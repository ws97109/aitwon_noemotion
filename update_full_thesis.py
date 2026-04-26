"""
全面更新論文＿李昇峰_v8.docx，修正事實錯誤、擴充各節學術內容。

修改項目：
A. 事實錯誤修正
   A1. para 137: MGT 情緒評分系統 → SACF 情感評分模型
   A2. para 383: 移除不精確的「SoTA 地位」宣稱，改為精確描述
   A3. para 131: 修正不當引用，改為正確的 MGT 引用

B. 各文獻探討節擴充
   B1. 2.1 智能體理論基礎：新增 LLM-based 生成式代理段落
   B2. 2.3 多智能體系統與社會互動：新增情感感知 MAS 應用段落
   B3. 2.4 大型語言模型在智能體中的應用：補充指令微調與記憶機制段落
   B4. 2.5 多智能體系統中的人格化與情感表達：修正 para 131 引用，補充連結

C. 參考文獻補充
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ════════════════════════════════════════════════════════════════
# 新段落內容
# ════════════════════════════════════════════════════════════════

# B1: 2.1 新增第三段
PARA_21_C = (
    '生成式 AI 智能體（Generative Agent）的概念隨著 LLM 能力的躍升而進入新的發展階段。'
    'Park 等人（2023）在 Generative Agents 框架中首次系統性地驗證：LLM 能夠驅動具備'
    '記憶流（Memory Stream）、反思（Reflection）與規劃（Planning）能力的擬人化虛擬居民，'
    '在開放式虛擬社區中自發湧現出晨間例行作息、社交關係形成乃至資訊八卦傳播等真實人類行為模式。'
    '該框架的核心貢獻在於將自然語言作為智能體認知狀態的統一表示媒介，'
    '使 LLM 的推理能力得以直接轉化為智能體的行為規劃能力，'
    '為大規模虛擬社會模擬開闢了嶄新的研究路徑。'
    '然而，Generative Agents 框架在情感建模方面仍採用純文字推理方式，'
    '未能對情感強度進行量化表示與系統性更新，這一不足促使本研究在其架構基礎上引入深度情感評分模組，'
    '以彌補情感感知能力的缺口。'
)

# B2: 2.3 新增第三段
PARA_23_C = (
    '情感感知能力的引入為多智能體模擬系統帶來了質的提升。'
    '在傳統 MAS 框架中，智能體的決策主要基於任務目標與環境狀態，'
    '缺乏對同伴情緒狀態的動態感知與回應機制。'
    '近年研究表明，將情感評分模型整合至智能體的感知-行動循環，'
    '能夠顯著提升虛擬互動的社交真實性：'
    '情緒高漲的智能體傾向於發起更多主動對話，情緒低落者則表現出社交迴避傾向，'
    '這與心理學研究中的情感-行為關聯理論高度一致（Frijda, 1986）。'
    '在問卷調查場景中，情緒狀態亦深刻影響智能體的回應風格與選項偏好，'
    '使基於情感記憶的虛擬受訪者能夠生成更具現實參照價值的調查數據。'
    '本研究在 Generative Agents 框架中系統性地整合了 SACF 情感評分層，'
    '使每位 AI 居民的情緒狀態得以量化更新，並動態影響其對話生成、互動決策與問卷回應，'
    '從而將情感感知能力從「定性描述」提升為「定量驅動」。'
)

# B3: 2.4 新增第三段（指令微調與工具使用）
PARA_24_C = (
    '在 LLM 的智能體化應用中，記憶管理與長期一致性是核心工程挑戰。'
    '由於 LLM 的上下文窗口有限，超出窗口的歷史記憶將遭截斷，'
    '導致智能體在長期互動中出現人格漂移（Persona Drift）與記憶錯亂現象。'
    '研究者提出多種外部記憶機制加以緩解：'
    '基於向量相似度的語義記憶檢索（Semantic Retrieval）將歷史記憶編碼為稠密向量並以餘弦相似度進行高效索引，'
    '使 LLM 能夠根據當前語境選擇性地回憶相關記憶片段；'
    '時效性衰減機制則以指數函數模擬人類記憶的遺忘曲線，確保近期事件具有更高的檢索優先級。'
    '本研究採用 BGE-M3 多語言向量模型對記憶條目進行編碼，'
    '結合時效性（衰減因子 0.995）、重要性（LLM 評分）與相關性（向量相似度）的三維加權公式，'
    '實現了對 9 位 AI 居民長達多輪互動的記憶一致性管理。'
    '此外，Qwen2.5:32b 通過 Ollama 框架在本地 GPU 上部署，'
    '既保障了繁體中文情境下的語義流暢性，也確保了多智能體並發運行的工程可行性。'
)

# B4: 2.5 修正第二段（移除錯誤引用，改為正確引用）
PARA_25_B_NEW = (
    '在將深度情感模型輸出整合至 LLM 生成管線的研究方面，'
    'Mai 等人（2025）的 MGT 框架展示了以語言模型為主幹、'
    '以多模態非語言特徵為輔助的注入架構，'
    '為情感感知語言生成提供了重要的技術參考。'
    '本研究採用類似的情感分數注入思路：'
    '以 SACF 模型產生的情感回歸分數（−3 至 +3）驅動智能體的情感狀態更新，'
    '再以情感標記（如「[情緒：快樂，強度：7]」）嵌入 qwen2.5:32b 的系統提示，'
    '引導後者生成與當前情感強度相符的對話內容。'
    '這種「情感評分→狀態標記→生成引導」的串接架構，'
    '相較於端對端的情感語言模型訓練，具有更高的模組化靈活性與可解釋性，'
    '同時允許情感模型（SACF）與語言生成模型（Qwen）分別在各自的領域資料上進行最優化訓練，'
    '避免了聯合訓練帶來的目標衝突與資料稀缺問題。'
    '這一設計選擇直接啟發自近期多模態情感注入研究的核心洞察：'
    '語言主幹的推理能力應受到保護，非語言情感訊號應以精確、可控的方式補充而非覆蓋語言語義。'
)

# A2: 結果討論第一段替換（修正 SoTA 宣稱）
PARA_DISCUSSION_NEW = (
    '本研究的實驗結果顯示，SACF 模型在 CMU-MOSI 多模態情感分析任務上取得全面競爭力，'
    '在 MAE 與 Corr 兩項指標上達到所有對比基線中的最優水準：'
    'MAE=0.587，相較最接近基線 ITHP（0.663）改善 11.5%；'
    'Corr=0.869，超越 MGT（0.764）達 +0.105，亦優於 ITHP（0.856）+0.013，'
    '顯示情感感知跨模態融合對情感強度的連續迴歸建模具有顯著優勢。'
    '在七分類準確率（Acc-7）方面，SACF 集成模型達到 52.19%，'
    '相較本研究比較表所引用的同一評估協議下的各傳統融合基線（最強為 MGT 的 50.44%）'
    '提升 +1.75 個百分點；'
    '需要指出的是，Mai 等人（2025）在其論文的特定實驗設定下另有報告 54.30% 的 Acc-7 成績，'
    '這一差異源於訓練資料分割策略（本研究採用 TrainVal 合併訓練）'
    '及評估協議的不同，不構成直接可比較關係。'
    '整體而言，SACF 在 MAE、Corr、Acc-2、F1 四項指標上的綜合表現優於同一評估框架內的所有基線，'
    '並在跨資料集的泛化性實驗中進一步驗證了模型特徵表示的廣泛適用性。'
)


def replace_run_text(para, old, new):
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new, 1)
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_full
    else:
        # XML fallback
        for t in para._p.findall('.//{%s}t' % W):
            if t.text and old in t.text:
                t.text = t.text.replace(old, new, 1)
                break
    return True


def set_para_text(para, new_text):
    """清空段落所有 run，填入新文字"""
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        for t in para._p.findall('.//{%s}t' % W):
            t.text = ''
        # 建立新 run
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = new_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        para._p.append(r)


def add_para_after(doc, anchor_para, text, style='Normal'):
    """在 anchor_para 之後插入新段落"""
    p_new = doc.add_paragraph(text, style=style)
    anchor_para._p.addnext(p_new._p)
    return p_new


def insert_para_before(anchor_para, doc, text, style='Normal'):
    p_new = doc.add_paragraph(text, style=style)
    anchor_para._p.addprevious(p_new._p)
    return p_new


def get_section_last_para(doc, heading_text):
    """取得指定 Heading 2 節下的最後一個 Normal 段落（在下一個 Heading 之前）"""
    found = False
    last = None
    for p in doc.paragraphs:
        if p.style.name == 'Heading 2' and heading_text in p.text:
            found = True
            continue
        if found:
            if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
                break
            if p.text.strip():
                last = p
    return last


def main():
    doc = Document('論文＿李昇峰_v8.docx')
    paras = doc.paragraphs

    # ════════════════════════════════════════════════════════════
    # A1. 修正 para 137 "MGT 情緒評分系統" → "SACF 情感評分模型"
    # ════════════════════════════════════════════════════════════
    for p in paras:
        if 'MGT 情緒評分系統' in p.text:
            ok = replace_run_text(p, 'MGT 情緒評分系統', 'SACF 情感評分模型')
            print(f'A1 [{"OK" if ok else "FAIL"}] MGT→SACF: {p.text[:80]}')

    # ════════════════════════════════════════════════════════════
    # A2. 修正結果討論第一段（移除不精確的 SoTA 宣稱）
    # ════════════════════════════════════════════════════════════
    for p in doc.paragraphs:
        if '確立了 SACF 在 CMU-MOSI 七分類情感分析任務上的 SoTA 地位' in p.text:
            set_para_text(p, PARA_DISCUSSION_NEW)
            print(f'A2 [OK] 結果討論第一段已更新')
            break

    # ════════════════════════════════════════════════════════════
    # A3. 修正 2.5 節 para 131：移除錯誤引用，改正為 MGT 引用
    # ════════════════════════════════════════════════════════════
    for p in doc.paragraphs:
        if ('Zadeh et al., 2017；Wang et al., 2021' in p.text and
                '多模態情感注入方法' in p.text):
            set_para_text(p, PARA_25_B_NEW)
            print(f'A3 [OK] 2.5節第二段引用已修正')
            break

    # ════════════════════════════════════════════════════════════
    # B1. 2.1 節新增第三段（LLM-based Generative Agents）
    # ════════════════════════════════════════════════════════════
    last_21 = get_section_last_para(doc, '智能體理論基礎')
    if last_21:
        next_heading = None
        found = False
        for p in doc.paragraphs:
            if p is last_21:
                found = True
                continue
            if found and p.style.name in ('Heading 1', 'Heading 2'):
                next_heading = p
                break
        if next_heading:
            p_new = doc.add_paragraph(PARA_21_C, style='Normal')
            next_heading._p.addprevious(p_new._p)
            print('B1 [OK] 2.1 節新增第三段（生成式代理）')
    else:
        print('B1 [FAIL] 找不到 2.1 節末段')

    # ════════════════════════════════════════════════════════════
    # B2. 2.3 節新增第三段（情感感知 MAS）
    # ════════════════════════════════════════════════════════════
    last_23 = get_section_last_para(doc, '多智能體系統與社會互動')
    if last_23:
        next_heading = None
        found = False
        for p in doc.paragraphs:
            if p is last_23:
                found = True
                continue
            if found and p.style.name in ('Heading 1', 'Heading 2'):
                next_heading = p
                break
        if next_heading:
            p_new = doc.add_paragraph(PARA_23_C, style='Normal')
            next_heading._p.addprevious(p_new._p)
            print('B2 [OK] 2.3 節新增第三段（情感感知 MAS）')
    else:
        print('B2 [FAIL] 找不到 2.3 節末段')

    # ════════════════════════════════════════════════════════════
    # B3. 2.4 節新增第三段（記憶管理與本地部署）
    # ════════════════════════════════════════════════════════════
    last_24 = get_section_last_para(doc, '大型語言模型在智能體中的應用')
    if last_24:
        next_heading = None
        found = False
        for p in doc.paragraphs:
            if p is last_24:
                found = True
                continue
            if found and p.style.name in ('Heading 1', 'Heading 2'):
                next_heading = p
                break
        if next_heading:
            p_new = doc.add_paragraph(PARA_24_C, style='Normal')
            next_heading._p.addprevious(p_new._p)
            print('B3 [OK] 2.4 節新增第三段（記憶管理與本地部署）')
    else:
        print('B3 [FAIL] 找不到 2.4 節末段')

    # ════════════════════════════════════════════════════════════
    # C. 補充參考文獻
    # ════════════════════════════════════════════════════════════
    ref_heading = None
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and '參考文獻' in p.text:
            ref_heading = p
            break

    new_refs = [
        ('Park', 'Park, J. S., O\'Brien, J. C., Cai, C. J., Morris, M. R., Liang, P., & Bernstein, M. S. (2023). '
         'Generative agents: Interactive simulacra of human behavior. '
         'In Proceedings of the 36th Annual ACM Symposium on User Interface Software and Technology (UIST \'23). '
         'ACM. https://doi.org/10.1145/3586183.3606763'),
        ('Frijda', 'Frijda, N. H. (1986). The emotions. Cambridge University Press.'),
        ('Han2021', 'Han, W., Chen, H., & Poria, S. (2021). Improving multimodal fusion with hierarchical mutual '
         'information maximization for multimodal sentiment analysis. '
         'In Proceedings of the 2021 Conference on Empirical Methods in Natural Language Processing (pp. 9180–9190). '
         'Association for Computational Linguistics.'),
        ('Yang2023', 'Yang, X., Feng, S., Wang, D., Zhang, Y., & Poria, S. (2023). '
         'Few-shot multimodal sentiment analysis based on multimodal probabilistic fusion prompts. '
         'In Proceedings of the 31st ACM International Conference on Multimedia (pp. 4425–4434). ACM.'),
        ('Hazarika2020', 'Hazarika, D., Zimmermann, R., & Poria, S. (2020). '
         'MISA: Modality-invariant and -specific representations for multimodal sentiment analysis. '
         'In Proceedings of the 28th ACM International Conference on Multimedia (pp. 1122–1131). ACM.'),
        ('Ekman1992', 'Ekman, P. (1992). An argument for basic emotions. '
         'Cognition & Emotion, 6(3–4), 169–200. https://doi.org/10.1080/02699939208411068'),
        ('Picard1997', 'Picard, R. W. (1997). Affective computing. MIT Press.'),
        ('Russell1980', 'Russell, J. A. (1980). A circumplex model of affect. '
         'Journal of Personality and Social Psychology, 39(6), 1161–1178.'),
    ]

    # 找到已存在的引用，避免重複
    existing_refs = set()
    for p in doc.paragraphs:
        for key, _ in new_refs:
            if key in p.text[:20]:
                existing_refs.add(key)

    added = 0
    if ref_heading:
        # 找第一個 Normal 參考文獻段後面依序插入
        first_ref_para = None
        for p in doc.paragraphs:
            if p is ref_heading:
                found = True
                continue
            if found and p.style.name == 'Normal' and p.text.strip():
                first_ref_para = p
                break

        insert_anchor = first_ref_para if first_ref_para else ref_heading
        for key, ref_text in new_refs:
            if key not in existing_refs:
                p_new = doc.add_paragraph(ref_text, style='Normal')
                insert_anchor._p.addnext(p_new._p)
                print(f'C [OK] 新增參考文獻: {key}')
                added += 1

    if added == 0:
        print('C [INFO] 所有參考文獻已存在，無需新增')

    # ════════════════════════════════════════════════════════════
    # 儲存
    # ════════════════════════════════════════════════════════════
    doc.save('論文＿李昇峰_v8.docx')
    print('\n完成！覆蓋儲存至 論文＿李昇峰_v8.docx')

    # 驗證
    doc2 = Document('論文＿李昇峰_v8.docx')
    print('\n=== 驗證：各節段落數 ===')
    sections = {
        '智能體理論基礎': 0,
        '情感計算與情感表達技術': 0,
        '多智能體系統與社會互動': 0,
        '大型語言模型在智能體中的應用': 0,
        '多智能體系統中的人格化與情感表達': 0,
    }
    current = None
    for p in doc2.paragraphs:
        if p.style.name == 'Heading 2' and p.text.strip() in sections:
            current = p.text.strip()
        elif current:
            if p.style.name in ('Heading 1', 'Heading 2'):
                current = None
            elif p.style.name == 'Normal' and p.text.strip():
                sections[current] += 1
    for sec, count in sections.items():
        print(f'  [{count}段] {sec}')

    print()
    print('=== 驗證：關鍵修正 ===')
    for p in doc2.paragraphs:
        if 'SACF 情感評分模型' in p.text and 'Generative' in p.text:
            print(f'  A1 ✓ 研究架構: {p.text[:80]}')
        if 'MGT（Mai et al., 2025）在其論文的特定實驗設定' in p.text:
            print(f'  A2 ✓ 結果討論修正: {p.text[:100]}')
        if 'Mai 等人（2025）的 MGT 框架' in p.text:
            print(f'  A3 ✓ 2.5節引用修正: {p.text[:80]}')


if __name__ == '__main__':
    main()
