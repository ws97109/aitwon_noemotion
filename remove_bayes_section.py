"""
從論文＿李昇峰_v5.docx 中刪除貝葉斯先驗修正相關內容：
1. 刪除 H4「貝葉斯先驗分布修正」及其說明段落（para 221-224）
2. 刪除 圖八（先驗分布）與 圖九（α搜尋）的 Caption 段落及空行（234-239）
3. 刪除圖目錄中 圖八、圖九 條目（para 57, 58）
4. 更新 圖六 Caption 與圖目錄條目，移除「貝葉斯先驗修正」字樣
5. 更新 para 161（正文引言）移除「貝葉斯先驗修正」
6. 更新 para 228 移除 alpha=3.0 相關說明
7. 更新 para 354（第四章）移除「貝葉斯先驗修正（α*=3.0）」
8. 圖十→圖八、圖十一→圖九、圖十二→圖十（圖目錄 + Caption）
9. 儲存為論文＿李昇峰_v6.docx
"""

from docx import Document
from docx.oxml.ns import qn

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def replace_wt_text(para_elem, old_str, new_str):
    """在 XML w:t 節點逐一替換字串。"""
    for t in para_elem.findall('.//{%s}t' % W):
        if t.text and old_str in t.text:
            t.text = t.text.replace(old_str, new_str, 1)
            return True
    return False


def replace_run_text(para, old_str, new_str):
    """在段落所有 run 中替換字串（合併 run 後替換）。"""
    full = para.text
    if old_str not in full:
        return False
    new_full = full.replace(old_str, new_str, 1)
    if para.runs:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = new_full
    return True


def delete_para(para):
    """從文件中刪除整個段落 XML 節點。"""
    p = para._p
    p.getparent().remove(p)


def main():
    doc = Document('論文＿李昇峰_v5.docx')
    paras = doc.paragraphs

    # ── Step 1: 刪除 貝葉斯先驗分布修正 文字段落（221-224） ──────
    # 找到目標段落（依內容比對，避免索引漂移）
    to_delete_keywords = [
        '貝葉斯先驗分布修正',
        '訓練驗證合併集（TrainVal）與測試集之間存在類別先驗分布差異',
        '驗證集驅動的 Alpha 最優化選擇',
        '先驗修正的縮放係數 α 決定了修正強度',
    ]
    deleted_count = 0
    # 需要找到這些段落並刪除（先收集，再刪除）
    to_delete = []
    for p in doc.paragraphs:
        for kw in to_delete_keywords:
            if kw in p.text:
                to_delete.append(p)
                break
    for p in to_delete:
        print(f'  [刪除] {p.text[:60]}')
        delete_para(p)
    deleted_count += len(to_delete)
    print(f'Step 1: 刪除 {deleted_count} 個先驗修正說明段落\n')

    # ── Step 2: 刪除 圖八 和 圖九 Caption 及周圍空行 ────────────
    # 注意：刪除後索引會變動，重新取 paragraphs
    paras = doc.paragraphs
    fig8_keywords = ['圖八 各資料集劃分', '圖 3.2.8']
    fig9_keywords = ['圖九 先驗修正縮放係數', '圖 3.2.9']

    to_delete = []
    for p in paras:
        txt = p.text
        for kw in fig8_keywords + fig9_keywords:
            if kw in txt:
                to_delete.append(p)
                break

    for p in to_delete:
        print(f'  [刪除 Caption] {p.text[:60]}')
        delete_para(p)
    print(f'Step 2: 刪除 {len(to_delete)} 個圖Caption段落\n')

    # ── Step 3: 刪除圖目錄中 圖八、圖九 條目 ──────────────────
    paras = doc.paragraphs
    tof_to_delete = []
    for p in paras:
        txt = p.text
        if ('圖九　　先驗修正' in txt or '圖八　　各資料集' in txt):
            tof_to_delete.append(p)
    for p in tof_to_delete:
        print(f'  [刪除 TOF] {p.text[:60]}')
        delete_para(p)
    print(f'Step 3: 刪除 {len(tof_to_delete)} 個圖目錄條目\n')

    # ── Step 4: 更新 圖六 Caption 移除「貝葉斯先驗修正」 ─────────
    paras = doc.paragraphs
    for p in paras:
        if '圖六 零洩漏推斷增強流程' in p.text and '貝葉斯先驗修正' in p.text:
            ok = replace_run_text(p, '、貝葉斯先驗修正', '')
            if not ok:
                replace_wt_text(p._p, '、貝葉斯先驗修正', '')
            print(f'  [更新 圖六 Caption] {p.text[:80]}')

    # ── Step 5: 更新圖目錄 圖六 條目 ────────────────────────────
    for p in paras:
        if '圖六　　零洩漏推斷增強流程' in p.text and '貝葉斯先驗修正' in p.text:
            ok = replace_run_text(p, '、貝葉斯先驗修正', '')
            if not ok:
                replace_wt_text(p._p, '、貝葉斯先驗修正', '')
            print(f'  [更新 TOF 圖六] {p.text[:80]}')
    print()

    # ── Step 6: 更新正文引言 para 161 ──────────────────────────
    for p in paras:
        if '結合測試時間增強（TTA）、多種子集成與貝葉斯先驗修正' in p.text:
            ok = replace_run_text(p, '、多種子集成與貝葉斯先驗修正', '與多種子集成')
            if not ok:
                replace_wt_text(p._p, '、多種子集成與貝葉斯先驗修正', '與多種子集成')
            print(f'  [更新 引言 TTA 段落] {p.text[:100]}')
    print()

    # ── Step 7: 更新 para 228（alpha=3.0 提及） ────────────────
    for p in paras:
        if 'v59 的三項改進（EMD Loss、TTA、alpha=3.0）' in p.text:
            ok = replace_run_text(p, '（EMD Loss、TTA、alpha=3.0）', '（EMD Loss、TTA 與多種子集成）')
            if not ok:
                replace_wt_text(p._p, '（EMD Loss、TTA、alpha=3.0）', '（EMD Loss、TTA 與多種子集成）')
            print(f'  [更新 para 228] {p.text[:100]}')
    print()

    # ── Step 8: 更新 第四章 para 354 ──────────────────────────
    for p in paras:
        if '貝葉斯先驗修正（α*=3.0）進一步校正類別分布偏差，系統性提升 Acc-7' in p.text:
            ok = replace_run_text(
                p,
                '，貝葉斯先驗修正（α*=3.0）進一步校正類別分布偏差，系統性提升 Acc-7',
                '，有效提升預測穩定性與準確率'
            )
            if not ok:
                replace_wt_text(
                    p._p,
                    '，貝葉斯先驗修正（α*=3.0）進一步校正類別分布偏差，系統性提升 Acc-7',
                    '，有效提升預測穩定性與準確率'
                )
            print(f'  [更新 第四章 推斷段落] {p.text[:120]}')
    print()

    # ── Step 9: 圖十→圖八、圖十一→圖九、圖十二→圖十（Caption + TOF） ──
    renumber_map = [
        # (舊中文數字, 新中文數字, 舊編號參考, 新編號參考)
        ('圖十一', '圖九',  '圖十一',  '圖九'),
        ('圖十二', '圖十',  '圖十二',  '圖十'),
        ('圖十 ',  '圖八 ', '圖十　',  '圖八　'),
        ('圖十　', '圖八　','圖十　',  '圖八　'),
    ]
    paras = doc.paragraphs

    # 先處理 圖十一、圖十二（避免 圖十 替換影響），再處理 圖十→圖八
    rename_targets = [
        # (search_text, old_zh, new_zh)
        ('圖十一', '圖十一', '圖九'),
        ('圖十二', '圖十二', '圖十'),
    ]
    for p in paras:
        txt = p.text
        for search, old, new in rename_targets:
            if search in txt:
                ok = replace_run_text(p, old, new)
                if not ok:
                    replace_wt_text(p._p, old, new)
                print(f'  [重編號 {old}→{new}] {p.text[:60]}')
                break

    paras = doc.paragraphs
    for p in paras:
        txt = p.text
        # 只替換圖十（後面不是一、二…），避免重複替換圖十一圖十二
        if '圖十 ' in txt or '圖十　' in txt or txt.strip().startswith('圖十 ') or txt.strip().startswith('圖十　'):
            if '圖十一' not in txt and '圖十二' not in txt:
                ok = replace_run_text(p, '圖十　', '圖八　')
                if not ok:
                    ok = replace_run_text(p, '圖十 ', '圖八 ')
                if not ok:
                    replace_wt_text(p._p, '圖十', '圖八')
                print(f'  [重編號 圖十→圖八] {p.text[:60]}')

    print('\nStep 9: 圖號重編完成\n')

    # ── 儲存 ────────────────────────────────────────────────────
    output = '論文＿李昇峰_v6.docx'
    doc.save(output)
    print(f'完成！儲存至 {output}')


if __name__ == '__main__':
    main()
