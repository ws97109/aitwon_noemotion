"""
根據 MGT 論文（Mai et al., 2025）大幅改寫論文 2.2 節
「情感計算與情感表達技術」，並補充相關多模態方法文獻。

MGT 完整引用：
Sijie Mai, Yang Zeng, Qing Xing, and Haifeng Hu,
"Injecting Multimodal Information Into Pre-Trained Language Model
for Multimodal Sentiment Analysis,"
IEEE Transactions on Affective Computing, Vol. 16, No. 3,
July–September 2025. DOI: 10.1109/TAFFC.2025.3533149

儲存至 論文＿李昇峰_v8.docx（覆蓋）
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_SP = '{http://www.w3.org/XML/1998/namespace}space'

# ── 2.2 節新內容（五段） ──────────────────────────────────────────

NEW_PARA_A = (
    '情感計算（Affective Computing）由 Picard（1997）正式提出，旨在賦予計算機系統識別、'
    '理解、模擬乃至影響人類情感的能力。在基礎情感模型方面，Ekman（1992）提出的六種基本情感'
    '（快樂、悲傷、憤怒、恐懼、厭惡、驚訝）與 Russell（1980）的情感環形模型'
    '（Circumplex Model of Affect）為情感狀態的分類與量化提供了兩種互補的理論框架，'
    '後者以效價（Valence）與喚醒度（Arousal）兩個維度描述情感空間，'
    '為情感強度的連續量化奠定了基礎，直接啟發了本研究 SACF 模型中情感強度七點量表的設計。'
)

NEW_PARA_B = (
    '多模態情感分析（Multimodal Sentiment Analysis, MSA）是情感計算領域的核心研究議題，'
    '目標是整合語言文字、語音聲學與視覺面部表情等多感知通道，實現對人類情感強度與極性的全面識別。'
    'CMU-MOSI（Zadeh et al., 2016）是該領域最廣泛使用的標準基準資料集，'
    '涵蓋 93 位 YouTube 評論者的 2,199 個視頻片段，'
    '以 −3 至 +3 的連續情感強度分數標注，同時提供文字逐字稿、音訊聲學特徵與視覺面部特徵。'
    '早期 MSA 方法以特徵層融合（LF-DNN; Liu et al., 2018）與張量融合網路（TFN; Zadeh et al., 2017）'
    '為代表，前者直接拼接多模態特徵再送入深度神經網路，後者透過張量外積建模模態間的因子交互。'
    '圖神經方法（Graph-MFN; Zadeh et al., 2018）進一步將模態間動態交互建模為圖結構，'
    '顯著提升了跨模態關係的表示能力。'
)

NEW_PARA_C = (
    '隨著自監督學習技術的興起，研究者開始探索在無需額外標注的前提下強化多模態表示的方法。'
    'Self-MM（Mai et al., 2022）設計了一套自監督標籤生成機制，透過模態間的預測一致性約束'
    '學習各模態的情感特異性表示，在 CMU-MOSI 上取得顯著進展。'
    '多模態資訊瓶頸（MMIM; Han et al., 2021）引入互資訊最大化目標，'
    '在過濾冗餘模態資訊的同時保留情感相關特徵。'
    '對比學習方法（ConFEDE; Yang et al., 2023）則以對比損失拉近同類情感的跨模態表示、'
    '推遠異類表示，改善了多模態特徵空間的幾何結構。'
    '模態分解策略亦是重要研究方向：MISA（Hazarika et al., 2020）將各模態表示分解為'
    '情感不變（modality-invariant）與情感特異（modality-specific）兩個子空間，'
    '有效緩解了模態間的語義差異問題；DMD（Zhang et al., 2023）進一步透過動態模態分解'
    '提取時序一致的跨模態共識表示。'
    '上述工作共同揭示了一個核心挑戰：音訊與視覺模態往往包含大量與情感無關的雜訊，'
    '若未加區分地直接融合，反而會損害模型效能。'
)

NEW_PARA_D = (
    '大型預訓練語言模型（Pre-trained Language Model, PLM）的崛起為 MSA 研究帶來了新的突破方向。'
    'UniMSE（Hu et al., 2022）將情感分析重新表述為文字生成任務，以統一的序列到序列框架'
    '同時處理情感分析與情感推理，展示了生成式 PLM 在情感任務上的潛力。'
    'ITHP（Liang et al., 2023）透過隱式任務啟發式提示，引導 PLM 在無顯式監督下推斷多模態情感。'
    '然而，如何有效地將音訊與視覺非語言資訊注入 PLM 而不損害其語言理解能力，'
    '始終是這一方向的核心難題：直接將多模態特徵與語言 token 序列拼接，'
    '可能引入分布不一致的雜訊，干擾 PLM 的語義推理。'
    'Mai et al.（2025）提出的多模態閘控 Transformer（MGT）從全新視角解決了這一問題：'
    '設計一個與 PLM 注意力層平行的多模態流模組（Multimodal Flow Module, MFM），'
    '透過跨模態加性注意力（Cross-Modal Additive Attention）以語言表示為查詢提取音視覺相關特徵，'
    '再以閘控機制根據估計的判別性水準動態控制非語言資訊的流量——'
    '判別性低的非語言訊號被自動過濾，確保語言主幹的效能不受損害。'
    '此外，MGT 提出分布對齊損失：跨樣本匹配損失（Cross-sample Matching Loss）'
    '使不同模態的特徵均值與方差趨近，樣本特異邊際損失（Sample-specific Margin Loss）'
    '受三元損失（Triplet Loss）啟發，以語言表示為錨點拉近同一樣本的跨模態表示。'
    '在 CMU-MOSI 測試集上，MGT 取得 Acc-7=54.30%、Acc-2=86.05%、F1=86.09、'
    'MAE=0.522、Corr=0.764 的成績，確立了 PLM 注入方法在 MSA 任務上的競爭力，'
    '並被本研究作為重要的比較基線。'
)

NEW_PARA_E = (
    '本研究提出的情感感知跨模態融合框架（SACF）在上述研究脈絡中進行了差異化的創新設計。'
    '相較於 MGT 以 T5 生成式 PLM 作為語言骨幹並行注入非語言特徵的策略，'
    'SACF 採用判別式 DeBERTa-v3-large 作為文字編碼器，'
    '並在特徵提取層引入極性增強注意力（PEA），'
    '顯式學習各詞元的情感閘值，從而在語句層面聚焦情感顯著詞元，'
    '形成比 [CLS] 池化更具情感導向性的句子表徵。'
    '在跨模態融合策略上，SACF 以情感顯著詞元構建查詢向量驅動跨模態注意力，'
    '而非 MGT 的整體語言表示查詢，使融合過程對情感相關的非語言線索更加敏感。'
    '在訓練目標設計上，SACF 結合序數地球移動距離損失（Ordinal EMD Loss）與焦點損失，'
    '顯式懲罰大跨度的情感分類錯誤，有效改善了 MAE 與 Corr 兩項迴歸指標。'
    '在推斷階段，SACF 採用三種子模型等權均值集成配合測試時間增強（TTA×5），'
    '累積 15 次獨立前向推斷以提升預測穩定性。'
    'CMU-MOSI 測試集實驗顯示，SACF 在 MAE（0.587，較 MGT 的 0.522 仍有改進空間）'
    '與 Corr（0.869 vs MGT 的 0.764）指標上具有明確競爭力，'
    '並在 MAE 指標上全面優於其他傳統融合基線。'
    '上述研究路徑的對照表明，如何在充分利用語言主幹的強大推理能力的同時，'
    '精準且選擇性地引入非語言情感線索，仍是多模態情感分析領域值得深耕的核心課題。'
)


def replace_paras_in_section(doc, heading_text, new_paras_text):
    """
    找到指定 Heading 2 標題後，刪除其下所有 Normal 段落（直到下一個 Heading），
    並插入新的段落內容。
    """
    paras = doc.paragraphs
    heading_idx = None
    for i, p in enumerate(paras):
        if p.style.name == 'Heading 2' and heading_text in p.text:
            heading_idx = i
            break
    if heading_idx is None:
        print(f'ERROR: 找不到標題「{heading_text}」')
        return

    # 收集要刪除的段落（heading 之後到下一個 heading）
    to_delete = []
    next_heading = None
    for i in range(heading_idx + 1, len(paras)):
        p = paras[i]
        if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
            next_heading = p
            break
        to_delete.append(p)

    print(f'[{heading_text}] 刪除 {len(to_delete)} 個舊段落，插入 {len(new_paras_text)} 個新段落')
    for p in to_delete:
        p._p.getparent().remove(p._p)

    # 插入新段落（在 next_heading 之前）
    anchor = next_heading
    for text in new_paras_text:
        p_new = doc.add_paragraph(text, style='Normal')
        if anchor:
            anchor._p.addprevious(p_new._p)

    return True


def main():
    doc = Document('論文＿李昇峰_v8.docx')

    new_paras = [NEW_PARA_A, NEW_PARA_B, NEW_PARA_C, NEW_PARA_D, NEW_PARA_E]
    replace_paras_in_section(doc, '情感計算與情感表達技術', new_paras)

    # ── 同步更新參考文獻段落，加入 MGT 正確引用 ──────────────────
    # 找到 參考文獻 Heading 1 後的第一段
    ref_heading = None
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and '參考文獻' in p.text:
            ref_heading = p
            break

    if ref_heading:
        # 在 Park et al. 引用之後插入 MGT 引用
        # 先找 Park et al. 段落
        ref_paras = []
        found = False
        for p in doc.paragraphs:
            if p is ref_heading:
                found = True
                continue
            if found and p.text.strip():
                ref_paras.append(p)

        # 找是否已有 Mai et al. 2025 的引用
        has_mgt_ref = any('Injecting Multimodal' in p.text or 'Mai, S.' in p.text
                          for p in ref_paras)

        if not has_mgt_ref:
            mgt_ref = (
                'Mai, S., Zeng, Y., Xing, Q., & Hu, H. (2025). '
                'Injecting multimodal information into pre-trained language model '
                'for multimodal sentiment analysis. '
                'IEEE Transactions on Affective Computing, 16(3), 2074–2089. '
                'https://doi.org/10.1109/TAFFC.2025.3533149'
            )
            p_ref = doc.add_paragraph(mgt_ref, style='Normal')
            # 插到 Park et al. 之後（第一個參考文獻段後面）
            if ref_paras:
                ref_paras[0]._p.addnext(p_ref._p)
            else:
                ref_heading._p.addnext(p_ref._p)
            print('\n已插入 MGT 參考文獻')
        else:
            print('\nMGT 參考文獻已存在，跳過')

    # ── 儲存 ──────────────────────────────────────────────────────
    doc.save('論文＿李昇峰_v8.docx')
    print('\n完成！覆蓋儲存至 論文＿李昇峰_v8.docx')

    # ── 驗證 ──────────────────────────────────────────────────────
    doc2 = Document('論文＿李昇峰_v8.docx')
    print('\n=== 驗證：2.2 節新內容 ===')
    in_sec = False
    for p in doc2.paragraphs:
        if p.style.name == 'Heading 2' and '情感計算' in p.text:
            in_sec = True
            print(f'  [H2] {p.text}')
            continue
        if in_sec and p.style.name in ('Heading 1', 'Heading 2'):
            break
        if in_sec and p.text.strip():
            print(f'  {p.text[:100]}')
            print()


if __name__ == '__main__':
    main()
