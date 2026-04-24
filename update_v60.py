"""
將論文＿李昇峰_v6.docx 更新至 v60 系統設定，儲存為 論文＿李昇峰_v7.docx

主要改動：
1. LLM: qwen3:14b → qwen2.5:32b
2. TTA×3 → TTA×5（含相關公式與計數）
3. 3×3=9 → 3×5=15 次獨立評估
4. SWA 步長 3→2（每 2 輪一次，共 10 快照）
5. 最終模型版本 v59 → v60，結果數值更新
6. 比較表格 SACF 集成列數值更新
7. 移除圖六 Caption 殘留的先驗修正描述
"""

from docx import Document
from docx.oxml.ns import qn

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ── 舊/新結果對照 ─────────────────────────────────────────────
OLD_ENS = {'acc7': '52.33', 'acc2': '86.44', 'f1': '86.43', 'mae': '0.591', 'corr': '0.868'}
NEW_ENS = {'acc7': '52.19', 'acc2': '86.30', 'f1': '86.28', 'mae': '0.587', 'corr': '0.869'}

# 新圖六 Caption 完整文字
NEW_FIG6_TEXT = (
    '圖六 零洩漏推斷增強流程：TTA×5、多種子集成\n'
    '圖 3.2.6 展示完整的零洩漏推斷流程，分為兩個主要層次。'
    '第一層（三個種子列）：三個獨立的 SWA 模型（種子分別為 42、123、2024）'
    '各自執行 5 次 MC Dropout 前向傳播（TTA），對 logits 取平均後得到各種子的推斷結果'
    'L₇^(s) ∈ R^(N×7)（N=686 個測試樣本）。'
    '第二層（集成）：三個種子的 logits 按均值融合為 L̄₇ = (1/3)·Σ L₇^(s)，'
    '最終以 argmax 解碼得到七分類預測。'
    '每個測試樣本共接受 3 個種子 × 5 次 TTA = 15 次獨立的隨機前向推斷，'
    '最大化推斷多樣性與預測穩定性。'
)

# 新圖八 Caption 完整文字
NEW_FIG8_TEXT = (
    '圖八  各種子結果、最終指標彙整與累積改進瀑布圖\n'
    '圖 3.2.10 左圖（最終集成結果）：v60 三種子 TTA×5 集成模型在 CMU-MOSI 測試集上'
    '的 Acc-7 = 52.19%（橙邊框標記最佳），超越 52% 目標線。'
    '中圖（最終指標橫條圖）：v60 集成模型的五項真實指標'
    '（Acc-7=52.19%、Acc-2=86.30%、F1=86.28%、MAE=0.5872、Corr=0.869），'
    '以不同顏色區分不同指標類型（綠=Acc-7，藍=Acc-2，青=F1，橙=1−MAE，紫=Corr×100）。'
    '右圖（累積改進瀑布圖）：從 v55 基線（47.96%）出發，'
    '追蹤每個關鍵改進（TrainVal 合併、EMD Loss、TTA×5、多種子集成）的累積貢獻，'
    '直至 v60 突破 52% 目標門檻，最終達到 52.19%。'
)


def replace_wt_all(para_elem, old_str, new_str):
    """在所有 w:t 中替換字串（只替換第一次出現）。"""
    for t in para_elem.findall('.//{%s}t' % W):
        if t.text and old_str in t.text:
            t.text = t.text.replace(old_str, new_str, 1)
            return True
    return False


def replace_run_text(para, old_str, new_str):
    """合併所有 run 後做字串替換，回填到第一個 run。"""
    full = para.text
    if old_str not in full:
        return False
    new_full = full.replace(old_str, new_str, 1)
    if para.runs:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = new_full
    else:
        replace_wt_all(para._p, old_str, new_str)
    return True


def set_wt_text(para_elem, new_text):
    """將段落第一個 w:t 設為 new_text，其餘清空。"""
    t_elems = para_elem.findall('.//{%s}t' % W)
    if not t_elems:
        return False
    t_elems[0].text = new_text
    if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
        t_elems[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    for t in t_elems[1:]:
        t.text = ''
    return True


def apply_text_replacements(doc, replacements):
    """對所有段落套用一系列 (old, new) 替換。"""
    for para in doc.paragraphs:
        for old, new in replacements:
            if old in para.text:
                ok = replace_run_text(para, old, new)
                if not ok:
                    replace_wt_all(para._p, old, new)


def main():
    doc = Document('論文＿李昇峰_v6.docx')
    paras = doc.paragraphs

    # ── Step 1: LLM qwen3:14b → qwen2.5:32b ──────────────────────
    print('=== Step 1: LLM 更新 ===')
    llm_replacements = [
        ('qwen3:14b', 'qwen2.5:32b'),
        ('Qwen3:14b', 'Qwen2.5:32b'),
    ]
    changed = 0
    for para in doc.paragraphs:
        for old, new in llm_replacements:
            if old in para.text:
                replace_run_text(para, old, new)
                print(f'  [LLM] {para.text[:80]}')
                changed += 1
    print(f'  共更新 {changed} 處\n')

    # ── Step 2: TTA×3 → TTA×5，及相關計數 ──────────────────────
    print('=== Step 2: TTA 次數更新 ===')
    tta_replacements = [
        ('TTA×3',            'TTA×5'),
        ('T_TTA=3',          'T_TTA=5'),
        ('執行 3 次 MC Dropout', '執行 5 次 MC Dropout'),
        ('對 3 次結果取均值',  '對 5 次結果取均值'),
        ('(1/3) · Σ_t',      '(1/5) · Σ_t'),
        ('3 次 TTA = 9 次',   '5 次 TTA = 15 次'),
        ('× 3 次 TTA = 9 次', '× 5 次 TTA = 15 次'),
        ('TTA = 9 次',        'TTA = 15 次'),
        ('9 次獨立前向推斷',   '15 次獨立前向推斷'),
        ('9 次獨立評估',       '15 次獨立評估'),
        ('共 9 次獨立',        '共 15 次獨立'),
    ]
    apply_text_replacements(doc, tta_replacements)

    # 圖目錄中的 TTA×3
    for para in doc.paragraphs:
        if 'TTA×3' in para.text or 'TTA×5' in para.text:
            print(f'  [TTA] {para.text[:80]}')
    print()

    # ── Step 3: 更新圖六 Caption（含移除先驗修正殘留） ────────────
    print('=== Step 3: 圖六 Caption 更新 ===')
    for para in paras:
        if '圖六 零洩漏推斷增強流程' in para.text:
            set_wt_text(para._p, NEW_FIG6_TEXT)
            print(f'  [圖六] {para.text[:80]}')
    print()

    # ── Step 4: 更新圖八 Caption ──────────────────────────────────
    print('=== Step 4: 圖八 Caption 更新 ===')
    for para in paras:
        if '圖八' in para.text and '各種子結果' in para.text:
            set_wt_text(para._p, NEW_FIG8_TEXT)
            print(f'  [圖八] {para.text[:80]}')
    print()

    # ── Step 5: 版本 v59 → v60 及結果數值 ────────────────────────
    print('=== Step 5: 版本與結果數值更新 ===')
    version_replacements = [
        ('最終模型（v59）', '最終模型（v60）'),
        ('v55 至 v59 五個版本', 'v55 至 v60 六個版本'),
        ('v59 最終突破目標線達到 52.33%', 'v60 最終突破目標線達到 52.19%'),
        ('v59 的三項改進（EMD Loss、TTA 與多種子集成）', 'v60 的三項改進（EMD Loss、TTA×5 與更密集 SWA）'),
        # 主要結果數值（確保由大到小替換避免衝突）
        ('+0.33 個百分點',   '+0.19 個百分點'),
        # 精確匹配集成結果
        ('集成 SACF 達 52.33%', '集成 SACF 達 52.19%'),
        ('集成（52.33%）',   '集成（52.19%）'),
        ('52.33% 的絕對數值', '52.19% 的絕對數值'),
        # v60 MAE/Corr 精確替換
        ('集成 SACF 的 MAE 為 0.591', '集成 SACF 的 MAE 為 0.587'),
        ('集成 Corr（0.868）', '集成 Corr（0.869）'),
        ('集成 MAE（0.591）', '集成 MAE（0.587）'),
    ]
    apply_text_replacements(doc, version_replacements)

    # 印出更新後確認
    for para in doc.paragraphs:
        if 'v60' in para.text or '52.19' in para.text:
            print(f'  [v60] {para.text[:100]}')
    print()

    # ── Step 6: 超參數表格更新 ──────────────────────────────────
    print('=== Step 6: 超參數表格更新 ===')
    for ti, table in enumerate(doc.tables):
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                key = cells[0].text.strip()
                if key == 'SWA 步長':
                    old_val = cells[1].text.strip()
                    for para in cells[1].paragraphs:
                        replace_run_text(para, '每 3 輪一次', '每 2 輪一次')
                    for para in cells[2].paragraphs:
                        replace_run_text(para, '共 7 個快照（42–60 輪）', '共 10 個快照（42–60 輪）')
                    print(f'  Table {ti} SWA 步長: {cells[1].text.strip()}')
                elif key == 'TTA 次數 T_TTA':
                    for para in cells[1].paragraphs:
                        replace_run_text(para, '3', '5')
                        replace_wt_all(para._p, '3', '5')
                    print(f'  Table {ti} TTA 次數: {cells[1].text.strip()}')
    print()

    # ── Step 7: 比較表格 SACF 集成列更新 ──────────────────────────
    print('=== Step 7: 比較表格 SACF 集成列更新 ===')
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = row.cells
            row_text = ''.join(c.text for c in cells)
            if '三種子集成' in row_text or ('SACF' in row_text and '52.33' in row_text):
                # 更新各欄位
                updates = {
                    1: NEW_ENS['acc7'],   # Acc7
                    2: NEW_ENS['acc2'],   # Acc2
                    3: NEW_ENS['f1'],     # F1
                    4: NEW_ENS['mae'],    # MAE
                    5: NEW_ENS['corr'],   # Corr
                }
                for col_idx, new_val in updates.items():
                    if col_idx < len(cells):
                        for para in cells[col_idx].paragraphs:
                            if para.runs:
                                for run in para.runs:
                                    run.text = ''
                                para.runs[0].text = new_val
                            else:
                                replace_wt_all(para._p, para.text, new_val)
                print(f'  Table {ti} Row {ri}: {[c.text.strip() for c in cells]}')
    print()

    # ── Step 8: 集成效益分析段落更新 ──────────────────────────────
    print('=== Step 8: 集成效益段落（Acc-7/MAE/Corr）更新 ===')
    # 第四章提到 Acc-7 和 MAE
    extra_replacements = [
        # Acc-2 and F1 in chapter 4
        ('分別達 86.44% 與 86.43%',  '分別達 86.30% 與 86.28%'),
        ('MGT（86.30% / 86.28%）',   'MGT（86.30% / 86.28%）'),  # no change needed
        # Update Corr discussion
        ('集成 SACF 達 0.868',        '集成 SACF 達 0.869'),
        ('ITHP（0.856）+0.012',       'ITHP（0.856）+0.013'),
        ('優於 MGT（0.822）+0.046',   '優於 MGT（0.822）+0.047'),
    ]
    apply_text_replacements(doc, extra_replacements)
    for para in doc.paragraphs:
        if '86.30% 與 86.28%' in para.text and '達' in para.text:
            print(f'  [Acc2/F1] {para.text[:100]}')
    print()

    # ── Step 9: 種子集成段落（update_seed_section 內容）───────────
    print('=== Step 9: 種子集成說明更新（隨機種子段落） ===')
    seed_replacements = [
        ('3 個種子 × 3 次 TTA = 9 次獨立的隨機評估', '3 個種子 × 5 次 TTA = 15 次獨立的隨機評估'),
        ('最大化了推斷多樣性與預測穩定性。\n三個模型的 TTA 平均 logits 以等權均值進行融合：L̄₇ = (1/3)·Σ_s L₇^(s)。\n實驗數據顯示（見圖十），集成後 Acc-7 = 52.33%，顯著優於各單一種子結果\n（種子 42：51.90%，種子 123：51.46%，種子 2024：50.58%），\n集成相對最優單種子提升 +0.43 個百分點',
         '最大化了推斷多樣性與預測穩定性。\n三個模型的 TTA 平均 logits 以等權均值進行融合：L̄₇ = (1/3)·Σ_s L₇^(s)。\n實驗數據顯示（見圖八），v60 集成後 Acc-7 = 52.19%，超越 52% 目標門檻'),
    ]
    for para in doc.paragraphs:
        for old, new in seed_replacements:
            if old in para.text:
                replace_run_text(para, old, new)
                replace_wt_all(para._p, old, new)
                print(f'  [種子] {para.text[:80]}')
    print()

    # ── 儲存 ────────────────────────────────────────────────────
    output = '論文＿李昇峰_v7.docx'
    doc.save(output)
    print(f'完成！儲存至 {output}')


if __name__ == '__main__':
    main()
