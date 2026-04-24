"""
更新論文＿李昇峰_v8.docx 中的「模型優勢分析」與「結果討論」，
加入跨資料集泛化性與參數效率的新發現。
儲存至論文＿李昇峰_v8.docx（覆蓋）。
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def insert_para_before(target_para, new_para):
    target_para._p.addprevious(new_para._p)


def add_para(doc, text, style='Normal', bold=False):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def replace_run_text(para, old, new):
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new, 1)
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_full
    return True


def main():
    doc = Document('論文＿李昇峰_v8.docx')
    paras = doc.paragraphs

    # ── 找到「模型優勢分析」Heading 的下一段落（para 365）
    # 在 （4）推斷增強策略 之後、「結果討論」之前，插入第五點
    anchor_discussion = None
    for p in paras:
        if p.style.name == 'Heading 2' and '結果討論' in p.text:
            anchor_discussion = p
            break

    if anchor_discussion is None:
        print('ERROR: 找不到「結果討論」段落')
        return

    # ── 在「結果討論」之前插入第（5）點（模型優勢分析的新內容）
    new_point5 = add_para(doc, (
        '（5）跨任務泛化能力：在 MMAFFBen 與 SemEval-2018 的系統性評估驗證，'
        'SACF 骨幹學習到的情感表示具備強烈的跨資料集泛化能力，'
        '以 0.435B 參數實現在兩個基準上均排名第三的效能（SemEval 整體 66.99、'
        'MMAFFBen 平均 55.50），在特定任務（英語情感強度、中文情感極性）上甚至超越規模大於其 16 倍的模型。'
        '此特性對需要跨語言、跨情境應用情感感知能力的多代理系統具有直接的實用意義。'
    ))
    insert_para_before(anchor_discussion, new_point5)

    # ── 更新「結果討論」的三個段落，加入新研究的引用
    # 找到結果討論的段落
    discussion_paras = []
    in_discussion = False
    for p in paras:
        if p.style.name == 'Heading 2' and '結果討論' in p.text:
            in_discussion = True
            continue
        if in_discussion:
            if p.style.name in ('Heading 1', 'Heading 2') and '結果討論' not in p.text:
                break
            if p.text.strip():
                discussion_paras.append(p)

    print(f'結果討論段落數：{len(discussion_paras)}')
    for i, p in enumerate(discussion_paras):
        print(f'  {i}: {p.text[:80]}')

    # ── 在「在計算效率方面」段落之後（即最後一個討論段之後），插入新的結語段
    if discussion_paras:
        last_discussion_para = discussion_paras[-1]
        # 找參考文獻 heading
        ref_heading = None
        for p in paras:
            if p.style.name == 'Heading 1' and '參考文獻' in p.text:
                ref_heading = p
                break

        if ref_heading:
            new_gen_para = add_para(doc, (
                '跨資料集泛化性實驗進一步完善了上述結論。'
                '在 MMAFFBen 多任務基準（共五個跨語言情感任務）與 SemEval-2018 三語言情感評估中，'
                'SACF-Text 均以遠低於競爭模型的參數量（0.435B）達到整體排名第三的成績，'
                '驗證了模型所習得的情感極性表示不僅限於 CMU-MOSI 域內，'
                '而是可廣泛遷移至中文情緒分類、多語言情感強度評估等多元場景。'
                '領域自適應預訓練消融實驗（MMAFFIn 預訓練）則揭示了一個重要的訓練動態：'
                '在主要 Acc-7 指標上統計不顯著的差異（p=0.391）背後，'
                '預訓練版本在種子穩定性（std 1.17→0.27，4.3× 提升）與極端情感識別（recall@+3 從 20% 至 30%）'
                '上呈現明確改善，為未來結合更長訓練週期或更低骨幹學習率的持續改進方案提供了實驗基礎。'
                '總體而言，本研究所提出的 SACF 模型在 CMU-MOSI 七分類任務上確立了當前最優效能，'
                '並透過多基準、多語言的系統驗證，展現了情感感知跨模態注意力設計的廣泛適用性與實用潛力。'
            ))
            insert_para_before(ref_heading, new_gen_para)

    # ── 同時更新 Table 2（版本對比表，仍顯示 v59）→ 更新為 v60
    print('\n=== 更新 Table 2（v59 → v60）===')
    for ti, table in enumerate(doc.tables):
        headers = [c.text.strip() for c in table.rows[0].cells]
        if 'v59 最終模型' in ' '.join(headers) or 'v59' in ' '.join(headers):
            print(f'  找到 Table {ti}，更新標題列')
            # 更新 header
            for cell in table.rows[0].cells:
                for para in cell.paragraphs:
                    if 'v59 最終模型（本研究）' in para.text:
                        replace_run_text(para, 'v59 最終模型（本研究）', 'v60 最終模型（本研究）')
            # 更新數值列
            updates = {
                'Acc-7 (%) ↑': ['47.96', '51.90', '52.19'],
                'Acc-2 (%) ↑': ['86.30', '86.88', '86.30'],
                'F1 (%) ↑':    ['86.26', '86.89', '86.28'],
                'MAE ↓':       ['0.5981','0.5659', '0.587'],
                'Corr ↑':      ['0.8646','0.8755', '0.869'],
            }
            for row in table.rows[1:]:
                cells = row.cells
                if len(cells) >= 4:
                    key = cells[0].text.strip()
                    if key in updates:
                        expected = updates[key]
                        for ci, val in enumerate(expected, 1):
                            if ci < len(cells):
                                for para in cells[ci].paragraphs:
                                    old_val = para.text.strip()
                                    if old_val == expected[ci-1] and ci == 3:
                                        # Only update last column (v59→v60)
                                        if para.runs:
                                            for r in para.runs:
                                                r.text = ''
                                            para.runs[0].text = val
                                        print(f'    {key}: {old_val} → {val}')

    # ── 儲存 ─────────────────────────────────────────────────
    doc.save('論文＿李昇峰_v8.docx')
    print('\n完成！覆蓋儲存至 論文＿李昇峰_v8.docx')


if __name__ == '__main__':
    main()
