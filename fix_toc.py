"""
重建論文＿李昇峰_v8.docx 的圖目錄與表目錄：
1. 刪除現有亂序的圖目錄條目
2. 按圖一→圖十正序重新插入圖目錄
3. 新增「表目錄」標題
4. 按表一→表五新增表目錄條目
儲存至 論文＿李昇峰_v8.docx（覆蓋）
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_SP = '{http://www.w3.org/XML/1998/namespace}space'

# ── 圖目錄條目（按順序） ──────────────────────────────────────────
FIGURE_ENTRIES = [
    '圖一\u3000\u3000SACF 整體架構圖',
    '圖二\u3000\u3000極性增強注意力（PEA）模組詳細示意圖',
    '圖三\u3000\u3000情感感知跨模態注意力（SACF）逐步計算示意圖',
    '圖四\u3000\u3000訓練策略全景：漸進解凍、EMA、SWA 與學習率排程',
    '圖五\u3000\u3000多工損失函數設計：組成結構、序數懲罰矩陣與 EMD 示意',
    '圖六\u3000\u3000零洩漏推斷增強流程：TTA×5、多種子集成',
    '圖七\u3000\u3000各版本模型性能演進對比（零洩漏條件下）',
    '圖八\u3000\u3000各種子結果、最終指標彙整與累積改進瀑布圖',
    '圖九\u3000\u3000記憶檢索關聯',
    '圖十\u3000\u3000人物介紹',
]

# ── 表目錄條目（按順序） ──────────────────────────────────────────
TABLE_ENTRIES = [
    '表一\u3000\u3000CMU-MOSI 多模態情感分析效能比較',
    '表二\u3000\u3000SACF-Text 在 MMAFFBen 五個純文字任務上的 ma-F1 比較',
    '表三\u3000\u3000SACF-Text 在 SemEval-2018 Task 1 三語言整體分數比較',
    '表四\u3000\u3000MMAFFIn 領域自適應預訓練消融：v60_baseline vs v60_MMAFFIn',
    '表五\u3000\u3000參數效率比較（效率分數 = 整體得分 ÷ 參數量（B））',
]


def make_tof_para(doc, text, style_name='table of figures'):
    """建立一個 table of figures 樣式的段落（只放簡潔標題）"""
    p_elem = OxmlElement('w:p')

    # 段落屬性
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name.replace(' ', '_'))
    pPr.append(pStyle)
    p_elem.append(pPr)

    # run
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(XML_SP, 'preserve')
    r.append(t)
    p_elem.append(r)

    # 包成 Paragraph 物件
    from docx.text.paragraph import Paragraph
    return Paragraph(p_elem, doc)


def make_title_para(doc, text):
    """建立 Title 樣式段落（用於「表目錄」標題）"""
    p_elem = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Title')
    pPr.append(pStyle)
    p_elem.append(pPr)

    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p_elem.append(r)

    from docx.text.paragraph import Paragraph
    return Paragraph(p_elem, doc)


def main():
    doc = Document('論文＿李昇峰_v8.docx')
    paras = doc.paragraphs

    # ── Step 1：收集並刪除所有 table of figures 段落 ─────────────
    tof_paras = [p for p in paras if p.style.name == 'table of figures']
    print(f'刪除 {len(tof_paras)} 個舊圖目錄條目...')
    for p in tof_paras:
        p._p.getparent().remove(p._p)

    # ── Step 2：找到「圖表目錄」Title 段落 ───────────────────────
    doc_paras = doc.paragraphs  # 重新取（刪除後索引變動）
    toc_title = None
    for p in doc_paras:
        if p.style.name == 'Title' and '圖表目錄' in p.text:
            toc_title = p
            break

    if toc_title is None:
        print('ERROR: 找不到「圖表目錄」Title 段落')
        return

    # ── Step 3：找到「導論」Heading 1 作為插入終點 ───────────────
    first_chapter = None
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and '導論' in p.text:
            first_chapter = p
            break

    if first_chapter is None:
        print('ERROR: 找不到「導論」段落')
        return

    # ── Step 4：在 first_chapter 之前，依序插入圖目錄條目 ─────────
    # 先建立「圖目錄」小標題
    fig_subtitle = make_title_para(doc, '圖目錄')
    first_chapter._p.addprevious(fig_subtitle._p)

    # 插入圖目錄條目：每次都插到 fig_subtitle 之後
    # 用 last_anchor 追蹤上一個插入的元素，讓新元素插到它後面
    last_anchor = fig_subtitle._p
    for entry_text in FIGURE_ENTRIES:
        p_new = make_tof_para(doc, entry_text)
        last_anchor.addnext(p_new._p)
        last_anchor = p_new._p

    print(f'已插入 {len(FIGURE_ENTRIES)} 個圖目錄條目')

    # ── Step 5：插入空行 ─────────────────────────────────────────
    empty = OxmlElement('w:p')
    first_chapter._p.addprevious(empty)

    # ── Step 6：插入「表目錄」小標題 ─────────────────────────────
    tbl_subtitle = make_title_para(doc, '表目錄')
    first_chapter._p.addprevious(tbl_subtitle._p)

    # ── Step 7：插入表目錄條目 ───────────────────────────────────
    last_anchor = tbl_subtitle._p
    for entry_text in TABLE_ENTRIES:
        p_new = make_tof_para(doc, entry_text, style_name='table of figures')
        last_anchor.addnext(p_new._p)
        last_anchor = p_new._p

    print(f'已插入 {len(TABLE_ENTRIES)} 個表目錄條目')

    # ── 儲存 ──────────────────────────────────────────────────────
    doc.save('論文＿李昇峰_v8.docx')
    print('\n完成！覆蓋儲存至 論文＿李昇峰_v8.docx')

    # ── 驗證 ──────────────────────────────────────────────────────
    doc2 = Document('論文＿李昇峰_v8.docx')
    print('\n=== 驗證：圖表目錄區段 ===')
    in_toc = False
    for p in doc2.paragraphs:
        if '圖表目錄' in p.text or '圖目錄' in p.text or '表目錄' in p.text:
            in_toc = True
        if p.style.name == 'Heading 1' and '導論' in p.text:
            break
        if in_toc and p.text.strip():
            print(f'  [{p.style.name:25s}] {p.text[:60]}')


if __name__ == '__main__':
    main()
