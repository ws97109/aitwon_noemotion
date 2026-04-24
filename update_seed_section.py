"""
更新論文＿李昇峰_v3.docx 中的隨機種子相關說明
1. 大幅擴充 3.2.5.2 多種子模型集成段落（para 220）
2. 更新超參數表格中「隨機種子」列的說明欄文字
"""

from docx import Document
from docx.oxml.ns import qn

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

# ─── 新版 3.2.5.2 段落文字 ─────────────────────────────────────────
NEW_SEED_TEXT = (
    '隨機種子（Random Seed）是初始化偽隨機數生成器（Pseudo-Random Number Generator, PRNG）'
    '狀態的整數值，決定了深度學習訓練中所有隨機性來源的起始狀態。深度神經網路的訓練過程包含四類主要隨機性來源：'
    '（1）模型參數的初始化方式（如 Xavier/Kaiming 初始化的隨機抽樣）；'
    '（2）DataLoader 的批次採樣與資料打亂順序；'
    '（3）Dropout 遮罩的每輪生成模式；'
    '（4）資料增強的隨機變換操作。'
    '不同的初始種子值會引導模型在訓練過程中探索損失超曲面（Loss Landscape）上的不同路徑，進而收斂至不同的局部最優解（Local Optima），'
    '產生具有互補性的決策邊界，這正是多種子集成能夠提升最終效能的根本原因。'
    '\n'
    '本研究設計了全面的 set_seed(seed) 函式，在每個種子訓練開始前同時固定六個層次的隨機狀態：'
    '① Python 內建 random 模組（random.seed(seed)）；'
    '② NumPy 隨機數生成器（np.random.seed(seed)）；'
    '③ PyTorch CPU 計算核心（torch.manual_seed(seed)）；'
    '④ CUDA GPU 計算（torch.cuda.manual_seed_all(seed)，涵蓋所有可用 GPU 裝置）；'
    '⑤ Python 雜湊隨機化（os.environ["PYTHONHASHSEED"] = str(seed)，防止字典與集合的不確定順序）；'
    '⑥ cuDNN 確定性模式（torch.backends.cudnn.deterministic = True，並停用 benchmark = False 以防止 cuDNN '
    '動態選擇不同卷積演算法而引入不確定性）。'
    '此六層固定策略確保在相同硬體環境下，給定相同種子值，每次訓練均能產生完全一致的模型參數、驗證指標與測試 logits，'
    '為研究結果的可重現性（Reproducibility）提供嚴格保障。'
    '\n'
    '本研究選用三個在量級上差異顯著的種子值：42、123 與 2024。'
    '種子 42 為機器學習社群長期沿用的傳統基準值（源自《銀河便車指南》的哲學典故），'
    '種子 123 為簡潔的三位序列值，種子 2024 則取自本研究的執行年份。'
    '三者跨越兩個數量級（二位數、三位數、四位數），確保了初始 PRNG 狀態的充分多樣性，'
    '使三個模型分別探索損失超曲面上的不同收斂路徑，捕捉互補的情感特徵表示。'
    '結合測試時間增強（TTA×3），每個測試樣本共接受 3 個種子 × 3 次 TTA = 9 次獨立的隨機評估，'
    '最大化了推斷多樣性與預測穩定性。'
    '三個模型的 TTA 平均 logits 以等權均值進行融合：L̄₇ = (1/3)·Σ_s L₇^(s)。'
    '實驗數據顯示（見圖十），集成後 Acc-7 = 52.33%，顯著優於各單一種子結果'
    '（種子 42：51.90%，種子 123：51.46%，種子 2024：50.58%），'
    '集成相對最優單種子提升 +0.43 個百分點，驗證了多種子集成在降低方差、提升統計穩健性方面的實際效益。'
)

# ─── 超參數表格隨機種子列的新說明文字 ─────────────────────────────
NEW_TABLE_DESC = (
    '三種子（42 / 123 / 2024）獨立訓練並集成；每種子設定六層 PRNG 固定（Python / NumPy / PyTorch CPU / CUDA / '
    'PYTHONHASHSEED / cuDNN 確定性模式），確保可重現性。'
)


def replace_para_text(para, new_text):
    """清除段落所有 run 文字，以第一個 run 存入新文字。"""
    if not para.runs:
        para.add_run(new_text)
        return
    # 保留第一個 run 的格式
    first = para.runs[0]
    bold = first.bold
    italic = first.italic
    # 清空所有 run
    for run in para.runs:
        run.text = ''
    para.runs[0].text = new_text
    para.runs[0].bold = bold
    para.runs[0].italic = italic


def main():
    doc = Document('論文＿李昇峰_v3.docx')

    # ── 1. 更新段落 220（3.2.5.2 多種子模型集成） ──────────────────
    print('=== 更新段落 220（3.2.5.2 多種子模型集成）===')
    para220 = doc.paragraphs[220]
    replace_para_text(para220, NEW_SEED_TEXT)
    print(f'  完成，字數：{len(para220.text)} 字')

    # ── 2. 更新超參數表格中「隨機種子」列的說明欄 ─────────────────
    print('\n=== 更新超參數表格隨機種子列說明 ===')
    updated = False
    for t_idx, table in enumerate(doc.tables):
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 3 and '隨機種子' in cells[0].text:
                # 更新第三欄（說明欄）
                desc_cell = cells[2]
                # 清空並重設文字
                for para in desc_cell.paragraphs:
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = NEW_TABLE_DESC
                    else:
                        para.add_run(NEW_TABLE_DESC)
                print(f'  Table {t_idx}，已更新「隨機種子」說明欄')
                print(f'  新說明：{NEW_TABLE_DESC[:60]}...')
                updated = True
    if not updated:
        print('  未找到隨機種子列，請確認表格結構')

    # ── 儲存 ────────────────────────────────────────────────────
    output = '論文＿李昇峰_v4.docx'
    doc.save(output)
    print(f'\n完成！儲存至 {output}')


if __name__ == '__main__':
    main()
