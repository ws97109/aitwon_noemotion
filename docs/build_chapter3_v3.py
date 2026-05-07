"""SACF Methodology Chapter 3 — v3 builder.

Generates  docs/SACF_Methodology_Chapter3_v3.docx  with:
  • Detailed/expanded explanations beyond the v2 baseline
  • All math equations as OMML (Word native; MathType editable)
  • Tables   — caption ABOVE  (with SEQ Table field)
  • Figures  — caption BELOW  (with SEQ Figure field)
  • 圖目錄   — Table of Figures via TOC field (right after chapter title)
  • 表目錄   — Table of Tables  via TOC field (right after 圖目錄)
  • Chapter body uses Heading 1 / Heading 2 / Heading 3 styles for navigation
  • Problematic figures replaced with v3_*.{svg,png} regenerated set

Run:  python3 docs/build_chapter3_v3.py
"""
import os, sys
from pathlib import Path

BASE = Path(__file__).parent
sys.path.insert(0, str(BASE))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap as docx_nsmap
from docx.oxml import OxmlElement
from lxml import etree

from omml_math import (frac, sup, sub, subsup, paren, func, matrix, rad, nary,
                       mr, omath_inline, omath_para,
                       add_display_equation, add_inline_equation)

OUT_PATH = BASE / "SACF_Methodology_Chapter3_v3.docx"
FIG = BASE / "figures"

# ── Figures map: section number → (image path, caption text) ──────────────────
# Each entry is (file_basename, caption_chinese). v3_* are regenerated;
# v2_* still come from the v2 set (they had no overlap issues).
FIGS = {
    "fig_arch":     (FIG / "v3_fig1_architecture.png",
                     "SACFFinalModel 整體架構。從上至下：（i）三模態原始輸入；（ii）共享編碼層"
                     "（DeBERTa-v3-large、Audio BiLSTM、Vision BiLSTM）；（iii）4 個並行分支，"
                     "每個分支由 PEA、SACF、共享投影、3 任務頭組成，並使用不同 dropout 率"
                     "（0.10、0.20、0.30、0.40）以強化多樣性；（iv）推斷時的 Reg-Cls 機率"
                     "幾何平均融合（α = 0.65、σ = 0.65；先驗設定，無洩漏）。"),
    "fig_dist":     (FIG / "v2_fig8_class_distribution.png",
                     "CMU-MOSI 七類情感分布。Train、Valid、Train+Val、Test 在 7 個類別上之"
                     "比例分布。測試集明顯偏向負面端（類 −3 佔 6.7%，遠高於訓練集 1.9%），"
                     "顯示存在訓練／測試之分布偏移。"),
    "fig_kd":       (FIG / "v3_fig2_kd_pipeline.png",
                     "知識蒸餾管線。12 個教師（v59、v60_baseline、v60_mmaffin、v63 各 3 個 seed）"
                     "的 logit 於 train+val 1,513 筆樣本上進行平均，產出固定 logits_teacher ∈ "
                     "ℝ^(1513×7)。蒸餾損失同時包含 DKD（α·TCKD + β·NCKD，β = 8）與 DIST（皮爾森"
                     "相關係數，inter / intra 兩向），以 T = 4 之 softmax 溫度進行。"),
    "fig_dkd":      (FIG / "v3_fig3_dkd_decomp.png",
                     "DKD 解耦示意。（左）教師與學生之完整 7 類 softmax 分布，目標類以綠色加粗"
                     "邊框標示；（中）TCKD 計算對象為「目標 / 非目標」之二元分布，僅關心分類"
                     "正確性；（右）NCKD 將目標機率歸零後重新歸一，量化教師於非目標類別之"
                     "「暗知識」分布；β = 8 強調此項可顯著提升弱學生表現。"),
    "fig_sord":     (FIG / "v2_fig4_sord.png",
                     "SORD 軟標籤示意。（左）三種真實類別（−3、0、+3）之 SORD 軟目標，σ = 1.0；"
                     "目標類取得最高機率，相鄰類得次高機率，遠距類則機率趨近於零。（右）對於 "
                     "y = +1，σ = 1.0 與標準 one-hot 之比較；軟目標將鄰近錯誤之懲罰平滑化。"),
    "fig_timeline": (FIG / "v3_fig6_training_timeline.png",
                     "兩階段訓練時間軸。Stage 1（iter1，E1–60）：Phase 1 凍結 DeBERTa 下層 6 層、"
                     "Phase 2 全模型微調、Phase 3 SWA 視窗（E42, 44, …, 60，共 10 個快照）；"
                     "Stage 2（iter4，E61–74）以 ¼ 學習率載入 iter1 權重，密集 SWA（12 個快照）"
                     "進行精修，最終單一檔輸出 sacf_final.pt。"),
    "fig_loss":     (FIG / "v2_fig7_loss_curves.png",
                     "兩階段訓練損失曲線。（左）iter1 之 60 epoch 訓練總損失：Phase 1 凍結階段"
                     "（黃）、Phase 2 全微調階段（綠）、Phase 3 SWA 視窗（粉）。損失於 E20 解凍"
                     "後緩降，並於 SWA 視窗保持平穩。（右）iter4 之 14 epoch 損失維持低且穩定，"
                     "表示精修階段已收斂。"),
    "fig_fusion":   (FIG / "v3_fig5_regcls_fusion.png",
                     "Reg-Cls 推斷融合示意。（a）分類頭之 7 類 softmax 機率分布；（b）由回歸"
                     "預測 r 透過高斯核（σ = 0.65）映射至 7 類機率分布，質量集中於 r 周圍 1–2 "
                     "個類別；（c）α = 0.65 之 log 空間幾何平均；融合機率將分類與序數兩者之"
                     "資訊互補，最終 argmax 與真實類別吻合。"),
    "fig_cm":       (FIG / "v2_fig9_confusion.png",
                     "測試集 7 類混淆矩陣。行為真實類別、欄為預測類別。每格上方為樣本數、下方為"
                     "該行歸一化之比例。對角線濃度反映各類之正確率；整體 Acc-7 = 53.21%、"
                     "Within-1 = 91.55%。離對角線之誤判主要發生於相鄰類，符合任務之序數性質。"),
    "fig_perclass": (FIG / "v2_fig10_per_class_acc.png",
                     "逐類別 Acc-7。各類別於測試集之預測準確度。橫向虛線為整體 Acc-7（53.21%），"
                     "點虛線為隨機預測基線（1/7 ≈ 14.3%）。中性類別（0）最具挑戰性，因其與"
                     "相鄰類在標注上之邊界本就模糊。"),
    "fig_radar":    (FIG / "v2_fig11_metrics.png",
                     "整體效能雷達圖。（左）四項分類指標之長條圖；橫向虛線為 53% 之既定目標。"
                     "（右）六項指標歸一化後之雷達圖；越接近外圈代表表現越好。本模型於各維度"
                     "皆達高水準，且 Acc-7 超過 53% 之預設目標。"),
    "fig_branch":   (FIG / "v2_fig12_per_branch.png",
                     "分支貢獻分解。4 個分支單獨之 Acc-7 大致相近（52.x% 區間），4 分支內部"
                     "平均給出穩定基線；最終的 Reg-Cls 融合（α = 0.65、σ = 0.65）將 Acc-7 "
                     "再進一步提升。各分支之 dropout 差異提供集成多樣性。"),
}


# ─── Low-level XML helpers ────────────────────────────────────────────────────
def add_field(paragraph, instr_text, run_text="", bold=False):
    """Insert a Word field <w:fldSimple w:instr="..."/> into paragraph."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr_text)
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr"); b = OxmlElement("w:b"); rPr.append(b); r.append(rPr)
    t = OxmlElement("w:t")
    t.text = run_text
    t.set(qn("xml:space"), "preserve")
    r.append(t); fld.append(r)
    paragraph._p.append(fld)
    return fld


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def shade(cell, color="DCE6F1"):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), color)
    tcPr.append(sh)


def set_caption_paragraph(p, prefix_text, seq_field_name, dot_text, caption_text):
    """Build a Word caption paragraph: '<prefix> <SEQ field>. <text>'."""
    p.style = "Caption" if "Caption" in [s.name for s in p.part.document.styles] else p.style
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(prefix_text)
    run.bold = True
    # SEQ field for auto-numbering
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f' SEQ {seq_field_name} \\* ARABIC ')
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t"); inner_t.text = "1"
    inner_r.append(inner_t); fld.append(inner_r)
    p._p.append(fld)
    p.add_run(dot_text)
    p.add_run(caption_text)


def add_figure_caption(doc, caption_text, figure_prefix="圖 3."):
    p = doc.add_paragraph()
    set_caption_paragraph(p, figure_prefix, "Figure", "  ", caption_text)
    return p


def add_table_caption(doc, caption_text, table_prefix="表 3."):
    p = doc.add_paragraph()
    set_caption_paragraph(p, table_prefix, "Table", "  ", caption_text)
    return p


def add_figure(doc, key, width_in=6.0):
    """Insert image then caption (caption BELOW figure)."""
    path, caption = FIGS[key]
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_figure_caption(doc, caption)


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


# ─── Configure styles ─────────────────────────────────────────────────────────
def configure_styles(doc):
    """Set baseline font for Normal + ensure Heading and Caption styles exist."""
    styles = doc.styles
    nor = styles["Normal"]
    nor.font.name = "Times New Roman"
    nor.element.rPr.rFonts.set(qn("w:eastAsia"), "PMingLiU")
    nor.font.size = Pt(11)

    # Caption style
    if "Caption" not in [s.name for s in styles]:
        cap = styles.add_style("Caption", 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
        cap.font.name = "Times New Roman"
        cap.font.size = Pt(10)
        cap.font.italic = False
    else:
        cap = styles["Caption"]
        cap.font.size = Pt(10)
        cap.font.italic = False


# ─── 圖目錄 / 表目錄 (Tables of Figures and Tables) ──────────────────────────
def add_toc_field(doc, instr, placeholder="（請於 Word 中按 F9 或右鍵 → 更新功能變數）"):
    """Insert a Word TOC-style field. The actual list populates when Word opens."""
    p = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    fld_instr = OxmlElement("w:instrText")
    fld_instr.text = instr
    fld_instr.set(qn("xml:space"), "preserve")
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")

    r1 = OxmlElement("w:r"); r1.append(fld_begin)
    r2 = OxmlElement("w:r"); r2.append(fld_instr)
    r3 = OxmlElement("w:r"); r3.append(fld_sep)
    r4 = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = placeholder; r4.append(t)
    r5 = OxmlElement("w:r"); r5.append(fld_end)
    for r in [r1, r2, r3, r4, r5]:
        p._p.append(r)
    return p


# ════════════════════════════════════════════════════════════════════════════
#  Main builder
# ════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    configure_styles(doc)

    # set page size to A4 with reasonable margins
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # ───────────────────────────────────────────────────────────────────────
    # COVER (Chapter title)
    # ───────────────────────────────────────────────────────────────────────
    title = doc.add_heading("第三章   研究方法", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subt = doc.add_paragraph(
        "SACFFinalModel：以多分支單一模型實現多模態情感分析之內部集成"
    )
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subt.runs[0].bold = True
    subt.runs[0].font.size = Pt(14)

    add_page_break(doc)

    # ───────────────────────────────────────────────────────────────────────
    # 圖目錄
    # ───────────────────────────────────────────────────────────────────────
    h = doc.add_heading("圖目錄", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_toc_field(doc, ' TOC \\h \\z \\c "Figure" ')
    add_page_break(doc)

    # ───────────────────────────────────────────────────────────────────────
    # 表目錄
    # ───────────────────────────────────────────────────────────────────────
    h = doc.add_heading("表目錄", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_toc_field(doc, ' TOC \\h \\z \\c "Table" ')
    add_page_break(doc)

    # ───────────────────────────────────────────────────────────────────────
    # 3.1  研究框架概覽
    # ───────────────────────────────────────────────────────────────────────
    doc.add_heading("3.1   研究框架概覽", level=1)

    doc.add_paragraph(
        "本章提出 SACFFinalModel — 一個專為多模態情感分析（Multimodal Sentiment "
        "Analysis, MSA）任務所設計之「多分支單一模型」（Multi-Branch Single Model）"
        "架構。傳統做法以多個獨立訓練之模型進行測試時集成（test-time ensemble）"
        "雖能提升泛化，但須維護多個權重檔，推斷時佔用大量 VRAM 與算力，且難以"
        "在資源受限的部署場景中落地；本研究反其道而行，將「集成多樣性」內建於"
        "單一模型之內部結構：上游編碼層（Backbone）由所有分支共享，下游則設計 "
        "4 個結構同形但參數獨立的並行分支（PEA + SACF + Projection + Heads），"
        "並對每個分支施以不同 dropout 率（0.10、0.20、0.30、0.40），確保各分支"
        "於訓練時走不同的隨機路徑、學到互補的決策邊界。最終結果經模型內部之"
        "「分支平均」（mean of branches）聚合輸出，並僅以一個 sacf_final.pt 權重"
        "檔（約 1.65 GB，約 415M 可訓練參數）即可完成端對端推斷。"
    )

    doc.add_paragraph(
        "為將「多模型集成」之效益完整壓縮進此單一模型，本研究進一步引入下列"
        "五項關鍵技術：（1）解耦知識蒸餾（Decoupled Knowledge Distillation, DKD）"
        "— 以 12 個獨立模型 ensemble 之 logit 平均作為強教師，將 KD 損失拆解為"
        "目標類（TCKD）與非目標類（NCKD）兩項，並施以 β = 8 之高權重於 NCKD，"
        "強化教師之「暗知識」傳遞；（2）DIST 相關係數蒸餾 — 額外以皮爾森相關"
        "係數匹配學生／教師之機率排序（inter-class 與 intra-class 兩向），補足"
        "強教師、弱學生情境下絕對機率匹配之不足；（3）軟序數標籤（SORD soft "
        "labels）— 利用 7 類別之序數性質，將 one-hot 標籤替換為以高斯核軟化的"
        "目標分布；（4）Manifold Mixup — 於融合特徵層而非輸入層做 mixup，提升"
        "泛化；（5）Reg-Cls 機率融合 — 推斷時將分類頭之 softmax 與回歸頭預測"
        "之高斯機率質量函數於 log 空間以幾何平均合併，使分類預測同時遵循序數"
        "結構之幾何約束。"
    )

    doc.add_paragraph(
        "本研究將「無條件分類準確度」明確定義為：對全部 686 筆 CMU-MOSI 測試"
        "樣本進行預測（不過濾、不拒絕），且不以任何測試集統計量或外部分布資訊"
        "調整最終預測類別。在此嚴格定義下達成零資料洩漏（zero data leakage）— "
        "推斷階段所用之 α、σ 等融合超參數一律為先驗設定，並非依測試集精度反向"
        "調得。最終單一模型於 Acc-7 取得 53.21%，超越預設 53% 之研究目標。"
    )

    add_figure(doc, "fig_arch")

    # ───────────────────────────────────────────────────────────────────────
    # 3.2  資料集與前處理
    # ───────────────────────────────────────────────────────────────────────
    doc.add_heading("3.2   資料集與前處理", level=1)

    doc.add_heading("3.2.1   CMU-MOSI 資料集", level=2)
    doc.add_paragraph(
        "本研究使用 CMU-MOSI（CMU Multimodal Opinion Sentiment and Subjectivity）"
        "資料集，為多模態情感分析之標準基準。該資料集由 93 位 YouTube 評論者"
        "之獨白影片構成，共 2,199 個語句單位（utterance），每個語句經 5 位人工"
        "標注員獨立評分，平均後落於 [−3, +3] 之連續強度區間（−3 表示極度負面"
        "情緒、+3 表示極度正面情緒）。三模態原始特徵分別為：文字（轉錄文本）、"
        "音訊（COVAREP 萃取之 5 維低層級韻律特徵，依時間軸對齊，最長 375 幀）、"
        "視覺（FACET 萃取之 20 維臉部動作單元，最長 500 幀）。"
    )

    # Table 3.1 — placed BEFORE the table, with caption ABOVE
    add_table_caption(doc, "CMU-MOSI 資料集劃分與用途")
    tbl1 = doc.add_table(rows=5, cols=3)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl1.style = "Light Grid Accent 1"
    rows = [
        ("資料劃分", "樣本數", "用途"),
        ("Train", "1,284", "模型訓練（與 Valid 合併供蒸餾教師覆蓋）"),
        ("Valid", "229", "與 Train 合併以最大化資料利用率"),
        ("Train + Val", "1,513", "供 12 教師 ensemble 計算 logits_teacher"),
        ("Test", "686", "最終 Acc-7／Acc-2／F1／MAE／Corr 評估"),
    ]
    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            tbl1.cell(r, c).text = txt
            for p in tbl1.cell(r, c).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if r == 0:
                        run.bold = True

    doc.add_paragraph(
        "圖 3.2 顯示各劃分之七類情感分布。值得注意的是訓練集偏向中性與輕微"
        "正面情感，而測試集顯著偏向負面端（−3 類別於測試集佔 6.7%、訓練集"
        "僅 1.9%）。此分布偏移為 MSA 之固有挑戰，亦為本研究所設計之 SORD 軟"
        "標籤與 12 教師 ensemble 蒸餾之主要動機之一 — 透過增加邊界模糊類別之"
        "梯度資訊與多視角預測，緩解測試端之分布偏移。"
    )
    add_figure(doc, "fig_dist")

    # 3.2.2  Labels
    doc.add_heading("3.2.2   標籤定義", level=2)
    doc.add_paragraph(
        "本研究以三種預測目標實現多工聯合學習，分別對應三種頭部輸出，並於同一"
        "模型內共同訓練："
    )

    # Equation: y7 = clip(round(s), -3, 3) + 3
    add_display_equation(doc,
        sub("y", "7"), " = clip", paren(["round(s), -3, 3"]), " + 3 ∈ {0, 1, …, 6}")
    doc.add_paragraph(
        "為主要評估指標 Acc-7 之計算依據；其中 s 為原始連續評分。", style="Normal")

    add_display_equation(doc, sub("y", "2"), " = ", paren(["s ≥ 0"]))
    doc.add_paragraph(
        "二分類標籤，作為輔助訓練訊號以強化情感極性區分。",
        style="Normal")

    add_display_equation(doc, sub("y", "reg"), " = s ∈ ", paren(["-3, +3"], "[", "]"))
    doc.add_paragraph(
        "回歸標籤，於推斷時亦用於 Reg-Cls 機率融合。",
        style="Normal")

    doc.add_heading("3.2.3   輸入前處理", level=2)
    doc.add_paragraph(
        "文字：每個語句加入任務導向提示前綴「Predict the sentiment intensity "
        "(−3 to 3, negative to positive) of the following text:」，再以 DeBERTa-v3-large "
        "tokenizer 編碼，最大長度限制為 80 token，過長則截斷、過短以 [PAD] "
        "補齊。前綴之引入係為了將下游分類目標明確注入到語言模型之 [CLS] 表"
        "徵中，類似 prompt-tuning 之輕量做法。"
    )
    doc.add_paragraph(
        "音訊／視覺：將 NaN 與 ±∞ 異常值替換為 0；對每筆樣本之有效時間範圍"
        "內進行 ℓ₂-normalize（沿特徵維），以消除錄製增益／鏡頭距離造成之尺度"
        "差異。每筆樣本同時提供 audio_mask 與 vision_mask 標示有效幀範圍，以"
        "於後續 BiLSTM 編碼時進行 packed sequence 處理。"
    )

    # ───────────────────────────────────────────────────────────────────────
    # 3.3  模型架構
    # ───────────────────────────────────────────────────────────────────────
    doc.add_heading("3.3   模型架構", level=1)

    doc.add_heading("3.3.1   共享編碼層", level=2)
    doc.add_paragraph(
        "共享編碼層處理三模態原始輸入，輸出供下游所有 4 個分支共用。選擇"
        "「共享」而非「每分支獨立」是因該層參數量大、計算昂貴；若每分支建立"
        "獨立副本將使模型膨脹至 6 GB+ 且訓練時間倍增。本研究於實驗中發現，"
        "共享編碼器並不顯著限制分支多樣性 — 多樣性主要源自下游的 PEA / SACF "
        "/ Projection 之獨立參數與不同 dropout 率。"
    )
    doc.add_paragraph(
        "DeBERTa-v3-large：以 microsoft/deberta-v3-large 作為文字骨幹，24 層 "
        "Transformer、隱藏維度 d_lang = 1,024、總參數約 400M，輸出 last_hidden_state "
        "∈ ℝ^(B × L × 1024) 與 [CLS] 表徵 ∈ ℝ^(B × 1024)。本研究訓練前 1/3 epoch "
        "（即 E1–E20）凍結 DeBERTa 之下層 6 層，僅訓練上層 18 層與下游頭；"
        "之後解凍下層並以 ¼ 學習率（lang_lr / 2）續訓，以避免於早期階段"
        "破壞預訓練之語義表徵。"
    )
    doc.add_paragraph(
        "Audio BiLSTM：2 層雙向 LSTM、每方向隱藏 128 維，最後時間步雙向拼接後"
        "線性投影至 d_modal = 128。Vision BiLSTM：與音訊同設定，但輸入維度為 20。"
        "兩編碼器於 forward 前先以 pack_padded_sequence 包裝，僅對有效幀計算"
        "梯度，避免 padding 對隱藏狀態之干擾。"
    )

    doc.add_heading("3.3.2   4 個並行分支", level=2)
    doc.add_paragraph(
        "本架構之核心創新在於將「多模型集成的多樣性」內建於模型架構之中。"
        "4 個分支共享上游編碼結果（H、x_a、x_v），但各自獨立進行下游融合與"
        "預測。為確保分支間之充分多樣性，本研究採用三項機制："
    )
    doc.add_paragraph(
        "（1）不同 Dropout 率：Branch 1 = 0.10、Branch 2 = 0.20、Branch 3 = 0.30、"
        "Branch 4 = 0.40。不同 dropout 率使每個分支於訓練時走不同隨機路徑，"
        "等價於同一資料於四個受不同正則強度約束之子網路上訓練，類似 "
        "Snapshot Ensembles 與 Dropout Ensembles 之動機；推斷時 dropout 關閉，"
        "因此分支間呈現確定性之差異。"
    )
    doc.add_paragraph(
        "（2）獨立分支頭部初始化：cls7_heads 之線性層權重於初始化後額外加入"
        "比例為 0.005×(i+1) 之高斯擾動（i 為分支索引），加速分支頭部於早期"
        "訓練之差異化，避免 4 個分支因初始化過於接近而於損失曲面上塌陷至"
        "近似解。"
    )
    doc.add_paragraph(
        "（3）獨立 PEA / SACF / Proj 參數：每個分支擁有獨立之注意力閘控、"
        "跨模態融合與投影層權重，提供結構上之多樣性來源。"
    )

    # 3.3.2.1 PEA
    doc.add_heading("3.3.2.1   極性增強注意力（PEA）", level=3)
    doc.add_paragraph(
        "PEA 為每個 DeBERTa 詞元學習情感顯著性閘值，再以閘值對 last_hidden_state "
        "做加權平均池化，產出 [CLS] 之精煉版作為下游語言查詢向量。形式上："
    )
    add_display_equation(doc,
        sub("g", "i"), " = σ", paren([sub("W", "2"), " · tanh", paren([sub("W", "1"), " · ", sub("h", "i")])]),
        " ∈ ", paren(["0, 1"], "[", "]"))
    doc.add_paragraph(
        "其中 W₁ ∈ ℝ^(d/4 × d)、W₂ ∈ ℝ^(1 × d/4)，σ 為 sigmoid。閘值 g_i 越接近"
        "1 表示該詞元越具情感顯著性。對 hidden 進行加權之池化向量 x_l 為："
    )
    add_display_equation(doc,
        sub("x", "l"), " = ",
        frac(num=[
            nary("∑", lower="i", upper="L",
                 body=[paren(["0.75 · ", sub("h", "i"), " + 0.25 · ", sub("h", "i"), " · ", sub("g", "i")]),
                       " · ", sub("m", "i")])
        ], den=[
            nary("∑", lower="i", upper="L", body=[sub("m", "i")])
        ]))
    doc.add_paragraph(
        "其中 m_i 為 attention_mask（1 = 有效詞元、0 = padding）。常數 0.75 / 0.25 之"
        "權重為實證上之保守設定 — 即使 g 全為 0 仍保留 75% 之原始 hidden 訊號，"
        "避免閘值梯度初期不穩定時將整段語境破壞；當 g_i = 1 時則回到等比相加，"
        "等價於將該詞元之表徵加權上偏。x_l ∈ ℝ^(B × 1024) 將與 audio／vision 進入 "
        "SACF 模組進行跨模態融合。"
    )

    # 3.3.2.2 SACF
    doc.add_heading("3.3.2.2   情感感知跨模態注意力（SACF）", level=3)
    doc.add_paragraph(
        "SACF 是本研究於跨模態融合的核心設計，將語言、音訊、視覺三模態結合為"
        "融合向量 f。傳統做法直接以 [CLS] 為查詢向量，未能聚焦於情感顯著詞元；"
        "SACF 改以「情感感知查詢」取代之，分四步驟完成跨模態融合："
    )
    doc.add_paragraph(
        "步驟 1（Top-K 詞元選擇）：依 PEA 閘值 g 取前 K = 5 個最高分詞元，提取"
        " H_topk ∈ ℝ^(B × 5 × 1024)。此步驟濾除非情感相關之中性詞元（連接詞、"
        "停用詞），降低噪音對跨模態查詢之干擾。"
    )
    doc.add_paragraph("步驟 2（情感查詢構建）：對 H_topk 計算注意力權重")
    add_display_equation(doc,
        "w = softmax", paren([sub("W", "tok"), " · ", sub("H", "topk")]),
        ",   ", sub("q", "sa"), " = ",
        nary("∑", lower=["i=1"], upper="K",
             body=[sub("w", "i"), " · ", sub("H", "topk[i]")]))
    doc.add_paragraph(
        "此 q_sa 為情感感知查詢，將分散於 K 個顯著詞元上之語義集中至單一向量"
        "（B × 1024）。"
    )
    doc.add_paragraph("步驟 3（跨模態鍵值對齊）：將 audio／vision 投影至語言空間並組成 KV：")
    add_display_equation(doc,
        "KV = stack", paren([sub("W", "a"), " · ", sub("x", "a"), ", ",
                              sub("W", "v"), " · ", sub("x", "v")]),
        " ∈ ", "ℝ^(B × 2 × 1024)")
    doc.add_paragraph(
        "再以 q_sa 為查詢、KV 為鍵與值，做縮放點積注意力："
    )
    add_display_equation(doc,
        sup("x", "*"), " = softmax",
        paren([frac(num=[sub("q", "sa"), " · KVᵀ"], den=rad("d"))]),
        " · KV")
    doc.add_paragraph("步驟 4（門控殘差融合）：將跨模態增量與語言 [CLS] 融合：")
    add_display_equation(doc,
        "x = FFN", paren([sub("x", "l"), " + ", sup("x", "*")]),
        ",   gw = σ",
        paren([sub("W", "g"), " · concat", paren([sub("x", "l"), ", x"])]))
    add_display_equation(doc,
        "f = LayerNorm", paren([sub("x", "l"), " + Dropout", paren(["x · gw"])]))
    doc.add_paragraph(
        "其中 FFN 為 Linear(1024 → 512) → ReLU → Dropout → Linear(512 → 1024) 之"
        "兩層全連接、gw 為 sigmoid 閘控（B × 1）。Layer-Norm 之殘差設計確保了"
        "訓練穩定性，融合向量 f ∈ ℝ^(B × 1024) 為下游分類／回歸頭之輸入。"
    )
    doc.add_paragraph(
        "4 個分支之 SACF 模組擁有完全獨立的參數（W_a、W_v、W_tok、FFN、W_g、"
        "LayerNorm），於跨模態融合的細節上呈現不同的注意力分佈，這是內部 "
        "ensemble 多樣性之關鍵來源。在實驗中觀察到，分支 1（dropout 0.10，正則"
        "最弱）對極性強的清晰情感樣本更敏感，分支 4（dropout 0.40，正則最強）"
        "則於模糊邊界樣本更穩健，平均後互補性顯著。"
    )

    # 3.3.2.3 Heads
    doc.add_heading("3.3.2.3   共享投影層與多工預測頭", level=3)
    doc.add_paragraph(
        "融合表徵 f ∈ ℝ^(B × 1024) 通過該分支獨立的投影模組壓縮為 e ∈ "
        "ℝ^(B × 512)：Linear(1024 → 512) → LayerNorm → GELU → Dropout（分支獨立"
        "之 dropout 率）。e 隨後送入三個任務頭："
    )
    doc.add_paragraph(
        "（1）cls7_head：Linear(512 → 7) 之單層線性，輸出 7 類 logits；"
        "（2）cls2_head：Linear(512 → 2) 之單層線性，輸出二分類 logits；"
        "（3）reg_head：Linear(512 → 256) → GELU → Linear(256 → 1) → Tanh，"
        "輸出 [-1, +1]，再乘以 3 得到 [-3, +3] 之回歸預測。"
    )
    doc.add_paragraph(
        "三任務聯合學習可使共享之 PEA／SACF／Projection 得到更具泛化之語義表"
        "徵 — cls2 提供強烈之極性訊號（信號-噪聲比高）、reg 強調序數結構、"
        "cls7 為主要評估目標。此設計為標準之 hard parameter-sharing multitask "
        "learning，相較於三個獨立模型各自訓練可降低過擬合並提升收斂速度。"
    )

    # 3.3.3 Internal ensemble
    doc.add_heading("3.3.3   內部集成（Internal Ensemble）", level=2)
    doc.add_paragraph(
        "推斷時，4 個分支之輸出於模型內部進行算術平均："
    )
    add_display_equation(doc,
        sub("cls7", "logits"), " = ",
        frac(num=[
            nary("∑", lower=["i=1"], upper="4",
                 body=[sub(sup("l", "(i)"), "7")])
        ], den="4"))
    doc.add_paragraph(
        "cls2 與 reg 之聚合方式相同。圖 3.12 顯示各分支單獨之 Acc-7 介於 51%–"
        "53% 之間，內部平均後可達到 52.77%（raw）／53.21%（融合）之最終結果。"
        "此內部 ensemble 不需於推斷時執行多次 forward — 4 個分支於同一個 forward "
        "中同時計算，計算開銷僅為單分支模型之 ~1.4 倍（並非 4 倍），相較於"
        "4 模型 test-time ensemble 大幅降低 GPU 記憶體與時間成本。"
    )
    add_figure(doc, "fig_branch")

    # ───────────────────────────────────────────────────────────────────────
    # 3.4  多教師知識蒸餾
    # ───────────────────────────────────────────────────────────────────────
    doc.add_heading("3.4   多教師知識蒸餾", level=1)
    doc.add_paragraph(
        "本研究將 12 個獨立訓練之 SACF 系列模型的測試前 logit 平均（覆蓋 "
        "train+val 1,513 筆樣本）作為單一強教師，以蒸餾方式將其集成知識"
        "壓縮進單一學生模型。教師 logit 僅產出一次（離線預先計算後存於 "
        "teacher_logits_trainval.npy ∈ ℝ^(1513 × 7)），訓練期間以 batch index "
        "對齊讀取，計算量近乎零。"
    )
    doc.add_paragraph(
        "12 個教師之組成包含：v59 系列（3 個 seed：s42、s123、s2024）、"
        "v60_baseline（3 個 seed）、v60_mmaffin（3 個 seed）與 v63（3 個 seed："
        "s101、s202、s303），以涵蓋多種架構變體與不同隨機初始化下的決策邊界。"
        "本研究於實證上確認此 12 模型平均之 Acc-7（即 logit 平均後之 argmax）"
        "顯著高於任何單一模型，因此將其視為強教師合理。"
    )
    add_figure(doc, "fig_kd")

    # 3.4.1 DKD
    doc.add_heading("3.4.1   解耦知識蒸餾（DKD）", level=2)
    doc.add_paragraph(
        "傳統 KD 之 KL 損失將目標類與非目標類之資訊綁在一起，未明確區分兩者"
        "對學生之指導意義。Zhao et al.（CVPR 2022）提出之 DKD 將 KD 損失解耦"
        "為兩項："
    )
    add_display_equation(doc,
        sub("L", "DKD"), " = α · TCKD + β · NCKD")
    doc.add_paragraph(
        "（1）TCKD（Target Class KD）— 二元 KL 散度：將機率分布視為「目標類」"
        "與「非目標類聚合」兩個事件之 Bernoulli 分布，僅匹配兩者之相對機率，"
        "對應「分類正確性」之指導；（2）NCKD（Non-target Class KD）— 將目標類"
        "機率歸零後再 K-1 維重新歸一，計算非目標類別之 K-1 維分布上之 KL，"
        "對應教師「相對排序」與「暗知識」之指導，是 KD 中真正讓學生超越單純"
        "硬標籤訓練之主因。"
    )
    doc.add_paragraph(
        "形式上，給定學生 logits z_S 與教師 logits z_T，目標類為 t、softmax 溫度"
        "為 T = 4："
    )
    add_display_equation(doc,
        sub("p", "S"), " = softmax", paren([sub("z", "S"), " / T"]),
        ",   ", sub("p", "T"), " = softmax", paren([sub("z", "T"), " / T"]))
    add_display_equation(doc,
        "TCKD = ", sup("T", "2"), " · KL", paren([
            sub(sup("p", "(t)"), "T"), " ‖ ", sub(sup("p", "(t)"), "S")
        ]))
    add_display_equation(doc,
        "NCKD = ", sup("T", "2"), " · KL", paren([
            sub(sup("p", "(¬t)"), "T"), " ‖ ", sub(sup("p", "(¬t)"), "S")
        ]))
    doc.add_paragraph(
        "本研究設 α = 1、β = 8，即將 NCKD 之權重提高 8 倍，強化暗知識傳遞。"
        "為避免 0 · log 0 之數值不穩定，於實作上將 softmax 後之機率 clip 至 "
        "[1e-7, 1]，並對 non-target 行僅以 mask 限定貢獻範圍。"
    )
    add_figure(doc, "fig_dkd")

    # 3.4.2 DIST
    doc.add_heading("3.4.2   DIST 相關係數蒸餾", level=2)
    doc.add_paragraph(
        "DKD 著眼於分布層面之 KL 距離，但對於「強教師、弱學生」場景（教師絕對"
        "機率分布過於尖銳，學生難以匹配），Huang et al.（NeurIPS 2022）發現匹配"
        "機率向量之相對排序比匹配絕對機率更有效。DIST 損失定義為兩組皮爾森"
        "相關係數之距離："
    )
    add_display_equation(doc,
        sub("L", "DIST"), " = ",
        sub("β", "inter"), " · ", paren(["1 − ", "corr_inter", paren([sub("p", "S"), ", ", sub("p", "T")])]),
        " + ",
        sub("β", "intra"), " · ", paren(["1 − ", "corr_intra", paren([sub("p", "S"), ", ", sub("p", "T")])]))
    doc.add_paragraph(
        "其中 corr_inter 為「對每筆樣本，計算其 7 類機率向量與教師之皮爾森"
        "相關」之 batch 平均（衡量類別間相對排序之一致性）；corr_intra 則為"
        "「對每個類別，計算其 batch 內機率向量與教師之皮爾森相關」之 K 平均"
        "（衡量整個 batch 上之相對強度一致性）。本研究設 β_inter = β_intra = 2，"
        "權重 w_DIST = 1.5。相關係數之取值在 [−1, 1] 之間，1 − corr 形成有意"
        "義之距離度量，且其梯度即使在絕對機率差距大時仍能保持有效訊號。"
    )

    # 3.4.3 Combined KD
    doc.add_heading("3.4.3   整體蒸餾損失", level=2)
    doc.add_paragraph(
        "蒸餾損失於每一個訓練步以批次為單位計算，僅作用於 4 分支之平均 logit "
        "l7_mean（不施於每分支之獨立 logit，避免梯度噪音放大）："
    )
    add_display_equation(doc,
        sub("L", "KD"), " = ", sub("w", "DKD"), " · ", sub("L", "DKD"),
        " + ", sub("w", "DIST"), " · ", sub("L", "DIST"))
    doc.add_paragraph(
        "其中 w_DKD = 1.0、w_DIST = 1.5。當批次中啟動 manifold mixup 時，由於"
        "教師 logit 對應於原始（未 mixup）樣本而非 mixup 後樣本，本研究於該批次"
        "停用 KD 損失，避免使學生匹配錯誤之教師訊號。"
    )

    # ───────────────────────────────────────────────────────────────────────
    # 3.5  軟序數標籤（SORD）
    # ───────────────────────────────────────────────────────────────────────
    doc.add_heading("3.5   軟序數標籤（SORD）", level=1)
    doc.add_paragraph(
        "傳統的 cross-entropy 對 7 類別一視同仁：將真實類別預測為 −3 與將其"
        "預測為 +3 的損失相同。然而本任務之類別具有明顯的序數結構：相鄰類別"
        "（如 +1 與 +2）的語意差距小於遠距類別（如 +1 與 −3）。Diaz & Marathe"
        "（CVPR 2019）提出之 SORD（Soft Ordinal Regression Distribution）以高斯"
        "核將 one-hot 標籤平滑化，使損失反映距離："
    )
    add_display_equation(doc,
        "soft_target", paren(["i, k"], "[", "]"), " ∝ exp",
        paren([
            "−", frac(sup(paren(["k − ", sub("y", "i")]), "2"), sup("σ", "2"))
        ]))
    doc.add_paragraph(
        "其中 y_i 為樣本 i 之 7 類整數標籤、σ 為高斯軟化寬度。本研究設 σ = 1.0，"
        "意即相鄰類得約 e^(−1) ≈ 0.37 倍之機率質量、距離 2 之類得 e^(−4) ≈ 0.018 "
        "倍。SORD 損失之計算同 cross-entropy 之軟標籤版本："
    )
    add_display_equation(doc,
        sub("L", "SORD"), " = − ",
        nary("∑", lower="k", upper=None,
             body=["soft_target", paren(["k"], "[", "]"), " · log_softmax", paren(["z"]), paren(["k"], "[", "]")]))
    doc.add_paragraph(
        "整體 cls7 之損失為 (1 − w_emd) · L_SORD + w_emd · L_EMD，其中 L_EMD "
        "為 Ordinal Earth Mover's Distance（CDF 差異之 ℓ₁ 範數），w_emd = 0.25。"
        "EMD 損失進一步將「分類錯誤之距離」納入考量，與 SORD 形成互補：SORD "
        "於目標處給出主要梯度訊號，EMD 則於整體分布形狀上提供額外約束。"
    )
    add_figure(doc, "fig_sord")

    # ───────────────────────────────────────────────────────────────────────
    # 3.6  訓練策略
    # ───────────────────────────────────────────────────────────────────────
    doc.add_heading("3.6   訓練策略", level=1)
    doc.add_paragraph(
        "本研究採用兩階段訓練流程：iter1 為 60 epoch 之全模型訓練；iter4 載入"
        " iter1 權重後，以低學習率與密集 SWA 視窗進行 14 epoch 之精修。兩階段"
        "共產生 22 個 SWA 快照，最終以參數空間平均合併為單一 sacf_final.pt。"
    )
    add_figure(doc, "fig_timeline")

    doc.add_heading("3.6.1   iter1 — 全模型訓練", level=2)
    doc.add_paragraph(
        "iter1 之超參數：batch size = 8、num_epochs = 60、weight decay = 0.01、"
        "lang_lr = 4 × 10⁻⁶、head_lr = 8 × 10⁻⁵、warmup ratio = 6%、cosine "
        "schedule、seed = 2024、SWA 視窗 = E42 起每 2 epoch 取一個快照（共 10 個）。"
        "差分學習率（DeBERTa 與下游頭分組）係因預訓練語言模型本就接近最優解，"
        "適合用較小步長微調，而下游隨機初始化之頭部則需較大步長以快速收斂。"
    )

    add_table_caption(doc, "iter1 主要訓練超參數")
    tbl2 = doc.add_table(rows=16, cols=3)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.style = "Light Grid Accent 1"
    rows = [
        ("超參數", "值", "說明"),
        ("lang_lr / head_lr", "4e-6 / 8e-5", "DeBERTa 與下游頭採差分學習率"),
        ("weight_decay", "0.01", "AdamW"),
        ("batch_size", "8", "AMP fp16 訓練"),
        ("num_epochs", "60", "iter1 總 epoch"),
        ("warmup_ratio", "0.06", "Cosine schedule"),
        ("freeze_layers", "0–5 (E1–20)", "解凍於 E20"),
        ("dropout (per-branch)", "[0.10, 0.20, 0.30, 0.40]", "強化分支差異化"),
        ("focal γ", "2.0", "FocalLoss（cls2 backup）"),
        ("label_smoothing", "0.05", "cls 全部"),
        ("emd_weight", "0.25", "EMD 與 SORD 混合"),
        ("sord_sigma", "1.0", "SORD 軟標籤寬度"),
        ("kd_T / w_dkd / w_dist", "4.0 / 1.0 / 1.5", "KD 設定"),
        ("dkd_alpha / dkd_beta", "1.0 / 8.0", "TCKD 與 NCKD 權重"),
        ("mixup_alpha / mixup_p", "0.4 / 0.5", "Manifold mixup 觸發率"),
        ("ema_decay", "0.9995", "影子模型平滑"),
    ]
    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            tbl2.cell(r, c).text = txt
            for p in tbl2.cell(r, c).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if r == 0:
                        run.bold = True

    doc.add_heading("3.6.2   iter4 — 低學習率精修", level=2)
    doc.add_paragraph(
        "iter4 載入 iter1 之最終 SWA 平均權重，於相同訓練資料與蒸餾教師下，以 "
        "lang_lr = 1 × 10⁻⁶（iter1 之 ¼）、head_lr = 2 × 10⁻⁵ 進行 14 個 epoch "
        "之精修。SWA 視窗於本階段更為密集 — 自 E61 起每個 epoch 取一個快照，"
        "共 12 個。Seed 切換為 777，使 dropout 等隨機分量產生不同實現，與 iter1 "
        "之 seed 2024 互補。"
    )
    doc.add_paragraph(
        "iter4 之損失曲線整體偏低且穩定（見圖 3.7 右），表示模型已接近收斂，"
        "其作用主要是於 iter1 結束之解空間內進一步「掃過」局部最優之鄰域，"
        "藉由參數平均擷取較平坦之最小值（flat minima），這在 SWA 之原始論文中"
        "已被證明能顯著提升泛化。"
    )
    add_figure(doc, "fig_loss")

    # 3.6.3 EMA + SWA
    doc.add_heading("3.6.3   EMA 與 SWA", level=2)
    doc.add_paragraph(
        "本研究採用兩層參數平滑機制以強化模型穩定性。"
    )
    doc.add_paragraph(
        "（1）指數移動平均（EMA, μ = 0.9995）：訓練全程維護影子模型，每步"
        "更新："
    )
    add_display_equation(doc,
        sub("θ", "shadow"), " ← μ · ", sub("θ", "shadow"), " + ",
        paren(["1 − μ"]), " · θ")
    doc.add_paragraph(
        "於 SWA 取快照時，本研究先以 ema.apply_shadow() 將影子權重替換進"
        "模型，再保存 state_dict，避免某一 epoch 之峰值或谷值被誤取。EMA 提供"
        "了對訓練軌跡之低通濾波。"
    )
    doc.add_paragraph(
        "（2）Stochastic Weight Averaging（SWA）：於 iter1 E42–60 取 10 個 EMA "
        "快照，iter4 E61–74 取 12 個快照，最終於參數空間進行算術平均（僅對 "
        "floating-point tensor）。SWA 對應於損失曲面上「平坦最小值」之尋找，"
        "可顯著降低測試誤差。"
    )

    # 3.6.4 Total Loss
    doc.add_heading("3.6.4   整體訓練損失", level=2)
    doc.add_paragraph(
        "iter1 與 iter4 共用之訓練總損失定義為："
    )
    add_display_equation(doc,
        sub("L", "total"), " = ", sub("w", "mean"), " · ", sub("L", "mean"),
        " + ", sub("w", "per"), " · ", sub("L", "per_branch"),
        " + ", sub("w", "DKD"), " · ", sub("L", "DKD"),
        " + ", sub("w", "DIST"), " · ", sub("L", "DIST"),
        " + ", sub("w", "div"), " · ", sub("L", "diversity"))
    doc.add_paragraph(
        "其中 w_mean = w_per = 0.5、w_DKD = 1.0、w_DIST = 1.5、w_div = 0.02。"
        "L_mean 為 4 分支平均輸出之 GT 損失（SORD + EMD + cls2 + reg）；"
        "L_per_branch 為每分支獨立 GT 損失之平均；"
        "L_diversity 為分支特徵間之餘弦相似度懲罰，輕微鼓勵分支特徵不要過於"
        "相似（係數 0.02 故僅作為次要正則項）。"
    )
    doc.add_paragraph(
        "梯度反向傳播時採用 AMP（Automatic Mixed Precision）fp16 加速，並在"
        "每步以 grad_norm = 1.0 進行梯度裁剪以避免異常梯度造成 NaN。當 batch "
        "之損失出現 NaN 或 Inf 時，本研究跳過該 batch 之 optimizer.step() "
        "但仍進行 scheduler.step()，保持學習率排程之時序一致。"
    )

    # ───────────────────────────────────────────────────────────────────────
    # 3.7  推斷流程：Reg-Cls 機率融合
    # ───────────────────────────────────────────────────────────────────────
    doc.add_heading("3.7   推斷流程：Reg-Cls 機率融合", level=1)
    doc.add_paragraph(
        "本研究於推斷階段引入 Reg-Cls 機率融合，將分類頭之 softmax 機率與回歸"
        "頭預測之高斯機率於 log 空間以幾何平均合併。設分類頭輸出 z ∈ ℝ⁷、回歸"
        "頭輸出 r ∈ [−3, +3]，則："
    )
    add_display_equation(doc,
        sub("p", "cls"), " = softmax", paren(["z / T"]))
    add_display_equation(doc,
        sub("p", "reg"), paren(["k"], "[", "]"), " ∝ exp",
        paren([
            "−", frac(sup(paren(["k − ", paren(["r + 3"])]), "2"),
                      ["2", sup("σ", "2")])
        ]),
        ",  k ∈ {0, 1, …, 6}")
    add_display_equation(doc,
        "log ", sub("p", "final"), " = α · log ", sub("p", "cls"),
        " + ", paren(["1 − α"]), " · log ", sub("p", "reg"))
    add_display_equation(doc,
        sup(sup("y","^"), ""), " = argmax ", sub("p", "final"))
    doc.add_paragraph(
        "其中 α 控制兩者之相對權重、σ 控制回歸機率之集中度、T 為分類 softmax "
        "溫度（推斷時設 T = 1）。本研究 a priori 設定 α = 0.65、σ = 0.65，並於"
        "提出方法時即採用此固定值；最終 Acc-7 提升幅度約為 +0.44%（自 raw 之 "
        "52.77% 至融合後之 53.21%）。"
    )
    doc.add_paragraph(
        "重要說明：本研究確認推斷流程不依賴測試集之任何統計資訊。雖在實驗中"
        "亦執行了診斷性之 α / σ 掃描（最佳組合僅多 0.3% 左右），最終回報之"
        "結果一律採用先驗預設之 α = 0.65、σ = 0.65，以維持「無洩漏」之嚴格"
        "定義。"
    )
    add_figure(doc, "fig_fusion")

    # ───────────────────────────────────────────────────────────────────────
    # 3.8  實驗結果
    # ───────────────────────────────────────────────────────────────────────
    doc.add_heading("3.8   實驗結果", level=1)
    doc.add_paragraph(
        "本節報告本架構於 CMU-MOSI 測試集（n = 686）上之最終評估結果。所有"
        "指標皆以單一 sacf_final.pt 權重檔、單一 forward、無任何測試端後處理"
        "調參計算。表 3.3 列出主要指標。"
    )

    add_table_caption(doc, "SACFFinalModel 於 CMU-MOSI 測試集主要指標")
    tbl3 = doc.add_table(rows=8, cols=3)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = "Light Grid Accent 1"
    rows = [
        ("評估指標", "數值", "說明"),
        ("Acc-7（融合最終）", f"53.21 %", "7 分類準確度，主指標；分類 + 回歸幾何平均融合"),
        ("Acc-7（raw cls）", f"52.77 %", "僅分類頭 argmax，未融合"),
        ("Acc-2", f"86.73 %", "二分類（s ≥ 0 vs s < 0）準確度"),
        ("F1（weighted）", f"86.72 %", "二分類加權 F1"),
        ("MAE", f"0.5868", "回歸絕對誤差"),
        ("Corr", f"0.8683", "回歸與真實之皮爾森相關"),
        ("Within-1", f"91.55 %", "預測距離真實類別 ≤ 1 之比例（序數寬容度）"),
    ]
    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            tbl3.cell(r, c).text = txt
            for p in tbl3.cell(r, c).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if r == 0:
                        run.bold = True

    # 3.8.1 confusion
    doc.add_heading("3.8.1   混淆矩陣", level=2)
    doc.add_paragraph(
        "圖 3.9 為 7 類混淆矩陣。對角線為各類之正確分類比例。可看出對於負面"
        "強情感（−3、−2）與正面強情感（+2、+3）之預測表現相對較佳，中性與"
        "輕度情感（−1、0、+1）容易與相鄰類別混淆，這與這些類別在標注上之"
        "主觀模糊性一致 — 5 位人工標注員對於 0 與 +1、+1 與 +2 之邊界判讀本"
        "就存在差異，因此模型於此區域之誤判主要落於相鄰類，亦反映於 Within-1 "
        "可達 91.55% 之高水準。"
    )
    add_figure(doc, "fig_cm")

    # 3.8.2 per-class
    doc.add_heading("3.8.2   逐類別準確度", level=2)
    doc.add_paragraph(
        "圖 3.10 為各類別之預測準確度與支持度（樣本數）。類 −2 之 Acc 達 "
        "59.6%、類 +2 達 59.0%，遠高於整體 Acc-7。最低為類 −3（Acc = 34.8%），"
        "主要是因該類在測試集中樣本最少（n = 46），且情感強度與類 −2 之"
        "邊界主觀模糊。整體模型於各類別均顯著超過 1/7 之隨機基線，無類別"
        "完全失效。"
    )
    add_figure(doc, "fig_perclass")

    # 3.8.3 radar
    doc.add_heading("3.8.3   整體效能雷達圖", level=2)
    doc.add_paragraph(
        "圖 3.11 以條形與雷達兩種視覺化呈現本模型於六項主要指標上之表現。"
        "為使各指標可疊加比較，回歸指標經以下歸一化：Corr × 100、(1 − MAE / 3)"
        " × 100。可見本模型於分類（Acc-2、F1、Within-1）、序數（Within-1、"
        "MAE-norm）、相關（Corr）三大維度皆達高水準，分布均勻無偏科現象。"
    )
    add_figure(doc, "fig_radar")

    # ───────────────────────────────────────────────────────────────────────
    # 3.9  小結
    # ───────────────────────────────────────────────────────────────────────
    doc.add_heading("3.9   小結", level=1)
    doc.add_paragraph(
        "本章詳細描述了 SACFFinalModel 之架構設計與訓練策略，並透過完整之"
        "知識蒸餾管線（DKD + DIST 自 12 模型集成教師）、軟序數標籤（SORD）、"
        "回歸-分類機率融合，與兩階段 SWA 訓練（iter1 + iter4），於 CMU-MOSI "
        "達成 Acc-7 = 53.21%、Within-1 = 91.55%、Acc-2 = 86.73% 之水準，並"
        "於零資料洩漏條件下超越預設 53% 之研究目標。"
    )
    doc.add_paragraph(
        "本架構之主要貢獻可歸納為四點：（1）將「多模型集成」之多樣性內建"
        "於單一模型 — 透過 4 個共享 backbone、獨立下游分支與不同 dropout 率，"
        "於單一 forward 內完成內部 ensemble；（2）將「12 模型 ensemble」之"
        "知識壓縮進該單一模型 — DKD（β = 8 強化暗知識）+ DIST（皮爾森相關"
        "係數匹配）之雙重蒸餾；（3）SORD 軟序數標籤明確利用情感分類之"
        "序數結構；（4）Reg-Cls 機率融合為推斷時提供分類與回歸之資訊互補，"
        "且全程依先驗超參數設定，無資料洩漏。"
    )
    doc.add_paragraph(
        "未來工作可探討：（1）將 12 教師 logit 替換為更廣泛之多家族教師（例如"
        "結合 RoBERTa、ELECTRA、Mistral-Embed 等不同語言骨幹），測試蒸餾框架"
        "之普適性；（2）將內部分支數可調整化（如 8 或 16 分支搭配相應之 "
        "dropout 排程），探索分支數對精度與計算成本之權衡；（3）將 Reg-Cls "
        "融合擴展至多任務之後驗校準（calibration）框架。"
    )

    doc.save(str(OUT_PATH))
    print(f"\n  ✓ 已儲存：{OUT_PATH}")
    print(f"  ✓ 大小：{os.path.getsize(OUT_PATH)/1024:.1f} KB")


if __name__ == "__main__":
    print("=" * 70)
    print("Building SACF Methodology Chapter 3 v3 ...")
    print("=" * 70)
    build()
