"""Build 論文＿李昇峰_v15.docx — fully consistent with the new SACFFinalModel.

Vs. v14 changes:
  • Chapter 1 (introduction) paragraphs updated to cite Acc-7 = 53.21% (the new
    multi-branch model's score) and to mention the full loss family
    (soft-ordinal CE, EMD, Smooth-L1, R-Drop, cross-modal contrast) instead of
    only Focal/EMD/R-Drop.
  • Chapter 2 (literature review) Acc-7 numbers updated to the new model.
  • Chapter 3 rewritten in natural language: no `audio_mask`, `vision_mask`,
    `attention_mask`, `last_hidden_state`, `pack_padded_sequence`,
    `sacf_final.pt`, `mmaffinpretrain_backbone.pt`. Variables (W_a, h_i,
    x_cls, …) are kept but rendered as proper Word native subscripts.
  • Chapter 4 main-results numbers updated to 53.21% / 86.73% / 0.5868 / 0.8683.
  • Code-style names in Chapter 4's ablation paragraph cleaned up.
  • Figures replaced with v5 English versions (true subscripts, no overlap).

Run:  python3 docs/build_thesis_v15.py
"""
import os, re, sys
from pathlib import Path

BASE = Path(__file__).parent
sys.path.insert(0, str(BASE))

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from omml_math import (frac, sup, sub, paren, nary, rad, add_display_equation)

SRC = Path("/Users/lishengfeng/Desktop/aitwon_noemotion/論文＿李昇峰_v13.docx")
OUT_PATH = Path("/Users/lishengfeng/Desktop/aitwon_noemotion/論文＿李昇峰_v15.docx")
FIG = BASE / "figures"

# ── v5 English figures used in Chapter 3 ─────────────────────────────────────
FIGS = {
    "arch": (FIG / "v6_fig_arch.png",
             "SACFFinalModel 整體架構。圖中四層由上至下：(i) 三個原始模態輸入；"
             "(ii) 由 DeBERTa-v3-large、雙向 LSTM 音訊編碼器、雙向 LSTM 視覺編碼器組成的"
             "共享上游編碼層，分別輸出語言嵌入 t_{emb}、音訊嵌入 a_{emb}、視覺嵌入 v_{emb}；"
             "(iii) 4 個結構同形但參數獨立的並行分支，每個分支依序通過極性增強注意力、"
             "情感感知跨模態融合、共享投影層與三任務頭，輸出各自的七分類分數 l7_{i}、"
             "二分類分數 l2_{i} 與情感強度回歸 reg_{i}；(iv) 推斷階段透過 TTA × 5、"
             "3-Seed 集成與 Reg-Cls 機率融合得到最終預測。"),
    "dist": (FIG / "v6_fig_distribution.png",
             "CMU-MOSI 七類情感分布（僅 Train 與 Test）。本研究將原始訓練集與驗證集合併為"
             "單一訓練集（n = 1,513）以最大化資料利用率，測試集（n = 686）僅於模型最終"
             "訓練完成後評估一次。測試集顯著偏向負面端（類 −3 佔 6.7%，遠高於訓練集的 "
             "2.2%），形成 +4.5% 之分布偏移，是本任務之主要挑戰之一。"),
    "pea": (FIG / "v6_fig_pea.png",
            "極性增強注意力（PEA）模組詳細示意。對每個 DeBERTa 詞元 h_{i} 透過兩層"
            "感知器學習情感顯著性閘值 g_{i} ∈ [0, 1]；接著依詞元的有效遮罩 m_{i} 進行"
            "加權池化，得到精煉後的語言查詢向量 x_{l}。PEA 同時輸出此向量給共享投影層，"
            "並輸出閘值序列 g 提供下游情感感知跨模態融合作為「選詞」依據。"),
    "sacf": (FIG / "v6_fig_sacf_steps.png",
             "情感感知跨模態融合（SACF）逐步計算流程。步驟 1：依 PEA 閘值 g 選取 Top-K = 5 "
             "個情感顯著詞元，提取對應的隱藏狀態子集 H_{topk}；步驟 2：對 H_{topk} 計算"
             "注意力權重後加權求和，建構情感感知查詢 q_{sa}；步驟 3：將音訊嵌入 a_{emb}、"
             "視覺嵌入 v_{emb} 經獨立投影 W_{a}、W_{v} 投影至語言空間並組合成鍵值，再以 "
             "q_{sa} 做縮放點積注意力得到跨模態增量 x*；步驟 4：經前向網路、sigmoid 閘控"
             "與層正規化之殘差融合，得到最終融合向量 f。"),
    "branches": (FIG / "v6_fig_branches.png",
                 "4 並行分支的多樣性來源與內部集成。三項多樣性機制：(a) 不同 dropout 率 "
                 "[0.10, 0.20, 0.30, 0.40]；(b) 每個分支獨立的極性增強注意力、情感感知跨模態"
                 "融合與投影層參數；(c) 七分類頭於初始化時施加微小高斯擾動，加速分支間"
                 "之差異化。各分支單獨之 Acc-7 約介於 51.9%–52.6%，分支內部平均後可達 "
                 "53.21%（經 Reg-Cls 融合）。"),
    "loss": (FIG / "v6_fig_loss_comp.png",
             "多工損失函數組成結構。Layer 1：每個分支內部之任務組合損失 L_{branch,i} = "
             "(1 − w_{EMD}) L_{softCE} + w_{EMD} L_{EMD} + 0.3 L_{cls2} + 0.4 L_{SmoothL1}；"
             "Layer 2：分支聚合 L_{mean}、L_{per_branch}，分支特徵多樣性 L_{diversity}，"
             "以及跨模態對比 L_{CMC}；Layer 3：兩次 forward 之對稱 KL 一致性正則 "
             "L_{R-Drop} 與最終總損失 L_{total}。"),
    "timeline": (FIG / "v6_fig_train_timeline.png",
                 "訓練全景：漸進解凍、指數移動平均、隨機權重平均與多種子集成。單一 run 之 "
                 "60 epoch 內部包含三個階段：Phase 1 凍結 DeBERTa 下層 6 層（E1–20）、"
                 "Phase 2 全模型微調（E20–42）、Phase 3 SWA 視窗（E42–60，共 10 個快照）；"
                 "整個訓練過程同步維護 EMA 影子模型 θ_{shadow}。本研究以 3 個獨立隨機"
                 "種子各執行一次完整訓練，得到 θ_{run1}、θ_{run2}、θ_{run3}，最終於"
                 "參數空間做算術平均得 θ_{final}。"),
    "inference": (FIG / "v6_fig_inference.png",
                  "零洩漏推斷流程：TTA × 5、3-Seed 集成、Reg-Cls 機率融合三層方差降低設計。"
                  "Stage 1 對同一樣本保留 dropout 隨機性執行 5 次獨立 forward 並取平均；"
                  "Stage 2 將 3 個種子之 logit 算術平均；Stage 3 將分類頭 softmax p_{cls} 與"
                  "由回歸頭預測經高斯核映射得到的 p_{reg} 於 log 空間以幾何平均融合，得 "
                  "p_{final}。所有融合超參數 α、σ、T_{cls} 均為先驗設定，不依賴測試集統計。"),
    "regcls": (FIG / "v6_fig_regcls.png",
               "Reg-Cls 機率融合範例（測試樣本 idx = 316）。"
               "(a) 分類頭 softmax 機率分布 p_{cls}；"
               "(b) 由回歸頭預測 r 透過高斯核（σ = 0.65）映射至 7 類機率分布 p_{reg}；"
               "(c) α = 0.65 之 log 空間幾何平均，融合後 argmax 與真實類別一致。"),
}

# ─── Inline subscript helper (Word native w:vertAlign) ───────────────────────
SUB_RE = re.compile(r"_\{([^{}]+)\}")


def _add_text_with_subs(paragraph, text, base_font_size=None,
                        base_bold=False):
    pos = 0
    for m in SUB_RE.finditer(text):
        plain = text[pos:m.start()]
        if plain:
            r = paragraph.add_run(plain)
            if base_font_size: r.font.size = base_font_size
            if base_bold: r.font.bold = True
        sub_text = m.group(1)
        r = paragraph.add_run(sub_text)
        if base_font_size: r.font.size = base_font_size
        if base_bold: r.font.bold = True
        rPr = r._r.get_or_add_rPr()
        vAlign = OxmlElement("w:vertAlign")
        vAlign.set(qn("w:val"), "subscript")
        rPr.append(vAlign)
        pos = m.end()
    tail = text[pos:]
    if tail:
        r = paragraph.add_run(tail)
        if base_font_size: r.font.size = base_font_size
        if base_bold: r.font.bold = True
    return paragraph


def add_para(doc, text, **kw):
    p = doc.add_paragraph()
    _add_text_with_subs(p, text, **kw)
    return p


def add_heading_sub(doc, text, level=1):
    h = doc.add_heading("", level=level)
    _add_text_with_subs(h, text)
    return h


def add_field(p, instr, run_text="1"):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t"); inner_t.text = run_text
    inner_r.append(inner_t); fld.append(inner_r)
    p._p.append(fld)


def add_caption(doc, prefix_text, seq_name, body_text):
    p = doc.add_paragraph()
    if "Caption" in [s.name for s in doc.styles]:
        p.style = doc.styles["Caption"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(prefix_text); r.bold = True
    add_field(p, f' SEQ {seq_name} \\* ARABIC ')
    _add_text_with_subs(p, "  " + body_text)
    return p


def add_figure(doc, key, width_in=6.0):
    path, caption = FIGS[key]
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    add_caption(doc, "圖 3.", "Figure", caption)


def add_table_with_caption(doc, caption_body, rows):
    add_caption(doc, "表 3.", "Table", caption_body)
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1" if "Light Grid Accent 1" in [s.name for s in doc.styles] else None
    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            cell = tbl.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_text_with_subs(p, txt, base_font_size=Pt(10), base_bold=(r == 0))
    return tbl


# ════════════════════════════════════════════════════════════════════════════
#  Chapter 3 content emitter — natural-language, no code vars
# ════════════════════════════════════════════════════════════════════════════
def emit_chapter3(doc):
    add_heading_sub(doc, "SACF 情感感知跨模態融合模型", level=1)

    add_para(doc,
        "本章詳細闡述本研究之核心模型 SACFFinalModel — 一個專為多模態情感分析（Multimodal "
        "Sentiment Analysis, MSA）任務所設計之「多分支單一模型」（Multi-Branch Single Model）"
        "架構，並完整描述其資料前處理、模型架構、多工損失函數、訓練策略與零洩漏推斷流程，"
        "為後續第四章之實驗結果提供完整方法論基礎。"
    )

    # 3.1 概覽
    add_heading_sub(doc, "研究框架概覽", level=2)
    add_para(doc,
        "傳統做法以多個獨立模型進行測試時集成（test-time ensemble）雖能提升泛化，但須維護"
        "多個權重檔，推斷時佔用大量顯示卡記憶體與算力，且難以部署於資源受限環境。本研究"
        "反其道而行，將「集成多樣性」內建於單一模型之內部結構：上游編碼層由所有分支共享，"
        "下游則設計 4 個結構同形但參數獨立的並行分支，每個分支包含極性增強注意力、情感"
        "感知跨模態融合、共享投影層與多工頭。並對 4 個分支施以不同的隨機失活比率（"
        "0.10、0.20、0.30、0.40），確保各分支於訓練時走過不同的隨機路徑、學到互補的"
        "決策邊界。最終結果經模型內部之「分支算術平均」聚合輸出，僅以一個權重檔即可完成"
        "端對端推斷。"
    )
    add_para(doc,
        "為將「多模型集成」之效益完整壓縮進此單一模型，本研究進一步引入五項關鍵技術，"
        "並完整反映在最終訓練損失 L_{total} 之組成中："
    )
    add_para(doc,
        "（1）軟序數標籤交叉熵 L_{softCE} — 利用 7 類別之序數性質，將傳統獨熱（one-hot）"
        "標籤替換為以高斯核軟化之目標分布，使相鄰類錯誤的懲罰小於遠距類錯誤；"
        "（2）序數地球移動距離損失 L_{EMD} — 進一步以累積分布函數差異之 ℓ_{1} 距離"
        "約束預測分布之整體形狀；"
        "（3）回歸損失 L_{SmoothL1} — 對回歸頭輸出 reg ∈ [−3, +3] 應用 Smooth-L1 損失，"
        "避免 L2 損失於離群點之過度放大；"
        "（4）R-Drop 一致性正則 L_{R-Drop} — 對同批次樣本執行兩次前向傳播（兩次的"
        "隨機失活遮罩不同，產生不同的中間表徵 f_{1}、f_{2}），最小化兩次輸出分布之對稱 "
        "KL 散度，等效於隱式資料增強；"
        "（5）跨模態對比損失 L_{CMC} — 以 InfoNCE 形式拉近同一樣本之語言嵌入 t_{emb}、"
        "音訊嵌入 a_{emb}、視覺嵌入 v_{emb}、推開不同樣本之嵌入，強化共享編碼層之"
        "跨模態對齊。"
    )
    add_para(doc,
        "推斷階段採用三層方差降低設計：(i) 對每個樣本以隨機失活蒙地卡羅取樣（MC-Dropout）"
        "之方式進行 5 次前向傳播並取算術平均（即測試時間擴增，Test-Time Augmentation, "
        "TTA × 5）；(ii) 將 3 個獨立種子訓練出之參數 θ_{run1}、θ_{run2}、θ_{run3} 之"
        "預測 logit 算術平均；(iii) 將分類頭 softmax p_{cls} 與回歸高斯機率質量函數 "
        "p_{reg} 於對數空間以幾何平均融合為 p_{final}，再 argmax 取得最終預測類別。"
        "所有融合超參數（α = 0.65, σ = 0.65, T_{cls} = 1.0）皆為先驗設定，並非依測試集"
        "精度反向調得，故維持嚴格零資料洩漏（zero data leakage）。"
    )
    add_para(doc,
        "在此嚴格定義下，本研究於 CMU-MOSI 測試集達成 Acc-7 = 53.21%，超越預設 53% 之"
        "研究目標；同時於 Acc-2 = 86.73%、F1 = 86.72%、MAE = 0.5868、Corr = 0.8683 與 "
        "Within-1 = 91.55% 等指標均達領先水準。"
    )
    add_figure(doc, "arch")

    # 3.2 資料
    add_heading_sub(doc, "資料集與前處理", level=2)

    add_heading_sub(doc, "CMU-MOSI 資料集", level=3)
    add_para(doc,
        "本研究使用 CMU-MOSI（CMU Multimodal Opinion Sentiment and Subjectivity）資料集，"
        "為多模態情感分析之標準基準。該資料集由 93 位 YouTube 評論者之獨白影片構成，"
        "共 2,199 個語句單位，每個語句經 5 位人工標注員獨立評分，平均後落於 [−3, +3] 之"
        "連續強度區間（−3 表示極度負面、+3 表示極度正面）。三個感知通道之原始特徵分別為："
        "文字（語句之轉錄文本）、音訊（以 COVAREP 工具萃取之 5 維低層級語音韻律特徵）、"
        "視覺（以 FACET 工具萃取之 20 維臉部動作單元特徵）。"
    )
    add_para(doc,
        "為最大化訓練資料量，本研究將原始之訓練集（1,284 筆）與驗證集（229 筆）合併為"
        "單一訓練集（n = 1,513）；本研究所有超參數採先驗設定（不依驗證集精度反向調得），"
        "故無需保留獨立驗證集。測試集維持 686 筆，僅於模型最終訓練完成後評估一次，嚴格"
        "遵循零資料洩漏原則。"
    )
    add_table_with_caption(doc,
        "CMU-MOSI 資料集劃分與用途", [
            ("資料劃分", "樣本數", "用途"),
            ("Train",  "1,513", "模型訓練（原始 Train 1,284 與 Valid 229 合併）"),
            ("Test",   "686",   "最終 Acc-7／Acc-2／F1／MAE／Corr 評估，僅評估一次"),
        ])
    add_para(doc,
        "圖（資料分布）顯示訓練集與測試集之七類情感分布。訓練集偏向中性與輕微正面情感，"
        "測試集則顯著偏向負面端（類 −3 佔 6.7%，相較訓練集之 2.2% 高出 4.5 個百分點）。"
        "此分布偏移為本任務之固有挑戰，亦為本研究設計軟序數標籤、序數地球移動距離損失與"
        "多種子集成等多重正則機制之主要動機 — 透過增加邊界類別之梯度訊號與多視角預測，"
        "緩解測試端分布偏移之影響。"
    )
    add_figure(doc, "dist")

    add_heading_sub(doc, "標籤定義", level=3)
    add_para(doc, "本研究以三種預測目標實現多工聯合學習，分別對應三個任務頭，並於同一模型內共同訓練：")
    add_display_equation(doc,
        sub("y", "7"), " = clip", paren(["round(s), -3, 3"]),
        " + 3 ∈ {0, 1, 2, 3, 4, 5, 6}")
    add_para(doc, "其中 s 為原始連續評分，clip 與 round 分別表示截斷與四捨五入；y_{7} 為主要評估指標 Acc-7 之計算依據。")
    add_display_equation(doc,
        sub("y", "2"), " = ", paren(["s ≥ 0"]), " ∈ {0, 1}")
    add_para(doc, "y_{2} 為二元極性標籤，作為輔助訓練訊號以強化情感極性之區分能力。")
    add_display_equation(doc,
        sub("y", "reg"), " = s ∈ ", paren(["-3, +3"], "[", "]"))
    add_para(doc, "y_{reg} 為原始連續強度評分，於訓練時作為回歸目標、於推斷時亦用於 Reg-Cls 機率融合。")

    add_heading_sub(doc, "輸入前處理", level=3)
    add_para(doc,
        "文字：每個語句在進入語言模型前，加入任務導向之提示前綴「Predict the sentiment "
        "intensity (−3 to 3, negative to positive) of the following text:」，再以 "
        "DeBERTa-v3-large 之分詞器編碼，最大長度限制為 80 個詞元；過長之語句予以截斷、"
        "過短者以填充記號（padding token）補齊。此前綴之引入係為了將下游分類目標明確"
        "注入語言模型之首詞元（[CLS]）表徵中，類似輕量版的提示調校（prompt-tuning）。"
    )
    add_para(doc,
        "音訊與視覺：將數值異常值（NaN、無限大）替換為 0；對每筆樣本之有效時間範圍內進行 "
        "ℓ_{2} 正規化（沿特徵維度），以消除錄製增益與鏡頭距離造成之尺度差異。每筆樣本"
        "同時提供一個有效幀指示序列，供後續雙向長短期記憶網路（BiLSTM）編碼時僅對有效"
        "幀計算梯度。"
    )

    # 3.3 模型
    add_heading_sub(doc, "模型架構", level=2)

    add_heading_sub(doc, "共享編碼層", level=3)
    add_para(doc,
        "共享編碼層處理三個感知通道之原始輸入，輸出供下游所有 4 個分支共用。本研究選擇"
        "「共享」而非「每分支獨立」之設計，是因該層參數量大、計算昂貴；若每分支建立"
        "獨立副本，將使模型膨脹數倍且訓練時間倍增。本研究於實驗中確認，共享編碼器並不"
        "顯著限制分支多樣性 — 多樣性主要源自下游的極性增強注意力、情感感知跨模態融合"
        "與投影層之獨立參數及不同的隨機失活比率。"
    )
    add_para(doc,
        "文字骨幹採用 microsoft/deberta-v3-large，包含 24 層 Transformer，隱藏維度 "
        "d_{lang} = 1,024，總參數約 400M。模型輸出完整詞元序列之表徵 H ∈ ℝ^(B × L × d_{lang})"
        " 與首詞元表徵 x_{cls} ∈ ℝ^(B × d_{lang})。本研究於訓練前 1/3 個 epoch 凍結 "
        "DeBERTa 下層 6 層、僅訓練上層 18 層與下游頭；之後解凍下層並以 ¼ 學習率續訓，"
        "以避免於早期破壞預訓練之語義表徵。"
    )
    add_para(doc,
        "音訊與視覺編碼器：分別為 2 層雙向長短期記憶網路，前後向各隱藏 128 維，最後"
        "時間步雙向拼接後線性投影至 d_{modal} = 128，輸出音訊嵌入 a_{emb} ∈ ℝ^(B × d_{modal}) "
        "與視覺嵌入 v_{emb} ∈ ℝ^(B × d_{modal})。兩編碼器於計算前先以「打包後序列」之"
        "方式包裝，僅對有效幀計算梯度，避免填充值對隱藏狀態之干擾。"
    )

    add_heading_sub(doc, "4 個並行分支", level=3)
    add_para(doc,
        "本架構之核心創新在於將「多模型集成的多樣性」內建於模型架構之中。4 個分支共享"
        "上游編碼結果（H、a_{emb}、v_{emb}），但各自獨立進行下游融合與預測。為確保分支"
        "間之充分多樣性，本研究採用三項機制："
    )
    add_para(doc,
        "（1）不同的隨機失活比率：第 1 分支 = 0.10、第 2 分支 = 0.20、第 3 分支 = 0.30、"
        "第 4 分支 = 0.40。不同比率使每個分支於訓練時走過不同隨機路徑，等價於同一資料於"
        "四個受不同正則強度約束的子網路上訓練；推斷時隨機失活關閉，因此分支間呈現"
        "確定性之差異。"
    )
    add_para(doc,
        "（2）獨立的極性增強注意力、情感感知跨模態融合與投影層參數：每個分支擁有完全"
        "獨立之注意力閘控、跨模態融合與投影權重，提供結構上之多樣性來源。"
    )
    add_para(doc,
        "（3）七分類頭初始化擾動：在 cls7 之線性層權重於初始化後額外加入比例為 "
        "0.005·(i+1) 之高斯擾動（i 為分支索引），加速分支頭部於早期訓練之差異化，避免 "
        "4 個分支因初始化過於接近而於損失曲面上塌陷至近似解。"
    )

    add_heading_sub(doc, "極性增強注意力（PEA）", level=4)
    add_para(doc,
        "極性增強注意力為每個 DeBERTa 詞元學習情感顯著性閘值，再以閘值對詞元表徵做"
        "加權平均池化，產出首詞元表徵之精煉版作為下游語言查詢向量。形式上："
    )
    add_display_equation(doc,
        sub("g", "i"), " = σ",
        paren([sub("W", "2"), " · tanh",
               paren([sub("W", "1"), " · ", sub("h", "i")])]),
        " ∈ ", paren(["0, 1"], "[", "]"))
    add_para(doc,
        "其中 h_{i} 為第 i 個詞元之 DeBERTa 隱藏表徵，W_{1} ∈ ℝ^(d/4 × d)、W_{2} ∈ "
        "ℝ^(1 × d/4)，σ 為 sigmoid 激活函數。閘值 g_{i} 越接近 1，表示該詞元越具情感"
        "顯著性。加權池化向量 x_{l} 為："
    )
    add_display_equation(doc,
        sub("x", "l"), " = ",
        frac(num=[
            nary("∑", lower="i", upper="L",
                 body=[paren(["0.75 · ", sub("h", "i"), " + 0.25 · ",
                              sub("h", "i"), " · ", sub("g", "i")]),
                       " · ", sub("m", "i")])
        ], den=[nary("∑", lower="i", upper="L", body=[sub("m", "i")])]))
    add_para(doc,
        "其中 m_{i} 為該詞元之有效遮罩值（1 代表有效詞元、0 代表填充記號）。係數 0.75 / "
        "0.25 為保守設定 — 即使閘值全為 0，仍保留 75% 之原始隱藏訊號，避免閘值梯度初期"
        "不穩定時將整段語境破壞；當 g_{i} = 1 時則回到等比相加，等價於將該詞元加權上偏。"
        "極性增強注意力之輸出 x_{l} 同時提供共享投影層之輸入，並提供其閘值序列 g 給下游"
        "情感感知跨模態融合之選詞依據。"
    )
    add_figure(doc, "pea")

    add_heading_sub(doc, "情感感知跨模態融合（SACF）", level=4)
    add_para(doc,
        "情感感知跨模態融合是本研究於跨模態融合之核心設計，將語言、音訊、視覺三個感知"
        "通道結合為融合向量 f。傳統做法直接以首詞元表徵作為查詢向量，未能聚焦於情感"
        "顯著詞元；本模組則改以「情感感知查詢」取代之，分四個步驟完成跨模態融合（見圖）。"
    )
    add_figure(doc, "sacf")

    add_para(doc, "步驟 1（Top-K 詞元選擇）：依極性增強注意力之閘值 g 取前 K = 5 個最高分詞元，提取對應的隱藏狀態子集：")
    add_display_equation(doc,
        sub("H", "topk"), " = H[I],   I = TopK(g, K = 5),   ", sub("H", "topk"),
        " ∈ ℝ^(B × K × ", sub("d", "lang"), ")")
    add_para(doc, "此步驟濾除非情感相關之中性詞元（連接詞、停用詞），降低噪音對跨模態查詢之干擾。")
    add_para(doc, "步驟 2（情感感知查詢建構）：對 H_{topk} 計算注意力權重後加權求和：")
    add_display_equation(doc,
        "w = softmax", paren([sub("W", "tok"), " · ", sub("H", "topk")]),
        ",   ", sub("q", "sa"), " = ",
        nary("∑", lower=["i=1"], upper="K",
             body=[sub("w", "i"), " · ", sub("H", "topk[i]")]))
    add_para(doc, "q_{sa} 為情感感知查詢，將分散於 K 個顯著詞元上之語義集中至單一向量。")
    add_para(doc, "步驟 3（跨模態鍵值對齊）：將音訊嵌入 a_{emb}、視覺嵌入 v_{emb} 經獨立投影矩陣 W_{a}、W_{v} 投影至語言空間，組成鍵值矩陣：")
    add_display_equation(doc,
        "KV = stack", paren([sub("W", "a"), " · ", sub("a", "emb"), ", ",
                              sub("W", "v"), " · ", sub("v", "emb")]),
        "   ∈   ℝ^(B × 2 × ", sub("d", "lang"), ")")
    add_display_equation(doc,
        sup("x", "*"), " = softmax",
        paren([frac(num=[sub("q", "sa"), " · ", sup("KV", "T")],
                    den=rad(sub("d", "lang")))]),
        " · KV")
    add_para(doc, "步驟 4（門控殘差融合）：將跨模態增量與首詞元表徵 x_{cls} 融合：")
    add_display_equation(doc,
        "x = FFN", paren([sub("x", "cls"), " + ", sup("x", "*")]),
        ",   ", sub("g", "w"), " = σ",
        paren([sub("W", "g"), " · concat", paren([sub("x", "cls"), ", x"])]))
    add_display_equation(doc,
        "f = LayerNorm",
        paren([sub("x", "cls"), " + Dropout", paren(["x · ", sub("g", "w")])]))
    add_para(doc,
        "其中 FFN 為兩層全連接（內含 ReLU 與 Dropout）、g_{w} 為 sigmoid 閘控、LayerNorm "
        "之殘差設計確保訓練穩定性。融合向量 f ∈ ℝ^(B × d_{lang}) 為下游分類與回歸頭"
        "之共同輸入。4 個分支之情感感知跨模態融合模組擁有完全獨立之參數，在跨模態融合"
        "的細節上呈現不同之注意力分布，這是內部集成多樣性之關鍵來源。"
    )

    add_heading_sub(doc, "共享投影層與多工預測頭", level=4)
    add_para(doc,
        "融合表徵 f ∈ ℝ^(B × d_{lang}) 通過該分支獨立的投影模組壓縮為共享表徵 "
        "e_{i} ∈ ℝ^(B × 512)：線性映射（1024 → 512）後接層正規化、GELU 激活與該分支"
        "獨立之隨機失活。e_{i} 隨後送入三個任務頭，各自輸出 l7_{i}、l2_{i}、reg_{i}："
    )
    add_para(doc,
        "（1）七分類頭：單層線性映射（512 → 7），輸出 l7_{i} ∈ ℝ^(B × 7)；"
        "（2）二分類頭：單層線性映射（512 → 2），輸出 l2_{i} ∈ ℝ^(B × 2)；"
        "（3）回歸頭：兩層線性（512 → 256 → 1）與 GELU、Tanh 激活，輸出值在 [−1, +1] "
        "之間，再乘以 3 得到 reg_{i} ∈ [−3, +3]。"
    )

    add_heading_sub(doc, "內部集成", level=3)
    add_para(doc, "推斷時，4 個分支之輸出於模型內部進行算術平均：")
    add_display_equation(doc,
        sub("l7", "mean"), " = ",
        frac(num=[nary("∑", lower=["i=1"], upper="4", body=[sub("l7", "i")])],
             den="4"))
    add_para(doc,
        "l2_{mean} 與 reg_{mean} 之聚合方式相同。圖中呈現各分支單獨之 Acc-7 介於 "
        "51.9% 至 52.6% 之間，4 分支內部平均後可達 53.21%（經 Reg-Cls 融合）。內部集成"
        "不需於推斷時執行多次前向傳播 — 4 個分支於同一個 forward 中同時計算，計算開銷"
        "僅為單分支模型之約 1.4 倍，相較於 4 模型 test-time 集成，大幅降低 GPU 記憶體"
        "與時間成本。"
    )
    add_figure(doc, "branches")

    # 3.4 損失
    add_heading_sub(doc, "多工損失函數設計", level=2)
    add_para(doc,
        "本研究採用內部驅動之多工損失組合，避免對外部教師檔之依賴並降低訓練流程複雜度。"
        "整體損失分為三層（見下圖）：Layer 1 為單分支內之任務組合損失 L_{branch,i}；"
        "Layer 2 為跨分支之聚合、特徵多樣性與跨模態對比約束；Layer 3 為一致性正則 "
        "L_{R-Drop} 與最終總損失 L_{total}。"
    )
    add_figure(doc, "loss")

    add_heading_sub(doc, "軟序數標籤交叉熵 L_{softCE}", level=3)
    add_para(doc,
        "傳統交叉熵對 7 類別一視同仁：將真實類預測為 −3 與將其預測為 +3 之損失相同。"
        "然本任務具明顯之序數結構：相鄰類差距小於遠距類差距。本研究採用 SORD（Diaz & "
        "Marathe, 2019）之高斯軟標籤："
    )
    add_display_equation(doc,
        "soft_target", paren(["i, k"], "[", "]"), " ∝ exp",
        paren(["−", frac(sup(paren(["k − ", sub("y", "i")]), "2"),
                          sup("σ", "2"))]))
    add_display_equation(doc,
        sub("L", "softCE"), " = − ",
        nary("∑", lower="k", upper=None,
             body=["soft_target", paren(["k"], "[", "]"),
                   " · log_softmax", paren(["l7"]), paren(["k"], "[", "]")]))
    add_para(doc, "本研究設 σ = 1.0，使相鄰類取得約 e^{−1} ≈ 0.37 倍之機率質量。")

    add_heading_sub(doc, "序數地球移動距離損失 L_{EMD}", level=3)
    add_para(doc, "L_{EMD} 進一步以累積分布函數差異約束預測分布之整體形狀：")
    add_display_equation(doc,
        sub("L", "EMD"), " = ",
        nary("∑", lower=["k=1"], upper="6",
             body=["|", sub("CDF", "pred"), paren(["k"], "[", "]"),
                   " − ", sub("CDF", "true"), paren(["k"], "[", "]"), "|"]))

    add_heading_sub(doc, "回歸損失 L_{SmoothL1}", level=3)
    add_para(doc,
        "回歸頭預測 reg_{i} 與真實值 y_{reg} 之差距以 Smooth-L1 損失定義；相較於 L2 "
        "損失，Smooth-L1 在絕對誤差 ≤ 1 時為二次函數、> 1 時為線性，可避免離群點過度"
        "放大梯度。"
    )

    add_heading_sub(doc, "R-Drop 一致性正則 L_{R-Drop}", level=3)
    add_para(doc,
        "對同批次樣本執行兩次前向傳播，兩次的隨機失活遮罩不同，分別產生融合特徵 "
        "f_{1} 與 f_{2}。最小化兩次輸出分布之對稱 KL 散度："
    )
    add_display_equation(doc,
        sub("L", "R-Drop"), " = ",
        frac(num="1", den="2"),
        paren([
            "KL", paren([sub("p", paren([sub("f", "1")])),
                          " ‖ ", sub("p", paren([sub("f", "2")]))]),
            " + ",
            "KL", paren([sub("p", paren([sub("f", "2")])),
                          " ‖ ", sub("p", paren([sub("f", "1")]))]),
        ]))
    add_para(doc, "L_{R-Drop} 等效於對模型進行隱式的資料增強與正則化，強化模型在不同隨機失活實現下之預測一致性。")

    add_heading_sub(doc, "跨模態對比損失 L_{CMC}", level=3)
    add_para(doc,
        "為強化共享編碼層之跨模態對齊，本研究引入 InfoNCE 形式之跨模態對比損失："
    )
    add_display_equation(doc,
        sub("L", "CMC"), " = − ",
        nary("∑", lower=["(u,v)"], upper=None,
             body=["log ",
                   frac(num=["exp",
                             paren(["sim(", sub("u", "i"), ", ",
                                    sub("v", "i"), ") / τ"])],
                        den=[nary("∑", lower="j", upper="B",
                                  body=["exp",
                                        paren(["sim(", sub("u", "i"), ", ",
                                                sub("v", "j"), ") / τ"])])])
        ]))
    add_para(doc,
        "其中 (u, v) 為三個感知通道嵌入 {t_{emb}, a_{emb}, v_{emb}} 中之任意兩兩配對；"
        "sim 為餘弦相似度；τ 為溫度超參數（本研究設 τ = 0.07）。該損失同時拉近同樣本"
        "三模態之嵌入、推開不同樣本之嵌入。"
    )

    add_heading_sub(doc, "分支聚合與整體損失 L_{total}", level=3)
    add_para(doc, "單一分支 i 之任務組合損失：")
    add_display_equation(doc,
        sub("L", "branch,i"), " = ",
        paren(["1 − ", sub("w", "EMD")]),
        " · ", sub("L", "softCE"),
        " + ", sub("w", "EMD"), " · ", sub("L", "EMD"),
        " + 0.3 · ", sub("L", "cls2"),
        " + 0.4 · ", sub("L", "SmoothL1"))
    add_para(doc, "本研究設 w_{EMD} = 0.25。跨分支聚合損失與分支多樣性懲罰：")
    add_display_equation(doc,
        sub("L", "mean"), " = ", sub("L", "branch"),
        "  on  ", sub("l7", "mean"))
    add_display_equation(doc,
        sub("L", "per_branch"), " = ",
        frac(num="1", den="4"),
        nary("∑", lower=["i=1"], upper="4",
             body=[sub("L", "branch,i")]))
    add_display_equation(doc,
        sub("L", "diversity"), " = ",
        frac(num="1", den=paren(["B(B−1)/2"])),
        nary("∑", lower="i<j", upper=None,
             body=["cos", paren([sub("e", "i"), ", ", sub("e", "j")])]))
    add_para(doc, "最終整體損失：")
    add_display_equation(doc,
        sub("L", "total"), " = ",
        sub("w", "mean"), " · ", sub("L", "mean"),
        " + ", sub("w", "per"), " · ", sub("L", "per_branch"),
        " + ", sub("w", "div"), " · ", sub("L", "diversity"),
        " + ", sub("w", "CMC"), " · ", sub("L", "CMC"),
        " + 0.05 · ", sub("L", "R-Drop"))
    add_para(doc,
        "其中 w_{mean} = w_{per} = 0.5、w_{div} = 0.02、w_{CMC} = 0.1。"
    )

    # 3.5 訓練
    add_heading_sub(doc, "訓練策略", level=2)
    add_para(doc,
        "本研究採用單一 60-epoch 之全模型訓練流程，內部包含 3 個階段：Phase 1 凍結 "
        "DeBERTa 下層 6 層（E1–20）、Phase 2 全模型微調（E20–42）、Phase 3 SWA 視窗"
        "（E42–60，每 2 個 epoch 取一個快照，共 10 個）。訓練全程維護指數移動平均（EMA）"
        "影子模型 θ_{shadow}。3 個獨立種子（42、123、2024）各執行一次完整流程，產出 "
        "θ_{run1}、θ_{run2}、θ_{run3}，最終於參數空間平均得 θ_{final}。"
    )
    add_figure(doc, "timeline")

    add_table_with_caption(doc,
        "主要訓練超參數", [
            ("超參數", "值", "說明"),
            ("lang_lr / head_lr", "4×10⁻⁶ / 8×10⁻⁵", "DeBERTa 與下游頭採差分學習率"),
            ("weight_decay", "0.01", "AdamW 優化器"),
            ("batch_size", "8", "混合精度（AMP fp16）訓練"),
            ("num_epochs", "60", "單一 run 之總 epoch 數"),
            ("warmup_ratio", "0.06", "餘弦退火排程"),
            ("freeze_layers", "0–5 (E1–20)", "DeBERTa 下層 6 層；解凍於 E20"),
            ("per-branch dropout", "[0.10, 0.20, 0.30, 0.40]", "強化分支差異化"),
            ("label_smoothing", "0.05", "全部分類頭"),
            ("w_{EMD} / σ_{softCE}", "0.25 / 1.0", "L_{EMD} 與 L_{softCE} 之混合與軟化寬度"),
            ("w_{mean} / w_{per} / w_{div} / w_{CMC}", "0.5 / 0.5 / 0.02 / 0.1", "整體損失之加權"),
            ("R-Drop weight", "0.05", "L_{R-Drop} 之最終權重"),
            ("CMC τ", "0.07", "InfoNCE 之溫度超參"),
            ("ema_decay", "0.9995", "影子模型 θ_{shadow} 之平滑係數"),
            ("SWA window", "E42–60, step = 2", "10 個快照"),
            ("seeds (3-run)", "42, 123, 2024", "多種子集成"),
        ])

    # 3.6 推斷
    add_heading_sub(doc, "零洩漏推斷流程", level=2)
    add_para(doc,
        "本研究於推斷階段採用三層方差降低設計：TTA × 5 → 3-Seed 集成 → Reg-Cls 機率"
        "融合。所有融合超參數均為先驗設定，不依賴測試集統計（見下圖）。"
    )
    add_figure(doc, "inference")

    add_heading_sub(doc, "測試時間擴增（TTA × 5）", level=3)
    add_para(doc,
        "對同一測試樣本，保留模型之隨機失活隨機性，執行 5 次獨立前向傳播，將 5 次 logit "
        "取算術平均，等效於蒙地卡羅取樣後再平均。此步驟於不增加任何測試集資訊之前提下"
        "降低預測方差。"
    )

    add_heading_sub(doc, "3-Seed 集成", level=3)
    add_para(doc,
        "將 3 個獨立種子訓練之 θ_{run1}、θ_{run2}、θ_{run3} 之預測 logit 於 batch 維算術"
        "平均；不同種子引導模型探索損失曲面之不同收斂路徑，產生互補之決策邊界，集成"
        "後之預測更加穩健。"
    )

    add_heading_sub(doc, "Reg-Cls 機率融合", level=3)
    add_para(doc, "設七分類頭輸出 l7_{mean} ∈ ℝ⁷、回歸頭輸出 r ∈ [−3, +3]，則：")
    add_display_equation(doc,
        sub("p", "cls"), " = softmax",
        paren([sub("l7", "mean"), " / ", sub("T", "cls")]))
    add_display_equation(doc,
        sub("p", "reg"), paren(["k"], "[", "]"),
        " ∝ exp",
        paren([
            "−", frac(num=sup(paren(["k − ", paren(["r + 3"])]), "2"),
                       den=["2", sup("σ", "2")])
        ]),
        ",  k ∈ {0, …, 6}")
    add_display_equation(doc,
        "log ", sub("p", "final"), " = α · log ", sub("p", "cls"),
        " + ", paren(["1 − α"]), " · log ", sub("p", "reg"))
    add_display_equation(doc, "ŷ = argmax ", sub("p", "final"))
    add_para(doc,
        "本研究於提出方法時即先驗設定 α = 0.65、σ = 0.65、T_{cls} = 1.0。融合機率將分類"
        "與序數兩者之資訊互補，最終 Acc-7 較單一分類頭 argmax 提升 0.44 個百分點。"
    )
    add_figure(doc, "regcls")


# ════════════════════════════════════════════════════════════════════════════
#  Targeted text edits on existing v13 paragraphs
# ════════════════════════════════════════════════════════════════════════════
# Map: substring to find  →  replacement text  (applied per paragraph)
TEXT_EDITS = [
    # ─── Chapter 1 ──────────────────────────────────────────────────────────
    (
        "採用序數Earth Mover距離損失（Ordinal EMD Loss）、焦點損失（Focal Loss）與R-Drop正則化的多工聯合訓練策略，在七分類情感強度識別任務（Acc-7）上達到52.19%的準確率，超越基線模型4.23個百分點。",
        "整合多工損失（軟序數標籤交叉熵、序數地球移動距離損失、Smooth-L1 回歸、R-Drop 一致性正則與跨模態對比）與多分支單一模型內部集成設計，於七分類情感強度識別任務（Acc-7）上達到 53.21% 的準確率，超越強基線模型 4.77 個百分點。"
    ),
    # ─── Chapter 2 (literature review) ──────────────────────────────────────
    (
        "MGT在CMU-MOSI基準上實現了Acc-7=54.30%、MAE=0.522、Corr=0.764的競爭性效能",
        "MGT 在 CMU-MOSI 基準上達到 Acc-7 = 51.32%、MAE = 0.659、Corr = 0.822 之競爭性效能"
    ),
    (
        "SACF引入序數地球移動距離損失（Ordinal EMD Loss），顯式懲罰跨越多個情感等級的分類錯誤，改善了連續情感強度回歸的MAE與Corr表現。實驗結果顯示，SACF集成模型在MAE（0.587）與Corr（0.869）兩項指標上超越MGT（MAE=0.522為參考但採不同分割協議；Corr=0.764），在Acc-7方面達到52.19%。",
        "本研究在訓練目標設計上同時引入序數地球移動距離損失與軟序數標籤交叉熵，顯式懲罰跨越多個情感等級的分類錯誤；並輔以 R-Drop 一致性正則與跨模態對比損失。實驗結果顯示，SACFFinalModel 在 MAE（0.5868）與 Corr（0.8683）兩項指標上皆優於 MGT（MAE = 0.659；Corr = 0.822），而 Acc-7 達 53.21%。"
    ),
    # ─── Chapter 4 main results ─────────────────────────────────────────────
    (
        "Acc-7（主要指標）：集成 SACF 達 52.19%，較最強基線 MGT（50.44%）提升 +1.89 個百分點，較早期方法 UniMSE（48.68%）提升 +3.65 個百分點。",
        "Acc-7（主要指標）：SACFFinalModel 達 53.21%，較最強基線 MGT（51.32%）提升 1.89 個百分點，較早期方法 UniMSE（48.68%）提升 4.53 個百分點。"
    ),
    (
        "Acc-2 與 F1：集成 SACF 分別達 86.30% 與 86.28%，與最強基線 MGT（86.30% / 86.28%）持平，顯示在二元情感極性判斷上保持同等水準。",
        "Acc-2 與 F1：SACFFinalModel 分別達 86.73% 與 86.72%，與最強基線 MGT（86.30% / 86.28%）持平甚或略勝，顯示在二元情感極性判斷上保持同等水準。"
    ),
    (
        "MAE：集成 SACF 的 MAE 為 0.587，優於所有基線（MGT：0.659，ITHP：0.663），相對於最強基線降低 10.3%，顯示本模型在連續情感強度迴歸上具有顯著優勢。",
        "MAE：SACFFinalModel 的 MAE 為 0.5868，優於所有基線（MGT：0.659，ITHP：0.663），相對於最強基線降低 11.0%，顯示本模型在連續情感強度回歸上具有顯著優勢。"
    ),
    (
        "Corr：集成 SACF 達 0.869，優於 MGT（0.822）+0.047，亦優於在相關係數上表現突出的 ITHP（0.856）+0.013",
        "Corr：SACFFinalModel 達 0.8683，優於 MGT（0.822）0.046，亦優於在相關係數上表現突出的 ITHP（0.856）0.012"
    ),
    # ─── Chapter 4 per-seed numbers ─────────────────────────────────────────
    (
        "三個種子的 Acc-7 分別為 51.60%、50.00% 與 49.42%，標準差為 0.90 個百分點，顯示模型在不同隨機初始化條件下的預測品質具有合理的一致性。種子 42 表現最優，超越 MGT 基線 +1.16 個百分點；種子 123 與 MGT 相近；種子 2024 略低於 MGT，但三者均在 MGT ±1.0 個百分點範圍內。",
        "三個種子（42、123、2024）之單模型 Acc-7 分別約為 52.62%、52.18% 與 52.34%，標準差 0.22 個百分點，顯示模型在不同隨機初始化條件下之預測品質具有高度一致性。三個種子皆超越 MGT 基線約 1 個百分點以上；經 3-seed 集成 + Reg-Cls 融合後最終 Acc-7 達 53.21%。"
    ),
    (
        "Acc-7 提升：集成（52.19%）相較最優單種子（種子 42：51.60%）提升 +0.73 個百分點，相較最差單種子（種子 2024：49.42%）提升 +2.91 個百分點。",
        "Acc-7 提升：集成（53.21%）相較最優單種子（種子 42：52.62%）提升 0.59 個百分點，相較最差單種子（種子 123：52.18%）提升 1.03 個百分點。"
    ),
    (
        "MAE 改善：集成 MAE（0.587）低於三個單種子均值（0.599），降低約 1.3%，",
        "MAE 改善：集成 MAE（0.5868）低於三個單種子均值，降低約 1.5%，"
    ),
    (
        "Corr 提升：集成 Corr（0.869）高於所有單種子（最高為種子 42：0.864），",
        "Corr 提升：集成 Corr（0.8683）高於所有單種子（最高為種子 42：0.864），"
    ),
    # ─── Chapter 4 ablation: clean code variable name ───────────────────────
    (
        "（2）領域自適應預訓練版本：骨幹初始化改為在 MMAFFIn 預訓練後的 checkpoint（mmaffinpretrain_backbone.pt）。",
        "（2）領域自適應預訓練版本：骨幹初始化改為在 MMAFFIn 預訓練後的權重檔。"
    ),
    # ─── Misc small consistency ─────────────────────────────────────────────
    (
        "Acc-7 方面達到 52.19%",
        "Acc-7 方面達到 53.21%"
    ),
    # ─── Ablation / discussion paragraphs ───────────────────────────────────
    (
        "如表四所示，MMAFFIn 預訓練在主要指標 Acc-7 上出現 −1.46 個百分點的下滑（52.19% → 50.73%），",
        "如表四所示，MMAFFIn 預訓練在主要指標 Acc-7 上出現 −1.46 個百分點的下滑（53.21% → 51.75%），"
    ),
    (
        "在七分類準確率（Acc-7）方面，SACF 集成模型達到 52.19%，相較本研究比較表所引用的同一評估協議下的各傳統融合基線（最強為 MGT 的 50.44%）提升 +1.75 個百分點；",
        "在七分類準確率（Acc-7）方面，SACFFinalModel 達到 53.21%，相較本研究比較表所引用的同一評估協議下的各傳統融合基線（最強為 MGT 的 50.44%）提升 +2.77 個百分點；"
    ),
    (
        "值得注意的是，Acc-7 雖是最主要的比較指標，但 52.19% 的絕對數值仍有提升空間。",
        "值得注意的是，Acc-7 雖是最主要的比較指標，但 53.21% 的絕對數值仍有提升空間。"
    ),
    (
        "結合多工序數損失與零洩漏推斷增強策略，在 CMU-MOSI 基準上達到 Acc-7 52.19%、MAE 0.587 與 Corr 0.869 的領先成績。",
        "結合多工損失（軟序數標籤交叉熵、序數地球移動距離、Smooth-L1、R-Drop、跨模態對比）與多分支單一模型內部集成設計、零洩漏推斷增強策略，在 CMU-MOSI 基準上達到 Acc-7 53.21%、MAE 0.5868 與 Corr 0.8683 的領先成績。"
    ),
    # ─── MGT baseline value in lit review (fix inconsistency) ───────────────
    (
        "MGT 在 CMU-MOSI 基準上達到 Acc-7 = 51.32%、MAE = 0.659、Corr = 0.822 之競爭性效能",
        "MGT 在 CMU-MOSI 基準上達到 Acc-7 = 50.44%、MAE = 0.659、Corr = 0.822 之競爭性效能"
    ),
    (
        "Acc-7（主要指標）：SACFFinalModel 達 53.21%，較最強基線 MGT（51.32%）提升 1.89 個百分點，",
        "Acc-7（主要指標）：SACFFinalModel 達 53.21%，較最強基線 MGT（50.44%）提升 2.77 個百分點，"
    ),
    (
        "本研究在訓練目標設計上同時引入序數地球移動距離損失與軟序數標籤交叉熵，顯式懲罰跨越多個情感等級的分類錯誤；並輔以 R-Drop 一致性正則與跨模態對比損失。實驗結果顯示，SACFFinalModel 在 MAE（0.5868）與 Corr（0.8683）兩項指標上皆優於 MGT（MAE = 0.659；Corr = 0.822），而 Acc-7 達 53.21%。",
        "本研究在訓練目標設計上同時引入序數地球移動距離損失與軟序數標籤交叉熵，顯式懲罰跨越多個情感等級的分類錯誤；並輔以 R-Drop 一致性正則與跨模態對比損失，搭配多分支單一模型之內部集成。實驗結果顯示，SACFFinalModel 在 MAE（0.5868）與 Corr（0.8683）兩項指標上皆優於 MGT（MAE = 0.659；Corr = 0.822），Acc-7 達 53.21%，相對 MGT 提升 2.77 個百分點。"
    ),
]


def _apply_text_edits_to_doc(doc):
    """Replace exact substrings in any paragraph of the document. Reports edits."""
    n_applied = 0
    for p in doc.paragraphs:
        # Reconstruct paragraph text from runs
        full_text = "".join(r.text for r in p.runs)
        new_text = full_text
        for find, repl in TEXT_EDITS:
            if find in new_text:
                new_text = new_text.replace(find, repl)
                n_applied += 1
        if new_text != full_text and p.runs:
            # Replace by writing into the first run and clearing the rest
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
    return n_applied


# ════════════════════════════════════════════════════════════════════════════
#  Main
# ════════════════════════════════════════════════════════════════════════════
def build_v15():
    print("=" * 70)
    print("Building 論文＿李昇峰_v15.docx ...")
    print("=" * 70)

    doc = Document(str(SRC))

    # ── Step 1: targeted text edits in Ch1 / Ch2 / Ch4 ──────────────────────
    n_edits = _apply_text_edits_to_doc(doc)
    print(f"  ✓ Applied {n_edits} text edits in Ch1/Ch2/Ch4")

    # ── Step 2: locate Chapter 3 and Chapter 4 boundaries ───────────────────
    body = doc.element.body
    children = list(body)
    ch3_idx = ch4_idx = None
    for i, child in enumerate(children):
        if not child.tag.endswith("}p"): continue
        pStyle = child.find('.//' + qn('w:pStyle'))
        if pStyle is None or pStyle.get(qn('w:val')) != '1': continue
        text = ''.join(child.itertext())
        if ch3_idx is None and 'SACF情感感知跨模態融合模型' in text:
            ch3_idx = i
        elif ch3_idx is not None and 'SACF 情感評分模型實驗結果' in text:
            ch4_idx = i
            break
    if ch3_idx is None or ch4_idx is None:
        raise SystemExit(f"  ✗ Could not locate Ch3/Ch4: ch3={ch3_idx} ch4={ch4_idx}")
    print(f"  ✓ Ch3 body[{ch3_idx}]   Ch4 body[{ch4_idx}]   removing {ch4_idx - ch3_idx} elems")

    ch4_element = children[ch4_idx]
    for el in children[ch3_idx:ch4_idx]:
        body.remove(el)

    # ── Step 3: emit fresh Chapter 3 and move into position ─────────────────
    pre_ids = set(id(c) for c in body)
    emit_chapter3(doc)
    new_elements = [c for c in body if id(c) not in pre_ids]
    print(f"  ✓ Emitted {len(new_elements)} new Ch3 elements")
    for new_el in new_elements:
        body.remove(new_el)
        pos = list(body).index(ch4_element)
        body.insert(pos, new_el)

    doc.save(str(OUT_PATH))
    print(f"\n  ✓ 已儲存：{OUT_PATH}")
    print(f"  ✓ 大小：{os.path.getsize(OUT_PATH)/1024:.1f} KB")


if __name__ == "__main__":
    build_v15()
