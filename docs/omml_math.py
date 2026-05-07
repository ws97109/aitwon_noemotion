"""OMML (Office MathML) builder for python-docx.

Word's native equation format is OMML. MathType edits/round-trips OMML elements,
so equations inserted as OMML are clickable and editable in MathType after the
file is opened in Word. This module provides a small DSL that converts Python
expressions into properly nested w:oMath elements.

Usage:
    from docx import Document
    from omml_math import add_equation, omath_inline
    doc = Document()
    p = doc.add_paragraph()
    add_equation(p, frac(num="x+1", den="x-1"))
    p.add_run("  for x ≠ 1.")
"""
from lxml import etree

W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NSMAP = {"m": M_NS, "w": W_NS}


def _e(tag, **attrs):
    """Create an m:* element."""
    el = etree.SubElement if False else etree.Element
    e = etree.Element(f"{{{M_NS}}}{tag}", nsmap={"m": M_NS})
    for k, v in attrs.items():
        e.set(f"{{{M_NS}}}{k}", str(v))
    return e


def _sub(parent, tag, **attrs):
    e = etree.SubElement(parent, f"{{{M_NS}}}{tag}")
    for k, v in attrs.items():
        e.set(f"{{{M_NS}}}{k}", str(v))
    return e


def _w_sub(parent, tag, text=None, **attrs):
    e = etree.SubElement(parent, f"{{{W_NS}}}{tag}")
    for k, v in attrs.items():
        e.set(f"{{{W_NS}}}{k}", str(v))
    if text is not None:
        e.text = text
    return e


# ─── Building blocks ──────────────────────────────────────────────────────────
def mr(text, italic=None):
    """An m:r run holding a literal math text token. Returns an XML element."""
    r = _e("r")
    if italic is not None:
        rPr = _sub(r, "rPr")
        sty = _sub(rPr, "sty")
        sty.set(f"{{{M_NS}}}val", "i" if italic else "p")
    t = _sub(r, "t")
    t.text = str(text)
    # preserve spaces
    t.set(f"{{http://www.w3.org/XML/1998/namespace}}space", "preserve")
    return r


def seq(*nodes):
    """Wrap a sequence of math elements into a single fragment via m:e content."""
    # used inside mfrac, etc.
    return list(nodes)


def normalize(node_or_str):
    """Convert a string/None into mr(...) element; passthrough for Element."""
    if node_or_str is None:
        return None
    if isinstance(node_or_str, str):
        return mr(node_or_str)
    if isinstance(node_or_str, list):
        return node_or_str
    return node_or_str


def frac(num, den):
    """m:f fraction. num/den can be strings or pre-built elements/lists."""
    f = _e("f")
    fPr = _sub(f, "fPr")
    _sub(fPr, "type").set(f"{{{M_NS}}}val", "bar")
    n = _sub(f, "num")
    d = _sub(f, "den")
    _attach(n, num)
    _attach(d, den)
    return f


def sup(base, exp):
    """m:sSup superscript."""
    s = _e("sSup")
    e = _sub(s, "e")
    sub = _sub(s, "sup")
    _attach(e, base); _attach(sub, exp)
    return s


def sub(base, sb):
    """m:sSub subscript."""
    s = _e("sSub")
    e = _sub(s, "e")
    sb_el = _sub(s, "sub")
    _attach(e, base); _attach(sb_el, sb)
    return s


def subsup(base, sb, sp):
    """m:sSubSup base with both subscript and superscript."""
    s = _e("sSubSup")
    e = _sub(s, "e")
    sb_el = _sub(s, "sub")
    sp_el = _sub(s, "sup")
    _attach(e, base); _attach(sb_el, sb); _attach(sp_el, sp)
    return s


def rad(rad_under, deg=None):
    """m:rad radical (square root by default)."""
    r = _e("rad")
    rPr = _sub(r, "radPr")
    if deg is None:
        _sub(rPr, "degHide").set(f"{{{M_NS}}}val", "1")
    deg_el = _sub(r, "deg")
    if deg is not None:
        _attach(deg_el, deg)
    e_el = _sub(r, "e")
    _attach(e_el, rad_under)
    return r


def nary(op, lower=None, upper=None, body=None):
    """m:nary big operator (sum/prod/integral etc.)."""
    n = _e("nary")
    nPr = _sub(n, "naryPr")
    chr_el = _sub(nPr, "chr"); chr_el.set(f"{{{M_NS}}}val", op)
    sub_el = _sub(n, "sub")
    sup_el = _sub(n, "sup")
    e_el = _sub(n, "e")
    if lower is not None: _attach(sub_el, lower)
    if upper is not None: _attach(upper_el := sup_el, upper)
    if body is not None: _attach(e_el, body)
    return n


def func(fname, body):
    """m:func — named function like sin, log, softmax."""
    f = _e("func")
    fName = _sub(f, "fName")
    _attach(fName, fname)
    e = _sub(f, "e")
    _attach(e, body)
    return f


def paren(body, lparen="(", rparen=")"):
    """m:d delimiter (parenthesis-like)."""
    d = _e("d")
    dPr = _sub(d, "dPr")
    if lparen != "(": _sub(dPr, "begChr").set(f"{{{M_NS}}}val", lparen)
    if rparen != ")": _sub(dPr, "endChr").set(f"{{{M_NS}}}val", rparen)
    e = _sub(d, "e")
    _attach(e, body)
    return d


def matrix(rows, lparen="[", rparen="]"):
    """Bracketed matrix from a list-of-lists of math tokens."""
    d = _e("d")
    dPr = _sub(d, "dPr")
    _sub(dPr, "begChr").set(f"{{{M_NS}}}val", lparen)
    _sub(dPr, "endChr").set(f"{{{M_NS}}}val", rparen)
    e = _sub(d, "e")
    m = _sub(e, "m")
    mPr = _sub(m, "mPr")
    _sub(mPr, "mcs").set(f"{{{M_NS}}}val", "")
    for row in rows:
        mr_el = _sub(m, "mr")
        for cell in row:
            e_cell = _sub(mr_el, "e")
            _attach(e_cell, cell)
    return d


def _attach(parent, content):
    """Append content (string / Element / list) to parent."""
    if content is None: return
    if isinstance(content, str):
        parent.append(mr(content))
    elif isinstance(content, list):
        for c in content:
            if isinstance(c, str): parent.append(mr(c))
            elif c is not None: parent.append(c)
    else:
        parent.append(content)


# ─── High-level inline / display equations ────────────────────────────────────
def omath_inline(*nodes):
    """Wrap nodes into an inline m:oMath element (returned as XML)."""
    om = etree.Element(f"{{{M_NS}}}oMath", nsmap={"m": M_NS})
    for n in nodes:
        if isinstance(n, str): om.append(mr(n))
        elif isinstance(n, list):
            for nn in n:
                if isinstance(nn, str): om.append(mr(nn))
                else: om.append(nn)
        else:
            om.append(n)
    return om


def omath_para(*nodes):
    """Wrap nodes into a centered m:oMathPara element (display equation)."""
    omp = etree.Element(f"{{{M_NS}}}oMathPara", nsmap={"m": M_NS, "w": W_NS})
    omp_pr = _sub(omp, "oMathParaPr")
    j = _sub(omp_pr, "jc"); j.set(f"{{{M_NS}}}val", "center")
    om = etree.SubElement(omp, f"{{{M_NS}}}oMath")
    for n in nodes:
        if isinstance(n, str): om.append(mr(n))
        elif isinstance(n, list):
            for nn in n:
                if isinstance(nn, str): om.append(mr(nn))
                else: om.append(nn)
        else:
            om.append(n)
    return omp


def add_display_equation(doc, *nodes):
    """Insert a centered display equation as a NEW paragraph in `doc`."""
    p = doc.add_paragraph()
    p.alignment = 1  # center
    omp = omath_para(*nodes)
    # python-docx paragraphs can hold m:oMathPara as a child of w:p
    p._p.append(omp)
    return p


def add_inline_equation(paragraph, *nodes):
    """Insert an inline m:oMath element into an existing paragraph."""
    om = omath_inline(*nodes)
    paragraph._p.append(om)
    return paragraph
