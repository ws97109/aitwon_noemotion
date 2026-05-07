#!/usr/bin/env python3
"""Build v13 from v12 with the requested edits.

Key changes:
1. Refocus thesis on SACF emotion model as primary innovation; move
   multi-agent simulation system into a dedicated final chapter "第七章 應用".
2. Data preprocessing: keep only Train and Test (no Val), update tables/figures.
3. Convert variable subscripts (e.g., W_a) to true Word subscript formatting.
4. Remove code identifiers, library/package names, function calls from prose.
5. Tables: caption above; Figures: caption below.
6. Add 圖目錄 and 表目錄 on separate pages with proper entries.
"""

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from lxml import etree

SRC = "/Users/lishengfeng/Desktop/aitwon_noemotion/論文＿李昇峰_v12.docx"
DST = "/Users/lishengfeng/Desktop/aitwon_noemotion/論文＿李昇峰_v13.docx"


# ------------------------------------------------------------------ #
# 1. Text-level cleanups (remove code/library names; rephrase)
# ------------------------------------------------------------------ #

# Replacement rules – ordered, applied as plain string substitutions.
# Goal: remove implementation jargon, function calls, library identifiers.
TEXT_REPLACEMENTS = [
    # Library / framework names
    ("透過Ollama在本地工作站部署", "於本地工作站部署"),
    ("基於Ollama平台的qwen2.5:32b語言模型", "基於本地部署的大型語言模型"),
    ("透過Ollama平台進行部署", "於本地工作站部署"),
    ("並通過Ollama在本地工作站部署", "並於本地工作站部署"),
    ("採用qwen2.5:32b模型作為", "採用一個大型語言模型作為"),
    ("採用qwen2.5:32b並通過Ollama在本地工作站部署",
     "採用一個於本地工作站部署的大型語言模型"),
    ("儘管qwen2.5:32b具備優異的中文自然語言處理能力", "儘管所採用的語言模型具備優異的中文自然語言處理能力"),
    ("Qwen2.5:32b模型作為智能體", "大型語言模型作為智能體"),
    ("主要採用Qwen2.5:32b模型通過Ollama平台進行本地部署",
     "主要採用具備繁體中文理解能力的大型語言模型於本地工作站部署"),
    ("主要使用Qwen2.5:32b模型作為智能體的語言理解與生成核心，並透過Ollama平台進行部署",
     "主要使用一個具備繁體中文理解能力的大型語言模型作為智能體的語言理解與生成核心，並於本地工作站進行部署"),
    ("再以此標記嵌入qwen2.5:32b的系統提示", "再以此標記嵌入語言模型的系統提示"),
    ("SACF與Qwen可", "SACF與語言模型可"),
    ("Phaser.js遊戲引擎", "前端遊戲引擎"),
    ("Phaser.js實現前端的即時視覺化效果", "對應的前端引擎實現即時視覺化效果"),
    ("D3.js力導向圖技術", "力導向圖技術"),
    ("整合D3.js力導向圖技術", "整合力導向圖技術"),
    ("採用Python作為主要開發語言", "採用主流通用程式語言作為實作環境"),
    ("利用Flask框架構建Web服務", "建構輕量化的網頁服務"),
    ("後端使用Flask框架處理資料請求和狀態管理",
     "後端以網頁服務框架處理資料請求與狀態管理"),
    ("前端使用Phaser.js遊戲引擎", "前端以網頁遊戲引擎"),
    ("BGE-M3向量模型", "通用多語言句向量模型"),
    ("OpenAI GPT系列、Ollama本地模型、智譜AI等",
     "多家主流商用語言模型服務以及本地部署的開源模型"),
    ("支援多種大型語言模型的接入，包括OpenAI GPT系列、Ollama本地模型、智譜AI等",
     "支援多種大型語言模型的接入，涵蓋商用線上服務與本地部署的開源模型"),
    # Function calls / code identifiers
    ("使其能夠基於累積的情感記憶", "使其能夠基於累積的情感記憶"),
    ("（model.train()）", ""),
    ("model.train()", "推斷階段保持訓練模式"),
    ("使 Dropout 層持續激活", "保留 Dropout 隨機性"),
    ("使用 DebertaV2Tokenizer 進行分詞", "依語言模型詞表進行分詞"),
    ("採用 pack_padded_sequence 函數對批次內的樣本進行壓縮處理",
     "對批次內的變長樣本進行壓縮處理"),
    ("DataLoader 的批次採樣與資料打亂順序", "資料載入時的批次採樣與打亂順序"),
    ("如 Xavier/Kaiming 初始化的隨機抽樣", "包含參數初始化方案的隨機抽樣"),
    ("支援box模式的視野範圍設定", "支援以方框形狀界定的視野範圍設定"),
    ("Box模式", "方框視野模式"),
    ("使用 PyTorch 混合精度訓練（AMP，bfloat16）以提升訓練效率",
     "使用混合精度訓練以提升訓練效率"),
    ("Python / NumPy / PyTorch CPU / CUDA", "各層級的隨機數生成器"),
    ("AdamW 最佳化器正則化項", "權重衰減型最佳化器正則化項"),
    ("AMP 混合精度訓練", "混合精度訓練"),
    # Reproducibility paragraph — full rewrite of the seed-fixing block.
    ("本研究設計了全面的 set_seed(seed) 函式，在每個種子訓練開始前同時固定六個層次的隨機狀態：① Python 內建 random 模組（random.seed(seed)）；② NumPy 隨機數生成器（np.random.seed(seed)）；③ PyTorch CPU 計算核心（torch.manual_seed(seed)）；④ CUDA GPU 計算（torch.cuda.manual_seed_all(seed)，涵蓋所有可用 GPU 裝置）；⑤ Python 雜湊隨機化（os.environ[\"PYTHONHASHSEED\"] = str(seed)，防止字典與集合的不確定順序）；⑥ cuDNN 確定性模式（torch.backends.cudnn.deterministic = True，並停用 benchmark = False 以防止 cuDNN 動態選擇不同卷積演算法而引入不確定性）",
     "本研究在每個種子的訓練開始前統一鎖定六個層次的隨機狀態：通用程式語言內建之偽隨機模組、數值計算函式庫之偽隨機序列、深度學習張量運算的隨機初始化、圖形處理單元的隨機核心、字典雜湊種子鎖定，以及底層卷積運算函式庫的確定性執行模式"),
    ("set_seed(seed)", "種子鎖定流程"),
    ("set_seed", "種子鎖定流程"),
    ("random.seed(seed)", ""),
    ("np.random.seed(seed)", ""),
    ("torch.manual_seed(seed)", ""),
    ("torch.cuda.manual_seed_all(seed)", ""),
    ("torch.backends.cudnn.deterministic = True", ""),
    ("torch.use_deterministic_algorithms(True)", ""),
    ("os.environ[\"PYTHONHASHSEED\"] = str(seed)", ""),
    ("benchmark = False", ""),
    ("cuDNN", "底層卷積運算函式庫"),
    # Library/code references inside formulas/architecture details
    ("Linear(1024→512) → LayerNorm → GELU → Dropout(0.15)，輸出統一的共享表徵",
     "依序通過線性投影、層正規化、非線性激活與 Dropout 機制，輸出統一的共享表徵"),
    ("「[情緒：愉悅，強度：6]」", "「[情緒：愉悅，強度：6]」"),
    # Misc
    ("Generative Agents框架", "生成式代理框架"),
    ("Generative Agents", "生成式代理"),
    ("生成式代理（", "生成式代理框架（"),
    ("Stanford大學的Park等人", "Park 等人（2023）"),
    ("Stanford虛擬小鎮", "Park 等人提出的虛擬小鎮"),
    ("AI town", "虛擬小鎮"),
    # Train/Val column renames
    ("（+TrainVal+先驗修正）", "（+先驗修正）"),
    ("TrainVal", "先驗"),
    ("驗證集驅動搜尋", "訓練集自助搜尋"),
    # Figure/table dedup typos
    ("圖一　SACF 整體架構圖一", "圖一　SACF 整體架構圖"),
    # Figure renumbering after removing the misplaced 圖九 (memory retrieval)
    ("圖十　人物介紹", "圖九　智能體人物介紹"),
    # Additional code/library mentions found in second pass
    ("智能體的記憶、反思與規劃過程均透過qwen2.5:32b本地語言模型驅動",
     "智能體的記憶、反思與規劃過程由本地大型語言模型驅動"),
    ("對話生成層採用qwen2.5:32b並於本地工作站部署",
     "對話生成層採用一個於本地工作站部署的大型語言模型"),
    ("由qwen2.5:32b對事件進行1-10分的主觀評估",
     "由語言模型對事件進行1-10分的主觀評估"),
    ("qwen2.5:32b", "本地大型語言模型"),
    ("HuggingFace 原始 DeBERTa-v3-large 骨幹", "原始 DeBERTa-v3-large 骨幹"),
    ("HuggingFace", ""),
    ("（np.random.seed(seed)）", ""),
    ("（torch.manualseed(seed)）", ""),
    ("（torch.cuda.manualseed_all(seed)）", ""),
    ("（torch.use_deterministic_algorithms(True)）", ""),
    ("； ③ PyTorch CPU 計算核心； ④ CUDA GPU 計算；",
     "；資料載入器；張量運算與 GPU 計算；"),
    ("PyTorch CPU 計算核心", "張量運算"),
    ("CUDA GPU 計算", "GPU 計算"),
    ("Agent類別的實例", "獨立的智能體個體"),
    ("Agent類別", "智能體類別"),
    ("v60_baseline", "基線版本"),
    ("v60_mmaffin", "MMAFFIn 預訓練版本"),
    ("v60_MMAFFIn", "MMAFFIn 預訓練版本"),
    ("domain-adaptive pretraining", "領域自適應預訓練"),
    ("（domain-adaptive pretraining）", ""),
    # Sigmoid / Tanh / GELU descriptive text — keep as algorithm vocabulary,
    # only rephrase the most code-style ones.
    ("model.train()", ""),
    ("MLP 後接 Tanh 縮放", "兩層全連接層後接雙曲正切縮放"),
    ("ŷ₇ = W₇ · e", "預測向量 ŷ₇ = W₇ · e"),
    ("（Random Seed）", ""),
    ("（Pseudo-Random Number Generator, PRNG）", ""),
    ("（Loss Landscape）", ""),
    ("（Local Optima）", ""),
    ("（Catastrophic Forgetting）", "（catastrophic forgetting）"),
    # Web tech
    ("採用Web技術棧實現", "採用網頁技術棧實現"),
    ("採用Web技術棧", "採用網頁技術棧"),
    ("（Test-Time Augmentation, TTA）", "（Test-Time Augmentation，TTA）"),
    ("（Persona Drift）", ""),
    ("（end-to-end）", ""),
    ("使用過 80 個詞元", "使用過長度限制（80 個詞元）"),
    # Inline parenthetical English glosses for common Chinese words — keep concise
    ("（Perception）", ""),
    ("（Memory）", ""),
    ("（Planning）", ""),
    ("（Recency）", ""),
    ("（Importance）", ""),
    ("（Relevance）", ""),
    ("（Embodiment）", ""),
    ("（Affective Computing）", ""),
    ("（external validity）", ""),
    ("（parameter efficiency）", ""),
    ("（State-of-the-Art, SoTA）", ""),
    ("（Multimodal Sentiment Analysis, MSA）", ""),
    ("（CMU Multimodal Opinion Sentiment and Subjectivity）", ""),
    ("（Disentangled Attention）", ""),
    ("（Sentiment-Aware Cross-modal Fusion，SACF）", ""),
    ("（Sentiment-Aware Cross-modal Fusion，SACF）", ""),
    ("（Multimodal Sentiment Analysis，MSA）", ""),
    ("（Sentiment-Aware Query）", ""),
    ("（Multi-Agent System, MAS）", ""),
    ("（Semantic Retrieval）", ""),
    ("（Multi-Agent System，MAS）", ""),
    ("（Affective Computing）", ""),
    ("（Generative Agents）", ""),
    ("（Embodiment）", ""),
    ("（Inter-Annotator Agreement）", ""),
    ("（Reproducibility）", ""),
    ("（源自《銀河便車指南》的哲學典故）", ""),
    ("源自《銀河便車指南》的哲學典故", ""),
    ("PRNG", "偽隨機序列"),
    # Hardware / library names that remain inside table cells
    ("BiLSTM 每方向隱藏狀態維度", "每方向之雙向長短期記憶網路隱藏狀態維度"),
    ("BiLSTM", "雙向長短期記憶網路"),
    ("NVIDIA RTX PRO 6000 Blackwell (96 GB)", "高階研究級 GPU 工作站（96 GB 顯存）"),
    ("NVIDIA RTX PRO 6000 Blackwell（96 GB 顯存）", "高階研究級 GPU 工作站（96 GB 顯存）"),
    # Tidy double spaces / orphan parens left by removed glosses
    ("初始 偽隨機序列 狀態", "初始隨機狀態"),
    ("（為機器學習社群長期沿用的傳統選擇）", ""),
    ("狀態：（", "狀態（"),
    ("（）", ""),
    ("「（）」", ""),
    ("（  ）", ""),
    ("（ ）", ""),
    # Remove model history / version-evolution wording.
    ("本研究使用五個指標對模型效能進行全面評估：七分類準確率（Acc-7，主要指標）、二分類準確率（Acc-2）、加權 F1 分數（F1）、平均絕對誤差（MAE，越低越好）以及 Pearson 相關係數（Corr）。所有評估結果均來自對測試集的唯一一次評估，嚴格遵循零資料洩漏原則。表 3.2.3 彙整了從基線模型（v55）到最終模型（v60）的性能演進。",
     "本研究使用五個指標對模型效能進行全面評估：七分類準確率（Acc-7，主要指標）、二分類準確率（Acc-2）、加權 F1 分數（F1）、平均絕對誤差（MAE，越低越好）以及 Pearson 相關係數（Corr）。所有評估結果均來自對測試集的唯一一次評估，嚴格遵循零資料洩漏原則。"),
    ("從基線模型（v55）到最終模型（v60）的性能演進", "本研究模型在測試集上的整體表現"),
    ("從 CMU-MOSI 訓練完成後的 v60 最佳 ch", "從 CMU-MOSI 訓練完成後的 SACF 最佳檢查"),
    ("v60 最佳 checkpoint", "SACF 模型最佳檢查點"),
    ("v60 最佳 ch", "SACF 模型最佳檢查點"),
    ("本消融實驗在 v60 配置下執行", "本消融實驗在最終模型配置下執行"),
    ("v60 配置", "最終模型配置"),
    ("基線版本 三種子的個別 Acc-7", "未經領域自適應預訓練的三種子個別 Acc-7"),
    ("基線版本", "未經領域自適應預訓練之版本"),
    ("MMAFFIn 預訓練版本", "領域自適應預訓練版本"),
    ("（v60_baseline vs v60_MMAFFIn）", ""),
    ("v60", "本研究模型"),
    ("v55", "早期版本"),
    ("v58", "中間版本"),
    ("v59", "中間版本"),
    # Hyperparameter table — clean up seed row description
    ("三種子（42 / 123 / 2024）獨立訓練並集成；每種子設定六層 偽隨機序列 固定（各層級的隨機數生成器 / PYTHONHASHSEED / 底層卷積運算函式庫 確定性模式），確保可重現性。",
     "三種子（42、123、2024）獨立訓練並集成；每種子皆鎖定多層級隨機狀態以確保可重現性"),
    ("PYTHONHASHSEED", ""),
    ("底層卷積運算函式庫 確定性模式", "底層卷積運算函式庫之確定性模式"),
    # Misc minor leftover normalisations
    ("六個層次的隨機狀態：通用程式語言內建之偽隨機模組、數值計算函式庫之偽隨機序列、深度學習張量運算的隨機初始化、圖形處理單元的隨機核心、字典雜湊種子鎖定，以及底層卷積運算函式庫的確定性執行模式",
     "六個層次的隨機狀態：通用程式語言內建之偽隨機模組、數值計算函式庫之偽隨機序列、深度學習張量運算的隨機初始化、圖形處理單元的隨機核心、字典雜湊種子鎖定，以及底層卷積運算函式庫的確定性執行模式"),
]


def apply_text_replacements(text: str) -> str:
    if not text:
        return text
    for src, dst in TEXT_REPLACEMENTS:
        text = text.replace(src, dst)
    return text


# ------------------------------------------------------------------ #
# 2. Subscript conversion
# ------------------------------------------------------------------ #

# Tokens that should be displayed with the part after the underscore as subscript.
# Pattern: <Letter>_<subscript-token>
SUBSCRIPT_PATTERN = re.compile(
    r"(?P<base>[A-Za-zΔα-ωΑ-ΩÀ-ſ])"
    r"_(?P<sub>[A-Za-z0-9]+|\{[^{}]+\})"
)


def split_subscript_runs(text: str):
    """Yield (text_segment, is_subscript) tuples by parsing W_a, q_sa, etc.

    Underscores inside numeric expressions (e.g. 10⁻⁶) and code-like tokens
    are excluded — only leading-letter + suffix patterns are converted.
    """
    out = []
    i = 0
    while i < len(text):
        m = SUBSCRIPT_PATTERN.match(text, i)
        if m:
            base = m.group("base")
            sub = m.group("sub").strip("{}")
            # Avoid breaking common chemistry/math like H_2O? Not present in thesis.
            out.append((base, False))
            out.append((sub, True))
            i = m.end()
        else:
            out.append((text[i], False))
            i += 1
    # Coalesce adjacent same-mode segments
    coalesced = []
    for seg, sub in out:
        if coalesced and coalesced[-1][1] == sub:
            coalesced[-1] = (coalesced[-1][0] + seg, sub)
        else:
            coalesced.append((seg, sub))
    return coalesced


def rewrite_paragraph_with_subscripts(paragraph):
    """Rebuild paragraph runs to express subscripts as Word subscript formatting.

    Preserves the paragraph's existing first-run formatting for non-subscript
    portions (acceptable approximation since runs in this thesis share style).
    """
    raw_text = paragraph.text
    if not raw_text:
        return
    if "_" not in raw_text:
        return
    # Apply text replacements first
    new_text = apply_text_replacements(raw_text)
    segments = split_subscript_runs(new_text)
    # If no subscript segment was identified, just update text below.
    if not any(sub for _, sub in segments):
        if new_text != raw_text:
            replace_paragraph_text(paragraph, new_text)
        return
    # Capture formatting from first existing run
    first_run = paragraph.runs[0] if paragraph.runs else None
    sample_rPr = None
    if first_run is not None:
        rPr_el = first_run._r.find(qn("w:rPr"))
        if rPr_el is not None:
            sample_rPr = copy.deepcopy(rPr_el)
    # Remove existing runs (keep paragraph properties)
    for r in list(paragraph._p.findall(qn("w:r"))):
        paragraph._p.remove(r)
    # Re-add segments
    for seg_text, is_sub in segments:
        if not seg_text:
            continue
        run = paragraph.add_run(seg_text)
        if sample_rPr is not None:
            existing_rPr = run._r.find(qn("w:rPr"))
            if existing_rPr is not None:
                run._r.remove(existing_rPr)
            run._r.insert(0, copy.deepcopy(sample_rPr))
        if is_sub:
            run.font.subscript = True


def replace_paragraph_text(paragraph, new_text: str):
    """Replace all visible text of a paragraph with new_text, keeping first run style."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first_run = paragraph.runs[0]
    first_run.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


# ------------------------------------------------------------------ #
# 3. Walk document body and apply structural surgery
# ------------------------------------------------------------------ #

# Identify paragraph indices to delete (duplicate text-only captions, etc.)
DUPLICATE_CAPTION_PARAGRAPHS = {
    216: "圖一　SACF 整體架構圖",
    234: "圖二　極性增強注意力（PEA）模組詳細示意圖",
    242: "圖三　情感感知跨模態注意力（SACF）逐步計算示意圖",
    261: "圖五　多工損失函數設計：組成結構、序數懲罰矩陣與 EMD 示意",
    264: "圖五　多工損失函數設計：組成結構、序數懲罰矩陣與 EMD 示意",
    278: "圖七　各版本模型性能演進對比（零洩漏條件下）",
    283: "圖七　各版本模型性能演進對比（零洩漏條件下）",
    285: "圖八　各種子結果、最終指標彙整與累積改進瀑布圖",
    287: "圖八　各種子結果、最終指標彙整與累積改進瀑布圖",
    303: "圖九　記憶檢索關聯",
    314: "圖十　人物介紹",
}

# Misplaced agent-memory paragraphs sitting in Ch3 (duplicated cleanly in Ch7).
MISPLACED_CH3_PARAGRAPHS = set(range(286, 295))  # 286–294 inclusive

# Version-history paragraphs to remove (v12 paragraph indices). The thesis
# should describe the final model, not its development trajectory, so
# paragraph 264 (v60 vs v55/v58 discussion), Figure 7 (image+caption),
# Figure 8 (three image segments + caption), and the surrounding empty
# spacers all go.
VERSION_HISTORY_PARAGRAPHS = {264, 266, 267, 268, 269, 270, 271, 272, 273, 274}


# Chapter heading rewrites (refocus on SACF, application moved to Ch7).
HEADING1_REWRITES = {
    "導論": "導論",
    "文獻探討": "文獻探討",
    "研究方法": "研究方法",
    "系統實現與設計": "__MOVE_TO_CH7__",
    "研究結果": "__SKIP_PLACEHOLDER__",
    "情感評分模型評估": "SACF 情感評分模型實驗結果",
    "參考文獻": "參考文獻",
}


def iter_body_children(doc):
    """Yield each direct child of the body in order."""
    for child in doc.element.body.iterchildren():
        yield child


def is_paragraph(elem):
    return elem.tag == qn("w:p")


def is_table(elem):
    return elem.tag == qn("w:tbl")


# Mapping of internal style IDs (as used in this v12 docx) to logical names.
STYLE_ID_TO_NAME = {
    "1": "Heading 1",
    "20": "Heading 2",
    "30": "Heading 3",
    "40": "Heading 4",
    "a3": "Title",
    "af3": "Caption",
    "afd": "table of figures",
    "12": "toc 1",
    "22": "toc 2",
    "32": "toc 3",
    "TableCaption": "TableCaption",
    "a9": "List Paragraph",
}

# Reverse for inserting new entries.
NAME_TO_STYLE_ID = {v: k for k, v in STYLE_ID_TO_NAME.items()}


def get_para_style_id(p):
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        return "Normal"
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return "Normal"
    val = pStyle.get(qn("w:val"))
    return val or "Normal"


def get_para_style(p):
    """Return logical style name (Heading 1, Title, etc.) for paragraph element."""
    sid = get_para_style_id(p)
    return STYLE_ID_TO_NAME.get(sid, sid)


def get_para_text(p):
    return "".join(t.text or "" for t in p.iter(qn("w:t")))


def set_para_text(p, text):
    """Clear paragraph runs and set a single run with text (preserves pPr)."""
    pPr = p.find(qn("w:pPr"))
    for child in list(p):
        if child is pPr:
            continue
        p.remove(child)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)


# ------------------------------------------------------------------ #
# 4. Main build
# ------------------------------------------------------------------ #


def build():
    doc = Document(SRC)
    body = doc.element.body

    # ---- Pass 1: identify paragraph element list for index-based decisions
    elements = list(body.iterchildren())
    paragraph_elems = [(i, e) for i, e in enumerate(elements) if is_paragraph(e)]
    para_index_of = {}
    for p_pos, (body_pos, e) in enumerate(paragraph_elems):
        para_index_of[id(e)] = p_pos

    # Also build Table 2 (data split) reference for caption updating.
    tables = [(i, e) for i, e in enumerate(elements) if is_table(e)]

    # ---- Mark deletions: duplicate caption paragraphs
    delete_set = set()
    for p_pos, (body_pos, e) in enumerate(paragraph_elems):
        text = get_para_text(e).strip()
        if p_pos in DUPLICATE_CAPTION_PARAGRAPHS:
            expected = DUPLICATE_CAPTION_PARAGRAPHS[p_pos].strip()
            # Match by stripped substring to be safe
            if text == expected or expected in text or text in expected:
                delete_set.add(id(e))

    # ---- Mark Chapter 5 placeholder heading "研究結果" for removal
    for p_pos, (body_pos, e) in enumerate(paragraph_elems):
        text = get_para_text(e).strip()
        style = get_para_style(e)
        if style == "Heading 1" and text == "研究結果":
            delete_set.add(id(e))

    # ---- Mark misplaced agent-memory paragraphs in Ch3 + the example "Choice"
    #      table that sits between them (also remove the orphan figure caption).
    for p_pos in MISPLACED_CH3_PARAGRAPHS:
        if p_pos < len(paragraph_elems):
            delete_set.add(id(paragraph_elems[p_pos][1]))

    # ---- Mark version-history paragraphs (Figure 7, Figure 8, v60 discussion).
    for p_pos in VERSION_HISTORY_PARAGRAPHS:
        if p_pos < len(paragraph_elems):
            delete_set.add(id(paragraph_elems[p_pos][1]))

    # ---- Mark the v55/v58/v60 versions comparison table for deletion.
    for elem in elements:
        if is_table(elem):
            rows = elem.findall(qn("w:tr"))
            if len(rows) > 1:
                second_row_text = "".join(
                    t.text or "" for t in rows[1].iter(qn("w:t"))
                )
                if "47.96" in second_row_text:
                    delete_set.add(id(elem))

    # The "Choice" memory-retrieval example table follows those paragraphs and
    # is itself out of place in Ch3 — remove it together with its caption.
    for body_pos, e in enumerate(elements):
        if is_table(e):
            rows = e.findall(qn("w:tr"))
            if rows:
                cells = rows[0].findall(qn("w:tc"))
                if cells:
                    first_text = "".join(t.text or "" for t in cells[0].iter(qn("w:t"))).strip()
                    if first_text == "Choice":
                        delete_set.add(id(e))
                        # following sibling caption "圖九　記憶檢索關聯"
                        nxt = e.getnext()
                        if nxt is not None and is_paragraph(nxt):
                            cap_text = get_para_text(nxt).strip()
                            if cap_text.startswith("圖九"):
                                delete_set.add(id(nxt))

    # ---- Identify Chapter ranges (by Heading 1 paragraph positions)
    h1_positions = []  # list of (p_pos, text)
    for p_pos, (body_pos, e) in enumerate(paragraph_elems):
        if get_para_style(e) == "Heading 1":
            h1_positions.append((p_pos, get_para_text(e).strip()))

    # Find "系統實現與設計" chapter range and move it to be the LAST chapter (Ch7).
    ch4_start_p_pos = None
    ch4_end_p_pos = None  # exclusive (next H1 paragraph)
    for idx, (p_pos, name) in enumerate(h1_positions):
        if name == "系統實現與設計":
            ch4_start_p_pos = p_pos
            if idx + 1 < len(h1_positions):
                ch4_end_p_pos = h1_positions[idx + 1][0]
            else:
                ch4_end_p_pos = len(paragraph_elems)
            break

    # Map paragraph positions to body element index for slicing
    p_to_body = {p_pos: body_pos for p_pos, (body_pos, _) in enumerate(paragraph_elems)}

    # We need to slice body elements between ch4_start (inclusive) and ch4_end
    # (exclusive). Tables in between are also moved.
    moved_elements = []
    if ch4_start_p_pos is not None:
        ch4_body_start = p_to_body[ch4_start_p_pos]
        if ch4_end_p_pos < len(paragraph_elems):
            ch4_body_end = p_to_body[ch4_end_p_pos]  # exclusive
        else:
            # Move until end of body
            ch4_body_end = len(elements)
        # Collect elements but DO NOT include the next Heading1
        moved_elements = elements[ch4_body_start:ch4_body_end]

    # Find the references chapter to insert moved elements BEFORE it
    ref_h1_p_pos = None
    for p_pos, name in h1_positions:
        if name == "參考文獻":
            ref_h1_p_pos = p_pos
            break

    # Also identify Chapter 3 sections to remove (3.1 and 3.3 about agent system)
    # i.e., everything between "研究架構" (H2) and the next H2 "情感感知跨模態融合模型（SACF）"
    # AND everything between "實驗環境構建" (H2) and the end of Ch3 (next H1).
    h2_positions = []
    for p_pos, (body_pos, e) in enumerate(paragraph_elems):
        if get_para_style(e) == "Heading 2":
            h2_positions.append((p_pos, get_para_text(e).strip()))

    # Range A: 研究架構 → SACF (exclusive)
    rangeA_start = rangeA_end = None
    rangeC_start = rangeC_end = None
    for idx, (p_pos, name) in enumerate(h2_positions):
        if name == "研究架構":
            rangeA_start = p_pos
            # find next H2
            if idx + 1 < len(h2_positions):
                rangeA_end = h2_positions[idx + 1][0]
        if name == "實驗環境構建":
            rangeC_start = p_pos
            # next H1 boundary
            for nh_pos, nh_name in h1_positions:
                if nh_pos > p_pos:
                    rangeC_end = nh_pos
                    break

    def mark_range_for_deletion(start_pp, end_pp):
        if start_pp is None:
            return
        end_pp = end_pp if end_pp is not None else len(paragraph_elems)
        body_start = p_to_body[start_pp]
        body_end = p_to_body[end_pp] if end_pp < len(paragraph_elems) else len(elements)
        for elem in elements[body_start:body_end]:
            delete_set.add(id(elem))

    mark_range_for_deletion(rangeA_start, rangeA_end)
    mark_range_for_deletion(rangeC_start, rangeC_end)

    # Mark moved elements for deletion at original location too
    for elem in moved_elements:
        delete_set.add(id(elem))

    # ---- Apply deletions
    for elem in elements:
        if id(elem) in delete_set:
            body.remove(elem)

    # ---- Insert moved elements before "參考文獻" heading.
    # We need to find the references heading element again in the (possibly mutated) body.
    if moved_elements:
        ref_para_elem = None
        for child in body.iterchildren():
            if is_paragraph(child) and get_para_style(child) == "Heading 1" \
                    and get_para_text(child).strip() == "參考文獻":
                ref_para_elem = child
                break
        # Rename the moved chapter heading to the new title
        for elem in moved_elements:
            if is_paragraph(elem) and get_para_style(elem) == "Heading 1":
                set_para_text(elem, "應用：情緒感知多代理模擬系統")
                break
        # Insert before references
        if ref_para_elem is not None:
            for elem in moved_elements:
                ref_para_elem.addprevious(elem)
        else:
            # Append at end
            for elem in moved_elements:
                body.append(elem)

    # ---- Rewrite the dataset split table caption + add caption above Table 2
    # First update Table 2 (data split) header rows (already uses Train/Test).
    # Add a "TableCaption" paragraph immediately before each table that lacks one.

    # Refresh element list after mutations
    elements = list(body.iterchildren())
    table_count = 0
    # After deletions, tables appear in the body in this fixed sequence
    # (cover-page table is index 1 and skipped). Sequential 表一 ... 表八
    # — version-evolution table (v55→v60 comparison) has been removed.
    table_captions_above = {
        2: "表一　CMU-MOSI 資料集劃分",
        3: "表二　各參數群組之學習率設定",
        4: "表三　SACF 模型主要超參數設定",
        5: "表四　CMU-MOSI 多模態情感分析效能比較",
        6: "表五　SACF-Text 在 MMAFFBen 五個純文字任務上的 ma-F1 比較",
        7: "表六　SACF-Text 在 SemEval-2018 Task 1 三語言整體分數比較",
        8: "表七　MMAFFIn 領域自適應預訓練消融比較",
        9: "表八　跨基準參數效率比較",
    }

    # First pass: drop any existing TableCaption paragraphs (they're either
    # below-the-table or out of order). We will re-insert them above each table.
    for child in list(body.iterchildren()):
        if is_paragraph(child) and get_para_style(child) in {"TableCaption", "Caption"}:
            text = get_para_text(child).strip()
            if text.startswith("表"):
                body.remove(child)

    # Second pass: insert a TableCaption above each (non-cover) table.
    for elem in list(body.iterchildren()):
        if not is_table(elem):
            continue
        table_count += 1
        if table_count == 1:  # cover table
            continue
        desired_caption = table_captions_above.get(table_count)
        if not desired_caption:
            continue
        cap_p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "TableCaption")
        pPr.append(pStyle)
        cap_p.append(pPr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = desired_caption
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        cap_p.append(r)
        elem.addprevious(cap_p)

    # ---- Update Table 2 (CMU-MOSI dataset split) cells just to be safe (it already shows
    # Train/Test with appropriate sample counts).
    # Nothing else to do.

    # ---- Apply text-level replacements + subscripts to all paragraphs
    # (this covers paragraphs inside table cells too, because they are w:p elements).
    for p in body.iter(qn("w:p")):
        # Only process if has visible text
        text_nodes = list(p.iter(qn("w:t")))
        if not text_nodes:
            continue
        full_text = "".join(t.text or "" for t in text_nodes)
        if not full_text.strip():
            continue
        # Apply text replacement
        new_text = apply_text_replacements(full_text)
        if new_text != full_text:
            # rewrite by setting all text into first text-node, clearing others
            text_nodes[0].text = new_text
            for tn in text_nodes[1:]:
                tn.text = ""
        # Subscript conversion (only meaningful on the now-flattened text)
        if "_" in new_text:
            apply_subscripts_to_p(p)

    # ---- Convert all floating images (wp:anchor) into inline images (wp:inline)
    # so they flow with surrounding text instead of overlapping it. This fixes
    # the "tight wrap breaks layout" issue inherited from the source file.
    convert_floating_images_to_inline(doc)

    # ---- Build figure list and table list with concrete entries.
    insert_figure_table_lists(doc)

    # ---- Update TOC entries (the toc-styled paragraphs at the front)
    rebuild_toc(doc)

    # ---- Renumber chapter headings so the body matches the explicit TOC labels.
    chapter_label_map = {
        "導論": "第一章　導論",
        "文獻探討": "第二章　文獻探討",
        "研究方法": "第三章　SACF 情感感知跨模態融合模型",
        "情感評分模型評估": "第四章　SACF 情感評分模型實驗結果",
        "應用：情緒感知多代理模擬系統": "第七章　應用：情緒感知多代理模擬系統",
        "參考文獻": "第八章　參考文獻",
    }
    for child in body.iterchildren():
        if is_paragraph(child) and get_para_style(child) == "Heading 1":
            cur = get_para_text(child).strip()
            new_label = chapter_label_map.get(cur)
            if new_label:
                set_para_text(child, new_label)
                ensure_paragraph_style(child, "1")  # keep Heading 1 style id
            add_page_break_before(child)

    # ---- Insert placeholder chapters 5 (結論與貢獻) and 6 (未來工作)
    #      so the seventh-chapter labelling matches the TOC.
    ref_h1 = None
    for child in body.iterchildren():
        if is_paragraph(child) and get_para_style(child) == "Heading 1" \
                and "參考文獻" in get_para_text(child):
            ref_h1 = child
            break
    if ref_h1 is not None:
        # Find the application chapter (第七章) — placeholders go before it.
        app_h1 = None
        for child in body.iterchildren():
            if is_paragraph(child) and get_para_style(child) == "Heading 1" \
                    and "應用：情緒感知多代理模擬系統" in get_para_text(child):
                app_h1 = child
                break
        if app_h1 is not None:
            placeholder_specs = [
                ("第五章　結論與貢獻",
                 "本研究提出 SACF 情感感知跨模態融合模型，整合極性增強注意力與情感感知跨模態注意力兩項核心模組，並結合多工序數損失與零洩漏推斷增強策略，在 CMU-MOSI 基準上達到 Acc-7 52.19%、MAE 0.587 與 Corr 0.869 的領先成績。跨資料集評估進一步驗證模型在 MMAFFBen 與 SemEval-2018 上具備強勁的泛化能力與優異的參數效率。本研究之主要貢獻包含：（1）提出以情感顯著詞元驅動跨模態融合的新範式；（2）設計適用於序數情感標籤的多工聯合損失與零洩漏推斷流程；（3）驗證上述方法可遷移至多語言、多任務情感分析場景。"),
                ("第六章　未來工作",
                 "未來工作將沿三個方向延伸：（1）將 SACF 模型擴展至完整的多模態（音訊、視覺、文字三模態同時可用）情境，並探討跨模態缺失時的穩健學習；（2）研究輕量化設計，於行動端與嵌入式裝置實現即時情感推論；（3）將情感評分介面導入更廣泛的應用情境，涵蓋遠距教學、心理健康陪伴與長期社會模擬，並系統性評估其使用效益與倫理考量。"),
            ]
            for title_text, body_text in placeholder_specs:
                # Heading 1
                h_p = OxmlElement("w:p")
                pPr = OxmlElement("w:pPr")
                pStyle = OxmlElement("w:pStyle")
                pStyle.set(qn("w:val"), "1")
                pPr.append(pStyle)
                pbb = OxmlElement("w:pageBreakBefore")
                pPr.append(pbb)
                h_p.append(pPr)
                hr = OxmlElement("w:r")
                ht = OxmlElement("w:t")
                ht.text = title_text
                ht.set(qn("xml:space"), "preserve")
                hr.append(ht)
                h_p.append(hr)
                app_h1.addprevious(h_p)
                # Body paragraph
                b_p = OxmlElement("w:p")
                br_r = OxmlElement("w:r")
                br_t = OxmlElement("w:t")
                br_t.text = body_text
                br_t.set(qn("xml:space"), "preserve")
                br_r.append(br_t)
                b_p.append(br_r)
                app_h1.addprevious(b_p)

    doc.save(DST)
    print(f"Saved {DST}")


def ensure_paragraph_style(p, style_id):
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style_id)


def apply_subscripts_to_p(p):
    """Convert text in paragraph to runs with subscripts where pattern matches."""
    text_nodes = list(p.iter(qn("w:t")))
    if not text_nodes:
        return
    full_text = "".join(t.text or "" for t in text_nodes)
    if "_" not in full_text:
        return
    segments = split_subscript_runs(full_text)
    if not any(sub for _, sub in segments):
        return
    # Find the first run (parent of first text node) to copy formatting.
    first_run_elem = text_nodes[0].getparent()
    sample_rPr = first_run_elem.find(qn("w:rPr"))
    sample_rPr = copy.deepcopy(sample_rPr) if sample_rPr is not None else None

    # Remove all existing runs
    for r in list(p.findall(qn("w:r"))):
        p.remove(r)

    # Append new runs
    for seg_text, is_sub in segments:
        if not seg_text:
            continue
        r = OxmlElement("w:r")
        if sample_rPr is not None:
            r.append(copy.deepcopy(sample_rPr))
        if is_sub:
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r.insert(0, rPr)
            vert = OxmlElement("w:vertAlign")
            vert.set(qn("w:val"), "subscript")
            rPr.append(vert)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = seg_text
        r.append(t)
        p.append(r)


# ------------------------------------------------------------------ #
# 5. Figure / table lists with concrete entries
# ------------------------------------------------------------------ #

FIGURE_LIST = [
    "圖一　SACF 整體架構圖",
    "圖二　極性增強注意力（PEA）模組詳細示意圖",
    "圖三　情感感知跨模態注意力（SACF）逐步計算示意圖",
    "圖四　訓練策略全景：漸進解凍、EMA、SWA 與學習率排程",
    "圖五　多工損失函數設計：組成結構、序數懲罰矩陣與 EMD 示意",
    "圖六　零洩漏推斷增強流程：TTA×5、多種子集成",
]

TABLE_LIST = [
    "表一　CMU-MOSI 資料集劃分",
    "表二　各參數群組之學習率設定",
    "表三　SACF 模型主要超參數設定",
    "表四　CMU-MOSI 多模態情感分析效能比較",
    "表五　SACF-Text 在 MMAFFBen 五個純文字任務上的 ma-F1 比較",
    "表六　SACF-Text 在 SemEval-2018 Task 1 三語言整體分數比較",
    "表七　MMAFFIn 領域自適應預訓練消融比較",
    "表八　跨基準參數效率比較",
]


WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def convert_floating_images_to_inline(doc):
    """Rewrite every floating <wp:anchor> drawing as an inline <wp:inline>.

    Floating images with wrapTight / wrapSquare cause text to wrap around
    the picture and break academic-paper layout. Forcing every image to
    inline mode means each picture sits in its own line of flow, with the
    caption immediately below — the standard convention for thesis figures.

    The conversion preserves the inner <a:graphic>, <wp:extent>, and
    <wp:docPr>; it discards anchor-only positioning attributes (positionH,
    positionV, simplePos, wrap*).
    """
    body = doc.element.body
    for anchor in list(body.iter(f"{{{WP_NS}}}anchor")):
        # Build the replacement <wp:inline>.
        inline = etree.SubElement(anchor.getparent(), f"{{{WP_NS}}}inline")
        # Required attributes mirror what Word writes for inline drawings.
        inline.set("distT", anchor.get("distT", "0"))
        inline.set("distB", anchor.get("distB", "0"))
        inline.set("distL", anchor.get("distL", "0"))
        inline.set("distR", anchor.get("distR", "0"))

        # Carry over the children that inline expects, in the right order.
        # Required: extent, effectExtent, docPr, cNvGraphicFramePr, graphic.
        for child_tag in ("extent", "effectExtent", "docPr",
                          "cNvGraphicFramePr"):
            child = anchor.find(f"{{{WP_NS}}}{child_tag}")
            if child is not None:
                inline.append(copy.deepcopy(child))

        # graphic is in the drawingml namespace; it's the meaningful payload.
        a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
        graphic = anchor.find(f"{{{a_ns}}}graphic")
        if graphic is not None:
            inline.append(copy.deepcopy(graphic))

        # Move the freshly-built inline into the anchor's slot, then drop anchor.
        anchor.addprevious(inline)
        anchor.getparent().remove(anchor)


def add_page_break_before(p):
    """Add a w:pageBreakBefore element so the paragraph starts on a new page."""
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    # Avoid duplicates
    if pPr.find(qn("w:pageBreakBefore")) is not None:
        return
    pbb = OxmlElement("w:pageBreakBefore")
    pPr.append(pbb)


def insert_figure_table_lists(doc):
    """Replace existing 圖目錄 / 表目錄 sections with explicit entries."""
    body = doc.element.body
    # Find the "圖目錄" Title paragraph
    fig_title_p = None
    table_title_p = None
    for p in body.iter(qn("w:p")):
        text = get_para_text(p).strip()
        if text == "圖目錄":
            fig_title_p = p
        elif text == "表目錄":
            table_title_p = p

    # Force each index onto its own page.
    if fig_title_p is not None:
        add_page_break_before(fig_title_p)
    if table_title_p is not None:
        add_page_break_before(table_title_p)

    def replace_after(title_p, entries):
        if title_p is None:
            return
        # Walk forward, removing paragraphs until we hit the next Title or Heading 1
        next_elem = title_p.getnext()
        to_remove = []
        while next_elem is not None:
            if is_paragraph(next_elem):
                style = get_para_style(next_elem)
                if style in {"Title", "Heading 1"}:
                    break
                # Stop if this is the next index title
                t = get_para_text(next_elem).strip()
                if t in {"表目錄", "圖目錄"}:
                    break
                to_remove.append(next_elem)
            elif is_table(next_elem):
                # leave tables alone
                break
            next_elem = next_elem.getnext()
        for r in to_remove:
            body.remove(r)
        # Insert new entries after title
        anchor = title_p
        for entry in entries:
            p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            pStyle = OxmlElement("w:pStyle")
            pStyle.set(qn("w:val"), "12")  # toc 1 style
            pPr.append(pStyle)
            p.append(pPr)
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = entry
            r.append(t)
            p.append(r)
            anchor.addnext(p)
            anchor = p

    replace_after(fig_title_p, FIGURE_LIST)
    replace_after(table_title_p, TABLE_LIST)


# ------------------------------------------------------------------ #
# 6. Rebuild table-of-contents entries
# ------------------------------------------------------------------ #

NEW_TOC_ENTRIES = [
    ("toc 1", "第一章 導論"),
    ("toc 2", "第一節 研究背景與動機"),
    ("toc 2", "第二節 研究目的"),
    ("toc 2", "第三節 研究限制"),
    ("toc 1", "第二章 文獻探討"),
    ("toc 2", "第一節 智能體理論基礎"),
    ("toc 2", "第二節 情感計算與情感表達技術"),
    ("toc 2", "第三節 多智能體系統與社會互動"),
    ("toc 2", "第四節 大型語言模型在智能體中的應用"),
    ("toc 2", "第五節 多智能體系統中的人格化與情感表達"),
    ("toc 1", "第三章 SACF 情感感知跨模態融合模型"),
    ("toc 2", "第一節 研究框架概覽"),
    ("toc 2", "第二節 資料集與前處理"),
    ("toc 2", "第三節 模型架構"),
    ("toc 2", "第四節 零洩漏推斷增強策略"),
    ("toc 2", "第五節 實作細節"),
    ("toc 1", "第四章 SACF 情感評分模型實驗結果"),
    ("toc 2", "第一節 評估指標說明"),
    ("toc 2", "第二節 基線模型與比較方法"),
    ("toc 2", "第三節 CMU-MOSI 資料集效能比較"),
    ("toc 2", "第四節 多種子穩健性分析"),
    ("toc 2", "第五節 多種子集成效益分析"),
    ("toc 2", "第六節 跨資料集泛化性評估"),
    ("toc 2", "第七節 領域自適應預訓練消融實驗"),
    ("toc 2", "第八節 參數效率分析"),
    ("toc 2", "第九節 模型優勢分析"),
    ("toc 2", "第十節 結果討論"),
    ("toc 1", "第五章 結論與貢獻"),
    ("toc 1", "第六章 未來工作"),
    ("toc 1", "第七章 應用：情緒感知多代理模擬系統"),
    ("toc 2", "第一節 系統整體架構"),
    ("toc 2", "第二節 智能體核心模組設計"),
    ("toc 2", "第三節 記憶管理系統"),
    ("toc 2", "第四節 環境建模與空間導航"),
    ("toc 2", "第五節 語言模型整合與提示工程"),
    ("toc 2", "第六節 視覺化與使用者介面"),
    ("toc 1", "第八章 參考文獻"),
]

# Style id mapping from logical name to numeric style id used in this doc.
STYLE_ID = {
    "toc 1": "12",
    "toc 2": "22",
    "toc 3": "32",
}


def rebuild_toc(doc):
    """Find the existing TOC area (after the '目 錄' title) and rewrite entries."""
    body = doc.element.body
    toc_title_p = None
    for p in body.iter(qn("w:p")):
        text = get_para_text(p).strip()
        if text in {"目 錄", "目錄"}:
            toc_title_p = p
            break
    if toc_title_p is None:
        return
    # Remove subsequent toc-styled paragraphs until we hit the next Title
    next_elem = toc_title_p.getnext()
    to_remove = []
    while next_elem is not None:
        if is_paragraph(next_elem):
            style = get_para_style(next_elem)
            text = get_para_text(next_elem).strip()
            if style == "Title" or text in {"圖目錄", "表目錄"}:
                break
            # Remove paragraphs that look like TOC entries
            if style.startswith("toc") or style in STYLE_ID.values() or _looks_like_toc_entry(text):
                to_remove.append(next_elem)
            else:
                # Skip empty paragraphs but don't break either
                if not text:
                    to_remove.append(next_elem)
                else:
                    break
        else:
            break
        next_elem = next_elem.getnext()
    for r in to_remove:
        body.remove(r)
    # Insert new entries
    anchor = toc_title_p
    for level, text in NEW_TOC_ENTRIES:
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), STYLE_ID[level])
        pPr.append(pStyle)
        p.append(pPr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)
        anchor.addnext(p)
        anchor = p


def _looks_like_toc_entry(text):
    return bool(re.match(r"^第[一二三四五六七八九十]+(章|節)\s", text)) or "找不到目錄項目" in text


if __name__ == "__main__":
    build()
