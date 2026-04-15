"""
Integrate SACF_Methodology_Chapter3.docx into 論文＿李昇峰.docx as a new
section 3.2, renumbering existing 3.2/3.3/3.4 to 3.3/3.4/3.5.

Writes 論文＿李昇峰_merged.docx at project root.
"""
import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.parts.image import ImagePart

ROOT = Path(__file__).parent.parent
THESIS = ROOT / "論文＿李昇峰.docx"
SACF   = ROOT / "docs" / "SACF_Methodology_Chapter3.docx"
OUT    = ROOT / "論文＿李昇峰_merged.docx"

thesis = Document(str(THESIS))
sacf   = Document(str(SACF))

# ── Step 1: copy image parts from SACF into thesis with UNIQUE partnames ─────
rid_map = {}
pkg = thesis.part.package
existing_partnames = {str(p.partname) for p in pkg.iter_parts()}

def unique_image_partname(ext: str) -> PackURI:
    idx = 100
    while f"/word/media/sacf_image{idx}.{ext}" in existing_partnames:
        idx += 1
    pn = PackURI(f"/word/media/sacf_image{idx}.{ext}")
    existing_partnames.add(str(pn))
    return pn

for rel_id, rel in sacf.part.rels.items():
    if rel.reltype == RT.IMAGE:
        old = rel.target_part
        ext = str(old.partname).rsplit(".", 1)[-1]
        new_partname = unique_image_partname(ext)
        new_part = ImagePart(new_partname, old.content_type, old.blob, pkg)
        new_rid = thesis.part.relate_to(new_part, RT.IMAGE)
        rid_map[rel_id] = new_rid

# ── Step 2: find target = the BODY Heading-2 "3.2 技術實現方法", skip TOC entries ──
target_elem = None
for p in thesis.paragraphs:
    if (p.text.strip().startswith("3.2 技術實現方法")
            and p.style.name == "Heading 2"):
        target_elem = p._element
        break
assert target_elem is not None, "insertion target not found"

# ── Step 3: build list of SACF body elements to insert (renumbered) ───────────
SEC_RE = re.compile(r"^(3)\.(\d+)(?:\.(\d+))?(\s|$)")

def renumber_in_text(text: str) -> str:
    """Prefix 3.X / 3.X.Y with 3.2 → 3.2.X / 3.2.X.Y (SACF sections under new 3.2)."""
    stripped = text.lstrip()
    lead = text[:len(text)-len(stripped)]
    m = SEC_RE.match(stripped)
    if not m:
        return text
    maj, sub, trail = m.group(2), m.group(3), m.group(4)
    if sub:
        new_prefix = f"3.2.{maj}.{sub}{trail}"
    else:
        new_prefix = f"3.2.{maj}{trail}"
    return lead + new_prefix + stripped[m.end():]

def renumber_figure_label(text: str) -> str:
    """Figure 3.X / 圖 3.X → Figure 3.2.X / 圖 3.2.X (both variants appear)."""
    t = re.sub(r"^Figure\s+3\.(\d+)", r"Figure 3.2.\1", text)
    t = re.sub(r"圖\s*3\.(\d+)", r"圖 3.2.\1", t)
    t = re.sub(r"表\s*3\.(\d+)", r"表 3.2.\1", t)
    return t

sacf_body = sacf.element.body
elems_to_insert = []
for child in sacf_body.iterchildren():
    tag = child.tag.split("}")[-1]
    if tag == "sectPr":
        continue
    # skip the "第三章 研究方法" heading from SACF (thesis already has it)
    text = "".join(t.text or "" for t in child.iter(qn("w:t")))
    if text.strip().startswith("第三章"):
        continue

    new_el = deepcopy(child)

    # remap image rels
    for blip in new_el.iter(qn("a:blip")):
        embed = blip.get(qn("r:embed"))
        if embed and embed in rid_map:
            blip.set(qn("r:embed"), rid_map[embed])

    # renumber text nodes: only modify text that starts with "3.X" section number
    # or that contains Figure 3.X / 圖 3.X labels
    for tnode in new_el.iter(qn("w:t")):
        old = tnode.text or ""
        new = renumber_in_text(old)
        new = renumber_figure_label(new)
        if new != old:
            tnode.text = new

    elems_to_insert.append(new_el)

# ── Step 3b: prepend a new H2 heading "3.2 情感感知跨模態融合模型 (SACF)" ──
from docx.oxml import OxmlElement
def make_heading_para(text: str, style: str) -> "CT_P":
    """Build a new w:p element with the given text and paragraph style."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    # Styles in python-docx use id-form of style name; Heading 2 → "Heading2"
    # map friendly name → actual style id used in thesis: "Heading 2" → "2"
    _map = {"Heading 1":"1","Heading 2":"2","Heading 3":"3","Heading 4":"4","Normal":"a"}
    pStyle.set(qn("w:val"), _map.get(style, style))
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p

intro_para_text = (
    "本節介紹本研究為智能體情感感知層所設計之核心技術——情感感知跨模態融合"
    "（Sentiment-Aware Cross-modal Fusion，SACF）模型。SACF 以 CMU-MOSI 多模態"
    "情感資料集進行訓練與驗證，作為多代理系統之情緒評分基礎，後續章節亦將據此"
    "模型之預測結果驅動智能體的情緒狀態更新與社交互動決策。"
)
elems_to_insert = (
    [make_heading_para("3.2 情感感知跨模態融合模型（SACF）", "Heading 2"),
     make_heading_para(intro_para_text, "Normal")]
    + elems_to_insert
)

# ── Step 3c: upgrade SACF sub-section paragraphs to proper Heading styles ─────
# Second-level (3.2.1 ~ 3.2.7) → Heading 3
# Third-level (3.2.X.Y) → Heading 4 (if exists) else leave as Normal bold
sec2_re = re.compile(r"^3\.2\.\d+\s")       # "3.2.1 "
sec3_re = re.compile(r"^3\.2\.\d+\.\d+\s")  # "3.2.1.1 "

for el in elems_to_insert:
    if el.tag != qn("w:p"):
        continue
    text = "".join(t.text or "" for t in el.iter(qn("w:t"))).lstrip()
    if sec3_re.match(text):
        style_id = "4"  # Heading 4
    elif sec2_re.match(text):
        style_id = "3"  # Heading 3
    else:
        continue
    # set paragraph style
    pPr = el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        el.insert(0, pPr)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style_id)

# ── Step 4: insert before target ──────────────────────────────────────────────
for el in elems_to_insert:
    target_elem.addprevious(el)

# ── Step 5: renumber thesis' own existing Ch3 headings (3.2 → 3.3 etc.) ───────
# Use concatenate-into-first-run strategy to avoid split-run issues.
renumber_map = [
    ("3.2 技術實現方法",          "3.3 技術實現方法"),
    ("3.2.1 記憶檢索與關聯機制",   "3.3.1 記憶檢索與關聯機制"),
    ("3.3 實驗環境構建",           "3.4 實驗環境構建"),
    ("3.3.1 智能體配置",           "3.4.1 智能體配置"),
    ("3.4 互動機制與對話生成",     "3.5 互動機制與對話生成"),
    ("3.4.1 智能體間互動決策",     "3.5.1 智能體間互動決策"),
    ("3.4.2 自然對話生成",         "3.5.2 自然對話生成"),
]

def rewrite_paragraph(p, new_text):
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""

# only inside Ch 3 BODY (Heading style — skip TOC entries)
inside_ch3 = False
for p in thesis.paragraphs:
    t = p.text.strip()
    sty = p.style.name
    if t.startswith("第三章") and sty == "Heading 1":
        inside_ch3 = True
        continue
    if t.startswith("第四章") and sty == "Heading 1":
        break
    if not inside_ch3 or not sty.startswith("Heading"):
        continue
    for old, new in renumber_map:
        if t == old:
            rewrite_paragraph(p, new)
            break

thesis.save(str(OUT))
print(f"✓ Saved: {OUT}")
print(f"  inserted {len(elems_to_insert)} elements from SACF methodology")
print(f"  remapped {len(rid_map)} image relationships")
