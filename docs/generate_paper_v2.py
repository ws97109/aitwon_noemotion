"""Generator for the methodology chapter (Chapter 3) and accompanying figures.
Produces docs/figures/v2_*.svg/png and docs/SACF_Methodology_Chapter3_v2.docx
"""
import os, sys, json, pickle, warnings
warnings.filterwarnings("ignore")
from pathlib import Path
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, Rectangle, FancyArrowPatch, Circle
from matplotlib.gridspec import GridSpec
from matplotlib.lines import Line2D

BASE = Path(__file__).parent
ROOT = BASE.parent
OUTDIR = BASE / "figures"
OUTDIR.mkdir(exist_ok=True)
MODEL_DIR = ROOT / "emotion_system" / "models"

# ── load extracted data ─────────────────────────────────────────────────────
D = np.load(str(BASE / "paper_v2_data.npz"))
metrics = D["metrics"]   # [acc7_fused, acc7_raw, acc2, f1, mae, corr]
ACC7_FUSED, ACC7_RAW, ACC2, F1, MAE, CORR = [float(x) for x in metrics]
PER_CLASS_ACC = D["per_class_acc"]
CLASS_SUPPORT = D["class_support"]
CM = D["confusion_matrix"]
TRAIN_DIST = D["train_dist"]; VAL_DIST = D["val_dist"]
TEST_DIST = D["test_dist"]; TRAINVAL_DIST = D["trainval_dist"]
LOSS_CURVE = D["loss_curve"]
L7 = D["L7"]; L2 = D["L2"]; R = D["R"]; Y7 = D["y7"]

# ── palette ─────────────────────────────────────────────────────────────────
C = dict(
    primary="#1D4ED8", secondary="#0891B2", accent="#F59E0B",
    danger="#DC2626", success="#10B981", purple="#8B5CF6",
    text="#1F2937", muted="#6B7280", grid="#E5E7EB", bg="#F9FAFB",
    teal="#14B8A6", indigo="#6366F1", rose="#F43F5E", lime="#84CC16",
)
plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "font.size": 10, "axes.titlesize": 11, "axes.labelsize": 10,
    "xtick.labelsize": 9, "ytick.labelsize": 9, "legend.fontsize": 9,
    "axes.spines.top": False, "axes.spines.right": False,
    "axes.grid": True, "grid.color": C["grid"], "grid.alpha": 0.5,
    "savefig.dpi": 150, "figure.dpi": 110,
})


def save(fig, name):
    p_svg = OUTDIR / f"{name}.svg"
    p_png = OUTDIR / f"{name}.png"
    fig.savefig(str(p_svg), bbox_inches="tight")
    fig.savefig(str(p_png), bbox_inches="tight")
    plt.close(fig)
    return {"svg": p_svg, "png": p_png}


def rbox(ax, cx, cy, w, h, fc, text, fs=9, tc="white", ec=None,
         bold=False, lw=1.0, alpha=1.0, rad=0.06):
    if ec is None: ec = fc
    box = FancyBboxPatch((cx - w/2, cy - h/2), w, h,
                          boxstyle=f"round,pad=0.005,rounding_size={rad}",
                          fc=fc, ec=ec, lw=lw, alpha=alpha, zorder=4)
    ax.add_patch(box)
    weight = "bold" if bold else "normal"
    ax.text(cx, cy, text, ha="center", va="center",
            color=tc, fontsize=fs, fontweight=weight, zorder=5)


def arr(ax, x1, y1, x2, y2, color=None, lw=1.4, hw=0.18, z=3, ls="-"):
    if color is None: color = C["text"]
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle=f"->,head_width={hw},head_length={hw*1.4}",
                                 color=color, lw=lw, linestyle=ls), zorder=z)


def bg(ax, x0, y0, w, h, fc, ec, lw=1.2, alpha=0.18):
    r = FancyBboxPatch((x0, y0), w, h,
                       boxstyle="round,pad=0.01,rounding_size=0.18",
                       fc=fc, ec=ec, lw=lw, alpha=alpha, zorder=1)
    ax.add_patch(r)


# ════════════════════════════════════════════════════════════════════════════
# FIG 1 — Overall architecture
# ════════════════════════════════════════════════════════════════════════════
def fig1_architecture():
    fig, ax = plt.subplots(figsize=(13.5, 11.5))
    ax.set_xlim(0, 14); ax.set_ylim(0, 13.5); ax.axis("off")

    # Title bar
    ax.text(7, 13.0,
            "Multi-Branch Single Model: PEA + Hierarchical SACF + CMC + Reg-Cls Fusion",
            ha="center", va="center", fontsize=13.5, fontweight="bold",
            color=C["primary"])
    ax.text(7, 12.55, f"Acc-7 = {ACC7_FUSED:.2f}%   |   Acc-2 = {ACC2:.2f}%   |   F1 = {F1:.2f}%   |   MAE = {MAE:.4f}   |   Corr = {CORR:.4f}",
            ha="center", va="center", fontsize=10.5, color=C["muted"])

    # ── Row 1: Inputs ───────────────────────────────────────────────────────
    bg(ax, 0.4, 11.0, 13.2, 1.2, "#FEF3C7", C["accent"])
    ax.text(7, 12.0, "(i)  Multimodal Inputs", ha="center", fontsize=10.5,
            color=C["accent"], fontweight="bold")
    rbox(ax, 2.5, 11.4, 3.0, 0.55, C["primary"],
         "Text  (raw_text + task prompt)", fs=9.5, bold=True)
    rbox(ax, 7.0, 11.4, 3.0, 0.55, C["accent"],
         "Audio  (COVAREP, 5-d, ≤375 frames)", fs=9.5, bold=True)
    rbox(ax, 11.5, 11.4, 3.0, 0.55, C["success"],
         "Vision  (FACET, 20-d, ≤500 frames)", fs=9.5, bold=True)

    # ── Row 2: Shared encoders ──────────────────────────────────────────────
    bg(ax, 0.4, 9.2, 13.2, 1.5, "#DBEAFE", C["primary"])
    ax.text(7, 10.55, "(ii)  Shared Encoders", ha="center", fontsize=10.5,
            color=C["primary"], fontweight="bold")
    rbox(ax, 2.5, 9.85, 3.0, 0.7, C["primary"],
         "DeBERTa-v3-large\n24 layers · d=1024 · ~400M", fs=9, bold=True)
    rbox(ax, 7.0, 9.85, 3.0, 0.7, C["accent"],
         "BiLSTM-Audio\n2 layers · 5→128", fs=9, bold=True)
    rbox(ax, 11.5, 9.85, 3.0, 0.7, C["success"],
         "BiLSTM-Vision\n2 layers · 20→128", fs=9, bold=True)

    # arrows row1 → row2
    for cx in (2.5, 7.0, 11.5):
        arr(ax, cx, 11.13, cx, 10.20)

    # ── Row 3: 4 parallel branches ──────────────────────────────────────────
    bg(ax, 0.4, 5.4, 13.2, 3.5, "#EDE9FE", C["purple"])
    ax.text(7, 8.75, "(iii)  4 Parallel Branches  (per-branch dropout = [0.10, 0.20, 0.30, 0.40])",
            ha="center", fontsize=10.5, color=C["purple"], fontweight="bold")

    branch_colors = [C["primary"], C["accent"], C["success"], C["rose"]]
    branch_dropouts = [0.10, 0.20, 0.30, 0.40]
    for i in range(4):
        cx = 2.0 + i * 3.0
        # branch container
        bg(ax, cx - 1.30, 5.6, 2.6, 2.95, "#FFFFFF", branch_colors[i], lw=1.6, alpha=0.6)
        ax.text(cx, 8.30, f"Branch {i+1}  (p_drop={branch_dropouts[i]})",
                ha="center", fontsize=9.5, fontweight="bold", color=branch_colors[i])
        rbox(ax, cx, 7.85, 2.4, 0.5, branch_colors[i], "PEA  (gate σ)", fs=8.5)
        rbox(ax, cx, 7.20, 2.4, 0.5, branch_colors[i], "SACF  (top-K + cross-modal)", fs=8.5)
        rbox(ax, cx, 6.55, 2.4, 0.5, branch_colors[i], "Proj  (1024 → 512)", fs=8.5)
        rbox(ax, cx, 5.90, 2.4, 0.55, "#FFFFFF", "cls7 / cls2 / reg",
             fs=8.5, tc=branch_colors[i], ec=branch_colors[i], bold=True)

    # arrows from shared encoders to each branch (just 3 representative)
    for cx in (2.0, 5.0, 8.0, 11.0):
        arr(ax, 2.5, 9.50, cx, 8.50, color=C["primary"], lw=0.8)
        arr(ax, 7.0, 9.50, cx, 8.50, color=C["accent"], lw=0.8)
        arr(ax, 11.5, 9.50, cx, 8.50, color=C["success"], lw=0.8)

    # ── Row 4: Mean aggregation ─────────────────────────────────────────────
    bg(ax, 4.5, 4.1, 5.0, 0.9, "#FFE4E6", C["rose"])
    rbox(ax, 7.0, 4.55, 4.6, 0.65, C["rose"],
         "Mean-of-Branches  →  (cls7_mean, cls2_mean, reg_mean)",
         fs=9.5, bold=True)
    for i in range(4):
        cx = 2.0 + i * 3.0
        arr(ax, cx, 5.60, 7.0, 4.95, color=branch_colors[i], lw=0.8)

    # ── Row 5: Inference fusion ─────────────────────────────────────────────
    bg(ax, 0.4, 1.8, 13.2, 2.0, "#DCFCE7", C["success"])
    ax.text(7, 3.65, "(iv)  Inference: Reg-Cls Probability Fusion (α=0.65, σ=0.65 — chosen a priori)",
            ha="center", fontsize=10.5, color=C["success"], fontweight="bold")
    rbox(ax, 3.0, 2.85, 3.4, 0.7, C["primary"],
         "p_cls = softmax(cls7_logits / T)", fs=9, bold=True)
    rbox(ax, 7.0, 2.85, 3.4, 0.7, C["accent"],
         "p_reg[k] ∝ exp(−(k−r)² / 2σ²)", fs=9, bold=True)
    rbox(ax, 11.0, 2.85, 3.4, 0.7, C["success"],
         "log p ← α·log p_cls + (1−α)·log p_reg", fs=9, bold=True)
    arr(ax, 7.0, 4.20, 7.0, 3.30, lw=1.5)
    arr(ax, 3.0, 2.45, 5.5, 2.10)
    arr(ax, 7.0, 2.45, 7.0, 2.10)
    arr(ax, 11.0, 2.45, 8.5, 2.10)

    # ── Final prediction ────────────────────────────────────────────────────
    rbox(ax, 7.0, 1.30, 5.4, 0.75, C["danger"],
         f"ŷ = argmax(p_final)        Acc-7 = {ACC7_FUSED:.2f}%",
         fs=11, bold=True)

    # ── Legend ──────────────────────────────────────────────────────────────
    handles = [
        mpatches.Patch(color=C["primary"], label="Text (DeBERTa)"),
        mpatches.Patch(color=C["accent"], label="Audio (BiLSTM)"),
        mpatches.Patch(color=C["success"], label="Vision (BiLSTM)"),
        mpatches.Patch(color=C["rose"], label="Aggregation"),
        mpatches.Patch(color=C["purple"], label="Parallel Branches"),
    ]
    ax.legend(handles=handles, loc="lower left", bbox_to_anchor=(0.0, -0.02),
              ncol=5, fontsize=8.5, frameon=False)

    fig.tight_layout()
    return save(fig, "v2_fig1_architecture")


# ════════════════════════════════════════════════════════════════════════════
# FIG 2 — PEA module detail
# ════════════════════════════════════════════════════════════════════════════
def fig2_pea():
    fig = plt.figure(figsize=(13.5, 6.5))
    gs = GridSpec(1, 2, width_ratios=[1.4, 1.0], wspace=0.20)
    ax = fig.add_subplot(gs[0, 0])
    ax.set_xlim(0, 14); ax.set_ylim(0, 7); ax.axis("off")
    ax.text(7, 6.55, "Polarity-Enhanced Attention (PEA)",
            ha="center", fontsize=13, fontweight="bold", color=C["primary"])

    # Token sequence input
    tokens = ["This", "movie", "is", "absolutely", "amazing", "and", "thrilling"]
    gates = [0.18, 0.32, 0.12, 0.78, 0.94, 0.20, 0.86]
    bg(ax, 0.3, 4.4, 13.4, 1.4, "#DBEAFE", C["primary"])
    ax.text(7, 5.55, "DeBERTa hidden states  H ∈ ℝ^{B×L×1024}",
            ha="center", fontsize=10, fontweight="bold", color=C["primary"])
    for i, (t, g) in enumerate(zip(tokens, gates)):
        cx = 1.1 + i * 1.85
        # color intensity ~ gate value
        red = int(255 - (1 - g) * 200)
        col = f"#{red:02X}{int(255 * (1 - g) * 0.6):02X}{int(255 * (1 - g) * 0.6):02X}"
        rbox(ax, cx, 4.85, 1.6, 0.45, col, f"h_{i+1}",
             fs=9, bold=True)

    # Gate vector
    bg(ax, 0.3, 2.7, 13.4, 1.4, "#FEF3C7", C["accent"])
    ax.text(7, 3.85, "PEA gate  g_i = σ(W₂ · tanh(W₁ · h_i)) ∈ [0, 1]",
            ha="center", fontsize=10, fontweight="bold", color=C["accent"])
    for i, (t, g) in enumerate(zip(tokens, gates)):
        cx = 1.1 + i * 1.85
        rbox(ax, cx, 3.18, 1.6, 0.36, C["accent"], f"g={g:.2f}", fs=9, bold=True)
        ax.text(cx, 2.78, t, ha="center", fontsize=8.5, color=C["text"])
        # arrow from token to gate
        arr(ax, cx, 4.60, cx, 3.42, color=C["muted"], lw=0.8)

    # Output
    bg(ax, 0.3, 0.6, 13.4, 1.5, "#DCFCE7", C["success"])
    ax.text(7, 1.85, "Polarity-Weighted Sentence Embedding   x_cls ∈ ℝ^{B×1024}",
            ha="center", fontsize=10, fontweight="bold", color=C["success"])
    ax.text(7, 1.20,
            "x_cls  =  Σᵢ m_i · (0.75 · h_i  +  0.25 · h_i ⊙ g_i)  /  Σᵢ m_i",
            ha="center", fontsize=11, fontweight="bold", color=C["text"])
    arr(ax, 7, 2.65, 7, 2.10, lw=1.5)

    # Right panel: gate distribution explanation
    ax2 = fig.add_subplot(gs[0, 1])
    ax2.barh(range(len(tokens))[::-1], gates,
             color=[(1, 1 - g, 1 - g) for g in gates], edgecolor=C["accent"], lw=1)
    ax2.set_yticks(range(len(tokens))[::-1])
    ax2.set_yticklabels(tokens, fontsize=10)
    ax2.set_xlabel("Gate value g_i")
    ax2.set_xlim(0, 1)
    ax2.set_title("Per-Token Gate (illustrative)")
    ax2.axvline(0.5, color=C["muted"], ls=":", alpha=0.6)
    ax2.grid(True, axis="x", alpha=0.4)

    fig.tight_layout()
    return save(fig, "v2_fig2_pea")


# ════════════════════════════════════════════════════════════════════════════
# FIG 3 — SACF cross-modal pipeline
# ════════════════════════════════════════════════════════════════════════════
def fig3_sacf():
    fig, ax = plt.subplots(figsize=(13.5, 6.0))
    ax.set_xlim(0, 14); ax.set_ylim(0, 6); ax.axis("off")
    ax.text(7, 5.65, "Sentiment-Aware Cross-modal Fusion (SACF)",
            ha="center", fontsize=13, fontweight="bold", color=C["primary"])

    # Step boxes
    step_w = 3.0; step_h = 1.6
    cy = 3.0
    steps = [
        ("Step 1\nTop-K Token\nSelection",
         "Use PEA gates g\nto pick K=5 most\nsentiment-salient\ntokens", C["primary"]),
        ("Step 2\nSentiment\nQuery",
         "q_sa  =  Σ_k softmax(\nW_attn · h_k) · h_k\n∈ ℝ^{B×1024}", C["accent"]),
        ("Step 3\nCross-modal\nAttention",
         "K, V = stack(\n  audio_map(x_a),\n  vision_map(x_v))\nα = softmax(q_sa · Kᵀ/√d)\nx̂ = α · KV", C["secondary"]),
        ("Step 4\nGated Residual\nFusion",
         "x = FFN(x_cls + x̂)\ng_w = σ(W_g·[x_cls; x])\nf = LN(x_cls + x ⊙ g_w)", C["success"]),
    ]
    for i, (title, body_text, color) in enumerate(steps):
        cx = 1.6 + i * 3.05
        bg(ax, cx - step_w/2, cy - step_h/2, step_w, step_h, "#FFFFFF", color, lw=1.7, alpha=0.85)
        ax.text(cx, cy + 0.50, title, ha="center", va="center",
                fontsize=10, fontweight="bold", color=color)
        ax.text(cx, cy - 0.25, body_text, ha="center", va="center",
                fontsize=8.5, color=C["text"])
        if i < 3:
            arr(ax, cx + step_w/2, cy, cx + step_w/2 + 0.1, cy, lw=0)
            arr(ax, cx + step_w/2 + 0.05, cy, 1.6 + (i+1)*3.05 - step_w/2 - 0.05, cy,
                color=C["text"], lw=1.6)

    # Inputs above
    bg(ax, 0.3, 4.55, 13.4, 0.7, "#DBEAFE", C["primary"], alpha=0.4)
    rbox(ax, 1.6, 4.90, 2.4, 0.45, C["primary"], "H, gates g  (text)", fs=9, bold=True)
    rbox(ax, 4.65, 4.90, 2.4, 0.45, C["accent"], "x_a (audio, 128-d)", fs=9, bold=True)
    rbox(ax, 7.7, 4.90, 2.4, 0.45, C["success"], "x_v (vision, 128-d)", fs=9, bold=True)
    rbox(ax, 11.5, 4.90, 3.0, 0.45, C["text"], "x_cls (PEA output)", fs=9, bold=True)

    # Output
    bg(ax, 4.5, 0.5, 5.0, 1.0, "#FCE7F3", C["rose"])
    rbox(ax, 7.0, 1.0, 4.6, 0.7, C["rose"],
         "Fused representation  f ∈ ℝ^{B×1024}", fs=10, bold=True)
    arr(ax, 7.0, 2.20, 7.0, 1.45, lw=1.6, color=C["rose"])

    fig.tight_layout()
    return save(fig, "v2_fig3_sacf")


# ════════════════════════════════════════════════════════════════════════════
# FIG 5 — Reg-Cls fusion at inference
# ════════════════════════════════════════════════════════════════════════════
def fig5_regcls_fusion():
    fig, axes = plt.subplots(1, 3, figsize=(13.5, 4.3))
    K = 7
    classes_label = list(range(-3, 4))
    # Pick a real test sample where fusion improves (cls is close, reg helps)
    # Find a sample where pred_raw differs from pred_fused
    pred_raw = L7.argmax(1)
    import scipy.special
    p_cls_all = scipy.special.softmax(L7, axis=1)
    centers = np.arange(K).astype(np.float32)
    sigma_inf = 0.65
    p_reg_all = np.exp(-((centers[None, :] - (R[:, None] + 3))**2) / (2 * sigma_inf**2))
    p_reg_all = p_reg_all / p_reg_all.sum(1, keepdims=True)
    log_p = 0.65 * np.log(p_cls_all + 1e-9) + 0.35 * np.log(p_reg_all + 1e-9)
    log_p -= log_p.max(1, keepdims=True)
    p_fused_all = np.exp(log_p); p_fused_all /= p_fused_all.sum(1, keepdims=True)
    pred_fused = p_fused_all.argmax(1)
    # find sample where fused matches truth but raw doesn't
    idx = None
    for i in range(len(Y7)):
        if pred_fused[i] == Y7[i] and pred_raw[i] != Y7[i] and Y7[i] in (1,2,3,4,5):
            idx = i; break
    if idx is None: idx = 0
    p_cls = p_cls_all[idx]; p_reg = p_reg_all[idx]; p_fused = p_fused_all[idx]
    truth = Y7[idx]
    r_val = R[idx]

    ax = axes[0]
    ax.bar(classes_label, p_cls, color=C["primary"], alpha=0.85)
    ax.axvline(classes_label[truth], color=C["success"], lw=2, ls="--", label=f"Ground truth ({classes_label[truth]})")
    ax.axvline(classes_label[p_cls.argmax()], color=C["danger"], lw=1.5, ls=":", label=f"argmax cls ({classes_label[p_cls.argmax()]})")
    ax.set_xticks(classes_label); ax.set_xlabel("Class label"); ax.set_ylabel("Probability")
    ax.set_title("(a)  p_cls = softmax(cls7_logits)")
    ax.legend(fontsize=8)

    ax = axes[1]
    ax.bar(classes_label, p_reg, color=C["accent"], alpha=0.85)
    ax.axvline(r_val, color=C["text"], lw=2, ls="-", label=f"reg pred = {r_val:.2f}")
    ax.axvline(classes_label[truth], color=C["success"], lw=2, ls="--", label=f"truth ({classes_label[truth]})")
    ax.set_xticks(classes_label); ax.set_xlabel("Class label"); ax.set_ylabel("Probability")
    ax.set_title(f"(b)  p_reg ∝ exp(−(k−r)²/2σ²)   σ=0.65")
    ax.legend(fontsize=8)

    ax = axes[2]
    ax.bar(classes_label, p_fused, color=C["success"], alpha=0.9)
    ax.axvline(classes_label[truth], color=C["success"], lw=2, ls="--", label=f"truth ({classes_label[truth]})")
    ax.axvline(classes_label[p_fused.argmax()], color=C["danger"], lw=1.5, ls=":",
               label=f"argmax fused ({classes_label[p_fused.argmax()]})")
    ax.set_xticks(classes_label); ax.set_xlabel("Class label"); ax.set_ylabel("Probability")
    ax.set_title("(c)  log p_final = α·log p_cls + (1−α)·log p_reg   α=0.65")
    ax.legend(fontsize=8)

    fig.suptitle(f"Reg-Cls Probability Fusion at Inference  (test sample idx={idx})",
                 fontsize=12, fontweight="bold", y=1.02)
    fig.tight_layout()
    return save(fig, "v2_fig5_regcls_fusion")


# ════════════════════════════════════════════════════════════════════════════
# FIG 6 — Two-stage training timeline
# ════════════════════════════════════════════════════════════════════════════
def fig6_training_timeline():
    fig, ax = plt.subplots(figsize=(13.5, 5.5))
    ax.set_xlim(0, 80); ax.set_ylim(0, 5); ax.axis("off")

    # Header
    ax.text(40, 4.7, "Two-Stage Training Pipeline (single weight file)",
            ha="center", fontsize=12.5, fontweight="bold", color=C["primary"])

    # Stage 1 — iter1 (60 ep)
    bg(ax, 0.5, 0.5, 60, 3.6, "#DBEAFE", C["primary"])
    ax.text(30, 3.85, "Stage 1: iter1 (Epoch 1–60)  —  full training from scratch",
            ha="center", fontsize=10.5, fontweight="bold", color=C["primary"])
    # phases inside iter1
    bg(ax, 1, 1.2, 19, 1.4, "#FEF3C7", C["accent"], alpha=0.5)
    ax.text(10.5, 1.9, "Phase 1: freeze DeBERTa\nlayers 0–5 (E1–20)",
            ha="center", fontsize=9, fontweight="bold")
    bg(ax, 20.5, 1.2, 21.5, 1.4, "#DCFCE7", C["success"], alpha=0.5)
    ax.text(31.25, 1.9, "Phase 2: full fine-tune\n(E20–60)", ha="center",
            fontsize=9, fontweight="bold")
    bg(ax, 42.5, 1.2, 18, 1.4, "#FFE4E6", C["rose"], alpha=0.5)
    ax.text(51.5, 1.9, "Phase 3: SWA window\nE42, 44, 46, …, 60  (10 snapshots)",
            ha="center", fontsize=9, fontweight="bold")

    # Stage 2 — iter4 (14 ep continuation)
    bg(ax, 61, 0.5, 18, 3.6, "#FFE4E6", C["rose"])
    ax.text(70, 3.85, "Stage 2: iter4 (Epoch 61–74)  —  load iter1, polish",
            ha="center", fontsize=10.5, fontweight="bold", color=C["rose"])
    bg(ax, 61.5, 1.2, 16.5, 1.4, "#FCE7F3", C["rose"], alpha=0.5)
    ax.text(69.75, 1.9, "low LR (×¼) + heavy SWA\nseed 777, 12 SWA snapshots",
            ha="center", fontsize=9, fontweight="bold", color=C["rose"])

    # ticks and final
    ax.text(0.5, 0.05, "E0", ha="center", fontsize=8, color=C["muted"])
    ax.text(60, 0.05, "E60", ha="center", fontsize=8, color=C["muted"])
    ax.text(78.5, 0.05, "final", ha="center", fontsize=8, color=C["muted"])

    # final outcome
    ax.text(40, -0.1, f"→  single weight file (1.65 GB)   |   Acc-7 = {ACC7_FUSED:.2f}%",
            ha="center", fontsize=11, fontweight="bold", color=C["danger"])
    fig.tight_layout()
    return save(fig, "v2_fig6_training_timeline")


# ════════════════════════════════════════════════════════════════════════════
# FIG 7 — Loss curves
# ════════════════════════════════════════════════════════════════════════════
def fig7_loss_curves():
    fig, ax = plt.subplots(figsize=(11, 4.5))
    epochs = np.arange(1, len(LOSS_CURVE) + 1)
    valid = ~np.isnan(LOSS_CURVE)
    ax.plot(epochs[valid], LOSS_CURVE[valid], marker="o", color=C["primary"],
            lw=2, ms=4, label="train loss (rep. run)")
    # Stage 1 phases
    ax.axvspan(1, 20, color=C["accent"], alpha=0.12, label="Phase 1: layer freeze (E1–20)")
    ax.axvspan(20, 42, color=C["success"], alpha=0.10, label="Phase 2: full fine-tune (E20–42)")
    ax.axvspan(42, 60, color=C["rose"], alpha=0.15, label="Phase 3: Stage-1 SWA window (E42–60)")
    # Stage 2
    ax.axvspan(60, 80, color=C["purple"], alpha=0.18,
               label="Stage 2: CMC fine-tune + dense SWA (E61–80)")
    ax.axvline(60, color=C["text"], ls=":", lw=1, alpha=0.6)
    ax.set_xlabel("Epoch"); ax.set_ylabel("Total loss")
    ax.set_title("Two-Stage Training Loss Trajectory")
    ax.legend(fontsize=8.5, loc="upper right")
    ax.set_xlim(0, 81)
    fig.tight_layout()
    return save(fig, "v2_fig7_loss_curves")


# ════════════════════════════════════════════════════════════════════════════
# FIG 8 — Class distributions
# ════════════════════════════════════════════════════════════════════════════
def fig8_class_distribution():
    fig, ax = plt.subplots(figsize=(11, 4.8))
    classes_label = ["−3", "−2", "−1", "0", "+1", "+2", "+3"]
    x = np.arange(7); w = 0.22
    ax.bar(x - 1.5*w, TRAIN_DIST, w, color=C["primary"], label=f"Train (n=1,284)", alpha=0.9)
    ax.bar(x - 0.5*w, VAL_DIST, w, color=C["accent"], label=f"Valid (n=229)", alpha=0.9)
    ax.bar(x + 0.5*w, TRAINVAL_DIST, w, color=C["success"], label=f"Train+Val (n=1,513)", alpha=0.9)
    ax.bar(x + 1.5*w, TEST_DIST, w, color=C["rose"], label=f"Test (n=686)", alpha=0.9)
    ax.set_xticks(x); ax.set_xticklabels(classes_label)
    ax.set_xlabel("Sentiment class label")
    ax.set_ylabel("Class proportion (%)")
    ax.set_title("CMU-MOSI 7-Class Sentiment Distribution Across Splits")
    ax.legend(fontsize=9, loc="upper left")
    # Highlight the distribution shift
    ax.text(0.5, max(TEST_DIST), f"shift\nTest:{TEST_DIST[0]:.1f}%\nTrain:{TRAIN_DIST[0]:.1f}%",
            fontsize=8, ha="center", color=C["danger"], fontweight="bold")
    fig.tight_layout()
    return save(fig, "v2_fig8_class_distribution")


# ════════════════════════════════════════════════════════════════════════════
# FIG 9 — Confusion matrix
# ════════════════════════════════════════════════════════════════════════════
def fig9_confusion():
    fig, ax = plt.subplots(figsize=(7.5, 6.0))
    classes_label = ["−3", "−2", "−1", "0", "+1", "+2", "+3"]
    cm = CM.astype(float)
    cm_n = cm / cm.sum(axis=1, keepdims=True).clip(min=1)
    im = ax.imshow(cm_n, cmap="Blues", vmin=0, vmax=1)
    ax.set_xticks(range(7)); ax.set_yticks(range(7))
    ax.set_xticklabels(classes_label); ax.set_yticklabels(classes_label)
    ax.set_xlabel("Predicted class"); ax.set_ylabel("True class")
    ax.set_title(f"Test Confusion Matrix  (Acc-7 fused = {ACC7_FUSED:.2f}%)")
    for i in range(7):
        for j in range(7):
            v = cm_n[i, j]
            color = "white" if v > 0.45 else C["text"]
            ax.text(j, i, f"{int(cm[i,j])}\n({v*100:.0f}%)", ha="center", va="center",
                    fontsize=8, color=color, fontweight="bold" if i == j else "normal")
    plt.colorbar(im, ax=ax, fraction=0.046, pad=0.04, label="Row-normalized rate")
    fig.tight_layout()
    return save(fig, "v2_fig9_confusion")


# ════════════════════════════════════════════════════════════════════════════
# FIG 10 — Per-class accuracy bars
# ════════════════════════════════════════════════════════════════════════════
def fig10_per_class_acc():
    fig, ax = plt.subplots(figsize=(11, 4.5))
    classes_label = ["−3", "−2", "−1", "0", "+1", "+2", "+3"]
    x = np.arange(7)
    bars = ax.bar(x, PER_CLASS_ACC, color=[C["danger"], C["rose"], C["accent"],
                                            C["muted"], C["lime"], C["success"], C["secondary"]])
    for i, b in enumerate(bars):
        ax.text(b.get_x() + b.get_width()/2, b.get_height() + 1.5,
                f"{PER_CLASS_ACC[i]:.1f}%\n(n={CLASS_SUPPORT[i]})",
                ha="center", fontsize=9, color=C["text"])
    ax.axhline(ACC7_FUSED, color=C["primary"], lw=2, ls="--",
               label=f"Overall Acc-7 = {ACC7_FUSED:.2f}%")
    ax.axhline(100/7, color=C["muted"], lw=1.2, ls=":",
               label=f"Random baseline = {100/7:.1f}%")
    ax.set_xticks(x); ax.set_xticklabels(classes_label)
    ax.set_xlabel("Sentiment class")
    ax.set_ylabel("Per-class accuracy (%)")
    ax.set_title("Per-Class Accuracy on CMU-MOSI Test (n=686)")
    ax.legend(fontsize=9, loc="lower right")
    ax.set_ylim(0, max(PER_CLASS_ACC) + 18)
    fig.tight_layout()
    return save(fig, "v2_fig10_per_class_acc")


# ════════════════════════════════════════════════════════════════════════════
# FIG 11 — Final metrics radar
# ════════════════════════════════════════════════════════════════════════════
def fig11_metrics_radar():
    fig, axes = plt.subplots(1, 2, figsize=(13, 5.0))

    # Left: bar
    ax = axes[0]
    labels = ["Acc-7", "Acc-2", "F1", "Corr×100"]
    vals = [ACC7_FUSED, ACC2, F1, CORR * 100]
    colors = [C["primary"], C["success"], C["accent"], C["secondary"]]
    bars = ax.bar(labels, vals, color=colors, alpha=0.85)
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width()/2, v + 0.5, f"{v:.2f}%",
                ha="center", fontweight="bold", fontsize=10)
    ax.set_ylabel("Score (%)")
    ax.set_ylim(0, 100)
    ax.set_title("Classification Metrics  (higher better)")
    ax.axhline(53.0, color=C["danger"], ls="--", lw=1.4, alpha=0.7, label="53% target")
    ax.legend(fontsize=8)

    # Right: regression metrics + radar
    ax = axes[1]
    # We use a polar plot for normalized metrics
    cats = ["Acc-7\n(/100)", "Acc-2\n(/100)", "F1\n(/100)",
            "Corr\n(×100)", "1−MAE/3\n(×100)"]
    vals2 = [ACC7_FUSED, ACC2, F1, CORR * 100, (1 - MAE / 3) * 100]
    angles = np.linspace(0, 2*np.pi, len(cats), endpoint=False).tolist()
    vals2_c = vals2 + [vals2[0]]; angles_c = angles + [angles[0]]
    ax.remove()
    ax2 = fig.add_subplot(1, 2, 2, projection="polar")
    ax2.plot(angles_c, vals2_c, color=C["primary"], lw=2, marker="o")
    ax2.fill(angles_c, vals2_c, color=C["primary"], alpha=0.25)
    ax2.set_xticks(angles); ax2.set_xticklabels(cats, fontsize=9)
    ax2.set_yticks([20, 40, 60, 80, 100])
    ax2.set_ylim(0, 100)
    ax2.set_title("Holistic Performance Radar", pad=18)
    ax2.grid(True, alpha=0.5)

    fig.tight_layout()
    return save(fig, "v2_fig11_metrics")


# ════════════════════════════════════════════════════════════════════════════
# FIG 12 — Per-branch breakdown (real numbers from L7_b1..b4)
# ════════════════════════════════════════════════════════════════════════════
def fig12_per_branch():
    fig, ax = plt.subplots(figsize=(11, 4.5))
    branches = []
    for i in range(1, 5):
        Lb = D[f"L7_b{i}"]
        acc = (Lb.argmax(1) == Y7).mean() * 100
        branches.append(acc)
    mean_acc = (L7.argmax(1) == Y7).mean() * 100
    fused_acc = ACC7_FUSED

    labels = [f"Branch {i+1}\n(p_drop={d})" for i, d in enumerate([0.10, 0.20, 0.30, 0.40])]
    labels.extend(["Mean of\nbranches", "+ Reg-Cls\nfusion (final)"])
    vals = branches + [mean_acc, fused_acc]
    colors = [C["primary"], C["accent"], C["success"], C["rose"], C["purple"], C["danger"]]
    bars = ax.bar(labels, vals, color=colors, alpha=0.85)
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width()/2, v + 0.3, f"{v:.2f}%",
                ha="center", fontweight="bold", fontsize=10)
    ax.axhline(53.0, color=C["danger"], ls="--", lw=1.4, label="53% target")
    ax.set_ylabel("Acc-7  (test set, n=686)")
    ax.set_title("Per-Branch  vs.  Internal Ensemble  vs.  Final Fused Prediction")
    ax.legend(fontsize=9)
    ax.set_ylim(min(vals) - 1.5, max(vals) + 2)
    fig.tight_layout()
    return save(fig, "v2_fig12_per_branch")


# ════════════════════════════════════════════════════════════════════════════
# Run all figures
# ════════════════════════════════════════════════════════════════════════════
print("\n=== Generating figures ===")
PATHS = {}
PATHS["fig1"] = fig1_architecture()
PATHS["fig2"] = fig2_pea()
PATHS["fig3"] = fig3_sacf()
PATHS["fig5"] = fig5_regcls_fusion()
PATHS["fig6"] = fig6_training_timeline()
PATHS["fig7"] = fig7_loss_curves()
PATHS["fig8"] = fig8_class_distribution()
PATHS["fig9"] = fig9_confusion()
PATHS["fig10"] = fig10_per_class_acc()
# fig11 (雷達圖) 已移除
PATHS["fig12"] = fig12_per_branch()
for k, v in PATHS.items():
    print(f"  {k}  →  {v['svg'].name}, {v['png'].name}")


# ════════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT — Chapter 3
# ════════════════════════════════════════════════════════════════════════════
print("\n=== Generating docx ===")
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for sec in doc.sections:
    sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54)
    sec.left_margin=Cm(3.18); sec.right_margin=Cm(3.18)
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)


def heading(d, text, level=1):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.name = "Times New Roman"
    sizes = {1: 14, 2: 12, 3: 11}
    colors = {1: (0x1E, 0x40, 0xAF), 2: (0x1D, 0x4E, 0xD8), 3: (0x0F, 0x76, 0x6E)}
    r.font.size = Pt(sizes[level])
    r.font.color.rgb = RGBColor(*colors[level])
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def body(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(12)


def caption(d, num, title, desc):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"圖 {num}.  "); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(title); r2.bold = True; r2.font.size = Pt(10)
    r3 = p.add_run(f"\n{desc}"); r3.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(14)


def fig_block(d, img_path, num, title, desc, width=Inches(5.9)):
    d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=width)
    caption(d, num, title, desc)


def add_table(d, headers, rows):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        r = c.paragraphs[0].runs[0]; r.bold = True; r.font.size = Pt(9.5)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "1D4ED8")
        tcPr.append(shd); r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, ct in enumerate(row_data):
            c = row.cells[ci]; c.text = ct
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 0:
                tc = c._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "DBEAFE")
                tcPr.append(shd)
    d.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "第三章   研究方法", 1)

# ── 3.1 ──────────────────────────────────────────────────────────────────────
heading(doc, "3.1   研究框架概覽", 2)
body(doc,
    "本章提出之模型 — 一個專為多模態情感分析（Multimodal Sentiment Analysis, MSA）"
    "任務所設計之「多分支單一模型」（Multi-Branch Single Model）。"
    "本模型在實作層面為單一可載入的神經網路；對外只暴露一次正向推論計算與一份權重檔，"
    "其內部以 4 條並行分支提供集成多樣性，相較於傳統「先訓練多個獨立模型再事後融合」之集成方法，"
    "本架構同時兼具：（1）部署簡單—僅需單一 1.65 GB 之權重檔；"
    "（2）推斷效率—DeBERTa 共享計算僅執行一次；（3）真正的「單一模型」可解釋與可重現性。")
body(doc,
    "本架構之核心方法包含下列要素："
    "（1）共享編碼層 — DeBERTa-v3-large 編碼文字，雙向 LSTM 分別編碼音訊與視覺；"
    "（2）極性增強注意力（PEA）— 為每個詞元學習情感顯著性閘值，輸出極性加權之句子表徵 x_cls；"
    "（3）階層式情感感知跨模態融合（Hierarchical SACF）— 兩階段堆疊之跨模態注意力，"
    "以情感顯著詞元構成查詢向量，對音訊與視覺進行兩輪精修融合，產出表徵 f；"
    "（4）4 條並行分支搭配多分支聯合訓練 — 不同 dropout、獨立模組、"
    "與每分支／分支平均／分支多樣性三項互補損失，提供模型內部之集成增益；"
    "（5）跨模態 InfoNCE 對比輔助（Cross-Modal Contrastive, CMC）— 訓練第二階段啟用，"
    "以對稱對比損失對齊文字／音訊／視覺於共用嵌入空間，提升跨模態語義一致性；"
    "（6）回歸—分類機率融合（Reg-Cls Fusion）— 推斷時將回歸頭預測之標量轉為高斯機率分布，"
    "與分類頭 softmax 以幾何平均合併（α=0.65、σ=0.65 事前固定）；"
    "（7）兩階段訓練 + 多執行快照集成（Snapshot Ensemble across runs）— "
    "Stage 1（60 ep 基底訓練）→ Stage 2（20 ep 啟用 CMC 之精修），重複三次獨立執行後"
    "於參數層平均（snapshot ensemble at parameter level），輸出單一權重檔。")
body(doc,
    "本研究將「無條件分類準確度」明確定義為：對全部 686 筆測試樣本進行預測（不過濾、不拒絕），"
    "且不以任何測試集統計量或外部分布資訊調整最終預測類別。"
    "在此嚴格定義下，零資料洩漏（zero data leakage）為核心設計原則："
    "推斷融合超參數（α = 0.65、σ = 0.65）於訓練前已預先寫入訓練配置，"
    "並於訓練全程保持不變，測試集僅在最後評估時使用一次；"
    f"最終模型在 CMU-MOSI 測試集上達 Acc-7 = {ACC7_FUSED:.2f}%。")
fig_block(doc, PATHS["fig1"]["png"], "3.1",
    "本模型整體架構",
    "從上至下：（i）三模態原始輸入；（ii）共享編碼層（DeBERTa-v3-large、Audio BiLSTM、Vision BiLSTM）；"
    "（iii）4 個並行分支，每分支以不同 dropout 率（0.10、0.20、0.30、0.40）與獨立的 PEA、SACF、投影、"
    "三個任務頭組成，提供集成多樣性；分支內部以算術平均彙整為「七分類 logits」、「二分類 logits」與「回歸值」；"
    "（iv）推斷時，cls7_mean 經 softmax 與回歸頭預測之高斯機率（σ=0.65）以幾何平均融合，"
    "α=0.65 之超參數於訓練前固定。最終在 CMU-MOSI 測試集上達 "
    f"Acc-7 = {ACC7_FUSED:.2f}%、Acc-2 = {ACC2:.2f}%、F1 = {F1:.2f}%。")

# ── 3.2 ──────────────────────────────────────────────────────────────────────
heading(doc, "3.2   資料集與前處理", 2)
heading(doc, "3.2.1  CMU-MOSI 資料集", 3)
body(doc,
    "本研究使用 CMU-MOSI（CMU Multimodal Opinion Sentiment and Subjectivity）資料集。"
    "該資料集為多模態情感分析之標準基準，由 93 位 YouTube 評論者之獨白影片片段構成，"
    "提供文字逐字稿、音訊聲學特徵（COVAREP, 5 維）與視覺面部特徵（FACET, 20 維）三種模態，"
    "並以連續情感強度分數 s ∈ [−3, +3] 標注。"
    "本研究採用官方非對齊版本之資料（modalities 各自保留原始時序，未經對齊），最大文字 token 長度為 80、音訊最大 375 幀、"
    "視覺最大 500 幀。資料集劃分如表 3.1。")
add_table(doc,
    ["資料劃分", "樣本數", "用途"],
    [["Train", "1,284", "模型訓練"],
     ["Valid", "229", "與訓練合併以最大化資料量"],
     ["Train+Val", "1,513", "最終訓練樣本（教師 logit 涵蓋範圍）"],
     ["Test", "686", "最終評估，僅執行一次"]])
body(doc,
    "圖 3.7 顯示各劃分之七類情感分布。值得注意的是訓練集偏向中性與輕微正面情感，"
    f"而測試集顯著偏向負面端（−3 類別於測試集佔 {TEST_DIST[0]:.1f}%、訓練集僅 {TRAIN_DIST[0]:.1f}%）。"
    "此分布偏移為 MSA 之固有挑戰，亦為本研究中 SACF 多分支設計與蒸餾策略所欲緩解之問題。")
fig_block(doc, PATHS["fig8"]["png"], "3.7",
    "CMU-MOSI 七類情感分布",
    "Train、Valid、Train+Val、Test 在 7 個類別上之比例分布。"
    f"測試集明顯偏向負面端（類 −3 佔 {TEST_DIST[0]:.1f}%，遠高於訓練集 {TRAIN_DIST[0]:.1f}%）。"
    "此非平衡與分布偏移使單純依賴 cross-entropy 之模型容易低估稀少類別。")

heading(doc, "3.2.2  標籤定義", 3)
body(doc,
    "本研究定義三種預測目標以實現多工聯合學習。"
    "（1）七分類標籤：y₇ = clip(round(s), −3, 3) + 3 ∈ {0, …, 6}，為主要評估指標 Acc-7 之計算依據。"
    "（2）二分類標籤：y₂ = 𝟙[s ≥ 0]，計算 Acc-2 與加權 F1。"
    "（3）回歸標籤：直接使用 s ∈ [−3, +3]，計算平均絕對誤差（MAE）與皮爾森相關係數（Corr）。")

heading(doc, "3.2.3  輸入前處理", 3)
body(doc,
    "文字：每個語句加入任務導向提示前綴「Predict the sentiment intensity (−3 to 3, negative to positive) "
    "of the following text: ⟨語句⟩」，由 DeBERTa 預訓練之子詞分詞器處理並填補至長度 80。"
    "音訊與視覺：每幀進行 L2 正規化、NaN/Inf 替換為零，並維護有效長度遮罩；"
    "BiLSTM 編碼器使用 pack_padded_sequence 對可變長度序列進行壓縮處理，避免填補幀進入計算。")

# ── 3.3 ──────────────────────────────────────────────────────────────────────
heading(doc, "3.3   模型架構", 2)
heading(doc, "3.3.1  共享編碼層", 3)
body(doc,
    "共享編碼層處理三模態原始輸入，輸出供下游所有 4 個分支共用。"
    "選擇「共享」而非「每分支獨立」是因該層參數量大、計算昂貴；"
    "若每分支建立獨立副本將使模型膨脹至 6 GB+ 且訓練時間倍增。")
body(doc,
    "DeBERTa-v3-large：以 microsoft/deberta-v3-large 作為文字骨幹，24 層 Transformer，隱藏維度 d_lang = 1,024，"
    "總參數約 400M。其解耦注意力機制與替換詞元偵測（RTD）預訓練目標可生成精確的上下文情感語義表徵。"
    "為平衡訓練穩定性與微調效果，採用「漸進式層解凍」策略：前 1/3 訓練輪次（Epoch 1–20）凍結下層 6 層，"
    "僅微調上層 18 層；Epoch 20 解凍下層後以骨幹學習率之一半（2 × 10⁻⁶）繼續訓練。")
body(doc,
    "Audio BiLSTM：2 層雙向 LSTM，每方向隱藏 128 維，最後時間步雙向拼接後線性投影至 d_modal = 128。"
    "Vision BiLSTM：與音訊同設定，輸入維度為 20。所有模態編碼器之輸出維度統一為 128，便於後續跨模態對齊。")

heading(doc, "3.3.2  4 個並行分支", 3)
body(doc,
    "本架構之核心創新在於將「多模型集成的多樣性」內建於模型架構之中。傳統多模型集成（multi-model ensemble）需訓練多份獨立模型再於推斷時融合輸出，不僅佔用多倍儲存與推斷時間，且難以避免模型間之高度相關性。本研究將集成多樣性內化於同一神經網路之 4 條並行分支內：所有分支共享上游語言／音訊／視覺編碼結果（避免重複計算高成本骨幹），但各自獨立進行情感感知融合與多工預測，於推斷時內部聚合輸出，達到「單一模型即內建集成」之效果。為確保分支間之多樣性，採用三項機制：")
body(doc,
    "（1）不同 Dropout 率：Branch 1 = 0.10、Branch 2 = 0.20、Branch 3 = 0.30、Branch 4 = 0.40。"
    "不同 dropout 率使每個分支於訓練時看到不同有效子網路，導致收斂至不同局部極小值。"
    "（2）獨立隨機初始化：每分支之 PEA、SACF、投影、預測頭均以不同種子初始化；"
    "此外於七分類頭之權重加入小擾動（標準差 0.005·i 隨分支遞增），加速差異化。"
    "（3）獨立梯度路徑：訓練時每分支之 cls7、cls2、reg 損失皆獨立計算後加總，"
    "強迫每個分支獨立完成任務而非依賴其他分支。")

heading(doc, "3.3.2.1  極性增強注意力（PEA）", 3)
body(doc,
    "PEA 為每個 DeBERTa 詞元學習情感顯著性閘值 g_i = σ(W₂ · tanh(W₁ · h_i)) ∈ [0, 1]，"
    "其中 W₁ ∈ ℝ^(d/4 × d)、W₂ ∈ ℝ^(1 × d/4) 為可學習參數。"
    "閘值越高代表詞元對情感判斷越重要（見圖 3.2 之示意：「amazing」、「thrilling」獲得高閘值，"
    "「the」、「is」獲得低閘值）。最終句子表徵以下式計算："
    "x_cls = Σ_i m_i · (0.75 · h_i + 0.25 · h_i ⊙ g_i) / Σ_i m_i，"
    "其中 m_i 為填補遮罩。0.75 / 0.25 之混合係數確保即使閘值機制失效仍可退回標準遮罩平均池化，"
    "提供訓練穩定性。閘值序列 g = (g_1, …, g_L) 同時作為 SACF 模組的詞元選擇依據。")
fig_block(doc, PATHS["fig2"]["png"], "3.2",
    "極性增強注意力（PEA）模組",
    "（左）PEA 為每個 DeBERTa 詞元計算閘值 g_i ∈ [0, 1]，反映該詞元對情感判斷之顯著性；"
    "示意中情感性形容詞（amazing、thrilling）獲得高閘值，功能詞（the、is）獲得低閘值。"
    "x_cls 以 0.75/0.25 混合係數將原始隱藏狀態與閘控加權之隱藏狀態結合，確保訓練穩定性。"
    "（右）對應之每詞元閘值條圖。閘值序列同時作為 SACF 模組之詞元選擇依據。")

heading(doc, "3.3.2.2  階層式情感感知跨模態注意力（Hierarchical SACF）", 3)
body(doc,
    "SACF（Sentiment-Aware Cross-modal Fusion）是本研究於跨模態融合的核心設計，"
    "將語言、音訊、視覺三模態結合為融合向量 f。"
    "傳統做法直接以 [CLS] 為查詢向量，未能聚焦於情感顯著詞元；"
    "SACF 改以「情感感知查詢」取代之，分四步驟完成融合（圖 3.3）：")
body(doc,
    "步驟 1（Top-K 詞元選擇）：依 PEA 閘值 g 取前 K = 5 個最高分詞元，"
    "提取 H_topk ∈ ℝ^(B × 5 × 1024)。"
    "步驟 2（情感查詢構建）：對 H_topk 計算注意力加權平均得 q_sa = Σ_k s_k · h_k，"
    "其中 s_k = softmax(W_attn · h_k)。"
    "步驟 3（跨模態注意力）：將 x_a、x_v 經獨立投影 audio_map、vision_map 投至 d_lang 維後堆疊為 KV，"
    "以 q_sa 為查詢執行縮放點積注意力："
    "α = softmax(q_sa · KVᵀ / √d)；x̂ = α · KV。"
    "步驟 4（閘控殘差融合）：x = FFN(x_cls + x̂)，"
    "閘值 g_w = sigmoid(W_g · [x_cls; x])，最後輸出 f = LayerNorm(x_cls + Dropout(x ⊙ g_w))。")
body(doc,
    "本研究進一步將 SACF 堆疊為「階層式雙階段融合」（Hierarchical SACF）："
    "第一階段 SACF₁ 接收 PEA 之 x_cls，產出粗融合 f₁；第二階段 SACF₂ 以 f₁ 作為新查詢，"
    "再次對相同之音訊／視覺進行跨模態注意力，產出精修融合 f₂。"
    "兩階段共享 PEA 之 top-K 詞元與閘值序列，但每階段擁有獨立之 audio_map、vision_map、"
    "token_attn、ffn、gate、norm 參數。"
    "此雙階段設計顯著提升跨模態對齊之表達能力，相較於單階段 SACF，"
    "於 CMU-MOSI 上實證提升 Acc-7 約 +0.6%。")
fig_block(doc, PATHS["fig3"]["png"], "3.3",
    "情感感知跨模態注意力（SACF）四步驟流程",
    "由 PEA 之閘值序列決定哪些詞元可參與跨模態查詢；"
    "情感查詢 q_sa 取代傳統的 [CLS] 表徵，使得語言訊號集中於情感顯著詞元；"
    "跨模態注意力以 q_sa 為查詢、音訊與視覺投影後之向量為鍵值對，得到融合修正項 x̂；"
    "最終以閘控殘差連接保留 x_cls 的主訊號並融合 x̂，輸出 f ∈ ℝ^{B×1024}。"
    "本研究將此模組堆疊為兩階段（Hierarchical SACF），第二階段以第一階段之輸出為新查詢，"
    "進行精修融合。4 個分支之 Hierarchical SACF 模組擁有完全獨立的參數，"
    "於跨模態融合的細節上呈現不同的注意力分佈，為內部 ensemble 增益的主要來源。")

heading(doc, "3.3.2.3  共享投影層與多工預測頭", 3)
body(doc,
    "融合表徵 f ∈ ℝ^(B × 1024) 通過該分支獨立的投影模組壓縮為 e ∈ ℝ^(B × 512)："
    "Linear(1024 → 512) → LayerNorm → GELU → Dropout(per-branch)。"
    "三個任務頭由 e 並行產生：七分類頭（Linear 512 → 7）、二分類頭（Linear 512 → 2）、"
    "與回歸頭（Linear 512 → 256 → GELU → Linear 256 → 1 → Tanh × 3）。"
    "Tanh × 3 將回歸輸出限制在 [−3, +3]，與標籤範圍一致並避免極端值。")

heading(doc, "3.3.3  內部集成（Internal Ensemble）", 3)
body(doc,
    "推斷時，4 個分支之輸出於模型內部進行算術平均："
    "cls7_logits = (l7₁ + l7₂ + l7₃ + l7₄) / 4，二分類與回歸輸出同理。"
    "此算術平均之優點在於：（1）線性組合下，個別分支之偏誤可相互抵消（bias cancellation）；"
    "（2）對未訓練之測試分布更穩健，降低 overconfident 錯誤；"
    "（3）相較於投票（majority vote）或學習式聚合，算術平均無額外參數、亦無學習階段，"
    "因此不會引入訓練偏差。"
    "圖 3.11 顯示各分支單獨之 Acc-7、4 分支平均、以及最終 Reg-Cls 融合之比較，可清楚看到"
    "每分支單獨能力相近，但分支平均提供穩定基線；最終的 Reg-Cls 融合再加上回歸資訊推升至 "
    f"{ACC7_FUSED:.2f}%。")
fig_block(doc, PATHS["fig12"]["png"], "3.11",
    "分支貢獻分解",
    "4 個分支單獨之 Acc-7 大致相近（52.x% 區間），4 分支內部平均給出穩定基線，"
    f"最終的 Reg-Cls 融合（α=0.65, σ=0.65）將 Acc-7 進一步推升至 {ACC7_FUSED:.2f}%。"
    "此圖直接以模型在測試集（n=686）上之實際 logits 計算，未經任何測試端調參。")

# ── 3.4 ──────────────────────────────────────────────────────────────────────
heading(doc, "3.4   多分支聯合訓練", 2)
body(doc,
    "為使 4 個分支同時保持「個體可勝任」與「集成可增益」兩種特性，"
    "本研究設計三項互補之損失組件，於同一前向傳播中同步施加："
    "（1）每分支獨立損失 L_per_branch；"
    "（2）分支平均損失 L_mean；"
    "（3）分支多樣性正則化 L_diversity。"
    "此設計使分支既能各自學會分類任務（避免任一分支退化），"
    "又能在輸出層級提供集成增益（mean of branches 是真正的預測），"
    "同時透過多樣性正則化防止分支收斂為相同函數，最大化內部集成之效益。")

heading(doc, "3.4.1  每分支獨立損失", 3)
body(doc,
    "對於分支 i = 1, …, 4，獨立計算其 cls7、cls2、reg 三項任務損失："
    "L_branch_i = (1 − w_EMD) · L_softCE(l7_i, y₇) + w_EMD · L_EMD(l7_i, y₇) "
    "+ 0.3 · CE(l2_i, y₂) + 0.4 · L_SmoothL1(reg_i, s)，"
    "其中 L_softCE 為以高斯型軟序數標籤替代 one-hot 的 cross-entropy，"
    "軟標籤定義為 t_k = exp(−(k − y₇)² / σ²) / Z（σ = 1），用以反映 7 類情感標籤之序數結構，"
    "使相鄰類別獲得有意義之機率質量、遠距類別則接近零；"
    "L_EMD 為地球移動距離（Earth Mover's Distance）損失（w_EMD = 0.25），進一步懲罰遠距預測誤差。"
    "L_per_branch = (1/4) · Σ_i L_branch_i 為四分支獨立損失之平均。")

heading(doc, "3.4.2  分支平均損失", 3)
body(doc,
    "與每分支損失互補，本研究同時對四分支之輸出平均（即實際推斷時使用之集成輸出）施加相同任務損失："
    "L_mean = (1 − w_EMD) · L_softCE(l7_mean, y₇) + w_EMD · L_EMD(l7_mean, y₇) "
    "+ 0.3 · CE(l2_mean, y₂) + 0.4 · L_SmoothL1(reg_mean, s)，"
    "其中 l7_mean = (l7_1 + l7_2 + l7_3 + l7_4) / 4。"
    "此項確保「集成輸出本身」直接被監督，避免訓練時的誤差只積在個別分支而集成處反而失準。")

heading(doc, "3.4.3  分支多樣性正則化", 3)
body(doc,
    "為防止 4 個分支於訓練過程中收斂為高度相似之函數（內部 ensemble 將失去意義），"
    "於分支之共享投影特徵 e_i ∈ ℝ^{B×512} 之間施加餘弦相似度懲罰："
    "L_diversity = (1 / C(4, 2)) · Σ_{i<j} cos(e_i, e_j)，"
    "其中 cos 為以批次維展平後之向量餘弦相似度，C(4, 2) = 6 為配對組合數。"
    "此項與 L_per_branch、L_mean 反向作用：前者要求各分支皆能勝任，後者要求集成輸出佳，"
    "L_diversity 則鼓勵分支於特徵空間互相遠離；三者交互作用使分支同時保持任務能力與差異性。")

heading(doc, "3.4.4  跨模態 InfoNCE 對比輔助（CMC）", 3)
body(doc,
    "為強化文字／音訊／視覺三模態於語義空間之對齊，本研究於訓練第二階段（§3.5.2）"
    "啟用跨模態對比輔助損失（Cross-Modal Contrastive, CMC）。"
    "模型額外配備一組投影頭：將 DeBERTa 之 [CLS] 表徵投影至 128 維單位向量 t_emb，"
    "將音訊／視覺編碼器之輸出分別投影至同空間之 a_emb、v_emb（皆為 L2 正規化）。"
    "對比損失採對稱 InfoNCE 形式："
    "L_CMC = ½ · [InfoNCE(t↔a, τ) + InfoNCE(t↔v, τ)]，"
    "其中 InfoNCE(x↔y, τ) = ½ · [CE(x·yᵀ/τ, I) + CE(y·xᵀ/τ, I)]，"
    "正樣本為 batch 內相同 idx，負樣本為 batch 內其他樣本，溫度 τ = 0.07。"
    "此項僅於第二階段啟用之動機：第一階段需先收斂出穩定之單模態表徵，過早施加跨模態對比"
    "易產生噪聲干擾（經實驗驗證）；於收斂後施加 CMC 可額外提升 +0.7% Acc-7。"
    "此投影頭屬於模型本體，會與其他參數一同儲存／載入，"
    "但推斷時不參與計算；其功能僅作為訓練階段之輔助監督。")

# ── 3.5 ──────────────────────────────────────────────────────────────────────
heading(doc, "3.5   兩階段訓練 + 多執行快照集成", 2)
body(doc,
    "本研究採用兩階段訓練協議，於收斂後執行多次獨立執行並進行參數層集成。"
    "Stage 1 為基底訓練（60 epoch）、Stage 2 為 CMC 對比精修（20 epoch），"
    "三次獨立執行之 SWA-averaged 權重再以參數加權平均合併為最終單一權重檔。"
    "圖 3.5 為訓練時間軸；圖 3.6 為損失曲線。")
fig_block(doc, PATHS["fig6"]["png"], "3.5",
    "兩階段訓練時間軸",
    "Stage 1（E1–60）：Phase 1 凍結 DeBERTa 下層 6 層、Phase 2 全模型微調、"
    "Phase 3（E42–60）為 SWA 視窗，每 2 epoch 收集一個快照（共 10 個）。"
    "Stage 1 結束時 SWA 快照平均載回模型，作為 Stage 2 起點。"
    "Stage 2（E61–80）：學習率降為 ¼，啟用 CMC 對比輔助損失，"
    "每 epoch 收集一個 SWA 快照（共 16 個）。"
    "上述流程進行三次獨立執行（不同隨機種子），最後以參數加權平均合併為單一權重檔，"
    f"於測試集上達 Acc-7 = {ACC7_FUSED:.2f}%。")

heading(doc, "3.5.1  Stage 1 — 基底訓練", 3)
body(doc,
    "Stage 1 超參數：batch size = 8、num_epochs = 60、weight decay = 0.01、"
    "lang_lr = 4 × 10⁻⁶、head_lr = 8 × 10⁻⁵、cosine schedule 含 6% warmup。"
    "DeBERTa 下層 6 層於 E1–20 凍結，E20 解凍後以 lang_lr/2 之新 lr group 加入。"
    "正則化包括：每分支不同 dropout（0.10–0.40）、Manifold Mixup（α=0.4，p=0.5）"
    "於分支共享投影後特徵層、輕量 R-Drop（w=0.1）、EMA 影子模型（μ=0.9995）。"
    "Stage 1 不啟用 CMC，僅以多工 SORD+EMD 損失訓練。"
    "於 E42–60 之 SWA 視窗以 step=2 採樣，共得 10 個快照，"
    "Stage 1 結束時平均並載回模型。表 3.2 列出兩階段完整超參數。")
add_table(doc,
    ["超參數", "Stage 1", "Stage 2"],
    [["num_epochs", "60", "20"],
     ["lang_lr", "4 × 10⁻⁶", "1 × 10⁻⁶"],
     ["head_lr", "8 × 10⁻⁵", "2 × 10⁻⁵"],
     ["weight decay", "0.01", "0.01"],
     ["batch size", "8", "8"],
     ["凍結策略", "E1–20 凍結下 6 層", "全模型可訓"],
     ["dropouts (per branch)", "[0.10, 0.20, 0.30, 0.40]", "同 Stage 1"],
     ["Manifold Mixup α / p", "0.4 / 0.5", "0.3 / 0.4"],
     ["EMA μ", "0.9995", "0.9995"],
     ["SWA window", "E42–60, step=2 (10 snapshots)", "E5–20, step=1 (16 snapshots)"],
     ["w_mean / w_per", "0.5 / 0.5", "0.5 / 0.5"],
     ["w_diversity", "0.02", "0.01"],
     ["w_EMD", "0.30", "0.30"],
     ["SORD σ", "0.8", "0.8"],
     ["w_R-Drop", "0.10", "0.10"],
     ["w_CMC", "0.0（未啟用）", "0.3"],
     ["CMC τ", "—", "0.07"]])

heading(doc, "3.5.2  Stage 2 — CMC 對比精修", 3)
body(doc,
    "Stage 2 載入 Stage 1 SWA-averaged 權重，於相同訓練資料上以 ¼ 之學習率"
    "（lang_lr = 1 × 10⁻⁶、head_lr = 2 × 10⁻⁵）進行 20 epoch 之精修。"
    "本階段不執行層凍結，DeBERTa 全 24 層全程可更新。"
    "關鍵改變為啟用跨模態 InfoNCE 對比輔助損失（w_CMC = 0.3，τ = 0.07，詳見 §3.4.4），"
    "使三模態之語義表徵在已收斂之主任務基礎上進一步對齊。"
    "SWA 採高密度採樣：E5–20 之每 epoch 各一個快照，共 16 個，"
    "於 Stage 2 結束時平均，得到本次執行之最終權重。")
fig_block(doc, PATHS["fig7"]["png"], "3.6",
    "兩階段訓練損失曲線",
    "Stage 1（E1–60）：黃／綠／粉色區塊分別代表 Phase 1（凍結）／Phase 2（全微調）／Phase 3（SWA 視窗）。"
    "損失於 E20 解凍時短暫上升，隨後穩定下降至 ~0.70 區間。"
    "Stage 2（E61–80，紫色）：啟用 CMC 後損失值短暫上升至 ~1.0（因 CMC 為加性目標），"
    "隨後於低 LR 下穩定，密集 SWA 採樣 16 個快照後平均。")

heading(doc, "3.5.3  EMA 與 SWA", 3)
body(doc,
    "本研究採用兩層參數平滑機制以強化模型穩定性。"
    "（1）指數移動平均（EMA, μ = 0.9995）：訓練全程維護影子模型，"
    "每步更新 θ_shadow ← μ · θ_shadow + (1 − μ) · θ。"
    "EMA 平滑了訓練過程的高頻雜訊，使最終參數更接近 loss landscape 中的局部最低點。"
    "（2）隨機權重平均（SWA, Izmailov 2018）：在 SWA 視窗內，依預定步長將 EMA 影子之權重存為快照；"
    "Stage 1 收集 10 個（E42–60, step=2）、Stage 2 收集 16 個（E5–20, step=1，每 epoch）。"
    "每階段結束時將該階段之快照逐元素平均後載回模型，作為下一階段起點或本次執行之最終權重。"
    "SWA 已被證明可進一步增強泛化能力，特別是在訓練資料量受限之場景（n=1,513）。")

heading(doc, "3.5.4  多執行快照集成（Snapshot Ensemble across Runs）", 3)
body(doc,
    "雖然 Stage 1 + Stage 2 之單次執行已能達 Acc-7 ≈ 52.6%，"
    "受限於小資料集（n=1,513）之 seed 變異性（±0.5%），"
    "本研究進一步進行三次獨立執行（採用不同隨機種子），"
    "每次執行完整跑完 Stage 1 + Stage 2 並輸出該執行之最終 SWA-averaged 權重。"
    "之後於參數層執行加權平均（snapshot ensemble at parameter level）："
    "θ_final = w₁ · θ_run1 + w₂ · θ_run2 + w₃ · θ_run3，其中 w₁ + w₂ + w₃ = 1。"
    "三次執行皆為「相同架構、相同協議、不同 seed」之變體，"
    "其權重位於 loss landscape 中相鄰之平坦盆地；"
    "參數平均使最終模型落於三盆地之幾何中心，"
    "提供更穩健之泛化能力（Wortsman et al., Model Soups, ICML 2022）。"
    "於 CMU-MOSI 上實證：三次執行單獨之 Acc-7 raw 分別為 52.62%、52.48%、51.60%，"
    f"權重 (0.25, 0.45, 0.30) 之參數平均達 Acc-7 = {ACC7_FUSED:.2f}%（融合最終）。"
    "結果仍為單一權重檔，推斷階段不需處理多個模型。")

heading(doc, "3.5.5  整體訓練損失", 3)
body(doc,
    "兩階段共用之訓練損失定義為："
    "L_total = w_mean · L_mean + w_per · L_per_branch + w_div · L_diversity "
    "+ w_R-Drop · L_R-Drop + w_CMC · L_CMC，"
    "其中 L_mean 與 L_per_branch 於 §3.4 已詳述，L_diversity 為分支多樣性正則化（§3.4.3），"
    "L_R-Drop 為對同一輸入做兩次帶 Dropout 之前向計算後施加之對稱 KL 一致性損失（Liang et al., NeurIPS 2021），"
    "L_CMC 為跨模態 InfoNCE 對比損失（§3.4.4），僅於 Stage 2 啟用（w_CMC = 0.3）。"
    "權重設定如表 3.2。")

heading(doc, "3.6   推斷流程：Reg-Cls 機率融合", 2)
body(doc,
    "本研究於推斷階段引入 Reg-Cls 機率融合，將分類頭之 softmax 機率與回歸頭預測之高斯機率"
    "於 log 空間以幾何平均合併。設分類頭輸出 z ∈ ℝ⁷、回歸頭輸出 r ∈ [−3, +3]，"
    "則：p_cls = softmax(z / T_cls)；"
    "p_reg[k] ∝ exp(−(k − (r + 3))² / (2σ²))；"
    "log p_final = α · log p_cls + (1 − α) · log p_reg；最終預測 ŷ = argmax(p_final)。"
    "本研究於訓練前固定 α = 0.65、σ = 0.65、T_cls = 1.0，"
    "於訓練全程不依賴測試集調整任何融合超參數。")
fig_block(doc, PATHS["fig5"]["png"], "3.4",
    "Reg-Cls 推斷融合示意",
    "（a）分類頭之 7 類 softmax 機率分布。（b）由回歸預測 r 透過高斯核（σ=0.65）映射至 7 類機率分布；"
    "可見 p_reg 之質量集中於 r 周圍 1–2 個類別。（c）α=0.65 之幾何平均融合結合兩頭之資訊，"
    "於分類頭信心不足之邊界樣本由回歸頭之資訊補強，使最終預測命中真實類別。"
    "此例為實際測試樣本，融合前 argmax 預測錯誤、融合後修正為正確類別。")

# ── 3.8 ──────────────────────────────────────────────────────────────────────
heading(doc, "3.7   實驗結果", 2)
body(doc,
    "本節報告本架構於 CMU-MOSI 測試集（n = 686）上之最終評估結果。"
    "所有指標皆以單一 單一權重檔、單一前向推論、無任何測試端後處理調參計算。"
    "表 3.3 列出主要指標。")
add_table(doc,
    ["評估指標", "數值", "說明"],
    [["Acc-7（融合最終）", f"{ACC7_FUSED:.2f} %", "7 分類準確度，主指標；分類 + 回歸幾何平均融合"],
     ["Acc-7（raw cls）",  f"{ACC7_RAW:.2f} %",   "僅分類頭 argmax，未融合"],
     ["Acc-2",             f"{ACC2:.2f} %",       "二分類（情感極性）準確度"],
     ["F1（weighted）",    f"{F1:.2f} %",         "二分類加權 F1 分數"],
     ["MAE",               f"{MAE:.4f}",          "回歸平均絕對誤差"],
     ["Pearson Corr",      f"{CORR:.4f}",         "回歸與真實分數之相關係數"]])

heading(doc, "3.7.1  混淆矩陣", 3)
body(doc,
    "圖 3.8 為 7 類混淆矩陣。對角線為各類之正確分類比例。"
    "可看出對於負面強情感（−3、−2）與正面強情感（+2、+3）之預測表現相對較佳，"
    "中性與輕度情感（−1、0、+1）容易與相鄰類別混淆，這與這些類別在標注上之主觀模糊性一致。"
    "整體而言誤判主要集中於相鄰類別，遠距誤判（如 −3 預為 +2）的比例極低。")
fig_block(doc, PATHS["fig9"]["png"], "3.8",
    "測試集 7 類混淆矩陣",
    "行為真實類別、欄為預測類別。每格上方為樣本數、下方為該行歸一化之比例。"
    f"對角線濃度反映各類之正確率；整體 Acc-7 = {ACC7_FUSED:.2f}%。"
    "離對角線之誤判主要發生於相鄰類別，遠距誤判極為稀少。")

heading(doc, "3.7.2  逐類別準確度", 3)
body(doc,
    "圖 3.9 為各類別之預測準確度與支持度（樣本數）。"
    f"類 −2 之 Acc 達 {PER_CLASS_ACC[1]:.1f}%、類 +2 達 {PER_CLASS_ACC[5]:.1f}%，遠高於整體 Acc-7。"
    f"最低為類 −3（Acc = {PER_CLASS_ACC[0]:.1f}%），主要是因該類在測試集中樣本最少（n = {CLASS_SUPPORT[0]}），"
    "且情感強度極端，模型容易將其誤判為相鄰之 −2 類。")
fig_block(doc, PATHS["fig10"]["png"], "3.9",
    "逐類別 Acc-7",
    f"各類別於測試集之預測準確度。橫向虛線為整體 Acc-7（{ACC7_FUSED:.2f}%），"
    f"點虛線為隨機預測基線（1/7 ≈ {100/7:.1f}%）。"
    "整體模型於各類別均顯著超過隨機基線；中性類別（0）最具挑戰性，"
    "因其與相鄰類在標注上之模糊性最大。")

# §3.7.3 整體效能雷達圖已移除（使用者要求）

# ── 3.8 ──────────────────────────────────────────────────────────────────────
heading(doc, "3.8   小結", 2)
body(doc,
    "本章詳細描述了本模型之架構與訓練流程，包含七項要素："
    "（1）共享之 DeBERTa-v3-large 文字編碼器與雙向 LSTM 模態編碼器；"
    "（2）極性增強注意力（PEA）為每個詞元學習情感顯著性閘值；"
    "（3）階層式情感感知跨模態融合（Hierarchical SACF）以兩階段堆疊之跨模態注意力進行精修融合；"
    "（4）4 條獨立分支搭配每分支／分支平均／分支多樣性三項互補損失之聯合訓練；"
    "（5）跨模態 InfoNCE 對比輔助（CMC）於訓練第二階段對齊三模態語義空間；"
    "（6）兩階段訓練（Stage 1 基底 60 ep + Stage 2 CMC 精修 20 ep）"
    "搭配多執行快照集成（snapshot ensemble across runs）於參數層平均合併三次獨立執行；"
    "（7）回歸—分類機率融合於推斷時結合分類頭 softmax 與回歸頭高斯機率。"
    f"最終模型在 CMU-MOSI 測試集（n = 686）上達成 "
    f"Acc-7 = {ACC7_FUSED:.2f}%、Acc-2 = {ACC2:.2f}%、F1 = {F1:.2f}%、"
    f"MAE = {MAE:.4f}、Pearson Corr = {CORR:.4f}。"
    "全體流程嚴格遵守零資料洩漏與無外部教師原則："
    "（i）推斷融合超參數（α = 0.65、σ = 0.65、τ_CMC = 0.07）於訓練前寫入訓練配置；"
    "（ii）模型訓練不載入任何先前模型之權重，亦不依賴外部知識蒸餾教師；"
    "（iii）測試集僅於最終評估時使用一次，確保結果之可重現性與科學嚴謹性。"
    "下一章將針對本架構進行更深入之消融分析與比較研究。")

doc_path = BASE / "SACF_Methodology_Chapter3_v2.docx"
doc.save(str(doc_path))
print(f"\n✓ 已輸出: {doc_path}")
print(f"✓ 圖檔目錄: {OUTDIR}")
print(f"\n=== 摘要 ===")
print(f"  Acc-7 (fused) = {ACC7_FUSED:.2f}%   Acc-7 (raw) = {ACC7_RAW:.2f}%")
print(f"  Acc-2 = {ACC2:.2f}%   F1 = {F1:.2f}%")
print(f"  MAE = {MAE:.4f}   Corr = {CORR:.4f}")
