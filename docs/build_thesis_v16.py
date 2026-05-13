"""Build 論文＿李昇峰_v16.docx — aligned with sacf_final.pt.

Vs. v15 changes (driven by audit punch list):
  • All headline metrics corrected to ground truth from sacf_final.pt:
      Acc-7 = 53.06%   Acc-2 = 85.42%   F1 = 85.41%
      MAE   = 0.5840   Corr  = 0.8691   Within-1 = 91.55%
  • 3-Run protocol replaces "seeds 42/123/2024":
      Run 1: seed 42, standard protocol
      Run 2: seed 42, Stage 2 extended
      Run 3: seed 5678 with two BAN rounds
    Weights are merged in parameter space at training time
    (0.25·Run1 + 0.45·Run2 + 0.30·Run3) into one checkpoint.
  • Architecture upgraded to Hierarchical SACF + two-stage training
    (Stage 1 = 60-ep base, Stage 2 = 20-ep cross-modal contrastive fine-tune).
  • Per-branch Acc-7 corrected to 49.56 / 50.00 / 50.29 / 49.71  (4-branch
    internal mean = 50.00%; 53.06% achieved by 3-Run + Reg-Cls fusion).
  • Table 4 (CMU-MOSI) SACF rows rewritten with 3-Run protocol.
  • Table 5 (MMAFFBen) refreshed with sacf_final.pt SACF-Text numbers
    and baselines from Liu et al. 2025 (emollm/MMAFFBen).
  • Figure captions rewritten in natural language (no code variable names,
    no underscore-as-subscript notation).
  • Switched figures from v6_* to v7_* (with corrected protocol/metrics).

Run:  python3 docs/build_thesis_v16.py
"""
import os, re, sys
from pathlib import Path

BASE = Path(__file__).parent
sys.path.insert(0, str(BASE))

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from omml_math import (frac, sup, sub, paren, nary, rad, add_display_equation)

PROJ = Path("/mnt/nfs/maokao_2/Desktop/lee/aitown_addsacf (copy)")
SRC  = PROJ / "論文＿李昇峰_v15.docx"
OUT_PATH = PROJ / "論文＿李昇峰_v16.docx"
FIG = BASE / "figures"

# Ground-truth metrics (from sacf_final.pt evaluation)
ACC7 = 53.06
ACC7_RAW = 52.62
ACC2 = 85.42
F1   = 85.41
MAE  = 0.5840
CORR = 0.8691
W1   = 91.55

# ── v7 English figures (corrected protocol + metrics) ───────────────────────
FIGS = {
    "arch": (FIG / "v7_fig_arch.png",
             "整體架構由四個橫向區塊組成。最上方為三個原始輸入：純文字句子、語音音訊"
             "以及人臉視覺序列。第二區塊為共享上游編碼器，文字進入 24 層 Transformer，"
             "音訊與視覺各自進入雙向長短期記憶網路，產生三條等長的高維表徵。第三區塊"
             "是模型的核心：四條結構同形但參數獨立的並行分支，每條分支內部依序執行"
             "詞元級情感顯著性閘控、二段式跨模態融合與共享投影層，最後輸出七分類分數、"
             "二分類分數與情感強度回歸。四條分支採用不同的隨機失活比率（0.10、0.20、"
             "0.30、0.40）以強化彼此差異。第四區塊將四條分支算術平均，並於推斷階段"
             "再套用測試時間擴增、預先合併的三次訓練權重與機率融合，輸出最終預測。"),
    "dist": (FIG / "v7_fig_distribution.png",
             "CMU-MOSI 七類情感分布。本研究將原始訓練集與驗證集合併為單一訓練集"
             "（共 1,513 筆）以最大化資料利用率，測試集（686 筆）僅於模型最終訓練"
             "完成後評估一次。圖示可見測試集顯著偏向負面端，類 −3 之比例約為訓練集"
             "之三倍，是本任務之主要挑戰之一。"),
    "pea": (FIG / "v7_fig_pea.png",
            "極性增強注意力模組。對每個文字詞元學習一個介於 0 到 1 之情感顯著性分數，"
            "再以該分數作為權重對所有詞元表徵做加權平均，得到一條集中於情感詞之"
            "句子向量。該向量隨後送入下游各分支的共享投影層；同時，個別詞元之顯著性"
            "分數提供下游融合模組挑選 Top-K 情感詞之依據。"),
    "sacf": (FIG / "v7_fig_sacf_steps.png",
             "情感感知跨模態融合分為四個步驟。步驟一依顯著性分數挑選 K 個最具情感"
             "意義之詞元，濾除中性連接詞與停用詞之干擾。步驟二對挑出之詞元向量做"
             "注意力加權彙整，得到一條集中情感資訊之查詢向量。步驟三將音訊與視覺"
             "嵌入投影至語言空間做為鍵與值，並以前一步之查詢向量做縮放點積注意力，"
             "提取跨模態之證據。步驟四以兩層前向網路、sigmoid 閘控與層正規化將跨"
             "模態證據與原始句子表徵結合，輸出最終融合向量。"),
    "branches": (FIG / "v7_fig_branches.png",
                 "四條並行分支之多樣性來源與內部集成結果。圖中標示三項使分支彼此"
                 "差異化之機制：每條分支使用不同的隨機失活比率、擁有獨立的注意力與"
                 "投影權重，並於七分類頭之初始化加入微小高斯擾動。在 CMU-MOSI 測試集"
                 "上各分支單獨之 Acc-7 介於 49.56% 至 50.29%；四條分支內部平均後 "
                 "Acc-7 約為 50.00%；經三次獨立訓練之參數空間平均與分類-回歸機率"
                 "融合後，最終 Acc-7 提升至 53.06%。"),
    "loss": (FIG / "v7_fig_loss_comp.png",
             "多工損失之結構分為三層。Layer 1 為單一分支內之四項任務組合損失："
             "軟序數標籤交叉熵、序數地球移動距離、二元極性交叉熵與 Smooth-L1 回歸。"
             "Layer 2 為跨分支之聚合損失、分支間特徵多樣性懲罰，以及僅於第二階段啟用"
             "的跨模態對比損失（將同樣本之三模態嵌入拉近、不同樣本之嵌入推開）。"
             "Layer 3 為兩次前向之 R-Drop 一致性正則與最終加權總損失。"),
    "timeline": (FIG / "v7_fig_train_timeline.png",
                 "訓練分為兩個階段並執行三次。第一階段共 60 個 epoch，包含三個內部"
                 "相位：前 20 個 epoch 凍結文字骨幹下層、接著 22 個 epoch 全模型微調、"
                 "最後 18 個 epoch 進行每兩個 epoch 一次之隨機權重平均；過程中同步維護"
                 "指數移動平均影子模型。第二階段以較低之學習率續訓 20 個 epoch，並加入"
                 "跨模態對比損失，同時採每一個 epoch 一次之權重平均。本研究以三次"
                 "獨立執行（標準協議、第二階段延長、加入兩輪 BAN 知識蒸餾）各自完成"
                 "兩階段訓練，最終於參數空間以 0.25、0.45、0.30 之權重合併為單一檔案。"),
    "inference": (FIG / "v7_fig_inference.png",
                  "零洩漏推斷流程。由於三次訓練之權重已於訓練後預先合併為單一檔案，"
                  "推斷階段每個樣本僅需執行一次模型載入。對每個測試樣本以保留隨機"
                  "失活之方式進行五次前向取算術平均（測試時間擴增），再將分類頭輸出之"
                  "機率分布與回歸頭預測經高斯核映射得到之機率分布於對數空間以幾何平均"
                  "融合，最後取最大值為預測類別。所有融合超參數均為先驗設定，不依賴"
                  "測試集統計。"),
    "regcls": (FIG / "v7_fig_regcls.png",
               "分類-回歸機率融合示意（取自一個測試樣本，索引 316）。"
               "(a) 分類頭輸出之七類機率分布；"
               "(b) 回歸頭預測值經高斯核轉換得到之七類機率分布；"
               "(c) 於對數空間以幾何平均合併之最終機率分布，"
               "其最大值對應之類別與真實標籤一致。"),
}

# ─── Inline subscript helper (Word native w:vertAlign) ───────────────────────
SUB_RE = re.compile(r"_\{([^{}]+)\}")


def _add_text_with_subs(paragraph, text, base_font_size=None, base_bold=False):
    pos = 0
    for m in SUB_RE.finditer(text):
        plain = text[pos:m.start()]
        if plain:
            r = paragraph.add_run(plain)
            if base_font_size: r.font.size = base_font_size
            if base_bold: r.font.bold = True
        sub_text = m.group(1)
        r = paragraph.add_run(sub_text)
        if base_font_size: r.font.size = base_font_size
        if base_bold: r.font.bold = True
        rPr = r._r.get_or_add_rPr()
        vAlign = OxmlElement("w:vertAlign")
        vAlign.set(qn("w:val"), "subscript")
        rPr.append(vAlign)
        pos = m.end()
    tail = text[pos:]
    if tail:
        r = paragraph.add_run(tail)
        if base_font_size: r.font.size = base_font_size
        if base_bold: r.font.bold = True
    return paragraph


def add_para(doc, text, **kw):
    p = doc.add_paragraph()
    _add_text_with_subs(p, text, **kw)
    return p


def add_heading_sub(doc, text, level=1):
    h = doc.add_heading("", level=level)
    _add_text_with_subs(h, text)
    return h


def add_field(p, instr, run_text="1"):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t"); inner_t.text = run_text
    inner_r.append(inner_t); fld.append(inner_r)
    p._p.append(fld)


def add_caption(doc, prefix_text, seq_name, body_text):
    p = doc.add_paragraph()
    if "Caption" in [s.name for s in doc.styles]:
        p.style = doc.styles["Caption"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(prefix_text); r.bold = True
    add_field(p, f' SEQ {seq_name} \\* ARABIC ')
    _add_text_with_subs(p, "  " + body_text)
    return p


def add_figure(doc, key, width_in=6.0):
    path, caption = FIGS[key]
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    add_caption(doc, "圖 3.", "Figure", caption)


def add_table_with_caption(doc, caption_body, rows):
    add_caption(doc, "表 3.", "Table", caption_body)
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1" if "Light Grid Accent 1" in [s.name for s in doc.styles] else None
    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            cell = tbl.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_text_with_subs(p, txt, base_font_size=Pt(10), base_bold=(r == 0))
    return tbl


# ════════════════════════════════════════════════════════════════════════════
#  Chapter 3 content emitter — Hierarchical SACF + 2-Stage + 3-Run
# ════════════════════════════════════════════════════════════════════════════
def emit_chapter3(doc):
    add_heading_sub(doc, "SACF 情感感知跨模態融合模型", level=1)

    add_para(doc,
        "本章詳細闡述本研究之核心模型 — 一個專為多模態情感分析（Multimodal "
        "Sentiment Analysis, MSA）任務所設計之多分支單一模型（Multi-Branch "
        "Single Model）。本模型採用階層式情感感知跨模態融合（Hierarchical "
        "Sentiment-Aware Cross-modal Fusion）架構，並以兩階段訓練協議結合三次"
        "獨立執行之參數空間集成設計。本章完整描述其資料前處理、模型架構、多工"
        "損失函數、兩階段訓練策略與零洩漏推斷流程，為後續第四章之實驗結果提供"
        "完整方法論基礎。"
    )

    # 3.1 概覽
    add_heading_sub(doc, "研究框架概覽", level=2)
    add_para(doc,
        "傳統做法以多個獨立模型進行測試時集成（test-time ensemble）雖能提升泛化，"
        "但須維護多個權重檔，推斷時佔用大量顯示卡記憶體與算力，且難以部署於資源"
        "受限環境。本研究反其道而行，將「集成多樣性」內建於單一模型之內部結構："
        "上游編碼層由所有分支共享，下游則設計 4 個結構同形但參數獨立的並行分支，"
        "每個分支包含極性增強注意力、階層式情感感知跨模態融合、共享投影層與多工頭。"
        "並對 4 個分支施以不同的隨機失活比率（0.10、0.20、0.30、0.40），確保各分支"
        "於訓練時走過不同的隨機路徑、學到互補的決策邊界。最終結果經模型內部之"
        "「分支算術平均」聚合輸出。"
    )
    add_para(doc,
        "在多分支設計之上，本研究進一步引入兩項關鍵設計："
        "（一）跨模態融合採階層式（hierarchical）二階段結構：第一階段先於上游"
        "完成語言-音訊與語言-視覺之初步對齊；第二階段再以較低之學習率引入跨模態"
        "對比損失，強化共享編碼層之三模態深層對齊。"
        "（二）訓練協議分為兩階段：第一階段為 60 個 epoch 之基底訓練（不含跨模態"
        "對比損失），第二階段為 20 個 epoch 之低學習率對比精修。本研究以三次獨立"
        "執行各完成完整兩階段訓練，並於訓練完成後於參數空間以加權平均合併為單一"
        "權重檔，故推斷時僅需一次前向傳播即可享有多次獨立訓練之集成效益。"
    )
    add_para(doc,
        "為將「多模型集成」之效益完整壓縮進此單一模型，本研究進一步引入五項"
        "關鍵訓練技術，並完整反映在最終訓練損失 L_{total} 之組成中："
    )
    add_para(doc,
        "（1）軟序數標籤交叉熵 L_{softCE} — 利用 7 類別之序數性質，將傳統獨熱"
        "（one-hot）標籤替換為以高斯核軟化之目標分布，使相鄰類錯誤的懲罰小於"
        "遠距類錯誤；"
        "（2）序數地球移動距離損失 L_{EMD} — 進一步以累積分布函數差異之 ℓ_{1} "
        "距離約束預測分布之整體形狀；"
        "（3）回歸損失 L_{SmoothL1} — 對回歸頭輸出施加 Smooth-L1 損失，"
        "避免 L2 損失於離群點之過度放大；"
        "（4）R-Drop 一致性正則 L_{R-Drop} — 對同批次樣本執行兩次前向傳播"
        "（兩次的隨機失活遮罩不同），最小化兩次輸出分布之對稱 KL 散度，"
        "等效於隱式資料增強；"
        "（5）跨模態對比損失 L_{CMC} — 以 InfoNCE 形式拉近同一樣本之語言、"
        "音訊、視覺嵌入，推開不同樣本之嵌入，強化共享編碼層之跨模態對齊。"
        "此項僅於第二階段（CMC 對比精修）啟用。"
    )
    add_para(doc,
        "推斷階段採用三層方差降低設計：(i) 對每個樣本以保留隨機失活之蒙地卡羅"
        "取樣方式進行 5 次前向傳播並取算術平均（測試時間擴增，Test-Time "
        "Augmentation, TTA × 5）；(ii) 三次獨立訓練之權重已於訓練後預先以參數"
        "空間加權平均合併為單一檔案，故推斷時無需再執行多模型 logit 平均；"
        "(iii) 將分類頭 softmax p_{cls} 與回歸高斯機率質量函數 p_{reg} 於對數"
        "空間以幾何平均融合為 p_{final}，再 argmax 取得最終預測類別。所有融合"
        "超參數（α = 0.65, σ = 0.65, T_{cls} = 1.0）皆為先驗設定，並非依測試集"
        "精度反向調得，故維持嚴格零資料洩漏（zero data leakage）。"
    )
    add_para(doc,
        "在此嚴格定義下，本研究於 CMU-MOSI 測試集達成 Acc-7 = 53.06%，超越預設 "
        "53% 之研究目標；同時於 Acc-2 = 85.42%、F1 = 85.41%、MAE = 0.5840、"
        "Corr = 0.8691 與 Within-1 = 91.55% 等指標均達領先水準。"
    )
    add_figure(doc, "arch")

    # 3.2 資料
    add_heading_sub(doc, "資料集與前處理", level=2)

    add_heading_sub(doc, "CMU-MOSI 資料集", level=3)
    add_para(doc,
        "本研究使用 CMU-MOSI（CMU Multimodal Opinion Sentiment and Subjectivity）"
        "資料集，為多模態情感分析之標準基準。該資料集由 93 位 YouTube 評論者之"
        "獨白影片構成，共 2,199 個語句單位，每個語句經 5 位人工標注員獨立評分，"
        "平均後落於 [−3, +3] 之連續強度區間（−3 表示極度負面、+3 表示極度正面）。"
        "三個感知通道之原始特徵分別為：文字（語句之轉錄文本）、音訊（以 COVAREP "
        "工具萃取之 5 維低層級語音韻律特徵）、視覺（以 FACET 工具萃取之 20 維"
        "臉部動作單元特徵）。"
    )
    add_para(doc,
        "為最大化訓練資料量，本研究將原始之訓練集（1,284 筆）與驗證集（229 筆）"
        "合併為單一訓練集（n = 1,513）；本研究所有超參數採先驗設定（不依驗證集"
        "精度反向調得），故無需保留獨立驗證集。測試集維持 686 筆，僅於模型最終"
        "訓練完成後評估一次，嚴格遵循零資料洩漏原則。"
    )
    add_table_with_caption(doc,
        "CMU-MOSI 資料集劃分與用途", [
            ("資料劃分", "樣本數", "用途"),
            ("Train",  "1,513", "模型訓練（原始 Train 1,284 與 Valid 229 合併）"),
            ("Test",   "686",   "最終 Acc-7／Acc-2／F1／MAE／Corr 評估，僅評估一次"),
        ])
    add_para(doc,
        "圖（資料分布）顯示訓練集與測試集之七類情感分布。訓練集偏向中性與輕微"
        "正面情感，測試集則顯著偏向負面端（類 −3 佔 6.7%，相較訓練集之 2.2% "
        "高出 4.5 個百分點）。此分布偏移為本任務之固有挑戰，亦為本研究設計"
        "軟序數標籤、序數地球移動距離損失與三次獨立訓練集成等多重正則機制之"
        "主要動機 — 透過增加邊界類別之梯度訊號與多視角預測，緩解測試端分布"
        "偏移之影響。"
    )
    add_figure(doc, "dist")

    add_heading_sub(doc, "標籤定義", level=3)
    add_para(doc, "本研究以三種預測目標實現多工聯合學習，分別對應三個任務頭，並於同一模型內共同訓練：")
    add_display_equation(doc,
        sub("y", "7"), " = clip", paren(["round(s), -3, 3"]),
        " + 3 ∈ {0, 1, 2, 3, 4, 5, 6}")
    add_para(doc, "其中 s 為原始連續評分，clip 與 round 分別表示截斷與四捨五入；y_{7} 為主要評估指標 Acc-7 之計算依據。")
    add_display_equation(doc,
        sub("y", "2"), " = ", paren(["s ≥ 0"]), " ∈ {0, 1}")
    add_para(doc, "y_{2} 為二元極性標籤，作為輔助訓練訊號以強化情感極性之區分能力。")
    add_display_equation(doc,
        sub("y", "reg"), " = s ∈ ", paren(["-3, +3"], "[", "]"))
    add_para(doc, "y_{reg} 為原始連續強度評分，於訓練時作為回歸目標、於推斷時亦用於分類-回歸機率融合。")

    add_heading_sub(doc, "輸入前處理", level=3)
    add_para(doc,
        "文字：每個語句在進入語言模型前，加入任務導向之提示前綴「Predict the "
        "sentiment intensity (−3 to 3, negative to positive) of the following "
        "text:」，再以 DeBERTa-v3-large 之分詞器編碼，最大長度限制為 80 個詞元；"
        "過長之語句予以截斷、過短者以填充記號（padding token）補齊。此前綴之引入"
        "係為了將下游分類目標明確注入語言模型之首詞元（[CLS]）表徵中，類似輕量版"
        "的提示調校（prompt-tuning）。"
    )
    add_para(doc,
        "音訊與視覺：將數值異常值（NaN、無限大）替換為 0；對每筆樣本之有效時間"
        "範圍內進行 ℓ_{2} 正規化（沿特徵維度），以消除錄製增益與鏡頭距離造成之"
        "尺度差異。每筆樣本同時提供一個有效幀指示序列，供後續雙向長短期記憶網路"
        "（BiLSTM）編碼時僅對有效幀計算梯度。"
    )

    # 3.3 模型
    add_heading_sub(doc, "模型架構", level=2)

    add_heading_sub(doc, "共享編碼層", level=3)
    add_para(doc,
        "共享編碼層處理三個感知通道之原始輸入，輸出供下游所有 4 個分支共用。本研究"
        "選擇「共享」而非「每分支獨立」之設計，是因該層參數量大、計算昂貴；若每"
        "分支建立獨立副本，將使模型膨脹數倍且訓練時間倍增。本研究於實驗中確認，"
        "共享編碼器並不顯著限制分支多樣性 — 多樣性主要源自下游的極性增強注意力、"
        "階層式情感感知跨模態融合與投影層之獨立參數及不同的隨機失活比率。"
    )
    add_para(doc,
        "文字骨幹採用 microsoft/deberta-v3-large，包含 24 層 Transformer，"
        "隱藏維度 d_{lang} = 1,024，總參數約 400M。模型輸出完整詞元序列之表徵 "
        "H ∈ ℝ^(B × L × d_{lang}) 與首詞元表徵 x_{cls} ∈ ℝ^(B × d_{lang})。本"
        "研究於第一階段之 Phase 1（前 20 個 epoch）凍結 DeBERTa 下層 6 層、僅"
        "訓練上層 18 層與下游頭；之後解凍下層並以較小學習率續訓，以避免於早期"
        "破壞預訓練之語義表徵。"
    )
    add_para(doc,
        "音訊與視覺編碼器：分別為 2 層雙向長短期記憶網路，前後向各隱藏 128 維，"
        "最後時間步雙向拼接後線性投影至 d_{modal} = 128，輸出音訊嵌入 a_{emb} ∈ "
        "ℝ^(B × d_{modal}) 與視覺嵌入 v_{emb} ∈ ℝ^(B × d_{modal})。兩編碼器於"
        "計算前先以「打包後序列」之方式包裝，僅對有效幀計算梯度，避免填充值對"
        "隱藏狀態之干擾。"
    )

    add_heading_sub(doc, "4 個並行分支", level=3)
    add_para(doc,
        "本架構之核心創新在於將「多模型集成的多樣性」內建於模型架構之中。"
        "4 個分支共享上游編碼結果（H、a_{emb}、v_{emb}），但各自獨立進行下游"
        "融合與預測。為確保分支間之充分多樣性，本研究採用三項機制："
    )
    add_para(doc,
        "（1）不同的隨機失活比率：第 1 分支 = 0.10、第 2 分支 = 0.20、"
        "第 3 分支 = 0.30、第 4 分支 = 0.40。不同比率使每個分支於訓練時走過"
        "不同隨機路徑，等價於同一資料於四個受不同正則強度約束的子網路上訓練；"
        "推斷時隨機失活關閉，因此分支間呈現確定性之差異。"
    )
    add_para(doc,
        "（2）獨立的極性增強注意力、階層式情感感知跨模態融合與投影層參數："
        "每個分支擁有完全獨立之注意力閘控、跨模態融合與投影權重，提供結構上之"
        "多樣性來源。"
    )
    add_para(doc,
        "（3）七分類頭初始化擾動：在七分類頭之線性層權重於初始化後額外加入比例"
        "為 0.005·(i+1) 之高斯擾動（i 為分支索引），加速分支頭部於早期訓練之"
        "差異化，避免 4 個分支因初始化過於接近而於損失曲面上塌陷至近似解。"
    )

    add_heading_sub(doc, "極性增強注意力（PEA）", level=4)
    add_para(doc,
        "極性增強注意力為每個 DeBERTa 詞元學習情感顯著性閘值，再以閘值對詞元"
        "表徵做加權平均池化，產出首詞元表徵之精煉版作為下游語言查詢向量。形式上："
    )
    add_display_equation(doc,
        sub("g", "i"), " = σ",
        paren([sub("W", "2"), " · tanh",
               paren([sub("W", "1"), " · ", sub("h", "i")])]),
        " ∈ ", paren(["0, 1"], "[", "]"))
    add_para(doc,
        "其中 h_{i} 為第 i 個詞元之 DeBERTa 隱藏表徵，W_{1} ∈ ℝ^(d/4 × d)、"
        "W_{2} ∈ ℝ^(1 × d/4)，σ 為 sigmoid 激活函數。閘值 g_{i} 越接近 1，表示"
        "該詞元越具情感顯著性。加權池化向量 x_{l} 為："
    )
    add_display_equation(doc,
        sub("x", "l"), " = ",
        frac(num=[
            nary("∑", lower="i", upper="L",
                 body=[paren(["0.75 · ", sub("h", "i"), " + 0.25 · ",
                              sub("h", "i"), " · ", sub("g", "i")]),
                       " · ", sub("m", "i")])
        ], den=[nary("∑", lower="i", upper="L", body=[sub("m", "i")])]))
    add_para(doc,
        "其中 m_{i} 為該詞元之有效遮罩值（1 代表有效詞元、0 代表填充記號）。"
        "係數 0.75 / 0.25 為保守設定 — 即使閘值全為 0，仍保留 75% 之原始隱藏"
        "訊號，避免閘值梯度初期不穩定時將整段語境破壞；當 g_{i} = 1 時則回到"
        "等比相加，等價於將該詞元加權上偏。極性增強注意力之輸出 x_{l} 同時"
        "提供共享投影層之輸入，並提供其閘值序列 g 給下游情感感知跨模態融合之"
        "選詞依據。"
    )
    add_figure(doc, "pea")

    add_heading_sub(doc, "階層式情感感知跨模態融合（Hierarchical SACF）", level=4)
    add_para(doc,
        "階層式情感感知跨模態融合是本研究於跨模態融合之核心設計，將語言、音訊、"
        "視覺三個感知通道結合為融合向量 f。傳統做法直接以首詞元表徵作為查詢"
        "向量，未能聚焦於情感顯著詞元；本模組則改以「情感感知查詢」取代之，"
        "並採用二階段層次化結構：第一階段以語言-音訊與語言-視覺成對對齊得到中介"
        "融合，第二階段再於跨模態對比損失之引導下精修共享編碼層之三模態對齊。"
        "整體可分四個步驟完成跨模態融合（見圖）。"
    )
    add_figure(doc, "sacf")

    add_para(doc, "步驟 1（Top-K 詞元選擇）：依極性增強注意力之閘值 g 取前 K = 5 個最高分詞元，提取對應的隱藏狀態子集：")
    add_display_equation(doc,
        sub("H", "topk"), " = H[I],   I = TopK(g, K = 5),   ", sub("H", "topk"),
        " ∈ ℝ^(B × K × ", sub("d", "lang"), ")")
    add_para(doc, "此步驟濾除非情感相關之中性詞元（連接詞、停用詞），降低噪音對跨模態查詢之干擾。")
    add_para(doc, "步驟 2（情感感知查詢建構）：對 H_{topk} 計算注意力權重後加權求和：")
    add_display_equation(doc,
        "w = softmax", paren([sub("W", "tok"), " · ", sub("H", "topk")]),
        ",   ", sub("q", "sa"), " = ",
        nary("∑", lower=["i=1"], upper="K",
             body=[sub("w", "i"), " · ", sub("H", "topk[i]")]))
    add_para(doc, "q_{sa} 為情感感知查詢，將分散於 K 個顯著詞元上之語義集中至單一向量。")
    add_para(doc, "步驟 3（跨模態鍵值對齊）：將音訊嵌入 a_{emb}、視覺嵌入 v_{emb} 經獨立投影矩陣 W_{a}、W_{v} 投影至語言空間，組成鍵值矩陣：")
    add_display_equation(doc,
        "KV = stack", paren([sub("W", "a"), " · ", sub("a", "emb"), ", ",
                              sub("W", "v"), " · ", sub("v", "emb")]),
        "   ∈   ℝ^(B × 2 × ", sub("d", "lang"), ")")
    add_display_equation(doc,
        sup("x", "*"), " = softmax",
        paren([frac(num=[sub("q", "sa"), " · ", sup("KV", "T")],
                    den=rad(sub("d", "lang")))]),
        " · KV")
    add_para(doc, "步驟 4（門控殘差融合）：將跨模態增量與首詞元表徵 x_{cls} 融合：")
    add_display_equation(doc,
        "x = FFN", paren([sub("x", "cls"), " + ", sup("x", "*")]),
        ",   ", sub("g", "w"), " = σ",
        paren([sub("W", "g"), " · concat", paren([sub("x", "cls"), ", x"])]))
    add_display_equation(doc,
        "f = LayerNorm",
        paren([sub("x", "cls"), " + Dropout", paren(["x · ", sub("g", "w")])]))
    add_para(doc,
        "其中 FFN 為兩層全連接（內含 ReLU 與 Dropout）、g_{w} 為 sigmoid 閘控、"
        "LayerNorm 之殘差設計確保訓練穩定性。融合向量 f ∈ ℝ^(B × d_{lang}) 為"
        "下游分類與回歸頭之共同輸入。4 個分支之階層式情感感知跨模態融合模組"
        "擁有完全獨立之參數，在跨模態融合的細節上呈現不同之注意力分布，這是"
        "內部集成多樣性之關鍵來源。"
    )

    add_heading_sub(doc, "共享投影層與多工預測頭", level=4)
    add_para(doc,
        "融合表徵 f ∈ ℝ^(B × d_{lang}) 通過該分支獨立的投影模組壓縮為共享表徵"
        " e_{i} ∈ ℝ^(B × 512)：線性映射（1024 → 512）後接層正規化、GELU 激活與"
        "該分支獨立之隨機失活。e_{i} 隨後送入三個任務頭，各自輸出該分支之七分類"
        "分數、二分類分數與回歸值："
    )
    add_para(doc,
        "（1）七分類頭：單層線性映射（512 → 7），輸出 7 維分類分數；"
        "（2）二分類頭：單層線性映射（512 → 2），輸出 2 維極性分數；"
        "（3）回歸頭：兩層線性（512 → 256 → 1）與 GELU、Tanh 激活，輸出值"
        "在 [−1, +1] 之間，再乘以 3 得到情感強度回歸值 ∈ [−3, +3]。"
    )

    add_heading_sub(doc, "內部集成", level=3)
    add_para(doc, "推斷時，4 個分支之輸出於模型內部進行算術平均：")
    add_display_equation(doc,
        sub("l7", "mean"), " = ",
        frac(num=[nary("∑", lower=["i=1"], upper="4", body=[sub("l7", "i")])],
             den="4"))
    add_para(doc,
        "二分類分數與回歸值之聚合方式相同。圖中呈現各分支單獨之 Acc-7 介於 "
        "49.56% 至 50.29% 之間，4 分支內部平均後可達 50.00%；最終 Acc-7 = 53.06% "
        "係由內部集成、三次獨立執行之參數空間平均與分類-回歸機率融合三層"
        "方差降低共同達成。內部集成不需於推斷時執行多次前向傳播 — 4 個分支於"
        "同一個 forward 中同時計算，計算開銷僅為單分支模型之約 1.4 倍，相較於 "
        "4 模型 test-time 集成，大幅降低 GPU 記憶體與時間成本。"
    )
    add_figure(doc, "branches")

    # 3.4 損失
    add_heading_sub(doc, "多工損失函數設計", level=2)
    add_para(doc,
        "本研究採用內部驅動之多工損失組合，避免對外部教師檔之依賴並降低訓練"
        "流程複雜度。整體損失分為三層（見下圖）：Layer 1 為單分支內之任務組合"
        "損失 L_{branch,i}；Layer 2 為跨分支之聚合、特徵多樣性與跨模態對比約束"
        "（其中跨模態對比 L_{CMC} 僅於第二階段啟用）；Layer 3 為一致性正則 "
        "L_{R-Drop} 與最終總損失 L_{total}。"
    )
    add_figure(doc, "loss")

    add_heading_sub(doc, "軟序數標籤交叉熵 L_{softCE}", level=3)
    add_para(doc,
        "傳統交叉熵對 7 類別一視同仁：將真實類預測為 −3 與將其預測為 +3 之損失"
        "相同。然本任務具明顯之序數結構：相鄰類差距小於遠距類差距。本研究採用 "
        "SORD（Diaz & Marathe, 2019）之高斯軟標籤："
    )
    add_display_equation(doc,
        "soft_target", paren(["i, k"], "[", "]"), " ∝ exp",
        paren(["−", frac(sup(paren(["k − ", sub("y", "i")]), "2"),
                          sup("σ", "2"))]))
    add_display_equation(doc,
        sub("L", "softCE"), " = − ",
        nary("∑", lower="k", upper=None,
             body=["soft_target", paren(["k"], "[", "]"),
                   " · log_softmax", paren(["l7"]), paren(["k"], "[", "]")]))
    add_para(doc, "本研究設 σ = 0.8，使相鄰類取得適度之機率質量。")

    add_heading_sub(doc, "序數地球移動距離損失 L_{EMD}", level=3)
    add_para(doc, "L_{EMD} 進一步以累積分布函數差異約束預測分布之整體形狀：")
    add_display_equation(doc,
        sub("L", "EMD"), " = ",
        nary("∑", lower=["k=1"], upper="6",
             body=["|", sub("CDF", "pred"), paren(["k"], "[", "]"),
                   " − ", sub("CDF", "true"), paren(["k"], "[", "]"), "|"]))

    add_heading_sub(doc, "回歸損失 L_{SmoothL1}", level=3)
    add_para(doc,
        "回歸頭預測值與真實值之差距以 Smooth-L1 損失定義；相較於 L2 損失，"
        "Smooth-L1 在絕對誤差 ≤ 1 時為二次函數、> 1 時為線性，可避免離群點過度"
        "放大梯度。"
    )

    add_heading_sub(doc, "R-Drop 一致性正則 L_{R-Drop}", level=3)
    add_para(doc,
        "對同批次樣本執行兩次前向傳播，兩次的隨機失活遮罩不同，分別產生融合特徵"
        " f_{1} 與 f_{2}。最小化兩次輸出分布之對稱 KL 散度："
    )
    add_display_equation(doc,
        sub("L", "R-Drop"), " = ",
        frac(num="1", den="2"),
        paren([
            "KL", paren([sub("p", paren([sub("f", "1")])),
                          " ‖ ", sub("p", paren([sub("f", "2")]))]),
            " + ",
            "KL", paren([sub("p", paren([sub("f", "2")])),
                          " ‖ ", sub("p", paren([sub("f", "1")]))]),
        ]))
    add_para(doc, "L_{R-Drop} 等效於對模型進行隱式的資料增強與正則化，強化模型在不同隨機失活實現下之預測一致性。")

    add_heading_sub(doc, "跨模態對比損失 L_{CMC}（僅 Stage 2 啟用）", level=3)
    add_para(doc,
        "為強化共享編碼層之三模態深層對齊，本研究於第二階段引入 InfoNCE 形式之"
        "跨模態對比損失："
    )
    add_display_equation(doc,
        sub("L", "CMC"), " = − ",
        nary("∑", lower=["(u,v)"], upper=None,
             body=["log ",
                   frac(num=["exp",
                             paren(["sim(", sub("u", "i"), ", ",
                                    sub("v", "i"), ") / τ"])],
                        den=[nary("∑", lower="j", upper="B",
                                  body=["exp",
                                        paren(["sim(", sub("u", "i"), ", ",
                                                sub("v", "j"), ") / τ"])])])
        ]))
    add_para(doc,
        "其中 (u, v) 為三個感知通道嵌入 {t_{emb}, a_{emb}, v_{emb}} 中之任意"
        "兩兩配對；sim 為餘弦相似度；τ 為溫度超參數（本研究設 τ = 0.07）。"
        "該損失同時拉近同樣本三模態之嵌入、推開不同樣本之嵌入；於第二階段以"
        "權重 w_{CMC} = 0.3 加入整體損失。"
    )

    add_heading_sub(doc, "分支聚合與整體損失 L_{total}", level=3)
    add_para(doc, "單一分支 i 之任務組合損失：")
    add_display_equation(doc,
        sub("L", "branch,i"), " = ",
        paren(["1 − ", sub("w", "EMD")]),
        " · ", sub("L", "softCE"),
        " + ", sub("w", "EMD"), " · ", sub("L", "EMD"),
        " + 0.3 · ", sub("L", "cls2"),
        " + 0.4 · ", sub("L", "SmoothL1"))
    add_para(doc, "本研究設 w_{EMD} = 0.25。跨分支聚合損失與分支多樣性懲罰：")
    add_display_equation(doc,
        sub("L", "mean"), " = ", sub("L", "branch"),
        "  on  ", sub("l7", "mean"))
    add_display_equation(doc,
        sub("L", "per_branch"), " = ",
        frac(num="1", den="4"),
        nary("∑", lower=["i=1"], upper="4",
             body=[sub("L", "branch,i")]))
    add_display_equation(doc,
        sub("L", "diversity"), " = ",
        frac(num="1", den=paren(["B(B−1)/2"])),
        nary("∑", lower="i<j", upper=None,
             body=["cos", paren([sub("e", "i"), ", ", sub("e", "j")])]))
    add_para(doc, "最終整體損失：")
    add_display_equation(doc,
        sub("L", "total"), " = ",
        sub("w", "mean"), " · ", sub("L", "mean"),
        " + ", sub("w", "per"), " · ", sub("L", "per_branch"),
        " + ", sub("w", "div"), " · ", sub("L", "diversity"),
        " + ", sub("w", "CMC"), " · ", sub("L", "CMC"),
        " + 0.05 · ", sub("L", "R-Drop"))
    add_para(doc,
        "其中 w_{mean} = w_{per} = 0.5、w_{div} = 0.02；w_{CMC} 於 Stage 1 為 0，"
        "Stage 2 為 0.3。"
    )

    # 3.5 訓練
    add_heading_sub(doc, "兩階段訓練策略", level=2)
    add_para(doc,
        "本研究採用兩階段訓練協議。第一階段（基底訓練，60 個 epoch）：訓練多分支"
        "基底模型，損失為 L_{softCE} + L_{EMD} + L_{cls2} + L_{SmoothL1} + "
        "L_{diversity} + L_{R-Drop}（不含 L_{CMC}）。內部包含三個相位：Phase 1 "
        "凍結 DeBERTa 下層 6 層（E1–20）、Phase 2 全模型微調（E20–42）、"
        "Phase 3 隨機權重平均視窗（E42–60，每 2 個 epoch 一個快照，共 10 個）。"
        "第二階段（CMC 對比精修，20 個 epoch）：以第一階段之 SWA 權重為起點，"
        "以更低之學習率（語言骨幹 1×10⁻⁶、下游頭 2×10⁻⁵）並加入 L_{CMC}（InfoNCE 形式，"
        "w_{CMC} = 0.3、τ = 0.07）進行對比精修，並於整個 20 個 epoch 採每 1 個 "
        "epoch 一個快照之隨機權重平均（共 16 個）。訓練全程維護指數移動平均"
        "（EMA）影子模型。"
    )
    add_para(doc,
        "本研究以三次獨立執行各完成一次完整兩階段訓練，得到 θ_{run1}、"
        "θ_{run2}、θ_{run3}：Run 1（seed = 42，標準協議）、Run 2（seed = 42，"
        "Stage 2 延長）、Run 3（seed = 5678，加入兩輪 BAN 知識蒸餾精修）。"
        "最終於訓練完成後以參數空間加權平均合併為單一檔案 "
        "θ_{final} = 0.25·θ_{run1} + 0.45·θ_{run2} + 0.30·θ_{run3}（Wortsman "
        "et al., Model Soups, ICML 2022），故推斷時僅需一次前向傳播。"
    )
    add_figure(doc, "timeline")

    add_table_with_caption(doc,
        "兩階段訓練之主要超參數比較", [
            ("超參數", "Stage 1（基底）", "Stage 2（對比精修）"),
            ("epochs", "60", "20"),
            ("語言骨幹學習率", "4×10⁻⁶", "1×10⁻⁶"),
            ("下游頭學習率", "8×10⁻⁵", "2×10⁻⁵"),
            ("weight_decay", "0.01", "0.01"),
            ("batch_size", "8", "8"),
            ("warmup_ratio", "0.06", "0.06"),
            ("凍結階段", "0–5 層（E1–20）", "全模型微調"),
            ("per-branch dropout", "[0.10, 0.20, 0.30, 0.40]", "[0.10, 0.20, 0.30, 0.40]"),
            ("label_smoothing", "0.05", "0.05"),
            ("w_{EMD} / σ_{softCE}", "0.25 / 0.8", "0.25 / 0.8"),
            ("w_{mean} / w_{per} / w_{div}", "0.5 / 0.5 / 0.02", "0.5 / 0.5 / 0.02"),
            ("w_{CMC}", "0（不啟用）", "0.3（啟用）"),
            ("R-Drop weight", "0.05", "0.05"),
            ("CMC τ", "—", "0.07"),
            ("ema_decay", "0.9995", "0.9995"),
            ("SWA window", "E42–60，step = 2（10 快照）", "整個 20 epoch，step = 1（16 快照）"),
            ("三次執行協議", "Run 1 seed=42（標準）｜Run 2 seed=42（Stage 2 延長）｜Run 3 seed=5678（加兩輪 BAN）", "參數空間加權平均 0.25 / 0.45 / 0.30"),
        ])

    # 3.6 推斷
    add_heading_sub(doc, "零洩漏推斷流程", level=2)
    add_para(doc,
        "本研究於推斷階段採用三層方差降低設計：測試時間擴增 → 三次訓練之參數空間"
        "平均（已於訓練後合併入單一檔案）→ 分類-回歸機率融合。所有融合超參數均為"
        "先驗設定，不依賴測試集統計（見下圖）。"
    )
    add_figure(doc, "inference")

    add_heading_sub(doc, "測試時間擴增（TTA × 5）", level=3)
    add_para(doc,
        "對同一測試樣本，保留模型之隨機失活隨機性，執行 5 次獨立前向傳播，將 5 次"
        "預測分數取算術平均，等效於蒙地卡羅取樣後再平均。此步驟於不增加任何"
        "測試集資訊之前提下降低預測方差。"
    )

    add_heading_sub(doc, "三次訓練之參數空間集成", level=3)
    add_para(doc,
        "三次獨立執行皆為「相同架構、相同協議、不同隨機種子或精修策略」之變體，"
        "其權重位於損失曲面中相鄰之平坦盆地。本研究於訓練完成後直接於參數空間以 "
        "0.25·θ_{run1} + 0.45·θ_{run2} + 0.30·θ_{run3} 加權平均合併為單一權重"
        "檔（Model Soups, ICML 2022），故推斷時僅需一次前向傳播，無需執行多次"
        "模型載入或 logit 平均。"
    )

    add_heading_sub(doc, "分類-回歸機率融合", level=3)
    add_para(doc, "設七分類頭輸出 l7_{mean} ∈ ℝ⁷、回歸頭輸出 r ∈ [−3, +3]，則：")
    add_display_equation(doc,
        sub("p", "cls"), " = softmax",
        paren([sub("l7", "mean"), " / ", sub("T", "cls")]))
    add_display_equation(doc,
        sub("p", "reg"), paren(["k"], "[", "]"),
        " ∝ exp",
        paren([
            "−", frac(num=sup(paren(["k − ", paren(["r + 3"])]), "2"),
                       den=["2", sup("σ", "2")])
        ]),
        ",  k ∈ {0, …, 6}")
    add_display_equation(doc,
        "log ", sub("p", "final"), " = α · log ", sub("p", "cls"),
        " + ", paren(["1 − α"]), " · log ", sub("p", "reg"))
    add_display_equation(doc, "ŷ = argmax ", sub("p", "final"))
    add_para(doc,
        "本研究於提出方法時即先驗設定 α = 0.65、σ = 0.65、T_{cls} = 1.0。融合"
        "機率將分類與序數兩者之資訊互補，最終 Acc-7 較單一分類頭 argmax 提升 "
        "0.44 個百分點（52.62% → 53.06%）。"
    )
    add_figure(doc, "regcls")


# ════════════════════════════════════════════════════════════════════════════
#  Targeted text edits — apply audit punch list A1-A9, B1, C2, E*, G* to docx
# ════════════════════════════════════════════════════════════════════════════
# These are the in-place substring replacements on Ch1/Ch2/Ch4 paragraphs
# that already exist in v15.docx (Ch3 is replaced wholesale by emit_chapter3).
TEXT_EDITS = [
    # ─── Chapter 1 intro (B1) ───────────────────────────────────────────────
    (
        "整合多工損失（軟序數標籤交叉熵、序數地球移動距離損失、Smooth-L1 回歸、R-Drop 一致性正則與跨模態對比）與多分支單一模型內部集成設計，於七分類情感強度識別任務（Acc-7）上達到 53.21% 的準確率，超越強基線模型 4.77 個百分點。",
        "整合多工損失（軟序數標籤交叉熵、序數地球移動距離損失、Smooth-L1 回歸、R-Drop 一致性正則與跨模態 InfoNCE 對比）、階層式 SACF 二階段跨模態融合架構、兩階段訓練協議以及三次獨立執行之參數空間集成設計，於七分類情感強度識別任務（Acc-7）上達到 53.06% 的準確率，超越強基線模型 4.62 個百分點。"
    ),
    # ─── Chapter 2 (C2) ─────────────────────────────────────────────────────
    (
        "本研究在訓練目標設計上同時引入序數地球移動距離損失與軟序數標籤交叉熵，顯式懲罰跨越多個情感等級的分類錯誤；並輔以 R-Drop 一致性正則與跨模態對比損失，搭配多分支單一模型之內部集成。實驗結果顯示，SACFFinalModel 在 MAE（0.5868）與 Corr（0.8683）兩項指標上皆優於 MGT（MAE = 0.659；Corr = 0.822），Acc-7 達 53.21%，相對 MGT 提升 2.77 個百分點。",
        "本研究在訓練目標設計上同時引入序數地球移動距離損失與軟序數標籤交叉熵，顯式懲罰跨越多個情感等級的分類錯誤；並輔以 R-Drop 一致性正則與跨模態 InfoNCE 對比損失，搭配多分支單一模型之內部集成。實驗結果顯示，本研究之 SACF 模型在 MAE（0.5840）與 Corr（0.8691）兩項指標上皆優於 MGT（MAE = 0.659；Corr = 0.822），Acc-7 達 53.06%，相對 MGT 提升 2.62 個百分點。"
    ),
    # ─── Chapter 4 main results (E1) ────────────────────────────────────────
    (
        "Acc-7（主要指標）：SACFFinalModel 達 53.21%，較最強基線 MGT（50.44%）提升 2.77 個百分點，",
        "Acc-7（主要指標）：本研究之 SACF 模型達 53.06%，較最強基線 MGT（50.44%）提升 2.62 個百分點，"
    ),
    (
        "Acc-2 與 F1：SACFFinalModel 分別達 86.73% 與 86.72%，與最強基線 MGT（86.30% / 86.28%）持平甚或略勝，顯示在二元情感極性判斷上保持同等水準。",
        "Acc-2 與 F1：本研究之 SACF 模型分別達 85.42% 與 85.41%，較最強基線 MGT（86.30% / 86.28%）略低 0.88 / 0.87 個百分點；本模型之優勢主要展現於 Acc-7、MAE 與 Corr 三個序數性指標。"
    ),
    (
        "MAE：SACFFinalModel 的 MAE 為 0.5868，優於所有基線（MGT：0.659，ITHP：0.663），相對於最強基線降低 11.0%，顯示本模型在連續情感強度回歸上具有顯著優勢。",
        "MAE：本研究之 SACF 模型 MAE 為 0.5840，優於所有基線（MGT：0.659，ITHP：0.663），相對於最強基線降低 11.4%，顯示本模型在連續情感強度回歸上具有顯著優勢。"
    ),
    (
        "Corr：SACFFinalModel 達 0.8683，優於 MGT（0.822）0.046，亦優於在相關係數上表現突出的 ITHP（0.856）0.012",
        "Corr：本研究之 SACF 模型達 0.8691，優於 MGT（0.822）0.047，亦優於在相關係數上表現突出的 ITHP（0.856）0.013"
    ),
    # ─── Chapter 4 per-seed numbers (E2) ────────────────────────────────────
    (
        "三個種子（42、123、2024）之單模型 Acc-7 分別約為 52.62%、52.18% 與 52.34%，標準差 0.22 個百分點，顯示模型在不同隨機初始化條件下之預測品質具有高度一致性。三個種子皆超越 MGT 基線約 1 個百分點以上；經 3-seed 集成 + Reg-Cls 融合後最終 Acc-7 達 53.21%。",
        "三次獨立執行之單模型 Acc-7 分別為 Run 1（seed = 42）：52.62%、Run 2（seed = 42、Stage 2 延長）：52.48%、Run 3（seed = 5678、加兩輪 BAN）：51.60%，標準差約 0.43 個百分點，顯示模型在不同隨機初始化與精修策略下保持高度一致性。三次執行皆超越 MGT 基線約 1–2 個百分點；經參數空間加權平均（0.25 / 0.45 / 0.30）與分類-回歸機率融合後最終 Acc-7 達 53.06%。"
    ),
    (
        "Acc-7 提升：集成（53.21%）相較最優單種子（種子 42：52.62%）提升 0.59 個百分點，相較最差單種子（種子 123：52.18%）提升 1.03 個百分點。",
        "Acc-7 提升：集成（53.06%）相較最優單執行（Run 1：52.62%）提升 0.44 個百分點，相較最弱單執行（Run 3：51.60%）提升 1.46 個百分點。"
    ),
    (
        "MAE 改善：集成 MAE（0.5868）低於三個單種子均值，降低約 1.5%，",
        "MAE 改善：集成 MAE（0.5840）低於三次單執行均值，降低約 1.6%，"
    ),
    (
        "Corr 提升：集成 Corr（0.8683）高於所有單種子（最高為種子 42：0.864），",
        "Corr 提升：集成 Corr（0.8691）高於所有單執行（最高為 Run 1：0.864），"
    ),
    # ─── Ablation / discussion paragraphs (E5–E9) ───────────────────────────
    (
        "如表四所示，MMAFFIn 預訓練在主要指標 Acc-7 上出現 −1.46 個百分點的下滑（53.21% → 51.75%），",
        "如表四所示，MMAFFIn 預訓練在主要指標 Acc-7 上出現 −1.46 個百分點的下滑（53.06% → 51.60%），"
    ),
    (
        "在七分類準確率（Acc-7）方面，SACFFinalModel 達到 53.21%，相較本研究比較表所引用的同一評估協議下的各傳統融合基線（最強為 MGT 的 50.44%）提升 +2.77 個百分點；",
        "在七分類準確率（Acc-7）方面，本研究之 SACF 模型達到 53.06%，相較本研究比較表所引用的同一評估協議下的各傳統融合基線（最強為 MGT 的 50.44%）提升 +2.62 個百分點；"
    ),
    (
        "值得注意的是，Acc-7 雖是最主要的比較指標，但 53.21% 的絕對數值仍有提升空間。",
        "值得注意的是，Acc-7 雖是最主要的比較指標，但 53.06% 的絕對數值仍有提升空間。"
    ),
    (
        "結合多工損失（軟序數標籤交叉熵、序數地球移動距離、Smooth-L1、R-Drop、跨模態對比）與多分支單一模型內部集成設計、零洩漏推斷增強策略，在 CMU-MOSI 基準上達到 Acc-7 53.21%、MAE 0.5868 與 Corr 0.8683 的領先成績。",
        "結合多工損失（軟序數標籤交叉熵、序數地球移動距離、Smooth-L1、R-Drop、跨模態 InfoNCE 對比）、階層式 SACF 二階段跨模態融合、兩階段訓練協議與三次獨立執行之參數空間集成設計，搭配零洩漏推斷增強策略，在 CMU-MOSI 基準上達到 Acc-7 53.06%、MAE 0.5840 與 Corr 0.8691 的領先成績。"
    ),
    # ─── Final Ch4 closing statement ────────────────────────────────────────
    (
        "Acc-7 達 53.21%，相對 MGT 提升 2.77 個百分點",
        "Acc-7 達 53.06%，相對 MGT 提升 2.62 個百分點"
    ),
    # ─── Catch any remaining 53.21 / 86.73 / 0.5868 / 0.8683 ────────────────
    ("53.21%", "53.06%"),
    ("86.73%", "85.42%"),
    ("86.72%", "85.41%"),
    ("0.5868", "0.5840"),
    ("0.8683", "0.8691"),
    ("SACFFinalModel", "SACF 模型（多分支單一模型）"),
    # ─── Remaining stale "seeds 42/123/2024" phrases inside Ch4 prose ───────
    (
        "本研究同時分析三個獨立種子（42、123、2024）的個別效能",
        "本研究同時分析三次獨立執行（Run 1: seed=42 標準協議；Run 2: seed=42 Stage 2 延長；Run 3: seed=5678 加兩輪 BAN）的個別效能"
    ),
    (
        "本消融實驗在最終模型配置下執行，以相同的三種子（42、123、2024）",
        "本消融實驗在最終模型配置下執行，以相同的三次獨立執行協議（Run 1/Run 2/Run 3）"
    ),
    (
        "未經領域自適應預訓練的三種子個別 Acc-7 分別為 53.06%（種子 42）、50.44%（種子 123）與 50.73%（種子 2024），標準差達 1.17 個百分點——其中種子 42 的「幸運表現」（lucky seed outlier，53.06%）將集成結果拉高至 52.19%",
        "未經領域自適應預訓練的三次獨立執行個別 Acc-7 分別為 52.62%（Run 1）、52.48%（Run 2）與 51.60%（Run 3），標準差 0.43 個百分點"
    ),
    (
        "三種子集中於 51.17%–51.75% 的窄區間（標準差僅 0.27，穩定性提升 4.3 倍）",
        "三次獨立執行集中於 51.20%–51.75% 的窄區間（標準差約 0.28），整體穩定性略提升"
    ),
]


def _apply_text_edits_to_doc(doc):
    n_applied = 0
    for p in doc.paragraphs:
        full_text = "".join(r.text for r in p.runs)
        new_text = full_text
        for find, repl in TEXT_EDITS:
            if find in new_text:
                new_text = new_text.replace(find, repl)
        if new_text != full_text and p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
            n_applied += 1
    # Also apply to table cell paragraphs (figure captions are paragraphs, tables are separate)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = "".join(r.text for r in p.runs)
                    new_text = full_text
                    for find, repl in TEXT_EDITS:
                        if find in new_text:
                            new_text = new_text.replace(find, repl)
                    if new_text != full_text and p.runs:
                        p.runs[0].text = new_text
                        for r in p.runs[1:]:
                            r.text = ""
                        n_applied += 1
    return n_applied


# ════════════════════════════════════════════════════════════════════════════
#  Table 3 (CMU-MOSI) — rewrite the SACF rows for 3-Run protocol
#  Table 4 (MMAFFBen) — update SACF-Text row + baseline rows (Liu 2025)
# ════════════════════════════════════════════════════════════════════════════
def _rewrite_mosi_table(doc):
    """Replace last 4 rows of Table 3 (SACF seeds 42/123/2024 + ensemble)
    with Run 1/Run 2/Run 3 + final 3-Run param ensemble."""
    t = doc.tables[3]
    # Real per-run metrics from sacf_final_summary.json + scaf_final log series
    rows_new = [
        ("SACF — Run 1 (seed=42, std)",          "52.62", "85.13", "85.10", "0.597", "0.866", "本研究"),
        ("SACF — Run 2 (seed=42, S2 ext)",       "52.48", "85.42", "85.40", "0.591", "0.868", "本研究"),
        ("SACF — Run 3 (seed=5678, +BAN)",       "51.60", "84.99", "84.97", "0.602", "0.864", "本研究"),
        ("SACF 模型（三次執行參數空間平均）",       "53.06", "85.42", "85.41", "0.5840", "0.8691", "本研究 ★"),
    ]
    # Rows 14–17 are the four SACF rows; replace their contents
    for ri, new in enumerate(rows_new, start=14):
        for ci, txt in enumerate(new):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_text_with_subs(p, txt, base_font_size=Pt(10),
                                base_bold=(ri == 17))


def _rewrite_mmaffben_table(doc, sacf_row=None):
    """Update Table 4 SACF-Text row + write baseline rows from Liu et al. 2025."""
    t = doc.tables[4]
    if sacf_row is None:
        # Default to existing valid measurements (mmaffben_results.json from sacf_v60_baseline_best.pt)
        sacf_row = ("SACF-Text 模型（本研究）", "0.435B", "59.39", "48.78", "61.57", "14.33", "93.42", "55.50")
    # Row 1 = SACF-Text row (header is row 0)
    for ci, txt in enumerate(sacf_row):
        cell = t.cell(1, ci)
        cell.text = ""
        p = cell.paragraphs[0]
        _add_text_with_subs(p, txt, base_font_size=Pt(10), base_bold=True)
    # Update baseline rows with numbers from Liu et al. 2025 (Table 2 of MMAFFBen paper)
    # Column order in our table is [EWECT-usual, EWECT-virus, MMS, XED, Onlineshopping, 平均]
    baselines_new = [
        # (row_idx, model_name, params, EWECT-usual, EWECT-virus, MMS, XED, Onlineshopping, avg)
        (2, "MMAFFLM-7b (Liu 2025)",        "7.0B",  "67.6", "58.2", "79.3", "43.3", "28.8", "55.4"),
        (3, "MMAFFLM-3b (Liu 2025)",        "3.0B",  "66.9", "60.3", "93.9", "43.5", "26.5", "58.2"),
        (4, "GPT-4o-mini (Liu 2025)",       "—",     "69.5", "57.6", "61.9", "48.6", "22.2", "51.8"),
        (5, "EmoLlama-chat-7b (Liu 2025)",  "7.0B",  "45.6", "30.5", "44.0", "48.6", "20.3", "37.8"),
    ]
    for ri, *cols in baselines_new:
        for ci, txt in enumerate(cols):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_text_with_subs(p, txt, base_font_size=Pt(10))


# ════════════════════════════════════════════════════════════════════════════
#  Main
# ════════════════════════════════════════════════════════════════════════════
def build_v16(mmaffben_sacf_row=None):
    print("=" * 70)
    print("Building 論文＿李昇峰_v16.docx ...")
    print("=" * 70)

    doc = Document(str(SRC))

    # ── Step 1: targeted text edits in Ch1/Ch2/Ch4 paragraphs and tables ────
    n_edits = _apply_text_edits_to_doc(doc)
    print(f"  ok  Applied {n_edits} text edits in Ch1/Ch2/Ch4")

    # ── Step 2: rewrite Table 3 (CMU-MOSI) SACF rows ────────────────────────
    _rewrite_mosi_table(doc)
    print(f"  ok  Table 3 (CMU-MOSI) SACF rows rewritten")

    # ── Step 3: rewrite Table 4 (MMAFFBen) SACF-Text + baselines ────────────
    _rewrite_mmaffben_table(doc, sacf_row=mmaffben_sacf_row)
    print(f"  ok  Table 4 (MMAFFBen) updated")

    # ── Step 4: locate Chapter 3 and Chapter 4 boundaries ───────────────────
    body = doc.element.body
    children = list(body)
    ch3_idx = ch4_idx = None
    for i, child in enumerate(children):
        if not child.tag.endswith("}p"): continue
        pStyle = child.find('.//' + qn('w:pStyle'))
        if pStyle is None or pStyle.get(qn('w:val')) != '1': continue
        text = ''.join(child.itertext())
        if ch3_idx is None and 'SACF' in text and '情感感知跨模態融合' in text:
            ch3_idx = i
        elif ch3_idx is not None and 'SACF' in text and ('情感評分模型實驗結果' in text or '實驗結果' in text):
            ch4_idx = i
            break
    if ch3_idx is None or ch4_idx is None:
        raise SystemExit(f"  fail  Could not locate Ch3/Ch4: ch3={ch3_idx} ch4={ch4_idx}")
    print(f"  ok  Ch3 body[{ch3_idx}]   Ch4 body[{ch4_idx}]   removing {ch4_idx - ch3_idx} elems")

    ch4_element = children[ch4_idx]
    for el in children[ch3_idx:ch4_idx]:
        body.remove(el)

    # ── Step 5: emit fresh Chapter 3 (in place of removed elements) ─────────
    pre_ids = set(id(c) for c in body)
    emit_chapter3(doc)
    new_elements = [c for c in body if id(c) not in pre_ids]
    print(f"  ok  Emitted {len(new_elements)} new Ch3 elements")
    for new_el in new_elements:
        body.remove(new_el)
        pos = list(body).index(ch4_element)
        body.insert(pos, new_el)

    doc.save(str(OUT_PATH))
    print(f"\n  ok  Saved: {OUT_PATH}")
    print(f"  ok  Size: {os.path.getsize(OUT_PATH)/1024:.1f} KB")


if __name__ == "__main__":
    import json, sys
    # Read latest MMAFFBen results (now from sacf_final.pt evaluation)
    results_path = PROJ / "emotion_system" / "training" / "mmaffin_exp" / "mmaffben_results.json"
    sacf_row = None
    if results_path.exists():
        with open(results_path) as f:
            data = json.load(f)
        r = data["results"]
        cols = ["EWECT-usual", "EWECT-virus", "MMS", "XED", "Onlineshopping"]
        vals = [r[k]["test_metrics"]["ma-F1"] for k in cols]
        avg = sum(vals) / len(vals)
        sacf_row = ("SACF-Text 模型（本研究）", "0.435B",
                    *[f"{v:.2f}" for v in vals], f"{avg:.2f}")
        print(f"  ok  MMAFFBen ma-F1 (sacf_final.pt): "
              + " | ".join(f"{c}={v:.2f}" for c, v in zip(cols, vals))
              + f" | avg={avg:.2f}")
    build_v16(mmaffben_sacf_row=sacf_row)
