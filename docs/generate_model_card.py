"""SACFFinalModel — Model Card generator (v3)

讀取 emotion_system/models/sacf_final_summary.json 中的最新訓練配置與測試指標，
並引用 docs/figures/v2_fig*.png 圖檔，輸出 docs/SACFFinalModel_Model_Card.docx。
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent
ROOT = BASE.parent
FIG = BASE / "figures"
MODELS = ROOT / "emotion_system" / "models"

SUMMARY = json.loads((MODELS / "sacf_final_summary.json").read_text())
CFG = SUMMARY["config"]
M = SUMMARY["metrics"]

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18); sec.right_margin = Cm(3.18)
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)


def heading(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.name = "Times New Roman"
    sizes = {1: 16, 2: 13, 3: 11}
    colors = {1: (0x1E, 0x40, 0xAF), 2: (0x1D, 0x4E, 0xD8), 3: (0x0F, 0x76, 0x6E)}
    r.font.size = Pt(sizes[level])
    r.font.color.rgb = RGBColor(*colors[level])
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def body(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(12)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(11)


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = "Consolas"; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)


def caption(num, title, desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"圖 {num}  "); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(title); r2.bold = True; r2.font.size = Pt(10)
    r3 = p.add_run(f"\n{desc}"); r3.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(12)


def fig_block(img_path, num, title, desc, width=Inches(6.0)):
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=width)
    caption(num, title, desc)


def table_caption(num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"表 {num}  "); r1.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(title); r2.bold = True; r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        run = c.paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "1D4ED8")
        tcPr.append(shd); run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, ct in enumerate(row_data):
            c = row.cells[ci]; c.text = ct
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9.5)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 0:
                tc = c._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "DBEAFE")
                tcPr.append(shd)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# 標題
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SACFFinalModel — 多分支單一模型介紹（v3）")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Multi-Branch Single Model · 444M parameters · Two-Stage SWA · "
              "Manifold Mixup · SORD · CMC · Reg-Cls Fusion")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

# ─────────────────────────────────────────────────────────────────────────────
heading("1.  模型摘要", 1)
body(
    "SACFFinalModel 是本研究最終提交的多模態情感分析模型，於 CMU-MOSI test 集達成 "
    f"Acc-7 = {M['Acc-7']:.2f}%（融合）／{M['Acc-7-raw']:.2f}%（raw cls），"
    f"Acc-2 = {M['Acc-2']:.2f}%，F1 = {M['F1']:.2f}%，MAE = {M['MAE']:.4f}，Corr = {M['Corr']:.4f}。"
    "在「架構層面」就是一個單一模型——對外只暴露一個 nn.Module、一個 forward 方法、一個 "
    "sacf_final.pt 檔案。其關鍵創新在於把「多模型集成」的精度增益內建於模型架構中："
    "共享一個 DeBERTa-v3-large 文字骨幹與兩個 BiLSTM 模態編碼器，"
    "但在融合與預測層級設置 4 個並行分支，每個分支有獨立的 PEA + SACF + 投影 + 任務頭。")

body(
    "相較於前一代採用 12 模型事後 logit-ensemble 的版本，v3 在訓練流程上新增三項技術："
    "（i）兩階段 SWA — 60 epoch 主訓練後切換種子再 20 epoch 微調，獲得更多樣的 SWA 快照；"
    "（ii）SORD（Soft Ordinal Regression Distribution）+ squared EMD — 強化情感序數結構；"
    "（iii）跨模態對比學習（CMC）+ Manifold Mixup — 提升融合表徵的判別力與泛化能力。"
    "推斷時新增 Reg-Cls 機率融合機制，將回歸分支的連續預測平滑分配至離散類別機率，"
    f"進一步將 Acc-7 從 raw {M['Acc-7-raw']:.2f}% 提升至 {M['Acc-7']:.2f}%。")

table_caption(1, "SACFFinalModel 規格摘要")
add_table(
    ["項目", "設計", "說明"],
    [
        ["模型類型", "多分支單一模型", "對外是單一 nn.Module，內部 4 個並行分支"],
        ["總參數量", "444.1M", "骨幹 ~400M（共享）+ 4 分支 ~44M（獨立）"],
        ["輸入模態", "文字 + 音訊 + 視覺", "同步處理三模態，遮罩可變長度序列"],
        ["輸出", "cls7 + cls2 + reg", "7 類情感、二分類正負面、連續強度回歸"],
        ["檔案大小", "~1.65 GB", "單一 sacf_final.pt（fp32 state_dict）"],
        ["推斷成本", "1 次 forward", "DeBERTa 共享計算 + 4 分支並行融合"],
        ["訓練時間", "~3 小時", "80 epochs（stage1 60 + stage2 20）"],
        ["Acc-7（融合）", f"{M['Acc-7']:.2f} %", "CMU-MOSI test 集（686 樣本）"],
        ["Acc-2 / F1", f"{M['Acc-2']:.2f} / {M['F1']:.2f} %", "二分類正負面"],
        ["MAE / Corr", f"{M['MAE']:.4f} / {M['Corr']:.4f}", "回歸誤差與 Pearson 相關"],
    ])

# ─────────────────────────────────────────────────────────────────────────────
heading("2.  整體架構", 1)
body(
    "圖 1 呈現完整的 SACFFinalModel 架構。資料流分為三個階段："
    "（i）共享編碼層處理三模態原始輸入；"
    "（ii）4 個並行分支各自進行情感感知融合與預測；"
    "（iii）內部 ensemble 將分支輸出平均後送出。"
    "整個過程僅需「一次」端對端 forward 呼叫，無外部後處理。"
    "與舊版（12 模型事後融合）相比，新架構推斷一次完成，部署檔案壓縮至單一 1.65 GB。")

fig_block(FIG / "v2_fig1_architecture.png", 1, "SACFFinalModel 整體架構",
          "從上至下：（i）三模態輸入；（ii）共享編碼層（DeBERTa-v3-large、Audio BiLSTM、Vision BiLSTM）；"
          "（iii）4 個並行分支，每個內含 PEA、SACF、共享投影、三個任務頭；"
          "（iv）內部 mean-of-logits ensemble；（v）三個輸出（cls7、cls2、reg）。"
          "四種顏色分別對應 4 個分支的獨立參數。")

# ─────────────────────────────────────────────────────────────────────────────
heading("3.  各組件詳解", 1)

heading("3.1   共享編碼層", 2)
body(
    "共享層處理三模態的原始輸入，輸出供下游所有分支使用。之所以選擇「共享」而非「每分支獨立」，"
    "是因為這部分的參數量大、計算昂貴。若為每個分支建立獨立副本將使模型膨脹至 1.6 GB×4，"
    "且分散的梯度訊號反而稀釋表徵學習。")

heading("3.1.1   DeBERTa-v3-large 文字骨幹", 3)
body(
    "本模型採用 microsoft/deberta-v3-large 作為文字編碼器（~400M 參數、24 層 Transformer、"
    "隱藏維度 d_lang = 1024）。DeBERTa 的解耦注意力（disentangled attention）將內容與相對位置編碼"
    "分離計算，在情感分析這類需要細緻語義理解的任務上明顯優於 BERT/RoBERTa 同等模型。"
    "輸入文字會被加上提示前綴「Predict the sentiment intensity (-3 to 3, ...) of the following text: 」"
    "以引導模型聚焦於情感判斷。")

body(
    "為了平衡訓練穩定性與微調效果，採用「漸進式層解凍」策略："
    "前 1/3 訓練輪次（Epoch 1–20）凍結下層 6 層（Layers 0–5），僅微調上層 18 層；"
    "Epoch 20 後解凍下層，以骨幹學習率的一半（lang_lr / 2 = 2 × 10⁻⁶）繼續訓練，"
    "並透過自定義 LambdaLR 排程器將解凍層的學習率延續主排程的 cosine 衰減曲線。")

heading("3.1.2   Audio / Vision BiLSTM 模態編碼器", 3)
body(
    "音訊模態（COVAREP 聲學特徵，維度 5）與視覺模態（FACET 面部表情特徵，維度 20）"
    "皆採用 2 層雙向 LSTM，每方向隱藏狀態 128 維，輸出在最後時間步雙向拼接後線性投影至 d_modal = 128 維。"
    "為了避免填補幀干擾，使用 pack_padded_sequence 對可變長度序列進行壓縮處理。"
    "Dropout 率為 0.2（BiLSTM 層間）。兩個模態編碼器在訓練時參數獨立，"
    "因為音訊與視覺特徵的時間動態差異很大。")

heading("3.2   4 個並行分支", 2)
body(
    "這是 SACFFinalModel 的核心創新：把「多模型集成的多樣性」內建於模型架構。"
    "4 個分支共享上游的編碼結果（H、x_a、x_v），但各自獨立進行下游的融合與預測。"
    "每個分支由四個子模組組成：PEA、SACF、共享投影層、三個任務頭。"
    "為了確保分支之間的多樣性，採用三項設計：")

bullet("不同的 Dropout 率：Branch 1 = 0.10、Branch 2 = 0.20、Branch 3 = 0.30、Branch 4 = 0.40。"
       "不同 dropout 率會讓每個分支在訓練時看到不同的有效子網路，導致收斂到不同的解空間。"
       "v3 版本將最大 dropout 從 0.225 提升至 0.40，顯著加大分支差異性。")
bullet("不同的隨機初始化：所有獨立模組（PEA、SACF、Proj、Heads）採用 PyTorch 預設的隨機初始化，"
       "每個分支從不同點開始；cls7 head 額外加入了不同強度的隨機擾動，加速分支差異化。")
bullet("Per-branch 損失 + 多樣性懲罰：每個分支獨立計算其 cls7/cls2/reg 損失，"
       "並對分支兩兩之間的特徵 cosine similarity 加 0.02 權重的懲罰項，"
       "強迫每個分支獨立勝任所有任務且彼此差異化。")

heading("3.2.1   極性增強注意力（PEA）", 3)
body(
    "PEA 模組為每個詞元學習一個情感顯著性閘值 g_i ∈ [0, 1]："
    "g_i = σ(W₂ · tanh(W₁ · h_i))，其中 W₁ ∈ ℝ^(d/4 × d)、W₂ ∈ ℝ^(1 × d/4) 為可學習參數。"
    "閘值越高代表該詞元對情感判斷越重要（如「精彩」、「糟糕」、「非常」、「完全」等情感詞）。"
    "句子表徵透過閘值加權池化計算："
    "x_cls = Σᵢ mᵢ · (0.75 · h_i + 0.25 · g_i · h_i)，其中 mᵢ 為 attention mask。")

fig_block(FIG / "v2_fig2_pea.png", 2, "Polarity-Enhanced Attention（PEA）模組",
          "PEA 為每個詞元學習情感顯著性閘值 g_i，並以 0.75:0.25 的比例融合原始表徵與閘控表徵，"
          "使模型把更多注意力分配給情感詞。")

heading("3.2.2   情感感知跨模態注意力（SACF）", 3)
body(
    "SACF 是本研究的核心創新模組，負責將語言訊號與音訊／視覺進行跨模態融合。"
    "傳統做法直接以語言模型 [CLS] 表徵作為固定查詢向量，但這混合了句子的全域語義，"
    "未能聚焦於情感相關內容。SACF 改以「情感感知查詢」取代 [CLS]，分四步驟完成融合：")

bullet("步驟 1（Top-K 詞元選擇）：根據 PEA 閘值 g 取出前 K = 5 個最具情感顯著性的詞元，"
       "提取對應隱藏狀態 H_topk ∈ ℝ^(B × 5 × 1024)。")
bullet("步驟 2（情感查詢構建）：對 H_topk 中的 5 個詞元計算注意力加權平均，"
       "生成情感查詢向量 q_sa = Σ_k softmax(W_attn · H_topk)_k · H_topk,k ∈ ℝ^(B × 1024)。")
bullet("步驟 3（縮放點積跨模態注意力）：以 q_sa 為查詢，"
       "音訊與視覺表徵線性投影後拼接為鍵值對 KV = [W_a · x_a ; W_v · x_v] ∈ ℝ^(B × 2 × 1024)，"
       "計算縮放點積注意力 x̂ = softmax(q_sa · KV^T / √d) · KV。")
bullet("步驟 4（閘值殘差融合）：z = FFN(x_cls + x̂)；融合閘值 g_w = σ(W_g · [x_cls ; z])；"
       "最終輸出 f = LayerNorm(x_cls + Dropout(g_w ⊙ z)) ∈ ℝ^(B × 1024)。")

body(
    "4 個分支的 SACF 模組擁有完全獨立的參數（audio_map、vision_map、token_attn、ffn、gate、norm），"
    "因此會在跨模態融合的細節上呈現不同的注意力分佈，這是 ensemble 增益的主要來源。")

fig_block(FIG / "v2_fig3_sacf.png", 3, "Sentiment-Aware Cross-modal Fusion（SACF）模組",
          "以情感感知查詢 q_sa 取代傳統 [CLS] 查詢，並透過閘值殘差融合控制跨模態訊號注入語言表徵的強度。")

heading("3.2.3   共享投影層", 3)
body(
    "融合表徵 f ∈ ℝ^(B × 1024) 通過共享投影模組進行維度壓縮與特徵精煉："
    "Linear(1024 → 512) → LayerNorm → GELU → Dropout(per-branch)，"
    "輸出共享表徵 e ∈ ℝ^(B × 512)。「共享投影」中的「共享」指同一分支內三個任務頭共用此 e，"
    "並非跨分支共享參數。每個分支的 Proj 是獨立模組。")

heading("3.2.4   三個任務預測頭", 3)
body("每個分支從 e ∈ ℝ^(B × 512) 分支出去三個獨立任務頭：")
bullet("cls7 頭：Linear(512 → 7)，輸出 7 類情感的 logits（−3 到 +3 強度的離散映射）。")
bullet("cls2 頭：Linear(512 → 2)，輸出正負面二分類的 logits。")
bullet("reg 頭：Linear(512 → 256) → GELU → Linear(256 → 1) → Tanh × 3，"
       "輸出限制在 [−3, +3] 範圍的連續情感強度。")
body(
    "多任務聯合學習能讓模型同時學習離散分類與連續強度，互相監督下提升表徵品質。"
    "推論時三個輸出可同時取得；cls7 預測由「raw argmax」或「Reg-Cls 機率融合」獲得（見第 6 節）。")

# ─────────────────────────────────────────────────────────────────────────────
heading("4.  訓練損失設計（v3）", 1)
body(
    "v3 採用更豐富的損失組合以充分利用情感序數結構與跨模態對齊訊號。"
    "總損失（每階段獨立加權）為："
    "L_total = w_mean · L_mean + w_per · L_per_branch + diversity_weight · L_diversity "
    "+ w_rdrop · L_R-Drop + w_cmc · L_CMC，"
    f"其中 w_mean = {CFG['w_mean']}、w_per = {CFG['w_per']}、"
    f"diversity_weight = {CFG['diversity_weight']}、w_rdrop = {CFG['w_rdrop']}、"
    f"w_cmc（stage2）= {CFG['w_cmc_stage2']}（stage1 = {CFG['w_cmc_stage1']}）。"
    "L_per_branch 與 L_mean 內部各自由 SORD + squared-EMD + CE + SmoothL1 組成。")

heading("4.1   L_per_branch（每分支獨立任務損失）", 2)
body(
    "對 i = 1..4 每個分支獨立計算："
    "L_branch_i = 0.7 · SORD(l7_i, y_7) + emd_weight · sq-EMD(l7_i, y_7) "
    "+ 0.3 · CE(l2_i, y_2) + 0.4 · SmoothL1(reg_i, y_reg)。")
bullet(f"SORD（Soft Ordinal Regression Distribution，σ = {CFG['sord_sigma']}）："
       "把整數標籤 y 轉換為以 y 為中心、寬度 σ 的軟標籤分佈，並以 KL 散度監督預測機率。"
       "相較於硬標籤交叉熵，SORD 顯式編碼了 7 類情感的序數結構，"
       "讓相鄰類別之間的混淆受到較輕的懲罰。")
bullet(f"Squared EMD Loss（權重 emd_weight = {CFG['emd_weight']}）："
       "L_EMD = (1/(N · 6)) · Σᵢ Σ_c=1^6 (F̂_i(c) − F_i(c))²，"
       "其中 F̂_i(c) 為預測 CDF、F_i(c) 為真實 CDF。"
       "平方版 EMD 比 L1 版對遠距誤差懲罰更嚴格，加強序數一致性。")
bullet("Per-branch 損失強迫每個分支自身就是個能勝任全部任務的好預測器，"
       "避免分支「退化」為僅學一部分樣本。")

heading("4.2   L_mean（分支平均輸出損失）", 2)
body(
    "先計算 4 分支輸出的平均："
    "l7_mean = (l7_1 + l7_2 + l7_3 + l7_4) / 4（cls2、reg 同理），"
    "再對平均輸出計算 SORD + sq-EMD + CE + SmoothL1 組合損失。"
    "此項是「集成輸出本身」被直接優化的關鍵——"
    "讓 4 分支「合作」最大化集成預測的正確性，而不只是各自為政。"
    f"權重也是 {CFG['w_mean']}，與 per-branch 損失等重以保持兩者平衡。")

heading("4.3   L_diversity（分支多樣性正則化）", 2)
body(
    "若沒有顯式的多樣性懲罰，4 個分支可能逐漸收斂到相同的函數，使集成失去意義。"
    "本模型用「分支間 cosine similarity 平均」作為輕度懲罰："
    "L_div = (1 / C(4, 2)) · Σ_{i<j} cos(l7_i, l7_j)，"
    "其中 cos(·, ·) 在 batch 維度上平均。"
    f"權重 {CFG['diversity_weight']}（輕度，避免主損失被壓制）。"
    "這項懲罰會在訓練早期較顯著（分支接近同質時）、訓練後期自然下降。")

heading("4.4   L_R-Drop（一致性正則化）", 2)
body(
    "對同一 batch 跑兩次 forward（不同 dropout 遮罩），計算 l7_a 與 l7_b 的對稱 KL 散度："
    "L_KL = ½ [KL(p_a ‖ stop_grad(p_b)) + KL(p_b ‖ stop_grad(p_a))]，"
    "其中 p = softmax(l7)。"
    f"權重 {CFG['w_rdrop']}（v3 從 0.05 提高至 0.1 以強化隨機一致性）。"
    "R-Drop 強化「同樣輸入兩次預測一致」，相當於對模型施加隱式資料增強。")

heading("4.5   L_CMC（跨模態對比學習，stage2 新增）", 2)
body(
    "v3 在 stage2 微調階段加入跨模態對比學習（Cross-Modal Contrastive），"
    "目的是讓「同一樣本」在文字／音訊／視覺三條路徑下產生的低維投影互相靠近，"
    "而與「其他樣本」的投影互相遠離。具體做法是引入 CMCProjection 模組將三模態表徵"
    "投影至 d_proj = 128 維的對比空間，再對三對組合（t-a、t-v、a-v）"
    f"分別計算 InfoNCE 損失（溫度 τ = {CFG['cmc_tau']}），加總後取平均。"
    f"權重 w_cmc = {CFG['w_cmc_stage2']}（僅 stage2 啟用，stage1 為 0），"
    "在不擴增訓練資料的前提下強化跨模態語義對齊。")

heading("4.6   Manifold Mixup（融合層 mixup）", 2)
body(
    f"以 p = {CFG['mixup_p']} 的機率對 batch 觸發 Manifold Mixup，"
    f"在 SACF 輸出的融合特徵層級進行：以 Beta(α, α)（α = {CFG['mixup_alpha']}）抽樣 λ，"
    "隨機排列 batch 索引 perm，將融合特徵 feat 替換為 λ · feat + (1−λ) · feat[perm]，"
    "並對對應標籤同步做凸組合。在融合層而非輸入層做 mixup 能在保留原始特徵語義的前提下，"
    "顯著增加決策邊界附近的訓練樣本，提升泛化能力。")

# ─────────────────────────────────────────────────────────────────────────────
heading("5.  兩階段訓練流程", 1)
body(
    f"v3 採用 TrainVal 合併（1,284 train + 229 val = 1,513 樣本）作為訓練集，"
    "test 集（686 樣本）僅在訓練完成後評估「一次」，嚴格遵循零資料洩漏原則。"
    f"總共 {CFG['stage1_epochs']} + {CFG['stage2_epochs']} = "
    f"{CFG['stage1_epochs'] + CFG['stage2_epochs']} epoch，分兩階段執行：")

fig_block(FIG / "v2_fig6_training_timeline.png", 4, "兩階段訓練時程",
          "Stage 1（Epoch 1–60）：種子 42、漸進式層解凍、SWA 收集 10 個快照。"
          "Stage 2（Epoch 61–80）：切換種子至 1234、降低學習率 ¼、啟用 CMC、再次 SWA 收集。"
          "最終參數為兩階段 SWA 快照集的算術平均。")

bullet(f"Stage 1（Epoch 1–{CFG['stage1_epochs']}）：種子 = {CFG['seed']}，"
       f"骨幹學習率 {CFG['stage1_lang_lr']:.0e}，任務模組學習率 {CFG['stage1_head_lr']:.0e}，"
       f"從 Epoch {CFG['stage1_swa_start']} 開始每 {CFG['stage1_swa_step']} epoch 收集一次 SWA 快照。"
       "此階段以 SORD + sq-EMD + per-branch + diversity + R-Drop + Mixup 為損失組合（不含 CMC）。")
bullet(f"Stage 2（Epoch {CFG['stage1_epochs']+1}–{CFG['stage1_epochs']+CFG['stage2_epochs']}）："
       f"切換種子至 {CFG['stage2_seed']}，骨幹學習率降至 {CFG['stage2_lang_lr']:.0e}（1/4），"
       f"任務模組學習率 {CFG['stage2_head_lr']:.0e}，"
       f"從 stage2 第 {CFG['stage2_swa_start']} epoch 開始每 {CFG['stage2_swa_step']} epoch 收集快照，"
       f"並啟用 CMC 損失（w_cmc = {CFG['w_cmc_stage2']}）強化跨模態對齊。")
bullet("最終模型參數：對兩階段所有 SWA 快照進行算術平均，並更新 BatchNorm 統計量，"
       "輸出單一 sacf_final.pt（~1.65 GB）。")

table_caption(2, "v3 訓練超參數")
add_table(
    ["超參數", "值", "備註"],
    [
        ["batch size", str(CFG["batch_size"]), "混合精度（AMP, bfloat16）"],
        ["num_epochs", f"{CFG['stage1_epochs']} + {CFG['stage2_epochs']}", "兩階段固定 epoch"],
        ["stage1 lang_lr", f"{CFG['stage1_lang_lr']:.0e}", "DeBERTa 主學習率（cosine + warmup）"],
        ["stage2 lang_lr", f"{CFG['stage2_lang_lr']:.0e}", "微調階段（降至 1/4）"],
        ["stage1 head_lr", f"{CFG['stage1_head_lr']:.0e}", "PEA / SACF / Proj / Heads"],
        ["stage2 head_lr", f"{CFG['stage2_head_lr']:.0e}", "微調階段（降至 1/4）"],
        ["weight_decay", str(CFG["weight_decay"]), "AdamW"],
        ["dropout（基準）", str(CFG["dropout"]), "分支實際使用 [0.10, 0.20, 0.30, 0.40]"],
        ["focal γ", str(CFG["focal_gamma"]), "焦點聚焦參數"],
        ["label smoothing", str(CFG["label_smoothing"]), "Focal/CE 通用"],
        ["sord_sigma", str(CFG["sord_sigma"]), "SORD 軟標籤寬度"],
        ["emd_weight", str(CFG["emd_weight"]), "sq-EMD 在 cls7 損失中占比"],
        ["diversity_weight", str(CFG["diversity_weight"]), "分支間相似度懲罰"],
        ["w_rdrop", str(CFG["w_rdrop"]), "R-Drop 一致性權重"],
        ["w_cmc（stage2）", str(CFG["w_cmc_stage2"]), "跨模態對比損失（stage1 = 0）"],
        ["cmc_tau", str(CFG["cmc_tau"]), "InfoNCE 溫度"],
        ["mixup_alpha", str(CFG["mixup_alpha"]), "Manifold Mixup Beta 抽樣參數"],
        ["mixup_p", str(CFG["mixup_p"]), "Mixup 觸發機率"],
        ["EMA decay", str(CFG["ema_teacher_decay"]), "Exponential Moving Average"],
        ["SWA（stage1）", f"start={CFG['stage1_swa_start']}, step={CFG['stage1_swa_step']}", "10 個快照"],
        ["SWA（stage2）", f"start={CFG['stage2_swa_start']}, step={CFG['stage2_swa_step']}", "額外快照"],
    ])

fig_block(FIG / "v2_fig7_loss_curves.png", 5, "兩階段訓練損失曲線",
          "縱軸為 epoch-wise 平均訓練損失。"
          f"Epoch {CFG['stage1_epochs']} 處的不連續對應 stage1→stage2 切換（重置 optimizer 與種子）。"
          "曲線於 SWA 收集區段平緩，反映權重已進入損失盆地。")

# ─────────────────────────────────────────────────────────────────────────────
heading("6.  推斷流程（含 Reg-Cls 融合）", 1)
body(
    "推斷時採用 emotion_system/sacf_final_loader.py 中的 load_sacf_final 函數載入模型：")

code_block(
    "from emotion_system.sacf_final_loader import load_sacf_final\n"
    "model = load_sacf_final('emotion_system/models/sacf_final.pt', device='cuda:0')\n"
    "cls7_logits, cls2_logits, reg = model(\n"
    "    input_ids, attention_mask, audio, audio_mask, vision, vision_mask)")

body(
    "v3 推斷的關鍵新機制是 Reg-Cls 機率融合："
    "將回歸分支輸出的連續強度 reg ∈ [−3, +3] 視為 7 類離散標籤上的高斯軟分佈，"
    "再與 cls7 的 softmax 機率以幾何平均融合，最後取 argmax 作為最終預測。"
    "具體公式：")
code_block(
    "p_reg(c) ∝ exp(−(reg − (c−3))² / (2σ²)),  c ∈ {0, ..., 6}\n"
    "p_cls(c) = softmax(cls7_logits / T)\n"
    "p_final(c) ∝ p_cls(c)^α · p_reg(c)^(1−α)\n"
    f"ŷ = argmax_c p_final(c),  α = {CFG['fuse_alpha']}, σ = {CFG['fuse_sigma']}, T = {CFG['fuse_T']}")
body(
    "回歸分支提供「全局序數結構」（連續強度的可信估計），"
    "分類分支提供「局部判別力」（鄰近類別之間的鋒銳邊界），"
    "兩者幾何平均後達到互補效果。"
    f"在本研究的設定下，Reg-Cls 融合將 Acc-7 從 raw {M['Acc-7-raw']:.2f}% 提升至 {M['Acc-7']:.2f}%。")

fig_block(FIG / "v2_fig5_regcls_fusion.png", 6, "Reg-Cls 機率融合機制",
          "左：回歸輸出視為以 reg 為中心、σ 為寬度的高斯軟分佈 p_reg(c)；"
          "中：分類輸出 softmax(p_cls)；右：幾何平均後 argmax。"
          "α 控制兩者比重；本模型取 α = 0.65（偏向分類）。")

# ─────────────────────────────────────────────────────────────────────────────
heading("7.  實驗結果", 1)

body(
    f"在 CMU-MOSI test 集（686 樣本）上，SACFFinalModel v3 取得："
    f"Acc-7（融合）= {M['Acc-7']:.2f}%、Acc-7（raw）= {M['Acc-7-raw']:.2f}%、"
    f"Acc-2 = {M['Acc-2']:.2f}%、F1 = {M['F1']:.2f}%、MAE = {M['MAE']:.4f}、"
    f"Pearson 相關 = {M['Corr']:.4f}。所有指標皆為單次評估（無 TTA、無多次平均）。")

fig_block(FIG / "v2_fig11_metrics.png", 7, "整體測試指標",
          "六項指標的條形圖呈現（Acc-7 融合 / Acc-7 raw / Acc-2 / F1 / MAE / Corr）。"
          "Acc-7 為 7 類情感強度精度（−3 到 +3），Acc-2 為二分類精度。")

fig_block(FIG / "v2_fig9_confusion.png", 8, "Acc-7 混淆矩陣",
          "7 × 7 混淆矩陣，列為真實標籤、行為預測標籤。"
          "對角線元素表示正確預測，鄰近對角線的非零元素反映情感序數結構下的「軟混淆」。")

fig_block(FIG / "v2_fig10_per_class_acc.png", 9, "每類別精度",
          "7 個情感強度等級各自的 recall。極端類別（−3、+3）樣本稀少（見 fig 8），"
          "中性類別（0）樣本最多但語義模糊，精度天然較低。")

fig_block(FIG / "v2_fig12_per_branch.png", 10, "4 分支獨立 vs. 集成精度比較",
          "4 個分支獨立評估 Acc-7（藍）與 mean-of-logits 集成（橙）。"
          "集成增益顯示分支多樣性的實際貢獻。")

# ─────────────────────────────────────────────────────────────────────────────
heading("8.  檔案與重現指令", 1)
body("完整訓練（產生新 sacf_final.pt）：")
code_block(
    "cd /mnt/nfs/maokao_2/Desktop/lee/aitown_addsacf\\ \\(copy\\)\n"
    "CUDA_VISIBLE_DEVICES=0 python emotion_system/training/scaf_final.py")

body("推斷評估（在 CMU-MOSI test 集上驗證已訓練的 sacf_final.pt）：")
code_block(
    "python emotion_system/sacf_final_loader.py "
    "--ckpt emotion_system/models/sacf_final.pt")

body("關鍵檔案位置：")
table_caption(3, "檔案位置與用途")
add_table(
    ["檔案", "路徑", "用途"],
    [
        ["訓練腳本", "emotion_system/training/scaf_final.py",
         "定義 SACFFinalModel 類別 + 兩階段訓練主程式"],
        ["推斷載入器", "emotion_system/sacf_final_loader.py",
         "load_sacf_final() 函數 + CLI 評估工具"],
        ["模型權重", "emotion_system/models/sacf_final.pt",
         "單一檔案（~1.65 GB），含 444M 參數的 state_dict"],
        ["訓練摘要", "emotion_system/models/sacf_final_summary.json",
         "訓練配置 + 測試集指標的 JSON 紀錄"],
        ["論文圖檔", "docs/figures/v2_fig*.png",
         "本文件引用之所有圖檔的 PNG/SVG 原始檔"],
        ["資料集", "emotion_system/data/mosi/unaligned_50.pkl",
         "CMU-MOSI 非對齊版（train/valid/test）"],
    ])

body(
    "Model Card 撰寫日期：2026-05-12。"
    "更新原因：紀錄 v3 模型架構與訓練流程，"
    "新增 SORD、CMC、Manifold Mixup、Reg-Cls 融合等技術細節，並更新所有實驗數據。", indent=False)

# ─────────────────────────────────────────────────────────────────────────────
out = BASE / "SACFFinalModel_Model_Card.docx"
doc.save(str(out))
print(f"Saved: {out}")
print(f"  Acc-7 (fused) = {M['Acc-7']:.2f}%, Acc-7 (raw) = {M['Acc-7-raw']:.2f}%, "
      f"Acc-2 = {M['Acc-2']:.2f}%, F1 = {M['F1']:.2f}%, MAE = {M['MAE']:.4f}, Corr = {M['Corr']:.4f}")
