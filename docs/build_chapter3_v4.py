"""SACF Methodology Chapter 3 — v4 builder (aligned with v13 final model).

Differences vs. v3:
  • Removed v2-era DKD / DIST / SORD / Manifold-Mixup vocabulary.
  • Adopted v13 model losses: L_softCE, L_EMD, L_SmoothL1, L_R-Drop, L_CMC,
    combined per branch as L_branch_i and aggregated to L_total with weights
    w_mean, w_per, w_div, w_CMC.
  • Data partition: only Train (n = 1,513) + Test (n = 686).  Train+Val merged.
  • Many more architectural figures added (PEA detail, SACF 4-step, branches,
    loss composition, training timeline, inference pipeline).
  • Inline subscripts: any text token written as  W_{a}  /  L_{mean}  /  θ_{run1}
    is rendered with Word's native w:vertAlign w:val="subscript" run, so the
    file looks identical to MathType subscripts in the running paragraph text.
  • Display equations stay as OMML (m:oMath / m:oMathPara), MathType editable.

Run:  python3 docs/build_chapter3_v4.py
Output:  docs/SACF_Methodology_Chapter3_v4.docx
"""
import os, re, sys
from pathlib import Path

BASE = Path(__file__).parent
sys.path.insert(0, str(BASE))

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from omml_math import (frac, sup, sub, subsup, paren, func, matrix, rad, nary,
                       mr, omath_inline, omath_para,
                       add_display_equation, add_inline_equation)

OUT_PATH = BASE / "SACF_Methodology_Chapter3_v4.docx"
FIG = BASE / "figures"

# ────────────────────────────────────────────────────────────────────────────
# Figures used in v4 (file_basename → caption)
# ────────────────────────────────────────────────────────────────────────────
FIGS = {
    "arch": (FIG / "v4_fig_arch.png",
             "SACFFinalModel 整體架構。從上至下：（i）三模態原始輸入；（ii）共享編碼層產生語言"
             "嵌入 t_{emb}、音訊嵌入 a_{emb}、視覺嵌入 v_{emb}；（iii）4 個獨立並行分支，每"
             "分支依序通過 PEA、SACF、共享投影層與多工頭，產出 (l7_{i}, l2_{i}, reg_{i})；"
             "（iv）零洩漏推斷：TTA×5 與 3-seed ensemble 後做 Reg-Cls 機率融合。"),
    "dist": (FIG / "v4_fig_distribution.png",
             "CMU-MOSI 七類情感分布（Train + Test）。本研究將原始 Train 與 Valid 合併為單一 "
             "Train（n = 1,513）以最大化資料利用率，並對 Test（n = 686）執行唯一一次評估。"
             "Test 集顯著偏向負面端（類 −3 佔 6.7%，遠高於 Train 之 2.2%），形成 +4.5% 之分布"
             "偏移，是本任務之主要挑戰之一。"),
    "pea": (FIG / "v4_fig_pea.png",
            "極性增強注意力（PEA）模組詳細示意。對每個 DeBERTa 詞元 h_{i} 透過兩層 MLP 學習"
            "情感顯著性閘值 g_{i} ∈ [0, 1]；隨後依 attention_mask m_{i} 進行加權池化得到精煉"
            "向量 x_{l}。PEA 輸出有兩個用途：x_{l} 作為各分支共享投影層之輸入、g 作為 Top-K=5 "
            "選詞之依據，提供下游 SACF 之查詢構建。"),
    "sacf": (FIG / "v4_fig_sacf_steps.png",
             "情感感知跨模態注意力（SACF）逐步計算流程。步驟 1：以 PEA 閘值 g 選取 Top-K = 5 "
             "個情感顯著詞元，提取隱藏狀態 H_{topk}；步驟 2：對 H_{topk} 加權平均構建情感"
             "查詢 q_{sa}；步驟 3：將音訊 a_{emb}、視覺 v_{emb} 經 W_{a}、W_{v} 投影後組成"
             " KV，與 q_{sa} 做縮放點積注意力得 x*；步驟 4：FFN + sigmoid 閘控殘差與 LayerNorm "
             "得最終融合向量 f。"),
    "branches": (FIG / "v4_fig_branches.png",
                 "4 並行分支的多樣性來源與內部集成。三項多樣性機制：(a) 不同 dropout 率 "
                 "[0.10, 0.20, 0.30, 0.40]；(b) PEA / SACF / Proj 之獨立參數；(c) cls7 頭"
                 "於初始化時施加 0.005·(i+1)·N(0,1) 之微小擾動。各分支單獨 Acc-7 介於 "
                 "51.9%–52.6%，分支內部算術平均 l7_{mean} 可達到 53.21%（融合後）。"),
    "loss": (FIG / "v4_fig_loss_comp.png",
             "整體多工損失函數組成結構。Layer 1：每個分支 i 的任務組合損失 L_{branch_i} = "
             "(1 − w_{EMD}) L_{softCE} + w_{EMD} L_{EMD} + 0.3 L_{cls2} + 0.4 L_{SmoothL1}；"
             "Layer 2：分支聚合 L_{mean}、L_{per_branch}、L_{diversity} 與跨模態對比 L_{CMC}；"
             "Layer 3：R-Drop 對稱 KL 一致性正則 L_{R-Drop} 與 L_{total}。"),
    "timeline": (FIG / "v4_fig_train_timeline.png",
                 "訓練全景：漸進解凍 + EMA + SWA + 3-種子集成。單一 run 之 60 epoch 包含 "
                 "Phase 1（凍結 DeBERTa 下層 6 層，E1–20）、Phase 2（全模型微調，E20–42）、"
                 "Phase 3（SWA 視窗，E42, 44, …, 60，共 10 個快照）；EMA shadow（μ = 0.9995）"
                 "於全程維護；3 個獨立 seed（42、123、2024）各產出一份 θ_{run}，最後將 3 個 "
                 "θ_{run} 平均得 θ_{final}。"),
    "inference": (FIG / "v4_fig_inference.png",
                  "零洩漏推斷流程：TTA×5 + 3-Seed Ensemble + Reg-Cls 融合。三層方差降低："
                  "Stage 1 以 MC-Dropout 對同一樣本 forward 5 次取平均；Stage 2 將 3 個種子之 "
                  "logit 算術平均；Stage 3 將分類 softmax p_{cls} 與回歸高斯 PMF p_{reg} 於 "
                  "log 空間以幾何平均融合得 p_{final}。所有融合超參數（α、σ、T_{cls}）皆為"
                  "先驗設定，不依賴測試集統計。"),
    "regcls": (FIG / "v4_fig_regcls.png",
               "Reg-Cls 機率融合（測試樣本 idx = 316）。（a）分類頭 softmax p_{cls}；"
               "（b）由回歸預測 r 透過高斯核 σ = 0.65 映射至 7 類機率分布 p_{reg}；"
               "（c）α = 0.65 之 log 空間幾何平均；融合後 argmax 與真實類別吻合。"),
    "cm": (FIG / "v2_fig9_confusion.png",
           "測試集 7 類混淆矩陣。行為真實類別、欄為預測類別。對角線濃度反映各類正確率；"
           "整體 Acc-7 = 53.21%、Within-1 = 91.55%。離對角線誤判主要落於相鄰類，符合任務之"
           "序數性質。"),
    "perclass": (FIG / "v2_fig10_per_class_acc.png",
                 "逐類別 Acc-7。各類別於測試集之預測準確度。橫向虛線為整體 Acc-7（53.21%），"
                 "點虛線為隨機預測基線（1/7 ≈ 14.3%）。中性類別最具挑戰性，因其與相鄰類在標"
                 "注上之邊界本就模糊。"),
    "radar": (FIG / "v4_fig_radar.png",
              "整體效能雷達圖。（左）四項分類指標之長條圖；橫向虛線為 53% 之既定目標。"
              "（右）六項指標歸一化後之雷達圖；越接近外圈代表表現越好。本模型於各維度皆達高"
              "水準。"),
}


# ─── 1. Inline subscript helper ──────────────────────────────────────────────
# A *single source of truth* for converting `_{...}` markers in plain text into
# Word-native subscript runs.  This makes the rendered file behave the same as
# if a human had typed the subscript with Word's superscript/subscript button —
# and crucially, MathType picks them up as proper subscript glyphs.
SUB_RE = re.compile(r"_\{([^{}]+)\}")


def _add_text_with_subs(paragraph, text, base_font_size=None,
                         base_bold=False, base_color=None):
    """Split `text` on `_{...}` markers and emit alternating normal + subscript
    Word runs (with `w:vertAlign w:val=\"subscript\"` set on the subscript runs).

    Examples:
        _add_text_with_subs(p, "W_{a} 與 W_{v} 之投影")
        →  W [sub:a] " 與 " W [sub:v] " 之投影"

    Returns the paragraph (chainable).
    """
    pos = 0
    for m in SUB_RE.finditer(text):
        plain = text[pos:m.start()]
        if plain:
            r = paragraph.add_run(plain)
            _apply_font(r, base_font_size, base_bold, base_color)
        sub_text = m.group(1)
        r = paragraph.add_run(sub_text)
        _apply_font(r, base_font_size, base_bold, base_color)
        # mark as subscript
        rPr = r._r.get_or_add_rPr()
        vAlign = OxmlElement("w:vertAlign")
        vAlign.set(qn("w:val"), "subscript")
        rPr.append(vAlign)
        pos = m.end()
    tail = text[pos:]
    if tail:
        r = paragraph.add_run(tail)
        _apply_font(r, base_font_size, base_bold, base_color)
    return paragraph


def _apply_font(run, size=None, bold=False, color=None):
    if size: run.font.size = size
    if bold: run.font.bold = True
    if color: run.font.color.rgb = color


def add_para(doc, text, **kw):
    """Add a paragraph with `_{...}` subscript markers parsed inline."""
    p = doc.add_paragraph()
    _add_text_with_subs(p, text, **kw)
    return p


def add_heading_sub(doc, text, level=1):
    """Heading with subscript-marker support."""
    h = doc.add_heading("", level=level)
    _add_text_with_subs(h, text)
    return h


# ─── 2. Captions / TOCs ──────────────────────────────────────────────────────
def add_field(p, instr, run_text="1"):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t"); inner_t.text = run_text
    inner_r.append(inner_t); fld.append(inner_r)
    p._p.append(fld)


def add_caption(doc, prefix_text, seq_name, body_text):
    p = doc.add_paragraph()
    p.style = "Caption" if "Caption" in [s.name for s in doc.styles] else p.style
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(prefix_text); r.bold = True
    add_field(p, f' SEQ {seq_name} \\* ARABIC ')
    # space + caption body
    _add_text_with_subs(p, "  " + body_text)
    return p


def add_figure(doc, key, width_in=6.0):
    path, caption = FIGS[key]
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    add_caption(doc, "圖 3.", "Figure", caption)


def add_table_with_caption(doc, caption_body, rows):
    """Add caption ABOVE then table. `rows` is list-of-tuples (cells)."""
    add_caption(doc, "表 3.", "Table", caption_body)
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, txt in enumerate(row):
            cell = tbl.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_text_with_subs(p, txt, base_font_size=Pt(10),
                                 base_bold=(r == 0))
    return tbl


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_toc_field(doc, instr, placeholder="（請於 Word 中按 F9 或右鍵 → 更新功能變數）"):
    p = doc.add_paragraph()
    for type_ in ("begin", None, "separate", "placeholder", "end"):
        if type_ == "begin":
            fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
            r = OxmlElement("w:r"); r.append(fld); p._p.append(r)
        elif type_ is None:
            r = OxmlElement("w:r")
            ins = OxmlElement("w:instrText"); ins.text = instr
            ins.set(qn("xml:space"), "preserve"); r.append(ins); p._p.append(r)
        elif type_ == "separate":
            fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "separate")
            r = OxmlElement("w:r"); r.append(fld); p._p.append(r)
        elif type_ == "placeholder":
            r = OxmlElement("w:r")
            t = OxmlElement("w:t"); t.text = placeholder
            r.append(t); p._p.append(r)
        elif type_ == "end":
            fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "end")
            r = OxmlElement("w:r"); r.append(fld); p._p.append(r)


def configure_styles(doc):
    nor = doc.styles["Normal"]
    nor.font.name = "Times New Roman"
    nor.element.rPr.rFonts.set(qn("w:eastAsia"), "PMingLiU")
    nor.font.size = Pt(11)
    if "Caption" in [s.name for s in doc.styles]:
        cap = doc.styles["Caption"]
        cap.font.size = Pt(10); cap.font.italic = False


# ════════════════════════════════════════════════════════════════════════════
# Main
# ════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # ─── Cover ───────────────────────────────────────────────────────────────
    title = doc.add_heading("第三章   研究方法", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subt = doc.add_paragraph(
        "SACFFinalModel：以多分支單一模型 + 多工損失（softCE / EMD / SmoothL1 / R-Drop / CMC）"
        "與 3-Seed Ensemble 實現多模態情感分析"
    )
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subt.runs[0].bold = True
    subt.runs[0].font.size = Pt(13)
    add_page_break(doc)

    # ─── 圖目錄 ──────────────────────────────────────────────────────────────
    h = doc.add_heading("圖目錄", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_toc_field(doc, ' TOC \\h \\z \\c "Figure" ')
    add_page_break(doc)

    # ─── 表目錄 ──────────────────────────────────────────────────────────────
    h = doc.add_heading("表目錄", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_toc_field(doc, ' TOC \\h \\z \\c "Table" ')
    add_page_break(doc)

    # ─── 3.1  研究框架概覽 ───────────────────────────────────────────────────
    doc.add_heading("3.1   研究框架概覽", level=1)
    add_para(doc,
        "本章提出 SACFFinalModel — 一個專為多模態情感分析（Multimodal Sentiment Analysis, "
        "MSA）任務所設計之「多分支單一模型」（Multi-Branch Single Model）架構。傳統做法以"
        "多個獨立模型進行測試時集成（test-time ensemble）雖能提升泛化，但須維護多個權重檔，"
        "推斷時佔用大量 VRAM 與算力，且難以在資源受限的部署場景中落地；本研究反其道而行，"
        "將「集成多樣性」內建於單一模型之內部結構：上游編碼層（Backbone）由所有分支共享，"
        "下游則設計 4 個結構同形但參數獨立的並行分支（PEA + SACF + Projection + Heads），"
        "並對每個分支施以不同 dropout 率（0.10、0.20、0.30、0.40），確保各分支於訓練時走"
        "不同的隨機路徑、學到互補的決策邊界。最終結果經模型內部之「分支平均」（mean of "
        "branches）聚合輸出，並僅以一個 sacf_final.pt 權重檔即可完成端對端推斷。"
    )
    add_para(doc,
        "為將「多模型集成」之效益完整壓縮進此單一模型，本研究進一步引入五項關鍵技術，"
        "並完整反映在最終訓練損失 L_{total} 之組成中："
    )
    add_para(doc,
        "（1）軟序數標籤交叉熵 L_{softCE} — 利用 7 類別之序數性質，將 one-hot 替換為以高斯"
        "核軟化之目標分布，使相鄰類錯誤的懲罰小於遠距類錯誤；"
        "（2）序數地球移動距離損失 L_{EMD} — 進一步以 CDF 差異之 ℓ_{1} 距離約束預測分布之"
        "整體形狀；"
        "（3）回歸 L_{SmoothL1} — 對回歸頭輸出 reg ∈ [−3, +3] 應用 Smooth-L1 損失，避免 "
        "L2 損失於離群點之過度放大；"
        "（4）R-Drop 一致性正則 L_{R-Drop} — 對同批次樣本執行兩次 forward（dropout 取兩"
        "個不同實現 f_{1}、f_{2}），最小化兩次輸出分布之對稱 KL 散度，等效於隱式資料增強；"
        "（5）跨模態對比損失 L_{CMC} — 以 InfoNCE 形式拉近同一樣本之 t_{emb}、a_{emb}、"
        "v_{emb} 三模態嵌入、推開不同樣本之嵌入，強化共享編碼器之跨模態對齊。"
    )
    add_para(doc,
        "推斷階段採用三層方差降低：(i) 對每個樣本以 MC-Dropout 進行 TTA × 5 次 forward；"
        "(ii) 將 3 個獨立 seed 之 θ_{run1}、θ_{run2}、θ_{run3} 之 logit 算術平均；"
        "(iii) 將分類 softmax p_{cls} 與回歸高斯 PMF p_{reg} 於 log 空間以幾何平均融合為 "
        "p_{final}，再 argmax 取得最終預測。所有融合超參數（α = 0.65, σ = 0.65, T_{cls} = 1）"
        "皆為先驗設定，不依賴測試集統計，故維持嚴格零資料洩漏（zero data leakage）。"
    )
    add_para(doc,
        "在此嚴格定義下，本研究於 CMU-MOSI 測試集達成 Acc-7 = 53.21%，超越預設 53% 之研究"
        "目標；同時於 Acc-2、F1、MAE、Corr、Within-1 等指標均達領先水準。"
    )
    add_figure(doc, "arch")

    # ─── 3.2  資料集與前處理 ─────────────────────────────────────────────────
    doc.add_heading("3.2   資料集與前處理", level=1)

    doc.add_heading("3.2.1   CMU-MOSI 資料集", level=2)
    add_para(doc,
        "本研究使用 CMU-MOSI（CMU Multimodal Opinion Sentiment and Subjectivity）資料集，"
        "為多模態情感分析之標準基準。該資料集由 93 位 YouTube 評論者之獨白影片構成，共 "
        "2,199 個語句單位（utterance），每個語句經 5 位人工標注員獨立評分，平均後落於 "
        "[−3, +3] 之連續強度區間（−3 表示極度負面、+3 表示極度正面）。三模態原始特徵分別"
        "為：文字（轉錄文本）、音訊（COVAREP 萃取之 5 維低層級韻律特徵）、視覺（FACET 萃"
        "取之 20 維臉部動作單元）。"
    )
    add_para(doc,
        "為最大化訓練資料量，本研究將原始 Train（1,284 筆）與 Valid（229 筆）合併為單一 "
        "Train（n = 1,513），不於 Valid 上進行超參數選擇 — 所有超參數採先驗設定，故無需"
        "驗證集；Test 維持 686 筆，唯一一次於模型最終訓練完成後進行評估，嚴格遵循零資料"
        "洩漏原則。表 3.1 與圖 3.2 顯示此資料劃分與類別分布。"
    )

    add_table_with_caption(doc,
        "CMU-MOSI 資料集劃分與用途", [
            ("資料劃分", "樣本數", "用途"),
            ("Train",  "1,513", "模型訓練（原始 Train 1,284 + Valid 229 合併）"),
            ("Test",   "686",   "最終 Acc-7／Acc-2／F1／MAE／Corr 評估，僅評估一次"),
        ])

    add_para(doc,
        "圖 3.2 顯示 Train 與 Test 之七類情感分布。Train 集偏向中性與輕微正面情感（類 +2 "
        "佔最大比例），而 Test 集顯著偏向負面端（−3 類別佔 6.7%，相較 Train 之 2.2% 高出 "
        "+4.5%）。此分布偏移為 MSA 任務之固有挑戰，亦為本研究設計 L_{softCE} 軟標籤、"
        "L_{EMD} 序數損失、3-Seed 集成等多重正則之主要動機 — 透過增加邊界類別之梯度信號"
        "與多視角預測，緩解測試端偏移之影響。"
    )
    add_figure(doc, "dist")

    # 3.2.2
    doc.add_heading("3.2.2   標籤定義", level=2)
    add_para(doc, "本研究以三種預測目標實現多工聯合學習，分別對應三種頭部輸出，並於同一模型內共同訓練：")
    add_display_equation(doc,
        sub("y", "7"), " = clip", paren(["round(s), -3, 3"]),
        " + 3 ∈ {0, 1, 2, 3, 4, 5, 6}")
    add_para(doc, "其中 s 為原始連續評分；y_{7} 為主要評估指標 Acc-7 之計算依據。")
    add_display_equation(doc,
        sub("y", "2"), " = ", paren(["s ≥ 0"]),
        " ∈ {0, 1}")
    add_para(doc, "y_{2} 為二分類標籤，作為輔助訓練訊號以強化情感極性區分。")
    add_display_equation(doc,
        sub("y", "reg"), " = s ∈ ",
        paren(["-3, +3"], "[", "]"))
    add_para(doc, "y_{reg} 為回歸標籤，於推斷時亦用於 Reg-Cls 機率融合。")

    doc.add_heading("3.2.3   輸入前處理", level=2)
    add_para(doc,
        "文字：每個語句加入任務導向提示前綴「Predict the sentiment intensity (−3 to 3, "
        "negative to positive) of the following text:」，再以 DeBERTa-v3-large tokenizer "
        "編碼，最大長度限制為 80 token。"
    )
    add_para(doc,
        "音訊／視覺：將 NaN 與 ±∞ 異常值替換為 0；對每筆樣本之有效時間範圍內進行 "
        "ℓ_{2}-normalize（沿特徵維），以消除錄製增益／鏡頭距離造成之尺度差異。每筆樣本"
        "同時提供 audio_mask 與 vision_mask 標示有效幀範圍，以於後續 BiLSTM 編碼時進行 "
        "packed sequence 處理。"
    )

    # ─── 3.3  模型架構 ───────────────────────────────────────────────────────
    doc.add_heading("3.3   模型架構", level=1)

    doc.add_heading("3.3.1   共享編碼層", level=2)
    add_para(doc,
        "共享編碼層處理三模態原始輸入，輸出供下游所有 4 個分支共用。選擇「共享」而非「每"
        "分支獨立」是因該層參數量大、計算昂貴；若每分支建立獨立副本將使模型膨脹至 6 GB+ "
        "且訓練時間倍增。共享編碼器並不顯著限制分支多樣性 — 多樣性主要源自下游的 PEA / "
        "SACF / Projection 之獨立參數與不同 dropout 率。"
    )
    add_para(doc,
        "DeBERTa-v3-large：以 microsoft/deberta-v3-large 作為文字骨幹，24 層 Transformer、"
        "隱藏維度 d_{lang} = 1,024、總參數約 400M，輸出 last_hidden_state ∈ ℝ^(B × L × 1024) "
        "與 [CLS] 表徵 x_{cls} ∈ ℝ^(B × 1024)。本研究訓練前 1/3 epoch（即 E1–E20）凍結 "
        "DeBERTa 下層 6 層，僅訓練上層 18 層與下游頭；之後解凍下層並以 ¼ 學習率續訓，避免"
        "於早期破壞預訓練之語義表徵。"
    )
    add_para(doc,
        "音訊／視覺編碼器：分別為 2 層雙向 LSTM，最後時間步雙向拼接後線性投影至 "
        "d_{modal} = 128，輸出 a_{emb} ∈ ℝ^(B × 128) 與 v_{emb} ∈ ℝ^(B × 128)。"
        "兩編碼器於 forward 前先以 pack_padded_sequence 包裝，僅對有效幀計算梯度。"
    )

    # 3.3.2  4 branches
    doc.add_heading("3.3.2   4 個並行分支", level=2)
    add_para(doc,
        "本架構之核心創新在於將「多模型集成的多樣性」內建於模型架構之中。4 個分支共享上"
        "游編碼結果（H、a_{emb}、v_{emb}），但各自獨立進行下游融合與預測。為確保分支間之"
        "充分多樣性，本研究採用三項機制（詳見圖 3.5）："
    )
    add_para(doc,
        "（1）不同 Dropout 率：Branch 1 = 0.10、Branch 2 = 0.20、Branch 3 = 0.30、"
        "Branch 4 = 0.40；（2）獨立 PEA / SACF / Proj 參數；（3）cls7 頭部初始化額外加入 "
        "0.005·(i+1)·N(0, 1) 之高斯擾動。"
    )

    # 3.3.2.1 PEA
    doc.add_heading("3.3.2.1   極性增強注意力（PEA）", level=3)
    add_para(doc,
        "PEA 為每個 DeBERTa 詞元學習情感顯著性閘值 g_{i}，再以閘值對 last_hidden_state 做"
        "加權平均池化，產出 [CLS] 之精煉版作為下游語言查詢向量。形式上："
    )
    add_display_equation(doc,
        sub("g", "i"), " = σ",
        paren([sub("W", "2"), " · tanh",
               paren([sub("W", "1"), " · ", sub("h", "i")])]),
        " ∈ ", paren(["0, 1"], "[", "]"))
    add_para(doc,
        "其中 W_{1} ∈ ℝ^(d/4 × d)、W_{2} ∈ ℝ^(1 × d/4)，σ 為 sigmoid。閘值 g_{i} 越接近 1 "
        "表示該詞元越具情感顯著性。加權池化向量 x_{l} 為："
    )
    add_display_equation(doc,
        sub("x", "l"), " = ",
        frac(num=[
            nary("∑", lower="i", upper="L",
                 body=[paren(["0.75 · ", sub("h", "i"), " + 0.25 · ",
                              sub("h", "i"), " · ", sub("g", "i")]),
                       " · ", sub("m", "i")])
        ], den=[
            nary("∑", lower="i", upper="L", body=[sub("m", "i")])
        ]))
    add_para(doc,
        "其中 m_{i} 為 attention_mask（1 = 有效詞元、0 = padding）。常數 0.75 / 0.25 為保守"
        "設定 — 即使 g 全為 0 仍保留 75% 之原始 hidden 訊號。x_{l} ∈ ℝ^(B × d_{lang}) 將與"
        "音訊／視覺進入 SACF 模組進行跨模態融合。圖 3.4 為 PEA 模組之詳細示意。"
    )
    add_figure(doc, "pea")

    # 3.3.2.2 SACF
    doc.add_heading("3.3.2.2   情感感知跨模態注意力（SACF）", level=3)
    add_para(doc,
        "SACF 是本研究於跨模態融合的核心設計，將語言、音訊、視覺三模態結合為融合向量 f。"
        "傳統做法直接以 [CLS] 為查詢向量，未能聚焦於情感顯著詞元；SACF 改以「情感感知查詢」"
        "取代之，分四步驟完成跨模態融合（見圖 3.5）。"
    )
    add_figure(doc, "sacf")

    add_para(doc, "步驟 1（Top-K 詞元選擇）：依 PEA 閘值 g 取前 K = 5 個最高分詞元，提取：")
    add_display_equation(doc,
        sub("H", "topk"), " = H[I]   ∈   ℝ^(B × K × ",
        sub("d", "lang"), ")")
    add_para(doc, "步驟 2（情感查詢構建）：對 H_{topk} 計算注意力權重後加權求和：")
    add_display_equation(doc,
        "w = softmax", paren([sub("W", "tok"), " · ", sub("H", "topk")]),
        ",   ", sub("q", "sa"), " = ",
        nary("∑", lower=["i=1"], upper="K",
             body=[sub("w", "i"), " · ", sub("H", "topk[i]")]))
    add_para(doc, "步驟 3（跨模態鍵值對齊）：將 a_{emb}、v_{emb} 投影至語言空間並組成 KV：")
    add_display_equation(doc,
        "KV = stack", paren([sub("W", "a"), " · ", sub("a", "emb"), ", ",
                              sub("W", "v"), " · ", sub("v", "emb")]),
        "   ∈   ℝ^(B × 2 × ", sub("d", "lang"), ")")
    add_display_equation(doc,
        sup("x", "*"), " = softmax",
        paren([frac(num=[sub("q", "sa"), " · ", sup("KV", "T")],
                    den=rad(sub("d", "lang")))]),
        " · KV")
    add_para(doc, "步驟 4（門控殘差融合）：將跨模態增量與 [CLS] 表徵 x_{cls} 融合：")
    add_display_equation(doc,
        "x = FFN", paren([sub("x", "cls"), " + ", sup("x", "*")]),
        ",   gw = σ",
        paren([sub("W", "g"), " · concat", paren([sub("x", "cls"), ", x"])]))
    add_display_equation(doc,
        "f = LayerNorm",
        paren([sub("x", "cls"), " + Dropout", paren(["x · gw"])]))
    add_para(doc,
        "FFN 為 Linear(d → d/2) → ReLU → Dropout → Linear(d/2 → d) 之兩層全連接；gw 為 "
        "sigmoid 閘控（B × 1）。Layer-Norm 之殘差設計確保訓練穩定性，融合向量 "
        "f ∈ ℝ^(B × d_{lang}) 為下游分類／回歸頭之輸入。"
    )

    # 3.3.2.3 Heads
    doc.add_heading("3.3.2.3   共享投影層與多工預測頭", level=3)
    add_para(doc,
        "融合表徵 f ∈ ℝ^(B × d_{lang}) 通過該分支獨立的投影模組壓縮為 e_{i} ∈ ℝ^(B × 512)："
        "Linear(1024 → 512) → LayerNorm → GELU → Dropout（每分支獨立 dropout 率）。"
        "e_{i} 隨後送入三個任務頭，各自輸出 l7_{i}、l2_{i}、reg_{i}。"
    )
    add_para(doc,
        "（1）cls7 head：Linear(512 → 7) 之單層線性，輸出 l7_{i} ∈ ℝ^(B × 7)；"
        "（2）cls2 head：Linear(512 → 2)，輸出 l2_{i} ∈ ℝ^(B × 2)；"
        "（3）reg head：Linear(512 → 256) → GELU → Linear(256 → 1) → Tanh，輸出 ∈ [−1, +1] "
        "再乘以 3，得 reg_{i} ∈ [−3, +3]。"
    )

    # 3.3.3 Internal ensemble
    doc.add_heading("3.3.3   內部集成（Internal Ensemble）與分支貢獻", level=2)
    add_para(doc,
        "推斷時，4 個分支之輸出於模型內部進行算術平均："
    )
    add_display_equation(doc,
        sub("l7", "mean"), " = ",
        frac(num=[
            nary("∑", lower=["i=1"], upper="4",
                 body=[sub("l7", "i")])
        ], den="4"))
    add_para(doc,
        "l2_{mean} 與 reg_{mean} 之聚合方式相同。圖 3.6 呈現各分支單獨之 Acc-7（介於 "
        "51.9%–52.6%）與分支平均後之表現（53.21%）。此內部 ensemble 不需於推斷時執行多次 "
        "forward — 4 個分支於同一個 forward 中同時計算，計算開銷僅為單分支模型之 ~1.4 倍。"
    )
    add_figure(doc, "branches")

    # ─── 3.4  損失函數設計 ───────────────────────────────────────────────────
    doc.add_heading("3.4   多工損失函數設計", level=1)
    add_para(doc,
        "本研究放棄 v2 階段使用之外部 12 教師蒸餾（DKD / DIST），改採內部驅動之多工損失"
        "組合，避免對外部教師檔之依賴並降低訓練流程複雜度。整體損失分為三層（見圖 3.7）："
        "Layer 1 為單分支內之任務組合損失 L_{branch_i}；Layer 2 為跨分支之聚合與多樣性、"
        "對比約束；Layer 3 為一致性正則 L_{R-Drop} 與最終 L_{total}。"
    )
    add_figure(doc, "loss")

    # 3.4.1 softCE
    add_heading_sub(doc, "3.4.1   軟序數標籤交叉熵 L_{softCE}", level=2)
    add_para(doc,
        "傳統 cross-entropy 對 7 類別一視同仁：將真實類預測為 −3 與將其預測為 +3 之損失"
        "相同。然本任務具明顯之序數結構：相鄰類差距小於遠距類差距。本研究採用 SORD 高斯"
        "軟標籤："
    )
    add_display_equation(doc,
        "soft_target", paren(["i, k"], "[", "]"), " ∝ exp",
        paren(["−",
               frac(sup(paren(["k − ", sub("y", "i")]), "2"), sup("σ", "2"))]))
    add_display_equation(doc,
        sub("L", "softCE"), " = − ",
        nary("∑", lower="k", upper=None,
             body=["soft_target", paren(["k"], "[", "]"),
                   " · log_softmax", paren(["l7"]), paren(["k"], "[", "]")]))
    add_para(doc, "本研究設 σ = 1.0，使相鄰類得 e^{-1} ≈ 0.37 倍機率質量，距離 2 之類得 0.018 倍。")

    # 3.4.2 EMD
    add_heading_sub(doc, "3.4.2   序數地球移動距離損失 L_{EMD}", level=2)
    add_para(doc,
        "L_{EMD} 進一步以累積分布函數（CDF）差異約束預測分布之整體形狀，將「預測誤差之距離」"
        "明確納入損失："
    )
    add_display_equation(doc,
        sub("L", "EMD"), " = ",
        nary("∑", lower=["k=1"], upper="6",
             body=["|",
                   sub("CDF", "pred"), paren(["k"], "[", "]"),
                   " − ",
                   sub("CDF", "true"), paren(["k"], "[", "]"),
                   "|"]))

    # 3.4.3 SmoothL1
    add_heading_sub(doc, "3.4.3   回歸損失 L_{SmoothL1}", level=2)
    add_para(doc,
        "回歸頭預測 reg_{i} ∈ [−3, +3] 與 y_{reg} 之差距以 Smooth-L1 損失定義；相較於 L2 "
        "損失，Smooth-L1 在絕對誤差 ≤ 1 時為二次函數、> 1 時為線性，可避免離群點過度放大"
        "梯度："
    )
    add_display_equation(doc,
        sub("L", "SmoothL1"), " = ",
        paren([
            "1{|reg − ", sub("y", "reg"),
            "| ≤ 1} · 0.5 · ", sup(paren(["reg − ", sub("y", "reg")]), "2"),
            "  +  1{|reg − ", sub("y", "reg"),
            "| > 1} · ", paren(["|reg − ", sub("y", "reg"), "| − 0.5"])
        ]))

    # 3.4.4 R-Drop
    add_heading_sub(doc, "3.4.4   R-Drop 一致性正則 L_{R-Drop}", level=2)
    add_para(doc,
        "對同批次樣本執行兩次隨機 forward，dropout 取兩個不同實現，分別產生融合特徵 "
        "f_{1} 與 f_{2}。最小化兩次輸出分布之對稱 KL 散度："
    )
    add_display_equation(doc,
        sub("L", "R-Drop"), " = ",
        frac(num="1", den="2"),
        paren([
            "KL", paren([sub("p", paren([sub("f", "1")])),
                          " ‖ ",
                          sub("p", paren([sub("f", "2")]))]),
            " + ",
            "KL", paren([sub("p", paren([sub("f", "2")])),
                          " ‖ ",
                          sub("p", paren([sub("f", "1")]))]),
        ]))
    add_para(doc,
        "L_{R-Drop} 等效於對模型進行隱式的資料增強與正則化，強化模型在 dropout 條件下之"
        "預測一致性。"
    )

    # 3.4.5 CMC
    add_heading_sub(doc, "3.4.5   跨模態對比損失 L_{CMC}", level=2)
    add_para(doc,
        "為強化共享編碼層之跨模態對齊，本研究引入 InfoNCE 形式之跨模態對比損失："
    )
    add_display_equation(doc,
        sub("L", "CMC"), " = − ",
        nary("∑", lower=["(u,v)"], upper=None,
             body=["log ",
                   frac(num=["exp",
                             paren(["sim(", sub("u", "i"), ", ", sub("v", "i"),
                                    ") / τ"])],
                        den=[nary("∑", lower="j", upper="B",
                                  body=["exp",
                                        paren(["sim(", sub("u", "i"), ", ",
                                                sub("v", "j"), ") / τ"])])])
        ]))
    add_para(doc,
        "其中 (u, v) 為 {t_{emb}, a_{emb}, v_{emb}} 中任意兩兩模態配對；sim 為 cosine 相似度；"
        "τ 為溫度超參數（本研究設 τ = 0.07）。該損失同時拉近同樣本三模態之嵌入、推開不同"
        "樣本之嵌入。"
    )

    # 3.4.6 Aggregated total loss
    add_heading_sub(doc, "3.4.6   分支聚合與整體損失 L_{total}", level=2)
    add_para(doc, "單一分支 i 之任務組合損失：")
    add_display_equation(doc,
        sub("L", "branch_i"), " = ",
        paren(["1 − ", sub("w", "EMD")]),
        " · ", sub("L", "softCE"),
        " + ", sub("w", "EMD"), " · ", sub("L", "EMD"),
        " + 0.3 · ", sub("L", "cls2"),
        " + 0.4 · ", sub("L", "SmoothL1"))
    add_para(doc, "本研究設 w_{EMD} = 0.25。跨分支聚合損失：")
    add_display_equation(doc,
        sub("L", "mean"), " = ", sub("L", "branch"),
        "  on  ", sub("l7", "mean"))
    add_display_equation(doc,
        sub("L", "per_branch"), " = ",
        frac(num="1", den="4"),
        nary("∑", lower=["i=1"], upper="4",
             body=[sub("L", "branch_i")]))
    add_para(doc,
        "L_{diversity} 為分支特徵 e_{i} 兩兩之間餘弦相似度之懲罰，輕微鼓勵分支特徵不要"
        "過於相似（係數小，僅作為次要正則項）："
    )
    add_display_equation(doc,
        sub("L", "diversity"), " = ",
        frac(num="1", den=paren(["B(B−1)/2"])),
        nary("∑", lower="i<j", upper=None,
             body=["cos", paren([sub("e", "i"), ", ", sub("e", "j")])]))
    add_para(doc, "最終整體損失：")
    add_display_equation(doc,
        sub("L", "total"), " = ",
        sub("w", "mean"), " · ", sub("L", "mean"),
        " + ", sub("w", "per"), " · ", sub("L", "per_branch"),
        " + ", sub("w", "div"), " · ", sub("L", "diversity"),
        " + ", sub("w", "CMC"), " · ", sub("L", "CMC"),
        " + 0.05 · ", sub("L", "R-Drop"))
    add_para(doc,
        "其中 w_{mean} = w_{per} = 0.5、w_{div} = 0.02、w_{CMC} = 0.1。當批次中啟動 R-Drop "
        "時，由於兩次 forward 之 dropout mask 不同，L_{mean} 與 L_{per_branch} 採兩次 forward "
        "之平均；L_{R-Drop} 直接以兩次 forward 之 logit 計算對稱 KL。"
    )

    # ─── 3.5  訓練策略 ───────────────────────────────────────────────────────
    doc.add_heading("3.5   訓練策略", level=1)
    add_para(doc,
        "本研究採用單一 60-epoch 之全模型訓練流程，內部包含 3 個階段：Phase 1 凍結 DeBERTa "
        "下層 6 層（E1–20）、Phase 2 全模型微調（E20–42）、Phase 3 SWA 視窗（E42–60，每 2 "
        "epoch 取一個快照，共 10 個）。訓練全程維護 EMA 影子模型 θ_{shadow}。3 個獨立 seed "
        "（42、123、2024）各執行一次完整流程，產出 θ_{run1}、θ_{run2}、θ_{run3}，最終於"
        "參數空間平均得 θ_{final} = sacf_final.pt。"
    )
    add_figure(doc, "timeline")

    # Hyperparameter table
    add_table_with_caption(doc,
        "iter1 主要訓練超參數", [
            ("超參數", "值", "說明"),
            ("lang_lr / head_lr", "4e-6 / 8e-5", "DeBERTa 與下游頭採差分學習率"),
            ("weight_decay", "0.01", "AdamW"),
            ("batch_size", "8", "AMP fp16 訓練"),
            ("num_epochs", "60", "單一 run 總 epoch"),
            ("warmup_ratio", "0.06", "Cosine schedule"),
            ("freeze_layers", "0–5 (E1–20)", "解凍於 E20"),
            ("dropout (per-branch)", "[0.10, 0.20, 0.30, 0.40]", "強化分支差異化"),
            ("label_smoothing", "0.05", "cls 全部"),
            ("w_{EMD} / σ_{softCE}", "0.25 / 1.0", "L_{EMD} 與 L_{softCE} 混合"),
            ("w_{mean} / w_{per} / w_{div} / w_{CMC}", "0.5 / 0.5 / 0.02 / 0.1", "整體損失加權"),
            ("R-Drop weight", "0.05", "L_{R-Drop} 權重"),
            ("CMC τ", "0.07", "InfoNCE 溫度"),
            ("ema_decay", "0.9995", "影子模型平滑 θ_{shadow}"),
            ("SWA window", "E42–60, step=2", "10 個快照"),
            ("seeds (3-run)", "42, 123, 2024", "3-seed ensemble"),
        ])

    add_para(doc,
        "於 SWA 取快照時，本研究先以 ema.apply_shadow() 將影子權重替換進模型，再保存 "
        "state_dict，避免某一 epoch 之峰值或谷值被誤取。SWA 平均對應於損失曲面上「平坦最"
        "小值」之尋找。"
    )
    add_para(doc,
        "梯度反向傳播時採用 AMP fp16 加速，每步以 grad_norm = 1.0 進行梯度裁剪以避免 NaN。"
        "若 batch 之損失出現 NaN 或 Inf 則跳過該 batch 之 optimizer.step()，但仍進行 "
        "scheduler.step()，保持學習率排程之時序一致。"
    )

    # ─── 3.6  零洩漏推斷流程 ─────────────────────────────────────────────────
    doc.add_heading("3.6   零洩漏推斷流程", level=1)
    add_para(doc,
        "本研究於推斷階段採用三層方差降低設計：TTA × 5 → 3-Seed Ensemble → Reg-Cls 機率"
        "融合。所有融合超參數均為先驗設定，不依賴測試集統計（圖 3.9）。"
    )
    add_figure(doc, "inference")

    doc.add_heading("3.6.1   測試時間增強 TTA × 5", level=2)
    add_para(doc,
        "對同一測試樣本，保留模型之 Dropout 隨機性，執行 T_{TTA} = 5 次獨立 forward，將 5 "
        "次 logit 取算術平均，等效於 MC-Dropout 之 sample-then-average，於不增加任何標籤"
        "資訊之前提下降低預測方差。"
    )

    doc.add_heading("3.6.2   3-Seed Ensemble", level=2)
    add_para(doc,
        "將 3 個獨立 seed（42、123、2024）訓練之 θ_{run1}、θ_{run2}、θ_{run3} 之 logit "
        "於 batch 維算術平均；不同 seed 引導模型探索損失曲面之不同收斂路徑，產生互補之"
        "決策邊界，集成後預測更穩健。"
    )

    doc.add_heading("3.6.3   Reg-Cls 機率融合", level=2)
    add_para(doc, "設 cls7 head 輸出 l7_{mean} ∈ ℝ⁷、reg head 輸出 r ∈ [−3, +3]，則：")
    add_display_equation(doc,
        sub("p", "cls"), " = softmax",
        paren([sub("l7", "mean"), " / ", sub("T", "cls")]))
    add_display_equation(doc,
        sub("p", "reg"), paren(["k"], "[", "]"),
        " ∝ exp",
        paren([
            "−", frac(num=sup(paren(["k − ", paren(["r + 3"])]), "2"),
                       den=["2", sup("σ", "2")])
        ]),
        ",  k ∈ {0, …, 6}")
    add_display_equation(doc,
        "log ", sub("p", "final"), " = α · log ", sub("p", "cls"),
        " + ", paren(["1 − α"]), " · log ", sub("p", "reg"))
    add_display_equation(doc, "ŷ = argmax ", sub("p", "final"))
    add_para(doc,
        "本研究 a priori 設定 α = 0.65、σ = 0.65、T_{cls} = 1.0。融合機率將分類與序數兩者"
        "之資訊互補，最終 Acc-7 較單一分類頭 argmax 提升 +0.44%。"
    )
    add_figure(doc, "regcls")

    # ─── 3.7  實驗結果 ───────────────────────────────────────────────────────
    doc.add_heading("3.7   實驗結果", level=1)
    add_para(doc,
        "本節報告本架構於 CMU-MOSI 測試集（n = 686）上之最終評估結果。所有指標皆以單一 "
        "sacf_final.pt 權重檔、無任何測試端後處理調參計算。表 3.3 列出主要指標。"
    )
    add_table_with_caption(doc,
        "SACFFinalModel 於 CMU-MOSI 測試集主要指標", [
            ("評估指標", "數值", "說明"),
            ("Acc-7（融合最終）", "53.21 %", "7 分類準確度，主指標；Reg-Cls 機率融合"),
            ("Acc-7（raw cls）", "52.77 %", "僅分類頭 argmax，未融合"),
            ("Acc-2", "86.73 %", "二分類（s ≥ 0 vs s < 0）準確度"),
            ("F1（weighted）", "86.72 %", "二分類加權 F1"),
            ("MAE", "0.5868", "回歸絕對誤差"),
            ("Corr", "0.8683", "回歸與真實之皮爾森相關"),
            ("Within-1", "91.55 %", "預測距離真實類別 ≤ 1 之比例（序數寬容度）"),
        ])

    doc.add_heading("3.7.1   混淆矩陣", level=2)
    add_para(doc,
        "圖 3.11 為 7 類混淆矩陣。對角線為各類之正確分類比例。可看出對於負面強情感（−3、"
        "−2）與正面強情感（+2、+3）之預測表現相對較佳，中性與輕度情感（−1、0、+1）容易"
        "與相鄰類別混淆 — 因 5 位人工標注員對於 0 與 +1、+1 與 +2 之邊界判讀本就存在差異，"
        "因此模型於此區域之誤判主要落於相鄰類，亦反映於 Within-1 可達 91.55% 之高水準。"
    )
    add_figure(doc, "cm")

    doc.add_heading("3.7.2   逐類別準確度", level=2)
    add_para(doc,
        "圖 3.12 為各類別之預測準確度與支持度。類 −2 之 Acc 達 59.6%、類 +2 達 59.0%，遠"
        "高於整體 Acc-7。最低為類 −3（Acc = 34.8%），主要是因該類在測試集中樣本最少（n = 46），"
        "且情感強度與類 −2 之邊界主觀模糊。"
    )
    add_figure(doc, "perclass")

    doc.add_heading("3.7.3   整體效能雷達圖", level=2)
    add_para(doc,
        "圖 3.13 以條形與雷達兩種視覺化呈現本模型於六項主要指標上之表現。為使各指標可"
        "疊加比較，回歸指標經以下歸一化：Corr × 100、(1 − MAE / 3) × 100。本模型於分類"
        "（Acc-2、F1、Within-1）、序數（Within-1、MAE-norm）、相關（Corr）三大維度皆達"
        "高水準，分布均勻無偏科現象。"
    )
    add_figure(doc, "radar")

    # ─── 3.8  小結 ───────────────────────────────────────────────────────────
    doc.add_heading("3.8   小結", level=1)
    add_para(doc,
        "本章詳細描述 SACFFinalModel 之架構設計與訓練策略：以 4 並行分支搭配獨立 dropout "
        "率與初始化擾動實現「內部 ensemble」；以 L_{softCE} + L_{EMD} + L_{SmoothL1} + "
        "L_{R-Drop} + L_{CMC} 多工損失組合提供充足之訓練訊號；以 EMA + SWA 平滑訓練軌跡；"
        "以 3-Seed Ensemble + TTA × 5 + Reg-Cls 機率融合三層方差降低於推斷時進一步逼近"
        "真實後驗。在嚴格零洩漏條件下，本架構於 CMU-MOSI 達成 Acc-7 = 53.21%、Within-1 = "
        "91.55%、Acc-2 = 86.73%，超越預設 53% 之研究目標。"
    )
    add_para(doc,
        "本架構之主要貢獻可歸納為四點：（1）將多模型集成之多樣性內建於單一模型 — 透過 4 "
        "並行分支於單一 forward 內完成內部 ensemble；（2）以 L_{softCE} + L_{EMD} 明確利用"
        "情感分類之序數結構；（3）以 L_{R-Drop} 與 L_{CMC} 強化模型於 dropout 條件下之"
        "一致性與跨模態對齊；（4）零洩漏推斷流程於不增加標籤資訊之前提下穩定提升 Acc-7。"
    )
    add_para(doc,
        "未來工作可探討：（1）將分支數可調整化（如 8 或 16 分支搭配相應之 dropout 排程），"
        "探索分支數對精度與計算成本之權衡；（2）將 L_{CMC} 擴展至多級別之對比（句子對比 + "
        "幀對比）；（3）將 Reg-Cls 融合擴展至多任務後驗校準框架。"
    )

    doc.save(str(OUT_PATH))
    print(f"\n  ✓ 已儲存：{OUT_PATH}")
    print(f"  ✓ 大小：{os.path.getsize(OUT_PATH)/1024:.1f} KB")


if __name__ == "__main__":
    print("=" * 70)
    print("Building SACF Methodology Chapter 3 v4 ...")
    print("=" * 70)
    build()
