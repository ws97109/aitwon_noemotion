"""
修正 論文＿李昇峰_v6.docx 圖號殘留問題：
- Para 64 TOF 有兩個 w:t：第一個已為圖九，第二個殘留圖十一 → 清除第二個
- Para 65 TOF 仍為圖十二 → 改為圖十
- Para 260 Caption 人物介紹被誤改為圖八 → 改回圖十
"""

from docx import Document

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'


def set_wt_text(para_elem, new_text):
    """將段落第一個 w:t 設為 new_text，其餘清空。"""
    t_elems = para_elem.findall('.//{%s}t' % W)
    if not t_elems:
        return False
    t_elems[0].text = new_text
    if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
        t_elems[0].set(XML_SPACE, 'preserve')
    for t in t_elems[1:]:
        t.text = ''
    return True


def main():
    doc = Document('論文＿李昇峰_v6.docx')
    paras = doc.paragraphs

    # ── Fix 1: Para 64 TOF 圖九（清除殘留的第二個 w:t） ──────────
    p64 = paras[64]
    t_elems = p64._p.findall('.//{%s}t' % W)
    print(f'Para 64 before: {[t.text for t in t_elems]}')
    if len(t_elems) >= 2:
        for t in t_elems[1:]:
            t.text = ''
    print(f'Para 64 after:  {p64.text[:60]}')

    # ── Fix 2: Para 65 TOF 圖十二 → 圖十 ─────────────────────────
    p65 = paras[65]
    print(f'\nPara 65 before: {p65.text[:60]}')
    set_wt_text(p65._p, '圖十\u3000\u3000人物介紹')
    print(f'Para 65 after:  {p65.text[:60]}')

    # ── Fix 3: Para 260 Caption 圖八→圖十（人物介紹） ────────────
    p260 = paras[260]
    print(f'\nPara 260 before: {p260.text[:60]}')
    set_wt_text(p260._p, '圖十 人物介紹')
    print(f'Para 260 after:  {p260.text[:60]}')

    # ── 驗證 ────────────────────────────────────────────────────
    print('\n=== 圖目錄驗證 ===')
    for p in doc.paragraphs:
        if p.style.name == 'table of figures':
            print(f'  {p.text[:60]}')

    print('\n=== Caption 驗證 ===')
    for p in doc.paragraphs:
        if p.style.name == 'Caption' and '圖' in p.text:
            print(f'  {p.text[:60]}')

    doc.save('論文＿李昇峰_v6.docx')
    print('\n修正完成，已覆蓋 論文＿李昇峰_v6.docx')


if __name__ == '__main__':
    main()
